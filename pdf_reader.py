# SPDX-License-Identifier: AGPL-3.0-or-later
"""
BoltPDF - Super Lightweight Multi-Core PDF Reader
Uses pypdfium2 (Google's PDFium / Chromium PDF engine) for maximum rendering
speed and PyQt6 for a minimal UI. Pages are rendered in a background thread
and fed to the UI as they complete.

OCR feature: Uses Windows built-in OCR (Windows.Media.Ocr) via PowerShell —
zero external installs needed. Overlays word-level selectable text. Click and
drag to select a contiguous range of words in reading order (like a normal
text editor). Hold Shift to extend the selection.

Tabs: Drag-and-drop multiple PDFs or open via file dialog — each opens in a
new navigable tab below the toolbar.

============================================================================
Copyright (C) 2026  BoltPDF

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU Affero General Public License as published
by the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU Affero General Public License for more details.

You should have received a copy of the GNU Affero General Public License
along with this program (see the LICENSE file).  If not, see
<https://www.gnu.org/licenses/>.

BoltPDF links the AGPL-licensed PyMuPDF and the GPL-3.0-licensed PyQt6;
the combined work is therefore distributed under the AGPL-3.0.  The
complete corresponding source code is available at
<https://github.com/HBDPN/BoltPDF>.  Third-party component licenses are
listed in THIRD_PARTY_LICENSES.txt.
============================================================================
"""

import sys
import os
import json
import re
import subprocess
import tempfile
import time
import copy
import hashlib

__version__ = "1.0.3"
_UPDATE_URL = "https://raw.githubusercontent.com/HBDPN/BoltPDF/main/version.json"
import pypdfium2 as pdfium
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QGraphicsView, QGraphicsScene,
    QGraphicsPixmapItem, QGraphicsRectItem, QFileDialog, QToolBar,
    QWidget, QLabel, QSizePolicy, QMessageBox, QProgressDialog,
    QTabWidget, QVBoxLayout, QDialog, QRadioButton, QLineEdit,
    QDialogButtonBox, QGroupBox, QHBoxLayout, QButtonGroup,
    QPushButton, QInputDialog, QGraphicsProxyWidget,
    QGraphicsSimpleTextItem, QGraphicsTextItem, QTextEdit,
    QFontComboBox, QSpinBox, QToolButton, QColorDialog, QFrame,
    QSplitter, QScrollArea, QGraphicsLineItem, QGraphicsEllipseItem,
)
from PyQt6.QtGui import (
    QPixmap, QImage, QAction, QKeySequence, QWheelEvent, QPainter,
    QPen, QColor, QBrush, QFont, QPolygonF, QIcon, QTextCharFormat,
    QTextCursor, QTransform, QCursor, QPainterPath,
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QSize, QRectF, QTimer, QPointF, QUrl


# ---------------------------------------------------------------------------
# Custom cursors for edit tools — drawn programmatically so there are no
# external image dependencies.
# ---------------------------------------------------------------------------
_edit_cursors: dict[str, QCursor] = {}


def _get_edit_cursor(tool: str) -> QCursor:
    """Return a cached QCursor for the given edit *tool* name."""
    if tool in _edit_cursors:
        return _edit_cursors[tool]
    sz = 32
    pm = QPixmap(sz, sz)
    pm.fill(QColor(0, 0, 0, 0))  # transparent
    p = QPainter(pm)
    p.setRenderHint(QPainter.RenderHint.Antialiasing)

    if tool == "stamp":
        # Stamp icon: handle + rectangular base
        p.setPen(QPen(QColor(180, 40, 40), 2))
        p.setBrush(QBrush(QColor(200, 60, 60, 200)))
        # Handle (vertical bar)
        p.drawRect(12, 2, 8, 10)
        # Stamp head (wider rectangle)
        p.drawRect(6, 12, 20, 8)
        # Base line
        p.setPen(QPen(QColor(180, 40, 40), 3))
        p.drawLine(4, 24, 28, 24)
        hot_x, hot_y = 16, 24

    elif tool == "note":
        # Sticky note icon: yellow square with folded corner and lines
        p.setPen(QPen(QColor(180, 160, 40), 1.5))
        p.setBrush(QBrush(QColor(255, 240, 120, 220)))
        p.drawRect(4, 4, 22, 22)
        # Folded corner
        fold = QPainterPath()
        fold.moveTo(20, 4)
        fold.lineTo(26, 10)
        fold.lineTo(20, 10)
        fold.closeSubpath()
        p.setBrush(QBrush(QColor(220, 200, 60, 200)))
        p.drawPath(fold)
        # Text lines
        p.setPen(QPen(QColor(140, 120, 20), 1.2))
        p.drawLine(8, 13, 20, 13)
        p.drawLine(8, 17, 18, 17)
        p.drawLine(8, 21, 16, 21)
        hot_x, hot_y = 4, 4

    elif tool == "redact":
        # Whiteout: white rectangle with dashed border
        p.setPen(QPen(QColor(160, 160, 160), 2, Qt.PenStyle.DashLine))
        p.setBrush(QBrush(QColor(255, 255, 255, 200)))
        p.drawRect(4, 8, 24, 16)
        # X through it
        p.setPen(QPen(QColor(200, 60, 60), 2))
        p.drawLine(8, 12, 24, 20)
        p.drawLine(8, 20, 24, 12)
        hot_x, hot_y = 4, 8

    elif tool == "add_text":
        # Text cursor: I-beam with "T"
        p.setPen(QPen(QColor(40, 40, 180), 2))
        font = QFont("Segoe UI", 18, QFont.Weight.Bold)
        p.setFont(font)
        p.drawText(4, 24, "T")
        # Small crosshair at bottom-left
        p.setPen(QPen(QColor(40, 40, 180), 1))
        p.drawLine(0, 28, 6, 28)
        p.drawLine(3, 25, 3, 31)
        hot_x, hot_y = 3, 28

    elif tool == "add_image":
        # Image icon: frame with mountain/sun
        p.setPen(QPen(QColor(40, 120, 40), 2))
        p.setBrush(QBrush(QColor(220, 240, 220, 180)))
        p.drawRect(4, 6, 24, 20)
        # Sun
        p.setBrush(QBrush(QColor(255, 200, 40)))
        p.setPen(QPen(Qt.PenStyle.NoPen))
        p.drawEllipse(8, 9, 7, 7)
        # Mountain
        p.setPen(QPen(QColor(60, 140, 60), 2))
        p.setBrush(QBrush(QColor(80, 160, 80, 180)))
        mountain = QPolygonF([
            QPointF(6, 24), QPointF(16, 14),
            QPointF(22, 19), QPointF(26, 15), QPointF(26, 24),
        ])
        p.drawPolygon(mountain)
        hot_x, hot_y = 4, 6

    elif tool.startswith("shape_"):
        # Crosshair cursor for shapes
        p.setPen(QPen(QColor(40, 40, 40), 1.5))
        cx, cy = 16, 16
        gap = 4
        arm = 10
        # Crosshair lines with gap in center
        p.drawLine(cx - arm, cy, cx - gap, cy)
        p.drawLine(cx + gap, cy, cx + arm, cy)
        p.drawLine(cx, cy - arm, cx, cy - gap)
        p.drawLine(cx, cy + gap, cx, cy + arm)
        # Small shape hint in corner
        stype = tool.replace("shape_", "")
        p.setPen(QPen(QColor(40, 100, 200), 1.5))
        if stype == "rect":
            p.drawRect(22, 22, 8, 6)
        elif stype == "circle":
            p.drawEllipse(22, 22, 8, 8)
        elif stype == "line":
            p.drawLine(22, 30, 30, 22)
        elif stype == "arrow":
            p.drawLine(22, 30, 30, 22)
            p.drawLine(30, 22, 26, 23)
            p.drawLine(30, 22, 29, 26)
        hot_x, hot_y = cx, cy

    else:
        # Fallback crosshair
        p.setPen(QPen(QColor(40, 40, 40), 1.5))
        p.drawLine(16, 6, 16, 26)
        p.drawLine(6, 16, 26, 16)
        hot_x, hot_y = 16, 16

    p.end()
    cursor = QCursor(pm, hot_x, hot_y)
    _edit_cursors[tool] = cursor
    return cursor


# ---------------------------------------------------------------------------
# A single OCR word overlay — positioned on top of the rendered page.
# ---------------------------------------------------------------------------
class WordOverlay(QGraphicsRectItem):
    """Transparent rectangle over one OCR'd word."""

    def __init__(self, x, y, w, h, text, line_id, word_index, parent=None):
        super().__init__(x, y, w, h, parent)
        self.word_text = text
        self.line_id = line_id
        self.word_index = word_index
        self.setPen(QPen(Qt.PenStyle.NoPen))
        self.setBrush(QBrush(Qt.BrushStyle.NoBrush))
        self.setAcceptHoverEvents(True)
        self.setCursor(Qt.CursorShape.IBeamCursor)
        self.setZValue(10)
        self._selected = False

    @property
    def selected(self):
        return self._selected

    def set_selected(self, sel: bool):
        if sel == self._selected:
            return
        self._selected = sel
        if sel:
            self.setBrush(QBrush(QColor(51, 153, 255, 80)))
        else:
            self.setBrush(QBrush(Qt.BrushStyle.NoBrush))


# ---------------------------------------------------------------------------
# OCR worker thread — uses Windows built-in OCR via PowerShell
# ---------------------------------------------------------------------------
class OCRWorker(QThread):
    page_ocr_done = pyqtSignal(int, list)
    all_done = pyqtSignal()
    error_occurred = pyqtSignal(str)
    status_update = pyqtSignal(str)

    def __init__(self, doc_path, page_indices, render_scale, parent=None):
        super().__init__(parent)
        self.doc_path = doc_path
        self.page_indices = page_indices
        self.render_scale = render_scale
        self._cancel = False
        if getattr(sys, 'frozen', False):
            base = sys._MEIPASS
            exe_dir = os.path.dirname(sys.executable)
        else:
            base = os.path.dirname(os.path.abspath(__file__))
            exe_dir = base
        # Locate ocr_helper.ps1 — try bundle dir first, then exe dir
        self._ps_script = os.path.join(base, "ocr_helper.ps1")
        if not os.path.isfile(self._ps_script):
            self._ps_script = os.path.join(exe_dir, "ocr_helper.ps1")
        # Use a user-writable temp directory for intermediate PNGs.
        # Do NOT write to exe_dir — when the app is installed under
        # C:\Program Files\BoltPDF, that directory is read-only for
        # normal users, which would cause OCR to crash.
        self._tmp_dir = os.path.join(
            os.environ.get("LOCALAPPDATA", tempfile.gettempdir()),
            "BoltPDF", "ocr")
        try:
            os.makedirs(self._tmp_dir, exist_ok=True)
        except Exception:
            self._tmp_dir = tempfile.gettempdir()

    def cancel(self):
        self._cancel = True

    def run(self):
        if not os.path.isfile(self._ps_script):
            self.error_occurred.emit(f"OCR helper script not found:\n{self._ps_script}")
            self.all_done.emit()
            return

        doc = None
        try:
            doc = pdfium.PdfDocument(self.doc_path)
            for page_idx in self.page_indices:
                if self._cancel:
                    break
                if page_idx < 0 or page_idx >= len(doc):
                    continue
                self.status_update.emit(f"OCR page {page_idx + 1}...")
                try:
                    page = doc[page_idx]
                    bitmap = page.render(scale=self.render_scale)
                    pil_image = bitmap.to_pil()
                    if pil_image.mode != "RGB":
                        pil_image = pil_image.convert("RGB")

                    with tempfile.NamedTemporaryFile(
                        suffix=".png", delete=False, dir=self._tmp_dir
                    ) as tmp:
                        tmp_path = tmp.name
                        pil_image.save(tmp, format="PNG")

                    try:
                        result = subprocess.run(
                            ["powershell", "-NoProfile", "-ExecutionPolicy",
                             "Bypass", "-File", self._ps_script,
                             "-ImagePath", tmp_path],
                            capture_output=True, text=True, timeout=30,
                            creationflags=(subprocess.CREATE_NO_WINDOW
                                           if sys.platform == "win32" else 0),
                        )
                        raw = result.stdout.strip()
                        if not raw:
                            self.page_ocr_done.emit(page_idx, [])
                            continue

                        raw = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', raw)
                        data = json.loads(raw)

                        if data.get("error"):
                            self.error_occurred.emit(
                                f"OCR page {page_idx + 1}: {data['error']}")
                            continue

                        lines = []
                        for line_data in data.get("lines", []):
                            words = line_data.get("words", [])
                            if words:
                                lines.append(words)

                        self.page_ocr_done.emit(page_idx, lines)

                    finally:
                        try:
                            os.unlink(tmp_path)
                        except OSError:
                            pass

                except Exception as e:
                    self.error_occurred.emit(f"OCR page {page_idx + 1}: {e}")

        except Exception as e:
            self.error_occurred.emit(f"OCR error: {e}")
        finally:
            if doc is not None:
                try:
                    doc.close()
                except Exception:
                    pass
            self.all_done.emit()


# ---------------------------------------------------------------------------
# Background page renderer — multi-process, viewport-priority.
#
# Rendering a PDF page is pure CPU work that holds Python's GIL via
# PDFium's C calls, so OS threads give zero speedup.  We use a pool of
# persistent worker *processes* instead — each opens its own
# PdfDocument once (no shared native state → no crashes) and renders
# pages on demand.  A small dispatcher QThread keeps ~2×N pages in
# flight so when the user scrolls, the new focus takes effect within
# one render cycle.
# ---------------------------------------------------------------------------
import threading
import queue as _queue_mod


def _render_worker_process(doc_path, scale, job_q, result_q):
    """Persistent multiprocessing worker.

    Opens the PDF document once and then renders pages as jobs arrive
    on *job_q*, shipping raw RGBA bytes back on *result_q*.  Exits
    cleanly when it receives a ``None`` sentinel.

    Must live at module level so it can be pickled for spawn start
    method (Windows / frozen exe).
    """
    try:
        import pypdfium2 as _pdfium
        doc = _pdfium.PdfDocument(doc_path)
    except Exception as e:
        try:
            result_q.put(("INIT_ERROR", -1, str(e), 0, 0))
        except Exception:
            pass
        return

    try:
        while True:
            try:
                msg = job_q.get()
            except (EOFError, KeyboardInterrupt, OSError):
                break
            if msg is None:
                break  # shutdown sentinel
            page_idx = msg
            try:
                page = doc[page_idx]
                try:
                    bitmap = page.render(scale=scale, draw_annots=False)
                    pil = bitmap.to_pil()
                finally:
                    page.close()
                if pil.mode != "RGBA":
                    pil = pil.convert("RGBA")
                raw = pil.tobytes("raw", "RGBA")
                try:
                    result_q.put(
                        ("OK", page_idx, raw, pil.width, pil.height))
                except Exception:
                    break
            except Exception as e:
                try:
                    result_q.put(("ERR", page_idx, str(e), 0, 0))
                except Exception:
                    break
    finally:
        try:
            doc.close()
        except Exception:
            pass


class PageRenderer(QThread):
    """Multi-process, viewport-aware page renderer.

    Maintains a pool of up to N worker processes (N ≈ CPU count,
    capped at 8) that render pages in parallel.  A dispatcher thread
    (this QThread) pushes jobs in *focus-priority* order, keeping ~2×N
    jobs in flight at a time so the queue reorders quickly when the
    user scrolls.

    Falls back to the single-threaded in-process path if spawning a
    worker pool fails for any reason.
    """
    page_ready = pyqtSignal(int, QImage)
    all_done = pyqtSignal()
    error_occurred = pyqtSignal(str)

    # Tune these if needed.  Beyond ~8 workers diminishing returns kick
    # in quickly because PDFium is already very fast per page and the
    # pickled RGBA bytes become the bottleneck.
    _MAX_WORKERS = 8

    # Only render pages within this many pages of the current focus.
    # Pages outside the window are never queued and are dropped from the
    # "done" set on focus change so they re-render if the user scrolls
    # back to them.  Mirrors DocumentTab.PAGE_BUFFER.
    _RENDER_WINDOW = 10

    def __init__(self, doc_path, scale, num_pages, parent=None):
        super().__init__(parent)
        self.doc_path = doc_path
        self.scale = scale
        self.num_pages = num_pages
        self._cancel = False

        # Light lock — only protects the queue / focus, never held
        # while waiting on IPC so it can never block rendering.
        self._lock = threading.Lock()
        self._queue: list[int] = []
        self._done: set[int] = set()
        self._focus = 0
        self._retries: dict[int, int] = {}
        # Flag toggled by set_focus() to tell the dispatcher to drain
        # stale jobs and re-prioritise.
        self._focus_dirty = False

        # Worker pool state (created in run())
        cpu = os.cpu_count() or 4
        self._n_workers = min(self._MAX_WORKERS, max(1, min(num_pages, cpu)))
        self._workers: list = []
        self._job_q = None
        self._result_q = None

    # -- called from the GUI thread (DocumentTab) --------------------------
    def set_focus(self, page_idx: int):
        """Re-center the render window on *page_idx*.

        Pages within ``_RENDER_WINDOW`` of *page_idx* get queued (if
        they are not already done).  Pages that have drifted outside
        the window are evicted from the ``_done`` set so they will be
        re-rendered from scratch if the user later scrolls back.

        Sets a dirty flag so the dispatcher will drain any stale jobs
        from the IPC queue before picking up new work — otherwise
        workers would chew through old pages before reaching the
        new focus window.
        """
        with self._lock:
            self._focus = page_idx
            lo = max(0, page_idx - self._RENDER_WINDOW)
            hi = min(self.num_pages - 1, page_idx + self._RENDER_WINDOW)
            # Drop out-of-window completions so scroll-back re-renders.
            self._done = {i for i in self._done if lo <= i <= hi}
            self._retries = {
                i: c for i, c in self._retries.items() if lo <= i <= hi
            }
            self._queue = sorted(
                [i for i in range(lo, hi + 1) if i not in self._done],
                key=lambda i: abs(i - page_idx),
            )
            self._focus_dirty = True

    def cancel(self):
        self._cancel = True

    # -- helpers -----------------------------------------------------------
    def _next_page(self) -> int | None:
        with self._lock:
            while self._queue:
                p = self._queue.pop(0)
                if p not in self._done:
                    return p
        return None

    def _requeue(self, page_idx: int):
        with self._lock:
            lo = max(0, self._focus - self._RENDER_WINDOW)
            hi = min(self.num_pages - 1, self._focus + self._RENDER_WINDOW)
            if page_idx not in self._done and lo <= page_idx <= hi:
                self._queue.append(page_idx)
                self._queue.sort(key=lambda i: abs(i - self._focus))

    def _rebuild_queue(self):
        with self._lock:
            lo = max(0, self._focus - self._RENDER_WINDOW)
            hi = min(self.num_pages - 1, self._focus + self._RENDER_WINDOW)
            self._queue = sorted(
                [i for i in range(lo, hi + 1) if i not in self._done],
                key=lambda i: abs(i - self._focus),
            )

    # -- worker-pool lifecycle --------------------------------------------
    def _start_workers(self) -> bool:
        """Spawn the worker processes.  Returns True on success."""
        try:
            ctx = multiprocessing.get_context("spawn")
            self._job_q = ctx.Queue()
            self._result_q = ctx.Queue()
            for _ in range(self._n_workers):
                p = ctx.Process(
                    target=_render_worker_process,
                    args=(self.doc_path, self.scale,
                          self._job_q, self._result_q),
                    daemon=True,
                )
                p.start()
                self._workers.append(p)
            return True
        except Exception as e:
            self.error_occurred.emit(f"Worker pool start failed: {e}")
            return False

    def _shutdown_workers(self):
        """Tear down the worker pool.

        On *cancel* we terminate immediately — we don't care about
        half-rendered pages and don't want tab-close to hang.  On
        normal completion a graceful sentinel shutdown is fine, but
        terminate() is still correct (workers are daemons with no
        external side effects).
        """
        for p in self._workers:
            try:
                if p.is_alive():
                    p.terminate()
            except Exception:
                pass
        for p in self._workers:
            try:
                p.join(timeout=0.5)
            except Exception:
                pass
        # Close queues — cancel the feeder thread so we don't block
        for q in (self._job_q, self._result_q):
            try:
                if q is not None:
                    q.close()
                    q.cancel_join_thread()
            except Exception:
                pass
        self._workers = []
        self._job_q = None
        self._result_q = None

    # -- fallback: in-process single-threaded render ----------------------
    def _run_in_process_fallback(self):
        """Used if multiprocessing can't be started.  Renders pages
        sequentially in this QThread using a local PdfDocument."""
        try:
            doc = pdfium.PdfDocument(self.doc_path)
        except Exception as exc:
            self.error_occurred.emit(f"Cannot open PDF: {exc}")
            self.all_done.emit()
            return
        try:
            while not self._cancel:
                page_idx = self._next_page()
                if page_idx is None:
                    with self._lock:
                        remaining = self.num_pages - len(self._done)
                    if remaining <= 0:
                        break
                    self._rebuild_queue()
                    with self._lock:
                        if not self._queue:
                            break
                    continue
                try:
                    page = doc[page_idx]
                    try:
                        bitmap = page.render(scale=self.scale, draw_annots=False)
                        pil = bitmap.to_pil()
                    finally:
                        page.close()
                    if pil.mode != "RGBA":
                        pil = pil.convert("RGBA")
                    data = pil.tobytes("raw", "RGBA")
                    qimg = QImage(
                        data, pil.width, pil.height,
                        QImage.Format.Format_RGBA8888).copy()
                    with self._lock:
                        self._done.add(page_idx)
                    if not self._cancel:
                        self.page_ready.emit(page_idx, qimg)
                except Exception as e:
                    with self._lock:
                        self._done.add(page_idx)
                    if not self._cancel:
                        self.error_occurred.emit(
                            f"Page {page_idx + 1}: {e}")
        finally:
            try:
                doc.close()
            except Exception:
                pass
            self.all_done.emit()

    # -- main dispatcher loop ---------------------------------------------
    def run(self):
        self.set_focus(self._focus)

        if not self._start_workers():
            self._run_in_process_fallback()
            return

        # One pending job per worker maximum — keeps the IPC queue
        # short so a focus change flushes stale work almost
        # immediately.  Setting this higher would queue more stale
        # pages ahead of the new focus window.
        MAX_IN_FLIGHT = self._n_workers
        in_flight: set[int] = set()

        try:
            while not self._cancel:
                # --- Handle focus changes: drain stale jobs ------------
                # set_focus() flips _focus_dirty from the GUI thread.
                # We pull every pending job out of the IPC queue so the
                # workers' next get() will immediately see the new
                # priority pages, and we forget them from in_flight so
                # the dispatcher refills for the new window.
                with self._lock:
                    dirty = self._focus_dirty
                    if dirty:
                        self._focus_dirty = False
                if dirty and self._job_q is not None:
                    drained = 0
                    try:
                        while True:
                            stale = self._job_q.get_nowait()
                            in_flight.discard(stale)
                            drained += 1
                    except _queue_mod.Empty:
                        pass
                    except Exception:
                        pass

                # --- Fill the in-flight pool with focus-priority jobs ---
                while len(in_flight) < MAX_IN_FLIGHT:
                    p = self._next_page()
                    if p is None:
                        break
                    try:
                        self._job_q.put_nowait(p)
                        in_flight.add(p)
                    except Exception:
                        break

                # --- Nothing dispatched and nothing pending → park.
                # We deliberately do NOT break here: the user may scroll
                # later, which will call set_focus() and refill the
                # queue with a new window of pages to render.  We sit
                # idle cheaply until that happens (or until cancel).
                if not in_flight:
                    self._rebuild_queue()
                    with self._lock:
                        queue_empty = not self._queue
                    if queue_empty:
                        # Idle-wait briefly so we don't spin, and so
                        # the cancel flag gets checked regularly.
                        if self._cancel:
                            break
                        time.sleep(0.02)
                        continue

                # --- Wait for the next finished page --------------------
                try:
                    result = self._result_q.get(timeout=1.0)
                except _queue_mod.Empty:
                    continue  # let the cancel flag be rechecked
                except (EOFError, OSError, BrokenPipeError) as e:
                    self.error_occurred.emit(f"Renderer IPC broken: {e}")
                    break

                if self._cancel:
                    break

                tag = result[0]
                if tag == "INIT_ERROR":
                    self.error_occurred.emit(
                        f"PDF open failed in worker: {result[2]}")
                    break

                page_idx = result[1]
                in_flight.discard(page_idx)

                # Drop results for pages that have drifted outside the
                # current render window (user scrolled during render).
                with self._lock:
                    lo = max(0, self._focus - self._RENDER_WINDOW)
                    hi = min(self.num_pages - 1,
                             self._focus + self._RENDER_WINDOW)
                if not (lo <= page_idx <= hi):
                    continue

                if tag == "ERR":
                    attempts = self._retries.get(page_idx, 0) + 1
                    self._retries[page_idx] = attempts
                    if not self._cancel:
                        self.error_occurred.emit(
                            f"Page {page_idx + 1} "
                            f"(attempt {attempts}): {result[2]}")
                    if attempts < 3:
                        self._requeue(page_idx)
                    else:
                        with self._lock:
                            self._done.add(page_idx)  # give up
                    continue

                # tag == "OK" — build QImage from raw RGBA bytes
                raw = result[2]
                w = result[3]
                h = result[4]
                try:
                    # QImage needs its own copy of the buffer so it's
                    # safe to reuse after the worker's bytes are freed.
                    qimg = QImage(
                        bytes(raw), w, h,
                        QImage.Format.Format_RGBA8888).copy()
                except Exception as e:
                    if not self._cancel:
                        self.error_occurred.emit(
                            f"Page {page_idx + 1} decode: {e}")
                    continue

                with self._lock:
                    self._done.add(page_idx)
                if not self._cancel:
                    self.page_ready.emit(page_idx, qimg)

        except Exception as exc:
            if not self._cancel:
                self.error_occurred.emit(
                    f"Render dispatcher crashed: {exc}")
        finally:
            self._shutdown_workers()
            self.all_done.emit()


# ---------------------------------------------------------------------------
# Multiprocessing helpers — top-level functions so they can be pickled.
# Each subprocess opens its own PdfDocument (separate process = safe).
# ---------------------------------------------------------------------------
import multiprocessing

def _mp_render_page(args):
    """Render one page to a temp JPEG file.  Returns (page_idx, tmp_path)."""
    doc_path, page_idx, scale, quality, tmp_dir = args
    import pypdfium2 as _pdf
    from PIL import Image as _Img
    doc = _pdf.PdfDocument(doc_path)
    page = doc[page_idx]
    bmp = page.render(scale=scale)
    pil = bmp.to_pil()
    page.close()
    doc.close()
    if pil.mode != "RGB":
        pil = pil.convert("RGB")
    tmp_path = os.path.join(tmp_dir, f"page_{page_idx:06d}.jpg")
    pil.save(tmp_path, "JPEG", quality=quality)
    return (page_idx, tmp_path)


def _mp_export_image(args):
    """Extract one embedded image to a JPEG file at native resolution.

    Strategy:
    1. Try extract() — this pulls the raw compressed stream straight out of the
       PDF (zero re-encoding if the image is already JPEG).  If the raw stream
       is JPEG we write it directly to disk to preserve the exact original.
    2. If the raw stream is not JPEG (e.g. PNG, JPEG2000, CCITT fax), open it
       with PIL and save as maximum-quality JPEG at the native pixel size.
    3. Fallback: render the image object to a bitmap at its native pixel
       dimensions and save as maximum-quality JPEG.
    """
    doc_path, page_idx, obj_index, filepath, _quality = args
    import io as _io
    import pypdfium2 as _pdf
    from PIL import Image as _Img

    doc = _pdf.PdfDocument(doc_path)
    page = doc[page_idx]
    img_objs = list(page.get_objects(
        filter=[_pdf.raw.FPDF_PAGEOBJ_IMAGE]))
    if obj_index >= len(img_objs):
        doc.close()
        return None
    obj = img_objs[obj_index]
    pdf_img = _pdf.PdfImage(raw=obj.raw, page=obj.page, pdf=obj.pdf)

    saved = False
    # --- Attempt 1: raw extract (preserves native bytes) -------------------
    try:
        buf = _io.BytesIO()
        pdf_img.extract(buf, fb_format="jpg")
        raw_bytes = buf.getvalue()
        if len(raw_bytes) > 0:
            # Check if it's already a JPEG — write raw bytes directly
            if raw_bytes[:2] == b'\xff\xd8':
                with open(filepath, "wb") as f:
                    f.write(raw_bytes)
                saved = True
            else:
                # Non-JPEG raw (PNG, etc) — open with PIL, save as JPEG
                buf.seek(0)
                pil = _Img.open(buf)
                if pil.mode not in ("RGB", "L"):
                    pil = pil.convert("RGB")
                pil.save(filepath, "JPEG", quality=100, subsampling=0)
                saved = True
    except Exception:
        pass

    # --- Attempt 2: render at native image dimensions ----------------------
    if not saved:
        try:
            # Get the native pixel size of the embedded image
            meta = pdf_img.get_metadata()
            native_w = meta.width if hasattr(meta, 'width') and meta.width > 0 else 0
            native_h = meta.height if hasattr(meta, 'height') and meta.height > 0 else 0

            if native_w > 0 and native_h > 0:
                bmp = pdf_img.get_bitmap(render=True)
            else:
                bmp = pdf_img.get_bitmap(render=True)
            pil = bmp.to_pil()
            if pil.mode not in ("RGB", "L"):
                pil = pil.convert("RGB")
            pil.save(filepath, "JPEG", quality=100, subsampling=0)
            saved = True
        except Exception:
            pass

    # --- Attempt 3: last resort, render full page crop ---------------------
    if not saved:
        try:
            bmp = pdf_img.get_bitmap(render=True)
            pil = bmp.to_pil()
            if pil.mode not in ("RGB", "L"):
                pil = pil.convert("RGB")
            pil.save(filepath, "JPEG", quality=100, subsampling=0)
        except Exception:
            pass

    doc.close()
    return filepath


# ---------------------------------------------------------------------------
# Combine worker — merges multiple PDFs into one using pypdf's PdfWriter.
# Each file is read into memory in parallel across CPU cores, then merged
# sequentially.  pypdf clones every page object and its full resource tree
# (fonts, images, MediaBox, CropBox, annotations, etc.) independently,
# so every page is byte-for-byte identical to the original.
# ---------------------------------------------------------------------------
def _mp_preread_pdf(args):
    """Read a PDF file into memory in a worker process.
    Returns (file_index, raw_bytes)."""
    file_index, pdf_path = args
    with open(pdf_path, "rb") as f:
        raw = f.read()
    return (file_index, raw)


class CombineWorker(QThread):
    progress = pyqtSignal(int)       # 0-100
    finished_ok = pyqtSignal(str)    # output path
    error_occurred = pyqtSignal(str)

    def __init__(self, pdf_paths, output_path, parent=None):
        super().__init__(parent)
        self.pdf_paths = pdf_paths
        self.output_path = output_path

    def run(self):
        try:
            from pypdf import PdfReader, PdfWriter
            import io

            n = len(self.pdf_paths)
            if n == 0:
                self.error_occurred.emit("No PDF files selected.")
                return

            # Phase 1: parallel read of all files into memory
            workers = max(1, min(n, os.cpu_count() or 1))
            args = [(i, p) for i, p in enumerate(self.pdf_paths)]

            file_data = [None] * n
            with multiprocessing.Pool(processes=workers) as pool:
                for file_idx, raw in pool.imap_unordered(
                        _mp_preread_pdf, args):
                    file_data[file_idx] = raw
                    pct = int(((file_idx + 1) / n) * 40)
                    self.progress.emit(pct)

            # Phase 2: sequential merge using pypdf
            # Count total pages first for progress tracking
            readers = []
            total_pages = 0
            for raw_bytes in file_data:
                reader = PdfReader(io.BytesIO(raw_bytes))
                readers.append(reader)
                total_pages += len(reader.pages)

            self.progress.emit(50)

            writer = PdfWriter()
            pages_done = 0

            for reader in readers:
                for page in reader.pages:
                    writer.add_page(page)
                    pages_done += 1
                    pct = 50 + int((pages_done / total_pages) * 50)
                    self.progress.emit(pct)

            with open(self.output_path, "wb") as out_f:
                writer.write(out_f)

            self.progress.emit(100)
            self.finished_ok.emit(self.output_path)

        except Exception as e:
            self.error_occurred.emit(str(e))


# ---------------------------------------------------------------------------
# Rebuild worker — snapshots every page as JPEG, rebuilds a new image PDF.
# Uses multiprocessing to render pages across all CPU cores in parallel.
# ---------------------------------------------------------------------------
class RebuildWorker(QThread):
    progress = pyqtSignal(int)
    finished_ok = pyqtSignal(str)
    error_occurred = pyqtSignal(str)

    def __init__(self, doc_path, output_path, render_scale=2.0, quality=90,
                 parent=None):
        super().__init__(parent)
        self.doc_path = doc_path
        self.output_path = output_path
        self.render_scale = render_scale
        self.quality = quality

    def run(self):
        tmp_dir = None
        try:
            from PIL import Image

            doc = pdfium.PdfDocument(self.doc_path)
            num_pages = len(doc)
            doc.close()

            tmp_dir = tempfile.mkdtemp(prefix="boltpdf_rebuild_")
            workers = max(1, os.cpu_count() or 1)
            args = [
                (self.doc_path, i, self.render_scale, self.quality, tmp_dir)
                for i in range(num_pages)
            ]

            # Render pages in parallel across all cores
            done = 0
            page_paths = [None] * num_pages
            with multiprocessing.Pool(processes=workers) as pool:
                for page_idx, tmp_path in pool.imap_unordered(
                        _mp_render_page, args):
                    page_paths[page_idx] = tmp_path
                    done += 1
                    self.progress.emit(done)

            # Assemble into a single PDF (sequential — must be page-order)
            page_images = []
            for p in page_paths:
                if p is None:
                    continue
                page_images.append(Image.open(p))

            if page_images:
                first = page_images[0]
                rest = page_images[1:]
                first.save(
                    self.output_path, "PDF", save_all=True,
                    append_images=rest, quality=self.quality,
                )
                # Close handles
                for img in page_images:
                    img.close()

            self.finished_ok.emit(self.output_path)

        except Exception as e:
            self.error_occurred.emit(str(e))
        finally:
            # Clean up temp render directory
            if tmp_dir and os.path.isdir(tmp_dir):
                import shutil
                try:
                    shutil.rmtree(tmp_dir)
                except Exception:
                    pass


# ---------------------------------------------------------------------------
# Image detection worker — scans pages for embedded image objects
# ---------------------------------------------------------------------------
class ImageDetectorWorker(QThread):
    """Finds all embedded images in the given pages and emits their bounds.

    Scans outward from a *focus* page so the images near the viewport are
    available for selection almost immediately.  Call ``set_focus()`` from
    the GUI thread to re-prioritise when the user scrolls.
    """
    page_images_found = pyqtSignal(int, list)
    all_done = pyqtSignal()
    error_occurred = pyqtSignal(str)

    def __init__(self, doc_path, page_indices, focus=0, parent=None):
        super().__init__(parent)
        self.doc_path = doc_path
        self._cancel = False
        self._lock = threading.Lock()
        self._done: set[int] = set()
        self._focus = focus
        self._all_pages = set(page_indices)
        # Build initial queue sorted by distance from focus
        self._queue: list[int] = sorted(
            page_indices, key=lambda i: abs(i - focus))

    def set_focus(self, page_idx: int):
        """Re-sort the remaining queue so *page_idx*'s neighbours go first."""
        with self._lock:
            self._focus = page_idx
            self._queue = sorted(
                [i for i in self._all_pages if i not in self._done],
                key=lambda i: abs(i - page_idx),
            )

    def cancel(self):
        self._cancel = True

    def _next_page(self) -> int | None:
        with self._lock:
            while self._queue:
                p = self._queue.pop(0)
                if p not in self._done:
                    return p
        return None

    def run(self):
        doc = None
        try:
            doc = pdfium.PdfDocument(self.doc_path)
            while not self._cancel:
                page_idx = self._next_page()
                if page_idx is None:
                    break
                try:
                    page = doc[page_idx]
                    page_h = page.get_height()
                    results = []
                    obj_seq = 0
                    for obj in page.get_objects(
                            filter=[pdfium.raw.FPDF_PAGEOBJ_IMAGE]):
                        left, bottom, right, top = obj.get_bounds()
                        img = pdfium.PdfImage(
                            raw=obj.raw, page=obj.page, pdf=obj.pdf)
                        px_w, px_h = img.get_px_size()
                        results.append({
                            "obj_index": obj_seq,
                            "left": left,
                            "bottom": bottom,
                            "right": right,
                            "top": top,
                            "page_h": page_h,
                            "px_w": px_w,
                            "px_h": px_h,
                        })
                        obj_seq += 1
                    with self._lock:
                        self._done.add(page_idx)
                    self.page_images_found.emit(page_idx, results)
                except Exception as e:
                    with self._lock:
                        self._done.add(page_idx)
                    self.error_occurred.emit(
                        f"Image detection page {page_idx + 1}: {e}")
        except Exception as e:
            self.error_occurred.emit(f"Image detection: {e}")
        finally:
            if doc is not None:
                try:
                    doc.close()
                except Exception:
                    pass
            self.all_done.emit()


# ---------------------------------------------------------------------------
# Image export worker — extracts selected images at native resolution
# ---------------------------------------------------------------------------
class ImageExportWorker(QThread):
    """Extracts specific embedded images from the PDF as JPEG files.
    Uses multiprocessing to extract across all CPU cores in parallel."""
    progress = pyqtSignal(int)
    finished_ok = pyqtSignal(str)
    error_occurred = pyqtSignal(str)

    def __init__(self, doc_path, image_list, output_dir, base_name,
                 quality=95, parent=None):
        """
        image_list: list of dicts with 'page_idx' and 'obj_index' keys.
        """
        super().__init__(parent)
        self.doc_path = doc_path
        self.image_list = image_list
        self.output_dir = output_dir
        self.base_name = base_name
        self.quality = quality

    def run(self):
        try:
            os.makedirs(self.output_dir, exist_ok=True)

            # Build argument list for parallel extraction
            args = []
            for seq, info in enumerate(self.image_list, 1):
                filename = f"{self.base_name}_export_{seq:03d}.jpg"
                filepath = os.path.join(self.output_dir, filename)
                args.append((
                    self.doc_path, info["page_idx"], info["obj_index"],
                    filepath, self.quality,
                ))

            workers = max(1, min(os.cpu_count() or 1, len(args)))
            done = 0
            with multiprocessing.Pool(processes=workers) as pool:
                for result in pool.imap_unordered(_mp_export_image, args):
                    done += 1
                    self.progress.emit(done)

            self.finished_ok.emit(self.output_dir)

        except Exception as e:
            self.error_occurred.emit(str(e))


# ---------------------------------------------------------------------------
# ImageOverlay — clickable rectangle drawn over a detected embedded image.
# Shows a thin blue border + tick when selected.
# ---------------------------------------------------------------------------
class ImageOverlay(QGraphicsRectItem):
    """Overlay for one detected image in the PDF.  Click to toggle selection."""

    def __init__(self, x, y, w, h, page_idx, obj_index, parent=None):
        super().__init__(x, y, w, h, parent)
        self.page_idx = page_idx
        self.obj_index = obj_index
        self._selected = False

        # Default: invisible border
        self.setPen(QPen(Qt.PenStyle.NoPen))
        self.setBrush(QBrush(Qt.BrushStyle.NoBrush))
        self.setAcceptHoverEvents(True)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setZValue(15)

        # Tick item (child)
        self._tick = _TickItem(self)
        self._tick.setVisible(False)
        self._position_tick()

    def _position_tick(self):
        """Place the tick in the top-right corner of the rect."""
        r = self.rect()
        size = min(24, r.width() * 0.15, r.height() * 0.15)
        size = max(size, 14)
        self._tick.set_size(size)
        self._tick.setPos(r.x() + r.width() - size - 2, r.y() + 2)

    @property
    def selected(self):
        return self._selected

    def toggle_selected(self):
        self._selected = not self._selected
        if self._selected:
            self.setPen(QPen(QColor(51, 153, 255), 2))
            self.setBrush(QBrush(QColor(51, 153, 255, 30)))
            self._tick.setVisible(True)
        else:
            self.setPen(QPen(Qt.PenStyle.NoPen))
            self.setBrush(QBrush(Qt.BrushStyle.NoBrush))
            self._tick.setVisible(False)

    def hoverEnterEvent(self, event):
        if not self._selected:
            self.setPen(QPen(QColor(51, 153, 255, 120), 1.5))
        super().hoverEnterEvent(event)

    def hoverLeaveEvent(self, event):
        if not self._selected:
            self.setPen(QPen(Qt.PenStyle.NoPen))
        super().hoverLeaveEvent(event)


class _TickItem(QGraphicsRectItem):
    """Small green tick badge drawn as a child of ImageOverlay."""

    def __init__(self, parent):
        super().__init__(parent)
        self._size = 20
        self.setPen(QPen(Qt.PenStyle.NoPen))
        self.setBrush(QBrush(Qt.BrushStyle.NoBrush))
        self.setZValue(16)

    def set_size(self, s):
        self._size = s
        self.setRect(0, 0, s, s)

    def paint(self, painter, option, widget=None):
        s = self._size
        # Green circle background
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        painter.setPen(QPen(Qt.PenStyle.NoPen))
        painter.setBrush(QBrush(QColor(34, 180, 85)))
        painter.drawEllipse(QRectF(0, 0, s, s))
        # White tick
        painter.setPen(QPen(QColor(255, 255, 255), max(1.5, s * 0.12)))
        p1 = QPointF(s * 0.22, s * 0.52)
        p2 = QPointF(s * 0.42, s * 0.72)
        p3 = QPointF(s * 0.78, s * 0.30)
        painter.drawLine(p1, p2)
        painter.drawLine(p2, p3)


# ===========================================================================
# Edit Mode — in-place PDF editing backed by PyMuPDF (fitz).
#
# Design:
#   - Edit Mode is a stateful toolbar — no overlays are drawn over the
#     page in the idle state.  Actions are triggered one at a time.
#   - Add Text    : click on a page to place a new text box.
#   - Add Image   : pick a file, then click on a page to place it.
#   - Replace Image: temporarily re-uses the image-select UI (the
#                    detector's orange highlight overlays) to let the
#                    user click an existing image in the PDF and swap
#                    it for a file from disk.
#   - Save As...  : writes every queued EditRecord to a new PDF.
#
# All PyMuPDF calls happen in SaveEditedPdfWorker so the Qt UI stays
# responsive.  The page renderer is paused while saving because PyMuPDF
# and pypdfium2 should not both have the same document open at once.
# ===========================================================================


class EditRecord:
    """One pending edit against the currently open PDF.

    Types:
        'text_move'    — existing text block moved / resized
        'text_edit'    — existing text block text replaced (and maybe moved)
        'text_add'     — brand new text box added by the user
        'text_delete'  — existing text block removed
        'image_move'   — existing image moved / resized
        'image_add'    — brand new image inserted from disk
        'image_replace'— existing image swapped for a new file
        'image_delete' — existing image removed
        'stamp_add'    — preset or custom stamp overlay
        'shape_add'    — rectangle, circle, line, or arrow
        'redact_add'   — white-out rectangle to hide content
        'note_add'     — sticky note annotation

    Rects are stored in **PDF points** (the unit PyMuPDF uses when
    writing back to the file), not scene coordinates.  Overlay items
    convert between the two via ``DocumentTab._pt_to_scene`` /
    ``_scene_to_pt`` helpers.
    """

    def __init__(self, kind: str, page_idx: int, *,
                 orig_rect=None, new_rect=None,
                 text=None, html=None,
                 font_size=None, font_name=None,
                 color=(0, 0, 0), background_color=None,
                 image_path=None, image_xref=None,
                 rotation=0.0,
                 # Shape-specific fields
                 shape_type=None,        # 'rect','circle','line','arrow'
                 stroke_color=None,      # (r,g,b) 0-255
                 stroke_width=2.0,
                 fill_color=None,        # (r,g,b) 0-255 or None
                 stroke_style="solid",   # 'solid' | 'dash' | 'dot'
                 fill_opacity=80,        # 0-255 (alpha for fill)
                 arrow_both=False,       # arrowheads on both ends
                 # Line/arrow endpoints in PDF points
                 line_start=None,        # (x, y)
                 line_end=None,          # (x, y)
                 # Note-specific
                 note_icon=None,         # fitz annotation icon name
                 ):
        self.kind = kind
        self.page_idx = page_idx
        self.orig_rect = orig_rect      # (x0, y0, x1, y1) in PDF points
        self.new_rect = new_rect
        self.text = text
        # Optional rich-text HTML captured by the inline editor's
        # formatting toolbar.  When present, the save worker prefers
        # ``page.insert_htmlbox`` so bold / italic / colours / alignment
        # survive the round-trip to the output PDF.
        self.html = html
        self.font_size = font_size
        self.font_name = font_name      # PyMuPDF builtin, e.g. 'helv-b'
        self.color = color
        self.background_color = background_color  # (r,g,b) or None
        self.image_path = image_path
        self.image_xref = image_xref
        # Rotation in degrees (clockwise) applied to the record's
        # ``new_rect``.  Set by the Transform tool's rotation handle.
        # PyMuPDF's ``insert_image`` / ``insert_textbox`` only accept
        # rotations in 90° steps, so the save worker snaps this to the
        # nearest multiple of 90 before writing.
        self.rotation = float(rotation)
        # Shape drawing fields
        self.shape_type = shape_type
        self.stroke_color = stroke_color or color
        self.stroke_width = stroke_width
        self.fill_color = fill_color
        self.stroke_style = stroke_style or "solid"
        self.fill_opacity = (80 if fill_opacity is None
                             else int(fill_opacity))
        self.arrow_both = bool(arrow_both)
        self.line_start = line_start
        self.line_end = line_end
        # Sticky note icon
        self.note_icon = note_icon or "Note"
        # Free-form rebuild metadata that isn't a first-class field —
        # used by the undo/redo renderer to faithfully recreate
        # existing-content previews (e.g. detected font family/bold/
        # italic for text_move / text_edit).  Deep-copied with the
        # record, so history snapshots stay self-contained.
        self.extra: dict = {}



def _pymupdf_base14(font_name, bold: bool, italic: bool) -> str:
    """Map a detected font (its PDF/internal name + style flags) to a
    valid PyMuPDF Base-14 builtin code.

    The detector stores names like 'tiro-bi' / 'helv-b' which are NOT
    accepted by insert_textbox (real codes are 4 letters: helv/hebo/
    heit/hebi, tiro/tibo/tiit/tibi, cour/cobo/coit/cobi).  Using this
    keeps serif/sans/mono + bold/italic on edited & moved text instead
    of everything collapsing to plain Helvetica.  Always returns a
    valid builtin, so Save As can never fail on the font."""
    fn = (font_name or "").lower()
    if "cour" in fn or "mono" in fn or "consol" in fn:
        reg, b, i, bi = "cour", "cobo", "coit", "cobi"
    elif ("tiro" in fn or "time" in fn or "serif" in fn or "roman" in fn
          or "georgia" in fn or "garamond" in fn or "minion" in fn
          or "cambria" in fn or "book" in fn):
        reg, b, i, bi = "tiro", "tibo", "tiit", "tibi"
    else:
        reg, b, i, bi = "helv", "hebo", "heit", "hebi"
    if ("bold" in fn or fn.endswith("-b") or fn.endswith("-bi")
            or "-bi" in fn):
        bold = True
    if ("italic" in fn or "oblique" in fn or fn.endswith("-i")
            or fn.endswith("-bi") or "-bi" in fn):
        italic = True
    if bold and italic:
        return bi
    if bold:
        return b
    if italic:
        return i
    return reg


class SaveEditedPdfWorker(QThread):
    """Applies every queued EditRecord against the source PDF using
    PyMuPDF and writes the result to a new file.

    The original file is never modified.  Redactions are used to erase
    moved / edited / deleted regions, then fresh text or images are
    inserted at the user's chosen destination rectangles.
    """
    progress = pyqtSignal(int)      # 0-100
    finished_ok = pyqtSignal(str)   # output path
    error_occurred = pyqtSignal(str)

    def __init__(self, source_path: str, output_path: str,
                 edits: list, parent=None):
        super().__init__(parent)
        self.source_path = source_path
        self.output_path = output_path
        self.edits = edits
        self._cancel = False

    def cancel(self):
        """Cooperative cancel — safer than QThread.terminate() which
        would kill the thread while fitz still holds C-level state."""
        self._cancel = True

    def run(self):
        try:
            import fitz  # PyMuPDF
        except Exception as e:
            self.error_occurred.emit(f"PyMuPDF not installed: {e}")
            return

        doc = None
        try:
            doc = fitz.open(self.source_path)
            total = max(1, len(self.edits))

            # Group by page so each page's redactions are applied once.
            by_page: dict = {}
            for ed in self.edits:
                by_page.setdefault(ed.page_idx, []).append(ed)

            processed = 0
            for page_idx in sorted(by_page.keys()):
                if self._cancel:
                    self.error_occurred.emit("Save cancelled")
                    return
                if page_idx < 0 or page_idx >= len(doc):
                    continue
                page = doc[page_idx]

                # Phase A: redact every region the user touched on an
                # existing element.  New additions obviously have no
                # original rect to redact.  When a background colour
                # was sampled for a text edit, hand it to the redactor
                # as fill so the resulting "erased" area matches the
                # surrounding page instead of going white.
                # Glyphs can extend slightly outside the text block's
                # reported bbox (ascenders / descenders / anti-aliased
                # fringes), so grow the redact rect by a small pad so
                # the original characters are guaranteed to be erased
                # instead of leaving faint ghosts on the output.
                redact_pad = 0.8  # PDF points
                has_redactions = False
                for ed in by_page[page_idx]:
                    if ed.orig_rect is None:
                        continue
                    try:
                        x0, y0, x1, y1 = ed.orig_rect
                        padded = fitz.Rect(
                            x0 - redact_pad, y0 - redact_pad,
                            x1 + redact_pad, y1 + redact_pad,
                        )
                        annot = page.add_redact_annot(padded)
                        has_redactions = True
                        if (ed.background_color is not None
                                and annot is not None):
                            try:
                                r, g, b = ed.background_color
                                annot.set_colors(
                                    fill=(r / 255.0, g / 255.0, b / 255.0))
                                annot.update()
                            except Exception:
                                pass
                    except Exception:
                        pass
                # Only apply redactions if we actually added any.
                # apply_redactions() removes ALL annotations on the page,
                # so we must save and restore non-redaction annotations
                # (like sticky notes) that should survive.
                if has_redactions:
                    # Collect existing non-redaction annotations to restore
                    saved_annots = []
                    for a in page.annots() or []:
                        try:
                            atype = a.type[0]
                            # Skip redaction annotations (they'll be consumed)
                            if atype == fitz.PDF_ANNOT_REDACT:
                                continue
                            saved_annots.append({
                                "type": atype,
                                "type_name": a.type[1],
                                "rect": a.rect,
                                "info": dict(a.info) if a.info else {},
                                "content": (a.info or {}).get("content", ""),
                                "colors": a.colors if hasattr(a, 'colors') else None,
                            })
                        except Exception:
                            pass

                    try:
                        page.apply_redactions(
                            images=fitz.PDF_REDACT_IMAGE_NONE
                            if hasattr(fitz, "PDF_REDACT_IMAGE_NONE") else 0,
                            graphics=fitz.PDF_REDACT_LINE_ART_REMOVE_IF_COVERED
                            if hasattr(fitz, "PDF_REDACT_LINE_ART_REMOVE_IF_COVERED")
                            else 1,
                            text=fitz.PDF_REDACT_TEXT_REMOVE
                            if hasattr(fitz, "PDF_REDACT_TEXT_REMOVE") else 1,
                        )
                    except Exception:
                        try:
                            page.apply_redactions()
                        except Exception:
                            pass

                    # Restore text annotations that were wiped by redactions
                    for sa in saved_annots:
                        try:
                            if sa["type"] == fitz.PDF_ANNOT_TEXT:
                                new_a = page.add_text_annot(
                                    fitz.Point(sa["rect"].x0, sa["rect"].y0),
                                    sa["content"])
                                if new_a and sa["info"]:
                                    try:
                                        new_a.set_info(sa["info"])
                                        new_a.update()
                                    except Exception:
                                        pass
                        except Exception:
                            pass

                # Phase B: insert replacement content.
                for ed in by_page[page_idx]:
                    try:
                        if ed.kind in ("text_add", "text_edit", "text_move"):
                            if ed.new_rect is None or not ed.text:
                                continue
                            rect = fitz.Rect(*ed.new_rect)
                            fs = ed.font_size or 11.0
                            color = ed.color or (0, 0, 0)
                            # Normalise the colour to a 0..1 tuple as
                            # fitz expects.
                            if any(c > 1 for c in color):
                                color = tuple(c / 255.0 for c in color)
                            # E8: pick a real Base-14 code matching the
                            # original style (serif/sans/mono + b/i) so
                            # edited & moved text blends in instead of
                            # collapsing to plain Helvetica.
                            _ex = getattr(ed, "extra", None) or {}
                            font_name = _pymupdf_base14(
                                ed.font_name,
                                bool(_ex.get("bold")),
                                bool(_ex.get("italic")))

                            # If the record carries rich HTML from the
                            # formatting toolbar, try insert_htmlbox
                            # first so bold / italic / colours /
                            # alignment survive.  Fall back to
                            # insert_textbox on any failure (older
                            # PyMuPDF versions lack insert_htmlbox).
                            # Snap an arbitrary rotation angle to the
                            # nearest multiple of 90° (PyMuPDF's
                            # insert_textbox / insert_image only accept
                            # 0, 90, 180, 270).
                            rot_deg = int(round(
                                (getattr(ed, "rotation", 0.0) or 0.0)
                                / 90.0) * 90) % 360

                            inserted_via_html = False
                            if ed.html:
                                try:
                                    if hasattr(page, "insert_htmlbox"):
                                        page.insert_htmlbox(rect, ed.html)
                                        inserted_via_html = True
                                except Exception:
                                    inserted_via_html = False
                            if inserted_via_html:
                                processed += 1
                                self.progress.emit(
                                    min(99, int(processed * 100 / total)))
                                continue

                            # insert_textbox wraps text to fit; if it
                            # reports overflow (-1) fall back to a
                            # single-line insert_text at the top of the
                            # rect so the user at least sees their edit.
                            try:
                                res = page.insert_textbox(
                                    rect, ed.text,
                                    fontsize=fs,
                                    fontname=font_name,
                                    color=color,
                                    align=0,
                                    rotate=rot_deg,
                                )
                            except Exception:
                                # Unknown font name — retry with helv.
                                res = page.insert_textbox(
                                    rect, ed.text,
                                    fontsize=fs,
                                    fontname="helv",
                                    color=color,
                                    align=0,
                                    rotate=rot_deg,
                                )
                            if isinstance(res, (int, float)) and res < 0:
                                try:
                                    page.insert_text(
                                        fitz.Point(rect.x0, rect.y0 + fs),
                                        ed.text,
                                        fontsize=fs, fontname=font_name,
                                        color=color,
                                        rotate=rot_deg,
                                    )
                                except Exception:
                                    page.insert_text(
                                        fitz.Point(rect.x0, rect.y0 + fs),
                                        ed.text,
                                        fontsize=fs, fontname="helv",
                                        color=color,
                                        rotate=rot_deg,
                                    )
                        elif ed.kind in ("image_add", "image_replace",
                                          "image_move"):
                            if ed.new_rect is None:
                                continue
                            rect = fitz.Rect(*ed.new_rect)
                            img_rot = int(round(
                                (getattr(ed, "rotation", 0.0) or 0.0)
                                / 90.0) * 90) % 360
                            if ed.image_path and os.path.isfile(ed.image_path):
                                try:
                                    page.insert_image(
                                        rect, filename=ed.image_path,
                                        keep_proportion=True,
                                        rotate=img_rot,
                                    )
                                except Exception:
                                    # Older PyMuPDF builds may not
                                    # accept the rotate kwarg on
                                    # insert_image — retry without.
                                    page.insert_image(
                                        rect, filename=ed.image_path,
                                        keep_proportion=True,
                                    )
                        elif ed.kind == "stamp_add":
                            if ed.new_rect is None or not ed.text:
                                continue
                            rect = fitz.Rect(*ed.new_rect)
                            fs = ed.font_size or 36.0
                            color = ed.color or (255, 0, 0)
                            if any(c > 1 for c in color):
                                color = tuple(c / 255.0 for c in color)
                            rot = int(round(
                                (getattr(ed, "rotation", 0.0) or 0.0)
                                / 90.0) * 90) % 360
                            # Draw stamp border rect
                            shape = page.new_shape()
                            shape.draw_rect(rect)
                            shape.finish(
                                color=color, width=2.0,
                                fill=None)
                            shape.commit()
                            # Insert stamp text centred in the rect
                            try:
                                page.insert_textbox(
                                    rect, ed.text,
                                    fontsize=fs,
                                    fontname="helv-b",
                                    color=color,
                                    align=1,  # centre
                                    rotate=rot,
                                )
                            except Exception:
                                page.insert_textbox(
                                    rect, ed.text,
                                    fontsize=fs,
                                    fontname="helv",
                                    color=color,
                                    align=1,
                                )

                        elif ed.kind == "shape_add":
                            if ed.new_rect is None:
                                continue
                            rect = fitz.Rect(*ed.new_rect)
                            s_color = ed.stroke_color or (0, 0, 0)
                            if any(c > 1 for c in s_color):
                                s_color = tuple(c / 255.0 for c in s_color)
                            f_color = None
                            if ed.fill_color is not None:
                                f_color = ed.fill_color
                                if any(c > 1 for c in f_color):
                                    f_color = tuple(c / 255.0 for c in f_color)
                            sw = ed.stroke_width or 2.0
                            sstyle = getattr(ed, "stroke_style",
                                             "solid")
                            stype = ed.shape_type or "rect"
                            fo = 1.0
                            if f_color is not None:
                                fo = max(0.0, min(
                                    1.0,
                                    (getattr(ed, "fill_opacity", 80)
                                     or 80) / 255.0))
                            ls = ed.line_start or (rect.x0, rect.y0)
                            le = ed.line_end or (rect.x1, rect.y1)

                            def _shape_fallback():
                                """Legacy baked-content path — used only
                                if the annotation API is unavailable, so
                                Save As can never fail outright."""
                                dz = {"dash": "[4 3] 0",
                                      "dot": "[1 3] 0"}.get(sstyle)
                                shp = page.new_shape()
                                if stype == "rect":
                                    shp.draw_rect(rect)
                                elif stype == "circle":
                                    shp.draw_oval(rect)
                                else:
                                    shp.draw_line(fitz.Point(*ls),
                                                  fitz.Point(*le))
                                shp.finish(
                                    color=s_color, width=sw,
                                    fill=f_color, dashes=dz,
                                    fill_opacity=fo,
                                    closePath=(stype in ("rect",
                                                         "circle")))
                                shp.commit()
                                if stype == "arrow":
                                    import math
                                    hl = min(12, sw * 4)
                                    ends = [(le, math.atan2(
                                        le[1] - ls[1], le[0] - ls[0]))]
                                    if getattr(ed, "arrow_both", False):
                                        ends.append((ls, math.atan2(
                                            ls[1] - le[1],
                                            ls[0] - le[0])))
                                    s2 = page.new_shape()
                                    for (tx, ty), ang in ends:
                                        for da in (math.pi * 0.85,
                                                   -math.pi * 0.85):
                                            s2.draw_line(
                                                fitz.Point(tx, ty),
                                                fitz.Point(
                                                    tx + hl
                                                    * math.cos(ang + da),
                                                    ty + hl
                                                    * math.sin(ang
                                                               + da)))
                                    s2.finish(color=s_color, width=sw)
                                    s2.commit()

                            # E7: write shapes as real, editable PDF
                            # annotations (Square / Circle / Line).
                            try:
                                annot = None
                                if stype == "rect":
                                    annot = page.add_rect_annot(rect)
                                elif stype == "circle":
                                    annot = page.add_circle_annot(rect)
                                elif stype in ("line", "arrow"):
                                    annot = page.add_line_annot(
                                        fitz.Point(*ls),
                                        fitz.Point(*le))
                                    if (stype == "arrow"
                                            and annot is not None):
                                        try:
                                            start_le = (
                                                fitz.PDF_ANNOT_LE_OpenArrow
                                                if getattr(
                                                    ed, "arrow_both",
                                                    False)
                                                else
                                                fitz.PDF_ANNOT_LE_None)
                                            annot.set_line_ends(
                                                start_le,
                                                fitz.
                                                PDF_ANNOT_LE_OpenArrow)
                                        except Exception:
                                            pass
                                if annot is None:
                                    _shape_fallback()
                                else:
                                    cols = {"stroke": s_color}
                                    if (f_color is not None
                                            and stype in ("rect",
                                                          "circle")):
                                        cols["fill"] = f_color
                                    annot.set_colors(**cols)
                                    try:
                                        dl = {"dash": [4, 3],
                                              "dot": [1, 3]}.get(sstyle)
                                        if dl:
                                            annot.set_border(
                                                width=sw, dashes=dl)
                                        else:
                                            annot.set_border(width=sw)
                                    except Exception:
                                        pass
                                    if (f_color is not None
                                            and fo < 1.0):
                                        try:
                                            annot.set_opacity(fo)
                                        except Exception:
                                            pass
                                    annot.update()
                            except Exception:
                                try:
                                    _shape_fallback()
                                except Exception:
                                    pass

                        elif ed.kind == "redact_add":
                            if ed.new_rect is None:
                                continue
                            rect = fitz.Rect(*ed.new_rect)
                            annot = page.add_redact_annot(rect)
                            if annot is not None:
                                try:
                                    annot.set_colors(fill=(1, 1, 1))
                                    annot.update()
                                except Exception:
                                    pass
                            try:
                                page.apply_redactions(
                                    images=fitz.PDF_REDACT_IMAGE_REMOVE
                                    if hasattr(fitz, "PDF_REDACT_IMAGE_REMOVE")
                                    else 1,
                                )
                            except Exception:
                                try:
                                    page.apply_redactions()
                                except Exception:
                                    pass

                        elif ed.kind == "highlight_add":
                            if ed.new_rect is None:
                                continue
                            rect = fitz.Rect(*ed.new_rect)
                            try:
                                annot = page.add_highlight_annot(rect)
                                if annot is not None:
                                    annot.set_colors(
                                        stroke=(1.0, 0.92, 0.23))
                                    try:
                                        annot.set_opacity(0.4)
                                    except Exception:
                                        pass
                                    annot.update()
                            except Exception:
                                pass

                        elif ed.kind == "note_add":
                            if ed.new_rect is None or not ed.text:
                                continue
                            rect = fitz.Rect(*ed.new_rect)
                            icon = ed.note_icon or "Note"
                            print(f"[BoltPDF] Saving note on page {page_idx+1}: "
                                  f"'{ed.text[:40]}' at ({rect.x0:.1f}, {rect.y0:.1f})",
                                  file=sys.stderr)
                            try:
                                annot = page.add_text_annot(
                                    fitz.Point(rect.x0, rect.y0),
                                    ed.text, icon=icon)
                                if annot:
                                    annot.update()
                                    print(f"[BoltPDF] Note saved successfully",
                                          file=sys.stderr)
                            except Exception as note_ex:
                                print(f"[BoltPDF] Note save with icon failed: {note_ex}",
                                      file=sys.stderr)
                                page.add_text_annot(
                                    fitz.Point(rect.x0, rect.y0),
                                    ed.text)

                        # text_delete / image_delete — redaction already
                        # erased the region; nothing more to insert.
                    except Exception as e:
                        # One broken edit should not abort the whole save.
                        self.error_occurred.emit(
                            f"Page {page_idx + 1}: {e}")

                    processed += 1
                    self.progress.emit(min(99, int(processed * 100 / total)))

            doc.save(self.output_path, deflate=True, garbage=3)
            self.progress.emit(100)
            self.finished_ok.emit(self.output_path)
        except Exception as e:
            self.error_occurred.emit(str(e))
        finally:
            if doc is not None:
                try:
                    doc.close()
                except Exception:
                    pass


# ---------------------------------------------------------------------------
# Custom QGraphicsView — zoom, pan, and text-flow drag-to-select.
# ---------------------------------------------------------------------------
class PDFGraphicsView(QGraphicsView):
    zoom_changed = pyqtSignal(float)

    def __init__(self, scene, parent=None):
        super().__init__(scene, parent)
        self.setRenderHints(QPainter.RenderHint.Antialiasing |
                            QPainter.RenderHint.SmoothPixmapTransform)
        self.setDragMode(QGraphicsView.DragMode.ScrollHandDrag)
        self.setTransformationAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.setAcceptDrops(True)
        self._current_scale = 1.0
        self._tint_mode = "none"

        # Fit-to-screen state
        self._fit_mode = False
        self._fit_current_page = 0

        # Selection state
        self._ocr_mode = False
        self._dragging = False
        self._edit_dragging = False
        self._anchor_idx: int | None = None
        self._extent_idx: int | None = None

    # -- helper: get the owning DocumentTab --------------------------------
    def _tab(self):
        """Return the DocumentTab that owns this view."""
        return self.parent()

    # -- Reading tint (paint-time composite — never re-renders pages) -----
    _TINTS = ("none", "night", "sepia", "warm", "dim")

    def set_tint(self, mode: str):
        mode = mode if mode in self._TINTS else "none"
        if mode == self._tint_mode:
            return
        self._tint_mode = mode
        vp = self.viewport()
        if vp:
            vp.update()

    def drawForeground(self, painter, rect):
        """Composite a single translucent (or Difference) fill over the
        exposed region after items are painted.  O(dirty pixels), GPU/
        raster friendly, and completely independent of the multi-process
        render pool and the on-disk page cache."""
        super().drawForeground(painter, rect)
        mode = getattr(self, "_tint_mode", "none")
        if mode == "none":
            return
        painter.save()
        if mode == "night":
            # white − pixel = full colour inversion of everything already
            # drawn (pages, text, dark-grey backdrop) in one fill.
            painter.setCompositionMode(
                QPainter.CompositionMode.CompositionMode_Difference)
            painter.fillRect(rect, Qt.GlobalColor.white)
        elif mode == "sepia":
            painter.setCompositionMode(
                QPainter.CompositionMode.CompositionMode_Multiply)
            painter.fillRect(rect, QColor(244, 232, 198))
        elif mode == "warm":
            painter.fillRect(rect, QColor(255, 188, 110, 38))
        elif mode == "dim":
            painter.fillRect(rect, QColor(0, 0, 0, 80))
        painter.restore()

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                if url.toLocalFile().lower().endswith(".pdf"):
                    event.acceptProposedAction()
                    return
        super().dragEnterEvent(event)

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            super().dragMoveEvent(event)

    def dropEvent(self, event):
        paths = []
        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if path.lower().endswith(".pdf"):
                paths.append(path)
        if paths:
            # Walk up to find the main window and open tabs there
            w = self.window()
            for path in paths:
                w.open_pdf_in_new_tab(path)
        else:
            super().dropEvent(event)

    def set_ocr_mode(self, enabled: bool):
        self._ocr_mode = enabled
        if enabled:
            self.setDragMode(QGraphicsView.DragMode.NoDrag)
            self.setCursor(Qt.CursorShape.IBeamCursor)
        else:
            self.setDragMode(QGraphicsView.DragMode.ScrollHandDrag)
            self.unsetCursor()
            self._dragging = False

    # -- zoom / scroll ---------------------------------------------------
    def wheelEvent(self, event: QWheelEvent):
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            delta = event.angleDelta().y()
            factor = 1.15 if delta > 0 else 1 / 1.15
            new_scale = self._current_scale * factor
            if 0.1 <= new_scale <= 10.0:
                self.scale(factor, factor)
                self._current_scale = new_scale
                self.zoom_changed.emit(self._current_scale)
        elif self._fit_mode:
            delta = event.angleDelta().y()
            if delta != 0:
                # Step relative to the page currently in view so a
                # prior scrollbar drag is respected (no jump back).
                self.fit_step(1 if delta < 0 else -1)
            event.accept()
        else:
            super().wheelEvent(event)

    def set_scale(self, target: float):
        if 0.1 <= target <= 10.0:
            factor = target / self._current_scale
            self.scale(factor, factor)
            self._current_scale = target
            self.zoom_changed.emit(self._current_scale)

    # -- fit-to-screen ---------------------------------------------------
    def set_fit_mode(self, enabled: bool):
        self._fit_mode = enabled
        if enabled:
            self._fit_current_page = self._detect_current_page()
            self._fit_to_page(self._fit_current_page)

    def _detect_current_page(self) -> int:
        tab = self._tab()
        if not tab or not tab._page_positions:
            return 0
        vp = self.mapToScene(self.viewport().rect()).boundingRect()
        center_y = vp.center().y()
        best, best_d = 0, float("inf")
        for idx in range(tab._num_pages):
            pc = tab._page_positions[idx] + tab._page_heights[idx] / 2
            d = abs(center_y - pc)
            if d < best_d:
                best_d = d
                best = idx
        return best

    def _fit_to_page(self, page_idx: int):
        tab = self._tab()
        if not tab or page_idx not in tab._page_positions:
            return
        page_y = tab._page_positions[page_idx]
        page_h = tab._page_heights[page_idx]
        # Use the authoritative per-page metrics (spread-aware) instead
        # of scanning the scene by Y — in two-page-spread layout two
        # pages share a Y, so the old scan grabbed the wrong page and
        # mis-centred right-hand pages.
        page_w = tab._page_widths.get(page_idx, 0)
        page_x = tab._page_x.get(page_idx, 0.0)
        if page_w == 0:
            return

        vw = self.viewport().width() - 20
        vh = self.viewport().height() - 20
        if page_w <= 0 or page_h <= 0:
            return

        scale_x = vw / page_w
        scale_y = vh / page_h
        target = min(scale_x, scale_y)

        self.resetTransform()
        self._current_scale = target
        self.scale(target, target)
        self.zoom_changed.emit(self._current_scale)

        page_center = QPointF(page_x + page_w / 2,
                              page_y + page_h / 2)
        self.centerOn(page_center)

    def fit_go_to_page(self, page_idx: int):
        tab = self._tab()
        if not tab:
            return
        page_idx = max(0, min(page_idx, tab._num_pages - 1))
        self._fit_current_page = page_idx
        self._fit_to_page(page_idx)

    def fit_step(self, direction: int):
        """Move one page in fit mode, relative to the page ACTUALLY in
        view — not a stored counter.  This keeps wheel/keyboard paging
        in sync after the user has dragged the scrollbar (which moves
        the view without touching the counter)."""
        self.fit_go_to_page(self._detect_current_page() + direction)

    # -- text-flow selection ---------------------------------------------
    def _nearest_word_index(self, scene_pos: QPointF) -> int | None:
        tab = self._tab()
        if not tab or not tab._all_words:
            return None

        best_idx = None
        best_dist = float("inf")
        for w in tab._all_words:
            center = w.rect().center()
            dx = scene_pos.x() - center.x()
            dy = scene_pos.y() - center.y()
            d = dx * dx + dy * dy
            if d < best_dist:
                best_dist = d
                best_idx = w.word_index
        return best_idx

    def _apply_selection(self):
        tab = self._tab()
        if not tab or self._anchor_idx is None or self._extent_idx is None:
            return
        lo = min(self._anchor_idx, self._extent_idx)
        hi = max(self._anchor_idx, self._extent_idx)
        for w in tab._all_words:
            w.set_selected(lo <= w.word_index <= hi)

    def mousePressEvent(self, event):
        # Image-select mode: toggle overlays on click
        tab = self._tab()
        if (tab and tab._image_select_mode
                and event.button() == Qt.MouseButton.LeftButton):
            scene_pos = self.mapToScene(event.pos())
            if tab.handle_image_overlay_click(scene_pos):
                event.accept()
                return

        # Edit mode: if the user has armed an "add" action, route the
        # click to DocumentTab.handle_edit_mode_click so a new text /
        # image gets queued at the clicked position.  Otherwise, a
        # click on existing text opens an inline editor so the user
        # can retype it live.
        if (tab and tab._edit_mode
                and event.button() == Qt.MouseButton.LeftButton):
            scene_pos = self.mapToScene(event.pos())
            drag_actions = ("redact", "highlight", "multiselect",
                            "shape_rect", "shape_circle",
                            "shape_line", "shape_arrow")
            if tab._edit_action in ('add_text', 'add_image',
                                     'stamp', 'note'):
                if tab.handle_edit_mode_click(scene_pos):
                    event.accept()
                    return
            elif tab._edit_action in drag_actions:
                # Start a drag for shape/redact — record the start
                tab.handle_edit_mode_click(scene_pos)
                self._edit_dragging = True
                event.accept()
                return
            elif tab._edit_action == 'transform':
                # If the click landed on one of the transform handles
                # or the body catcher of the active selection, let Qt
                # deliver it to that item normally so drag events flow.
                item_under = self.itemAt(event.pos())
                if tab._active_transform is not None:
                    sel = tab._active_transform
                    if (item_under is sel.body
                            or item_under in sel.handles.values()):
                        super().mousePressEvent(event)
                        return
                # Otherwise try to start (or replace) a selection at
                # the clicked position.
                if tab.try_begin_transform_at(scene_pos):
                    # Don't accept the event — let it propagate so the
                    # just-created TransformBody can receive the same
                    # press and immediately enter "move" drag mode.
                    super().mousePressEvent(event)
                    return
                # Empty click → dismiss any lingering selection.
                if tab._active_transform is not None:
                    try:
                        tab._active_transform.dismiss()
                    except Exception:
                        pass
                    tab._active_transform = None
                event.accept()
                return
            elif tab._edit_action is None:
                # Don't hijack clicks that landed inside an already-open
                # inline editor's proxy widget — let Qt route them to
                # the QLineEdit so the user can position their caret.
                item_under = self.itemAt(event.pos())
                if (tab._inline_editor is not None
                        and item_under is tab._inline_editor):
                    super().mousePressEvent(event)
                    return
                # Clicks on an already-edited text block should reopen
                # the inline editor on THAT record so the user can keep
                # refining it (instead of getting a brand-new editor on
                # the hidden original).
                if tab.try_begin_inline_text_reedit_at(scene_pos):
                    event.accept()
                    return
                # Clicks on other existing live-preview items (image
                # replacements, etc.) should NOT reopen a fresh editor
                # on the hidden original span — just ignore them.
                if item_under is not None and item_under in tab._live_edit_items:
                    super().mousePressEvent(event)
                    return
                if tab.try_begin_inline_text_edit_at(scene_pos):
                    event.accept()
                    return

        if self._ocr_mode and event.button() == Qt.MouseButton.LeftButton:
            scene_pos = self.mapToScene(event.pos())
            idx = self._nearest_word_index(scene_pos)
            if idx is not None:
                shift = bool(event.modifiers() & Qt.KeyboardModifier.ShiftModifier)
                if shift and self._anchor_idx is not None:
                    self._extent_idx = idx
                else:
                    self._anchor_idx = idx
                    self._extent_idx = idx
                self._dragging = True
                self._apply_selection()
            event.accept()
            return
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self._dragging and self._ocr_mode:
            scene_pos = self.mapToScene(event.pos())
            idx = self._nearest_word_index(scene_pos)
            if idx is not None and idx != self._extent_idx:
                self._extent_idx = idx
                self._apply_selection()
            event.accept()
            return
        # Edit-mode drag for shapes / redact: show rubber-band preview
        if getattr(self, '_edit_dragging', False):
            tab = self._tab()
            if tab and tab._edit_mode:
                scene_pos = self.mapToScene(event.pos())
                tab.update_drag_preview(scene_pos)
                event.accept()
                return
        # Edit-mode hover: flip cursor to an I-beam whenever the
        # pointer is over editable text so the user knows a single
        # click will open the inline editor there.  Only active when
        # no other tool (add_text / add_image / transform) is armed
        # and no inline edit is already open.
        tab = self._tab()
        if (tab and tab._edit_mode
                and tab._edit_action is None
                and tab._inline_editor is None):
            try:
                scene_pos = self.mapToScene(event.pos())
                if tab.scene_point_over_editable_text(scene_pos):
                    if self.cursor().shape() != Qt.CursorShape.IBeamCursor:
                        self.setCursor(Qt.CursorShape.IBeamCursor)
                else:
                    if self.cursor().shape() != Qt.CursorShape.ArrowCursor:
                        self.setCursor(Qt.CursorShape.ArrowCursor)
            except Exception:
                pass
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if self._dragging and event.button() == Qt.MouseButton.LeftButton:
            self._dragging = False
            event.accept()
            return
        # Finish shape / redact drag
        if (getattr(self, '_edit_dragging', False)
                and event.button() == Qt.MouseButton.LeftButton):
            self._edit_dragging = False
            tab = self._tab()
            if tab and tab._edit_mode:
                scene_pos = self.mapToScene(event.pos())
                # Remove rubber band
                rubber = getattr(tab, '_drag_rubber', None)
                if rubber is not None:
                    try:
                        tab._scene.removeItem(rubber)
                    except Exception:
                        pass
                    tab._drag_rubber = None
                tab.handle_edit_drag_release(scene_pos)
            event.accept()
            return
        super().mouseReleaseEvent(event)

    def clear_selection(self):
        tab = self._tab()
        if tab:
            for w in tab._all_words:
                w.set_selected(False)
        self._anchor_idx = None
        self._extent_idx = None


# ---------------------------------------------------------------------------
# InlineFormatToolbar — floats above the inline text editor so users can
# apply font/size/bold/italic/underline/strikethrough/alignment/colours
# while editing a paragraph.  Clicks on it don't steal edit-mode focus.
# ---------------------------------------------------------------------------
class InlineFormatToolbar(QFrame):
    """Compact formatting toolbar for the inline text editor.

    Parent must be the QGraphicsView's viewport (or any widget that
    overlays the PDF view), so it appears above whatever is being
    edited.  Call ``bind(editor)`` right after creation so the toolbar
    knows which QTextEdit to drive, then ``sync_from_editor()`` to
    pull the current selection's formatting back onto the controls.
    """

    # Emitted when the user hits the "Done" button — the host
    # DocumentTab uses this to commit the edit.
    done_clicked = pyqtSignal()
    cancel_clicked = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self._editor: QTextEdit | None = None
        self._suppress = False  # guard against sync→control→sync loops

        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setFrameShadow(QFrame.Shadow.Raised)
        self.setAutoFillBackground(True)
        # Clicks on the toolbar (including its padding / gaps) must not
        # bleed through to the QGraphicsView underneath — otherwise
        # they'd be treated as "click outside the inline editor" and
        # commit / cancel the edit prematurely.
        self.setAttribute(Qt.WidgetAttribute.WA_NoMousePropagation, True)
        self.setStyleSheet("""
            InlineFormatToolbar {
                background: #f4f4f4;
                border: 1px solid #888;
                border-radius: 4px;
            }
            QToolButton, QPushButton {
                background: #ffffff;
                border: 1px solid #bbb;
                border-radius: 3px;
                padding: 2px 6px;
                min-width: 20px;
            }
            QToolButton:checked, QPushButton:checked {
                background: #c0d8f5;
                border: 1px solid #5a8fd0;
            }
            QToolButton:hover, QPushButton:hover {
                background: #e8f0fc;
            }
            QFontComboBox, QSpinBox {
                background: #ffffff;
                border: 1px solid #bbb;
                padding: 1px 2px;
            }
            QLabel { color: #222; }
        """)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(4, 3, 4, 3)
        layout.setSpacing(3)

        # Font family
        self.font_combo = QFontComboBox()
        self.font_combo.setMaximumWidth(150)
        self.font_combo.setEditable(True)
        layout.addWidget(self.font_combo)

        # Font size
        self.size_spin = QSpinBox()
        self.size_spin.setRange(4, 144)
        self.size_spin.setValue(12)
        self.size_spin.setFixedWidth(50)
        layout.addWidget(self.size_spin)

        def _sep():
            s = QFrame()
            s.setFrameShape(QFrame.Shape.VLine)
            s.setFrameShadow(QFrame.Shadow.Sunken)
            layout.addWidget(s)

        _sep()

        # Bold / Italic / Underline / Strike
        self.bold_btn = QToolButton()
        self.bold_btn.setText("B")
        f = QFont(); f.setBold(True)
        self.bold_btn.setFont(f)
        self.bold_btn.setCheckable(True)
        self.bold_btn.setToolTip("Bold (Ctrl+B)")
        self.bold_btn.setShortcut(QKeySequence("Ctrl+B"))
        layout.addWidget(self.bold_btn)

        self.italic_btn = QToolButton()
        self.italic_btn.setText("I")
        f = QFont(); f.setItalic(True)
        self.italic_btn.setFont(f)
        self.italic_btn.setCheckable(True)
        self.italic_btn.setToolTip("Italic (Ctrl+I)")
        self.italic_btn.setShortcut(QKeySequence("Ctrl+I"))
        layout.addWidget(self.italic_btn)

        self.under_btn = QToolButton()
        self.under_btn.setText("U")
        f = QFont(); f.setUnderline(True)
        self.under_btn.setFont(f)
        self.under_btn.setCheckable(True)
        self.under_btn.setToolTip("Underline (Ctrl+U)")
        self.under_btn.setShortcut(QKeySequence("Ctrl+U"))
        layout.addWidget(self.under_btn)

        self.strike_btn = QToolButton()
        self.strike_btn.setText("S")
        f = QFont(); f.setStrikeOut(True)
        self.strike_btn.setFont(f)
        self.strike_btn.setCheckable(True)
        self.strike_btn.setToolTip("Strikethrough")
        layout.addWidget(self.strike_btn)

        _sep()

        # Alignment
        self.align_left = QToolButton()
        self.align_left.setText("L")
        self.align_left.setToolTip("Align left")
        self.align_left.setCheckable(True)
        layout.addWidget(self.align_left)

        self.align_center = QToolButton()
        self.align_center.setText("C")
        self.align_center.setToolTip("Align center")
        self.align_center.setCheckable(True)
        layout.addWidget(self.align_center)

        self.align_right = QToolButton()
        self.align_right.setText("R")
        self.align_right.setToolTip("Align right")
        self.align_right.setCheckable(True)
        layout.addWidget(self.align_right)

        self.align_just = QToolButton()
        self.align_just.setText("J")
        self.align_just.setToolTip("Justify")
        self.align_just.setCheckable(True)
        layout.addWidget(self.align_just)

        _sep()

        # Font colour
        self.fg_btn = QPushButton("A")
        fg_f = QFont(); fg_f.setBold(True)
        self.fg_btn.setFont(fg_f)
        self.fg_btn.setToolTip("Font colour")
        self._fg_color = QColor(0, 0, 0)
        self._refresh_fg_swatch()
        layout.addWidget(self.fg_btn)

        # Highlight colour
        self.hl_btn = QPushButton("H")
        self.hl_btn.setToolTip("Highlight colour")
        self._hl_color = QColor(255, 255, 0)
        self._refresh_hl_swatch()
        layout.addWidget(self.hl_btn)

        # Background colour
        self.bg_btn = QPushButton("Bg")
        self.bg_btn.setToolTip("Paragraph background")
        self._bg_color = QColor(255, 255, 255)
        self._refresh_bg_swatch()
        layout.addWidget(self.bg_btn)

        # Transparent background toggle
        self.trans_btn = QToolButton()
        self.trans_btn.setText("Trans")
        self.trans_btn.setCheckable(True)
        self.trans_btn.setToolTip("Transparent background")
        layout.addWidget(self.trans_btn)

        _sep()

        # Done / Cancel
        self.done_btn = QPushButton("Done")
        self.done_btn.setStyleSheet(
            "QPushButton { background: #2e7d32; color: white;"
            " font-weight: bold; border: 1px solid #1b5e20; }"
            "QPushButton:hover { background: #388e3c; }")
        self.done_btn.setToolTip("Apply the edit (Ctrl+Enter)")
        layout.addWidget(self.done_btn)

        self.cancel_btn = QPushButton("Cancel")
        self.cancel_btn.setStyleSheet(
            "QPushButton { background: #c62828; color: white;"
            " font-weight: bold; border: 1px solid #8e1515; }"
            "QPushButton:hover { background: #d32f2f; }")
        self.cancel_btn.setToolTip("Discard the edit (Esc)")
        layout.addWidget(self.cancel_btn)

        self.adjustSize()

        # Wire button signals — editor hook-up happens in bind().
        self.done_btn.clicked.connect(self.done_clicked.emit)
        self.cancel_btn.clicked.connect(self.cancel_clicked.emit)

    # -- editor integration -----------------------------------------------
    def bind(self, editor: QTextEdit):
        self._editor = editor
        self.font_combo.currentFontChanged.connect(self._on_family)
        self.size_spin.valueChanged.connect(self._on_size)
        self.bold_btn.toggled.connect(self._on_bold)
        self.italic_btn.toggled.connect(self._on_italic)
        self.under_btn.toggled.connect(self._on_under)
        self.strike_btn.toggled.connect(self._on_strike)
        self.align_left.clicked.connect(
            lambda: self._set_align(Qt.AlignmentFlag.AlignLeft))
        self.align_center.clicked.connect(
            lambda: self._set_align(Qt.AlignmentFlag.AlignHCenter))
        self.align_right.clicked.connect(
            lambda: self._set_align(Qt.AlignmentFlag.AlignRight))
        self.align_just.clicked.connect(
            lambda: self._set_align(Qt.AlignmentFlag.AlignJustify))
        self.fg_btn.clicked.connect(self._pick_fg)
        self.hl_btn.clicked.connect(self._pick_hl)
        self.bg_btn.clicked.connect(self._pick_bg)
        self.trans_btn.toggled.connect(self._on_transparent)
        editor.cursorPositionChanged.connect(self.sync_from_editor)
        editor.selectionChanged.connect(self.sync_from_editor)

    def sync_from_editor(self):
        """Pull the current selection's character format onto the
        controls so users see what's under the caret."""
        if self._editor is None:
            return
        self._suppress = True
        try:
            cursor = self._editor.textCursor()
            fmt = cursor.charFormat()
            fnt = fmt.font()
            if fnt.family():
                self.font_combo.setCurrentFont(fnt)
            if fnt.pointSize() > 0:
                self.size_spin.setValue(fnt.pointSize())
            self.bold_btn.setChecked(fnt.bold())
            self.italic_btn.setChecked(fnt.italic())
            self.under_btn.setChecked(fnt.underline())
            self.strike_btn.setChecked(fnt.strikeOut())
            align = self._editor.alignment()
            self.align_left.setChecked(
                bool(align & Qt.AlignmentFlag.AlignLeft))
            self.align_center.setChecked(
                bool(align & Qt.AlignmentFlag.AlignHCenter))
            self.align_right.setChecked(
                bool(align & Qt.AlignmentFlag.AlignRight))
            self.align_just.setChecked(
                bool(align & Qt.AlignmentFlag.AlignJustify))
        finally:
            self._suppress = False

    # -- slash helpers for character formats ------------------------------
    def _merge_char(self, fmt: QTextCharFormat):
        ed = self._editor
        if ed is None:
            return
        cursor = ed.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.SelectionType.Document)
        cursor.mergeCharFormat(fmt)
        ed.mergeCurrentCharFormat(fmt)
        ed.setFocus()

    def _on_family(self, fnt: QFont):
        if self._suppress:
            return
        f = QTextCharFormat()
        f.setFontFamilies([fnt.family()])
        self._merge_char(f)

    def _on_size(self, v: int):
        if self._suppress:
            return
        f = QTextCharFormat()
        f.setFontPointSize(float(v))
        self._merge_char(f)

    def _on_bold(self, on: bool):
        if self._suppress:
            return
        f = QTextCharFormat()
        f.setFontWeight(QFont.Weight.Bold if on else QFont.Weight.Normal)
        self._merge_char(f)

    def _on_italic(self, on: bool):
        if self._suppress:
            return
        f = QTextCharFormat()
        f.setFontItalic(on)
        self._merge_char(f)

    def _on_under(self, on: bool):
        if self._suppress:
            return
        f = QTextCharFormat()
        f.setFontUnderline(on)
        self._merge_char(f)

    def _on_strike(self, on: bool):
        if self._suppress:
            return
        f = QTextCharFormat()
        f.setFontStrikeOut(on)
        self._merge_char(f)

    def _set_align(self, flag):
        if self._editor is None:
            return
        self._editor.setAlignment(flag)
        self._editor.setFocus()

    # -- colour pickers ---------------------------------------------------
    def _refresh_fg_swatch(self):
        c = self._fg_color
        self.fg_btn.setStyleSheet(
            "QPushButton {{ color: rgb({r},{g},{b}); font-weight: bold;"
            " background: #ffffff; border: 1px solid #bbb;"
            " border-radius: 3px; padding: 2px 6px; }}"
            .format(r=c.red(), g=c.green(), b=c.blue()))

    def _refresh_hl_swatch(self):
        c = self._hl_color
        self.hl_btn.setStyleSheet(
            "QPushButton {{ background: rgb({r},{g},{b}); color: black;"
            " border: 1px solid #bbb; border-radius: 3px; padding: 2px 6px; }}"
            .format(r=c.red(), g=c.green(), b=c.blue()))

    def _refresh_bg_swatch(self):
        c = self._bg_color
        self.bg_btn.setStyleSheet(
            "QPushButton {{ background: rgb({r},{g},{b}); color: black;"
            " border: 1px solid #bbb; border-radius: 3px; padding: 2px 6px; }}"
            .format(r=c.red(), g=c.green(), b=c.blue()))

    def _pick_fg(self):
        col = QColorDialog.getColor(self._fg_color, self, "Font colour")
        if col.isValid():
            self._fg_color = col
            self._refresh_fg_swatch()
            f = QTextCharFormat()
            f.setForeground(QBrush(col))
            self._merge_char(f)

    def _pick_hl(self):
        col = QColorDialog.getColor(
            self._hl_color, self, "Highlight colour")
        if col.isValid():
            self._hl_color = col
            self._refresh_hl_swatch()
            f = QTextCharFormat()
            f.setBackground(QBrush(col))
            self._merge_char(f)

    def _pick_bg(self):
        col = QColorDialog.getColor(
            self._bg_color, self, "Paragraph background")
        if col.isValid():
            self._bg_color = col
            self._refresh_bg_swatch()
            self.trans_btn.setChecked(False)
            ed = self._editor
            if ed is not None:
                ed.setStyleSheet(
                    "QTextEdit {{ background: rgb({r},{g},{b});"
                    " border: 2px solid #3399ff; padding: 1px; }}"
                    .format(r=col.red(), g=col.green(), b=col.blue()))

    def _on_transparent(self, on: bool):
        if self._suppress:
            return
        ed = self._editor
        if ed is None:
            return
        if on:
            ed.setStyleSheet(
                "QTextEdit { background: transparent;"
                " border: 2px dashed #3399ff; padding: 1px; }")
        else:
            c = self._bg_color
            ed.setStyleSheet(
                "QTextEdit {{ background: rgb({r},{g},{b});"
                " border: 2px solid #3399ff; padding: 1px; }}"
                .format(r=c.red(), g=c.green(), b=c.blue()))

    # -- public helpers for host --------------------------------------------
    def bg_color(self) -> QColor | None:
        """Current paragraph background (or None if transparent)."""
        if self.trans_btn.isChecked():
            return None
        return QColor(self._bg_color)

    def apply_initial_style(self, family: str, size_pt: float,
                            bold: bool, italic: bool,
                            fg: tuple, bg: tuple | None):
        """Seed the toolbar and editor with the detected paragraph
        formatting so the initial look matches the original PDF."""
        self._suppress = True
        try:
            self.font_combo.setCurrentFont(QFont(family))
            self.size_spin.setValue(max(4, int(round(size_pt))))
            self.bold_btn.setChecked(bold)
            self.italic_btn.setChecked(italic)
            self._fg_color = QColor(*fg)
            self._refresh_fg_swatch()
            # Seed the Bg picker with the sampled paragraph background
            # when we have one; the editor itself is already filled
            # opaquely so "transparent" starts OFF.  The user can still
            # flip it on from the toolbar if they want the underlying
            # page to show through.
            if bg is not None:
                self._bg_color = QColor(*bg)
                self._refresh_bg_swatch()
                self.trans_btn.setChecked(False)
            else:
                self.trans_btn.setChecked(False)
        finally:
            self._suppress = False


# ---------------------------------------------------------------------------
# Transform tool — move / scale / rotate handles drawn over a selected
# text box or image so the user can drag it around, resize it from any
# corner or edge, or rotate it via the top knob.
# ---------------------------------------------------------------------------
class TransformHandle(QGraphicsRectItem):
    """One small square drawn at a corner / edge / rotation spot.

    Handles their own mouse drag events and forwards them to the parent
    ``TransformSelection`` which owns the rect math.
    """

    HANDLE_SIZE = 11

    def __init__(self, role: str, selection):
        half = self.HANDLE_SIZE / 2.0
        super().__init__(-half, -half, self.HANDLE_SIZE, self.HANDLE_SIZE)
        self._role = role
        self._selection = selection
        self.setBrush(QBrush(QColor(255, 255, 255)))
        self.setPen(QPen(QColor(51, 153, 255), 1.5))
        self.setZValue(202)
        self.setAcceptedMouseButtons(Qt.MouseButton.LeftButton)
        cursors = {
            "nw": Qt.CursorShape.SizeFDiagCursor,
            "se": Qt.CursorShape.SizeFDiagCursor,
            "ne": Qt.CursorShape.SizeBDiagCursor,
            "sw": Qt.CursorShape.SizeBDiagCursor,
            "n": Qt.CursorShape.SizeVerCursor,
            "s": Qt.CursorShape.SizeVerCursor,
            "e": Qt.CursorShape.SizeHorCursor,
            "w": Qt.CursorShape.SizeHorCursor,
            "rot": Qt.CursorShape.CrossCursor,
            "move": Qt.CursorShape.SizeAllCursor,
        }
        self.setCursor(cursors.get(role, Qt.CursorShape.ArrowCursor))
        if role == "rot":
            # Visually distinct rotation knob — green fill
            self.setBrush(QBrush(QColor(102, 204, 102)))

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self._selection.begin_drag(self._role, event.scenePos())
            event.accept()
            return
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        self._selection.update_drag(event.scenePos())
        event.accept()

    def mouseReleaseEvent(self, event):
        self._selection.end_drag()
        event.accept()


class TransformBody(QGraphicsRectItem):
    """Invisible rectangle covering the whole selection area.

    Lives at a lower Z than the corner handles so clicks in the middle
    of the selection become move-drags while clicks on the corners hit
    the scale handles.
    """

    def __init__(self, selection):
        super().__init__()
        self._selection = selection
        self.setBrush(QBrush(QColor(51, 153, 255, 28)))
        self.setPen(QPen(QColor(51, 153, 255), 1.2, Qt.PenStyle.DashLine))
        self.setZValue(200)
        self.setAcceptedMouseButtons(Qt.MouseButton.LeftButton)
        self.setCursor(Qt.CursorShape.SizeAllCursor)

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self._selection.begin_drag("move", event.scenePos())
            event.accept()
            return
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        self._selection.update_drag(event.scenePos())
        event.accept()

    def mouseReleaseEvent(self, event):
        self._selection.end_drag()
        event.accept()


class TransformSelection:
    """Owns the handles + dashed outline for one selected edit record.

    ``rect`` is kept in scene coordinates (axis-aligned) and mirrored to
    the associated ``EditRecord.new_rect`` (in PDF points) on each drag
    update so Save As always sees the latest geometry.  Rotation is
    stored in degrees and written to ``EditRecord.rotation``.
    """

    def __init__(self, tab, page_idx: int, rect_scene: QRectF,
                 record: "EditRecord",
                 visuals: list,
                 rotation: float = 0.0):
        self.tab = tab
        self.scene = tab._scene
        self.page_idx = page_idx
        self.record = record
        self.rect = QRectF(rect_scene)
        self.rotation = float(rotation)
        # Visual items that should move / scale / rotate with the
        # selection (the live text or pixmap preview).  Cover rects are
        # kept SEPARATE because they stay anchored to the original
        # position.
        self.visuals = list(visuals)
        # Store each visual's initial offset and size relative to the
        # selection rect so multi-item groups (notes, arrows) move/scale
        # correctly.
        self._visual_offsets: list[QPointF] = []
        self._visual_sizes: list[tuple] = []  # (w, h) of each visual
        self._visual_init_lines: list = []    # initial QLineF for line items
        self._initial_rect = QRectF(rect_scene)  # original selection rect
        self._is_group = len(visuals) > 1
        for v in self.visuals:
            try:
                vp = v.pos() if hasattr(v, 'pos') else QPointF(0, 0)
                self._visual_offsets.append(
                    QPointF(vp.x() - rect_scene.x(),
                            vp.y() - rect_scene.y()))
                vbr = v.boundingRect() if hasattr(v, 'boundingRect') else QRectF(0, 0, 0, 0)
                self._visual_sizes.append((vbr.width(), vbr.height()))
                # Store initial line for QGraphicsLineItem
                if isinstance(v, QGraphicsLineItem):
                    self._visual_init_lines.append(v.line())
                else:
                    self._visual_init_lines.append(None)
            except Exception:
                self._visual_offsets.append(QPointF(0, 0))
                self._visual_sizes.append((0, 0))
                self._visual_init_lines.append(None)

        # Invisible body catcher for "click in middle → move"
        self.body = TransformBody(self)
        self.scene.addItem(self.body)

        self.handles: dict = {}
        for role in ("nw", "n", "ne", "e", "se", "s", "sw", "w", "rot"):
            h = TransformHandle(role, self)
            self.scene.addItem(h)
            self.handles[role] = h

        # Drag state
        self._drag_role = None
        self._drag_start_pt: QPointF | None = None
        self._drag_start_rect: QRectF | None = None
        self._drag_start_rot: float = 0.0

        self._sync_overlay()

    # -- geometry helpers -------------------------------------------------
    def _sync_overlay(self):
        import math
        r = QRectF(self.rect)
        # Body rectangle uses its own transform origin at the centre so
        # setRotation rotates about the middle.
        self.body.setRect(0, 0, r.width(), r.height())
        self.body.setPos(r.x(), r.y())
        self.body.setTransformOriginPoint(r.width() / 2, r.height() / 2)
        self.body.setRotation(self.rotation)

        # Handle positions in UN-rotated rect space
        pts = {
            "nw": QPointF(r.left(), r.top()),
            "n": QPointF(r.center().x(), r.top()),
            "ne": QPointF(r.right(), r.top()),
            "e": QPointF(r.right(), r.center().y()),
            "se": QPointF(r.right(), r.bottom()),
            "s": QPointF(r.center().x(), r.bottom()),
            "sw": QPointF(r.left(), r.bottom()),
            "w": QPointF(r.left(), r.center().y()),
            "rot": QPointF(r.center().x(), r.top() - 30),
        }
        # Rotate each point around the rect centre for display
        centre = r.center()
        a = math.radians(self.rotation)
        cos_a = math.cos(a)
        sin_a = math.sin(a)
        for role, pt in pts.items():
            dx = pt.x() - centre.x()
            dy = pt.y() - centre.y()
            nx = centre.x() + dx * cos_a - dy * sin_a
            ny = centre.y() + dx * sin_a + dy * cos_a
            self.handles[role].setPos(nx, ny)

        # Update the visual content items (image / text) to the new
        # rect and rotation.
        if self._is_group:
            self._apply_group_visuals(r)
        else:
            for v in self.visuals:
                self._apply_to_visual(v, r)

    def _apply_group_visuals(self, r: QRectF):
        """Position grouped visuals (notes, stamps) relative to the
        selection rect, scaling their offsets and sizes proportionally."""
        ir = self._initial_rect
        sx = r.width() / max(1, ir.width())
        sy = r.height() / max(1, ir.height())
        # Translation delta from original rect origin to new rect origin
        dx = r.x() - ir.x()
        dy = r.y() - ir.y()
        for i, v in enumerate(self.visuals):
            try:
                off = self._visual_offsets[i]
                ow, oh = self._visual_sizes[i]
                # New position = selection origin + scaled offset
                nx = r.x() + off.x() * sx
                ny = r.y() + off.y() * sy
                # New size = original size * scale
                nw = ow * sx
                nh = oh * sy
                if isinstance(v, QGraphicsPixmapItem):
                    pm = v.pixmap()
                    if pm and not pm.isNull():
                        pw = max(1, pm.width())
                        ph = max(1, pm.height())
                        t = QTransform()
                        t.scale(nw / pw, nh / ph)
                        v.setTransform(t)
                    v.setPos(nx, ny)
                    v.setRotation(self.rotation)
                elif isinstance(v, QGraphicsTextItem):
                    v.setTextWidth(max(10, nw))
                    v.setPos(nx, ny)
                    v.setRotation(self.rotation)
                elif isinstance(v, QGraphicsRectItem):
                    v.setRect(0, 0, max(1, nw), max(1, nh))
                    v.setPos(nx, ny)
                    v.setRotation(self.rotation)
                elif isinstance(v, QGraphicsSimpleTextItem):
                    v.setPos(nx, ny)
                    v.setRotation(self.rotation)
                elif isinstance(v, QGraphicsLineItem):
                    # Lines (arrows): transform initial endpoints relative
                    # to the initial rect, then map to the new rect.
                    init_line = self._visual_init_lines[i]
                    if init_line is not None:
                        x1 = r.x() + (init_line.x1() - ir.x()) * sx
                        y1 = r.y() + (init_line.y1() - ir.y()) * sy
                        x2 = r.x() + (init_line.x2() - ir.x()) * sx
                        y2 = r.y() + (init_line.y2() - ir.y()) * sy
                        v.setLine(x1, y1, x2, y2)
                    v.setRotation(self.rotation)
                else:
                    v.setPos(nx, ny)
                    v.setRotation(self.rotation)
            except Exception as e:
                print(f"[BoltPDF] group visual apply failed: {e}",
                      file=sys.stderr)

    def _apply_to_visual(self, visual, r: QRectF):
        try:
            if isinstance(visual, QGraphicsPixmapItem):
                pm = visual.pixmap()
                if pm is None or pm.isNull():
                    return
                pw = max(1, pm.width())
                ph = max(1, pm.height())
                sx = r.width() / pw
                sy = r.height() / ph
                t = QTransform()
                t.scale(sx, sy)
                visual.setTransform(t)
                visual.setPos(r.x(), r.y())
                visual.setTransformOriginPoint(pw / 2, ph / 2)
                visual.setRotation(self.rotation)
            elif isinstance(visual, QGraphicsTextItem):
                visual.setTextWidth(r.width())
                visual.setPos(r.x(), r.y())
                visual.setTransformOriginPoint(
                    r.width() / 2, r.height() / 2)
                visual.setRotation(self.rotation)
            elif isinstance(visual, QGraphicsLineItem):
                # Line: set endpoints to match the selection rect diagonal
                init_line = self._visual_init_lines[0] if self._visual_init_lines else None
                if init_line is not None:
                    ir = self._initial_rect
                    sx = r.width() / max(1, ir.width())
                    sy = r.height() / max(1, ir.height())
                    x1 = r.x() + (init_line.x1() - ir.x()) * sx
                    y1 = r.y() + (init_line.y1() - ir.y()) * sy
                    x2 = r.x() + (init_line.x2() - ir.x()) * sx
                    y2 = r.y() + (init_line.y2() - ir.y()) * sy
                    visual.setLine(x1, y1, x2, y2)
                visual.setRotation(self.rotation)
            elif isinstance(visual, QGraphicsEllipseItem):
                visual.setRect(0, 0, r.width(), r.height())
                visual.setPos(r.x(), r.y())
                visual.setTransformOriginPoint(
                    r.width() / 2, r.height() / 2)
                visual.setRotation(self.rotation)
            elif isinstance(visual, QGraphicsRectItem):
                # A generic rect item — resize + reposition
                visual.setRect(0, 0, r.width(), r.height())
                visual.setPos(r.x(), r.y())
                visual.setTransformOriginPoint(
                    r.width() / 2, r.height() / 2)
                visual.setRotation(self.rotation)
        except Exception as e:
            print(f"[BoltPDF] transform apply_to_visual failed: {e}",
                  file=sys.stderr)

    # -- drag handlers ----------------------------------------------------
    def begin_drag(self, role: str, scene_pt: QPointF):
        self._drag_role = role
        self._drag_start_pt = QPointF(scene_pt)
        self._drag_start_rect = QRectF(self.rect)
        self._drag_start_rot = self.rotation
        self._moved = False

    def update_drag(self, scene_pt: QPointF):
        if (self._drag_role is None or self._drag_start_pt is None
                or self._drag_start_rect is None):
            return
        self._moved = True
        import math
        r0 = QRectF(self._drag_start_rect)
        dx = scene_pt.x() - self._drag_start_pt.x()
        dy = scene_pt.y() - self._drag_start_pt.y()
        role = self._drag_role
        r = QRectF(r0)

        if role == "move":
            r.translate(dx, dy)
        elif role == "nw":
            r.setTopLeft(QPointF(r0.left() + dx, r0.top() + dy))
        elif role == "ne":
            r.setTopRight(QPointF(r0.right() + dx, r0.top() + dy))
        elif role == "sw":
            r.setBottomLeft(QPointF(r0.left() + dx, r0.bottom() + dy))
        elif role == "se":
            r.setBottomRight(QPointF(r0.right() + dx, r0.bottom() + dy))
        elif role == "n":
            r.setTop(r0.top() + dy)
        elif role == "s":
            r.setBottom(r0.bottom() + dy)
        elif role == "w":
            r.setLeft(r0.left() + dx)
        elif role == "e":
            r.setRight(r0.right() + dx)
        elif role == "rot":
            c = r0.center()
            ang = math.degrees(math.atan2(
                scene_pt.y() - c.y(), scene_pt.x() - c.x())) + 90.0
            # Snap to 15° increments when the user holds Shift (via Qt's
            # global modifier state); otherwise free rotation.
            from PyQt6.QtWidgets import QApplication
            mods = QApplication.keyboardModifiers()
            if mods & Qt.KeyboardModifier.ShiftModifier:
                ang = round(ang / 15.0) * 15.0
            self.rotation = ang
            r = r0  # rect unchanged for rotation

        # Normalise + enforce a minimum size so the user can't collapse
        # the selection into nothingness.
        r = r.normalized()
        if r.width() < 14:
            r.setWidth(14)
        if r.height() < 14:
            r.setHeight(14)
        # E3: snap to nearby objects / page margins on a move drag
        # (only translation — resize/rotate are left exact).
        if role == "move":
            try:
                r = self.tab.apply_snap(r, self.page_idx, self.record)
            except Exception:
                pass
        self.rect = r
        self._sync_overlay()
        self._commit_to_record()

    def end_drag(self):
        self._drag_role = None
        self._drag_start_pt = None
        self._drag_start_rect = None
        self._commit_to_record()
        try:
            self.tab._clear_snap_guides()
        except Exception:
            pass
        # One undo step per completed move/resize/rotate gesture (not
        # per mouse-move, and not for a no-op click that never dragged).
        if getattr(self, "_moved", False):
            self._moved = False
            try:
                self.tab._edit_checkpoint()
            except Exception:
                pass

    def _commit_to_record(self):
        if self.record is None:
            return
        try:
            # The object may have been dragged onto a DIFFERENT page.
            # Re-home it to that page so its coordinates and its
            # page link stay correct — e.g. a sticky note moved to
            # another page now jumps there from the Notes panel.
            pg = self.tab._page_at_scene_rect_center(self.rect)
            if pg is not None:
                self.page_idx = pg
                self.record.page_idx = pg
            new_rect_pt = self.tab._scene_rect_to_pt(
                self.page_idx, self.rect)
            self.record.new_rect = new_rect_pt
            self.record.rotation = self.rotation
            self.tab.status_changed.emit()
        except Exception as e:
            print(f"[BoltPDF] transform commit_to_record failed: {e}",
                  file=sys.stderr)

    # -- teardown ---------------------------------------------------------
    def dismiss(self):
        try:
            self.scene.removeItem(self.body)
        except Exception:
            pass
        for h in list(self.handles.values()):
            try:
                self.scene.removeItem(h)
            except Exception:
                pass
        self.handles.clear()
        self.visuals = []
        try:
            self.tab._clear_snap_guides()
        except Exception:
            pass


# ---------------------------------------------------------------------------
# DocumentTab — one tab per open PDF.  Holds all per-document state.
# ---------------------------------------------------------------------------
class DocumentTab(QWidget):
    """Contains a PDFGraphicsView and all state for one open PDF document."""

    # Emitted whenever this tab's status changes so the main window can
    # update the toolbar (only acts if this tab is the active one).
    status_changed = pyqtSignal()

    # How many pages either side of the viewport to keep in RAM.
    # Must match ``PageRenderer._RENDER_WINDOW`` so the renderer and the
    # disk cache stay in lock-step: anything further than this is
    # unloaded from both memory and disk.
    PAGE_BUFFER = 10

    def __init__(self, parent=None):
        super().__init__(parent)
        self._doc_path = None
        self._num_pages = 0
        self._render_scale = 2.0
        self._renderer = None
        self._ocr_worker = None
        self._page_items: dict[int, QGraphicsPixmapItem] = {}
        self._page_positions: dict[int, float] = {}
        # Per-page scene X origin (0 for continuous/single; non-zero for
        # the right-hand page of a two-page spread).  Centralising it
        # here means _pt_to_scene / _scene_to_pt — which every overlay
        # subsystem already uses — handle spread automatically.
        self._page_x: dict[int, float] = {}
        self._reading_mode: str = "continuous"
        self._page_heights: dict[int, float] = {}
        self._page_widths: dict[int, float] = {}

        # Disk-backed page cache — rendered pages are written here so they
        # can be evicted from RAM and reloaded cheaply on scroll.
        self._cache_dir: str | None = None
        self._cached_pages: set[int] = set()   # pages written to disk

        # OCR state
        self._ocr_active = False
        self._ocr_done_pages: set[int] = set()
        self._page_word_overlays: dict[int, list[WordOverlay]] = {}
        self._all_words: list[WordOverlay] = []
        self._next_word_index = 0
        self._line_breaks: set[int] = set()
        self._scroll_ocr_timer = None

        # Image export selection state
        self._image_select_mode = False
        self._image_overlays: list[ImageOverlay] = []
        self._image_overlays_by_page: dict[int, list[ImageOverlay]] = {}
        # Pages whose embedded-image overlays have been built.  Used by
        # ``_ensure_image_detector_covers`` to decide whether newly
        # visible pages need a fresh detection pass (the detector
        # otherwise finishes and never runs again, leaving late-loaded
        # pages un-clickable in image-select mode).
        self._image_detected_pages: set[int] = set()
        self._img_detector = None
        self._export_worker = None

        # --- Edit Mode ---------------------------------------------------
        # Edit Mode is a stateful toolbar with no overlay layer drawn
        # over the page.  Pending edits go into ``_edit_records`` and
        # get flushed to disk by SaveEditedPdfWorker when the user hits
        # Save As.
        self._edit_mode = False
        self._edit_action = None          # None | 'add_text' | 'add_image'
        self._pending_add_image_path = None
        self._edit_records: list = []     # EditRecord objects
        # -- Undo/redo history (Stage A scaffolding) ---------------------
        # Each entry is a deep-copied snapshot of _edit_records.  The
        # current state lives at _edit_history[_edit_history_index].
        # _edit_restoring suppresses checkpointing while we rebuild the
        # scene during an undo/redo so the restore can't recurse.
        self._edit_history: list = []
        self._edit_history_index: int = -1
        self._edit_restoring: bool = False
        self._EDIT_HISTORY_CAP = 100
        # Debounced crash-recovery journal writer (E5).
        self._journal_timer = QTimer(self)
        self._journal_timer.setSingleShot(True)
        self._journal_timer.setInterval(1200)
        self._journal_timer.timeout.connect(self._write_journal)
        self._save_worker = None
        self._save_progress = None
        # When the user clicks "Replace Image" in edit mode, we flip
        # into a temporary pick state that re-uses the image-select
        # detector's orange overlays so the user can click a real image
        # in the PDF.  This flag tells handle_image_overlay_click to
        # treat the next click as a "pick for replace" instead of the
        # default toggle-for-export behaviour.
        self._edit_image_pick_mode = None  # None | 'replace'
        # Scene items used to render live previews of pending edits so
        # the user sees their change immediately instead of waiting for
        # Save As.  Each entry is a QGraphicsItem already added to the
        # scene — removing them from the scene also dismisses them
        # visually.
        self._live_edit_items: list = []
        # Inline text editor (QGraphicsProxyWidget wrapping a QLineEdit)
        # that appears over a clicked text run.  Only one can be open at
        # a time.
        self._inline_editor = None
        self._inline_editor_ctx = None  # dict with page_idx, orig_rect, etc.
        # Floating formatting toolbar that appears alongside the inline
        # text editor.  Parented to the view's viewport so it doesn't
        # zoom or scroll with the scene content.
        self._inline_editor_toolbar = None
        # Transform tool — currently active selection (or None) and a
        # map from live-preview items back to the EditRecord they
        # represent, so clicks on a previously-edited item re-select
        # the existing record instead of creating a duplicate.
        self._active_transform: "TransformSelection | None" = None
        # E2 multi-select (dedicated "Select" tool — independent of the
        # single-object Transform path).  _multi_sel holds records;
        # _multi_hl holds the dashed highlight items drawn around them.
        self._multi_sel: list = []
        self._multi_hl: list = []
        self._next_group_id: int = 0
        # E3 snapping: snap a moved object's edges/centre to other
        # objects and the page margins, with live guide lines.
        self._snap_enabled: bool = True
        self._snap_guide_items: list = []
        self._item_to_record: dict = {}
        # All visual scene items belonging to a single record, so the
        # Transform tool can move them as a group and Delete can remove
        # them all.  Keyed by id(record).
        self._record_visuals: dict[int, list] = {}
        # Cover rects drawn under text_move / image_move records so the
        # original content appears to be removed.  Keyed by record so
        # exit_edit_mode(discard=True) can sweep them.
        self._record_covers: dict = {}
        # Lazy per-page cache of text block bounding boxes (in PDF
        # points).  Used by the view's hover handler to decide whether
        # the cursor is over editable text so it can flip to an I-beam.
        # Populated on demand via ``_get_edit_text_bboxes``.
        self._edit_text_bboxes: dict = {}

        # --- Comment annotations ---------------------------------------------
        # Per-page list of annotation dicts extracted from the PDF using
        # fitz.  Keyed by 0-based page index.  Used by the NotesPanel in
        # the main window to display comments for the current page.
        self._annotations_by_page: dict[int, list[dict]] = {}
        self._has_annotations: bool = False
        self._annotation_overlays: dict[int, list] = {}

        # --- Search ----------------------------------------------------------
        self._search_highlights: list = []   # QGraphicsRectItems on the scene
        self._search_matches: list = []      # (page_idx, (x0,y0,x1,y1)) tuples
        self._search_current: int = -1       # index into _search_matches
        self._search_query: str = ""
        self._search_worker: SearchWorker | None = None
        self._search_done: bool = True
        self._search_highlight_hwm: int = 0
        # Page → list of indices into _search_matches for O(1) viewport
        # filtering instead of iterating every match.
        self._search_page_index: dict[int, list[int]] = {}
        # Track which pages currently have drawn highlights so scroll
        # refreshes only add/remove the delta.
        self._search_drawn_pages: set[int] = set()
        # Cached visible page to avoid recomputing per-item.
        self._search_vis_page: int = 0
        # Throttle status_changed during active search (emit at most
        # every 300 ms to avoid flooding _sync_toolbar).
        self._search_status_timer: QTimer | None = None
        self._search_status_pending: bool = False
        # Per-page plain text cache — populated once by the first search,
        # reused by every subsequent search so we can skip pages that
        # definitely don't contain the query without calling fitz.
        self._search_text_cache: dict[int, str] | None = None

        # Scene / View
        self._scene = QGraphicsScene(self)
        self._view = PDFGraphicsView(self._scene, self)
        self._view.setBackgroundBrush(Qt.GlobalColor.darkGray)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self._view)

    # -- public accessors for toolbar sync --------------------------------
    @property
    def doc_path(self):
        return self._doc_path

    @property
    def num_pages(self):
        return self._num_pages

    @property
    def ocr_active(self):
        return self._ocr_active

    @property
    def image_select_mode(self):
        return self._image_select_mode

    @property
    def edit_mode(self):
        return self._edit_mode

    @property
    def edit_action(self):
        return self._edit_action

    @property
    def view(self):
        return self._view

    # -- File loading -----------------------------------------------------
    # Target pixel area — scale is chosen so pages stay near this size.
    # 1500×1000 ≈ 1.5 Mpx, sharp enough for a 1080p display.
    _TARGET_PIXELS = 1_500_000

    def _choose_render_scale(self, pt_w, pt_h):
        """Pick a render scale that keeps the rasterised page near
        ``_TARGET_PIXELS`` pixels.  Clamps to [0.5, 3.0]."""
        raw = (self._TARGET_PIXELS / max(pt_w * pt_h, 1)) ** 0.5
        return max(0.5, min(round(raw, 2), 3.0))

    def load_pdf(self, path):
        self._stop_workers()
        self._scene.clear()
        self._page_items.clear()
        self._page_positions.clear()
        self._page_x.clear()
        self._page_heights.clear()
        self._page_widths.clear()
        self._page_word_overlays.clear()
        self._all_words.clear()
        self._line_breaks.clear()
        self._next_word_index = 0
        self._ocr_done_pages.clear()
        self._ocr_active = False
        self._view.set_ocr_mode(False)
        self._image_select_mode = False
        self._image_overlays.clear()
        self._image_overlays_by_page.clear()
        self._image_detected_pages.clear()
        # Reset edit mode state — any queued edits from a previous
        # document are thrown away when a new file is opened.
        self._edit_mode = False
        self._edit_action = None
        self._pending_add_image_path = None
        self._edit_records.clear()
        self._edit_image_pick_mode = None
        # Live preview items from the previous document are gone because
        # the scene was cleared above; just drop the tracking list.
        self._live_edit_items.clear()
        self._inline_editor = None
        self._inline_editor_ctx = None
        if self._inline_editor_toolbar is not None:
            try:
                self._inline_editor_toolbar.deleteLater()
            except Exception:
                pass
        self._inline_editor_toolbar = None
        if self._active_transform is not None:
            try:
                self._active_transform.dismiss()
            except Exception:
                pass
        self._active_transform = None
        self._item_to_record.clear()
        self._record_visuals.clear()
        self._record_covers.clear()
        self._edit_text_bboxes.clear()
        self._annotations_by_page.clear()
        self._annotation_overlays.clear()
        self._has_annotations = False
        self._search_text_cache = None

        # Reset disk cache
        self._cleanup_cache()
        self._cache_dir = tempfile.mkdtemp(prefix="boltpdf_cache_")
        self._cached_pages.clear()

        self._doc_path = path

        doc = pdfium.PdfDocument(path)
        self._num_pages = len(doc)

        # --- Fast dimension estimation -----------------------------------
        # Querying every page's size is very slow on large files (60 s+ for
        # a 700 MB engineering PDF).  Instead, sample a handful of pages and
        # assume all pages share the most-common size.  Any mismatch is
        # corrected when the page actually renders.
        sample_indices = sorted(set(
            [0]
            + [self._num_pages // 4]
            + [self._num_pages // 2]
            + [3 * self._num_pages // 4]
            + [self._num_pages - 1]
        ))
        sample_indices = [i for i in sample_indices if 0 <= i < self._num_pages]

        sample_sizes = []
        for i in sample_indices:
            page = doc[i]
            sample_sizes.append((page.get_width(), page.get_height()))
            page.close()

        # Pick the most common size as the default
        from collections import Counter
        default_size = Counter(
            [(round(w, 1), round(h, 1)) for w, h in sample_sizes]
        ).most_common(1)[0][0]
        default_pt_w, default_pt_h = default_size
        doc.close()

        # Adaptive scale based on page dimensions
        self._render_scale = self._choose_render_scale(
            default_pt_w, default_pt_h)

        # Build the scene layout using the estimated size for every page
        gap = 20
        y_offset = 0.0
        est_w = default_pt_w * self._render_scale
        est_h = default_pt_h * self._render_scale
        spread = (self._reading_mode == "spread"
                  and self._num_pages > 1)
        col2_x = est_w + gap  # x-origin of the right-hand page
        for i in range(self._num_pages):
            if spread:
                # Simple pairs: 0|1, 2|3, …  Left page at x=0, right
                # page at col2_x, both sharing the same row (y).
                if i % 2 == 0:
                    page_x, row_y = 0.0, y_offset
                else:
                    page_x, row_y = col2_x, y_offset
            else:
                page_x, row_y = 0.0, y_offset
            self._page_positions[i] = row_y
            self._page_x[i] = page_x
            self._page_heights[i] = est_h
            self._page_widths[i] = est_w
            rect = self._scene.addRect(page_x, row_y, est_w, est_h)
            rect.setBrush(Qt.GlobalColor.white)
            rect.setPen(QPen(Qt.PenStyle.NoPen))
            rect.setZValue(-1)
            # Advance to the next row after the right page (or after a
            # lone last page) in spread; every page in single column.
            if not spread or i % 2 == 1 or i == self._num_pages - 1:
                y_offset += est_h + gap

        scene_w = (col2_x + est_w) if spread else est_w
        self._scene.setSceneRect(
            QRectF(-20, -20, scene_w + 40, y_offset + 40))
        self._view.set_scale(1.0)
        vsb = self._view.verticalScrollBar()
        if vsb:
            vsb.setValue(0)

        self._renderer = PageRenderer(
            self._doc_path, self._render_scale, self._num_pages, self)
        self._renderer.page_ready.connect(self._on_page_rendered)
        self._renderer.error_occurred.connect(self._on_render_error)
        self._renderer.start()

        # Re-prioritise rendering when the user scrolls
        self._connect_render_scroll_listener()

        # Extract and display clickable comment annotations
        self._load_annotations()

        self.status_changed.emit()
        # Offer to restore any unsaved edits from a previous session.
        QTimer.singleShot(0, self._maybe_recover_edits)

    # -- Comment annotations ------------------------------------------------
    def _load_annotations(self):
        """Scan every page for comment annotations and store them in
        ``_annotations_by_page`` so the NotesPanel can display them."""
        if not self._doc_path:
            return
        try:
            import fitz
        except ImportError:
            print("[BoltPDF] PyMuPDF (fitz) not installed — cannot load annotations",
                  file=sys.stderr)
            return
        try:
            doc = fitz.open(self._doc_path)
        except Exception as exc:
            print(f"[BoltPDF] Could not open PDF for annotations: {exc}",
                  file=sys.stderr)
            return

        # Skip only Link and Widget (form-field) annotations — everything
        # else that carries text is interesting.
        _SKIP_TYPES = set()
        for name in ("PDF_ANNOT_LINK", "PDF_ANNOT_WIDGET",
                      "PDF_ANNOT_POPUP"):
            val = getattr(fitz, name, None)
            if val is not None:
                _SKIP_TYPES.add(val)

        for page_idx in range(len(doc)):
            page = doc[page_idx]
            try:
                annots = page.annots()
            except Exception:
                continue
            if annots is None:
                continue
            for annot in annots:
                try:
                    atype = annot.type[0]
                    aname = annot.type[1]
                except Exception:
                    continue
                if atype in _SKIP_TYPES:
                    continue
                # Collect every text field the annotation exposes
                info = {}
                try:
                    info = annot.info or {}
                except Exception:
                    pass
                content = (info.get("content", "") or "").strip()
                subject = (info.get("subject", "") or "").strip()
                author = (info.get("title", "") or "").strip()
                # Also check the raw /Contents key via annot.info and
                # the popup text via get_text() as a fallback.
                if not content:
                    try:
                        content = (annot.get_text("text") or "").strip()
                    except Exception:
                        pass
                if not content and not subject:
                    continue
                # Store annotation rect for overlay positioning
                try:
                    arect = annot.rect
                    rect_tuple = (arect.x0, arect.y0, arect.x1, arect.y1)
                except Exception:
                    rect_tuple = None
                self._annotations_by_page.setdefault(page_idx, []).append({
                    "author": author,
                    "content": content,
                    "subject": subject,
                    "type": aname,
                    "rect": rect_tuple,
                })
        doc.close()
        self._has_annotations = bool(self._annotations_by_page)
        total = sum(len(v) for v in self._annotations_by_page.values())
        print(f"[BoltPDF] Annotations: {total} notes across "
              f"{len(self._annotations_by_page)} pages",
              file=sys.stderr)

    def _on_page_rendered(self, idx, qimg):
        if qimg.isNull():
            print(f"[BoltPDF] Page {idx + 1}: received null QImage, skipping",
                  file=sys.stderr)
            return

        # If the user has already scrolled this page more than
        # PAGE_BUFFER away, drop it — don't waste disk I/O or RAM
        # persisting something that's already outside the window.
        vis = self.get_visible_page() if self._num_pages else 0
        if abs(idx - vis) > self.PAGE_BUFFER:
            return

        # Persist to disk cache (JPEG — ~15× smaller than BMP, fast I/O)
        if self._cache_dir:
            cache_path = os.path.join(self._cache_dir, f"{idx}.jpg")
            if qimg.save(cache_path, "JPEG", 90):
                self._cached_pages.add(idx)
            else:
                print(f"[BoltPDF] Page {idx + 1}: cache write failed",
                      file=sys.stderr)

        # Show in scene
        self._show_page_pixmap(idx, qimg=qimg)

    def _show_page_pixmap(self, idx, *, qimg=None):
        """Add a QGraphicsPixmapItem for *idx*.  Uses the provided QImage or
        loads from the disk cache.  No-op if the item is already in the scene."""
        if idx in self._page_items:
            return
        if qimg is None:
            qimg = self._load_cached_image(idx)
            if qimg is None:
                return
        pixmap = QPixmap.fromImage(qimg)
        if pixmap.isNull():
            print(f"[BoltPDF] Page {idx + 1}: QPixmap conversion failed",
                  file=sys.stderr)
            return
        item = QGraphicsPixmapItem(pixmap)
        item.setPos(self._page_x.get(idx, 0.0),
                    self._page_positions.get(idx, 0))
        item.setZValue(0)
        self._scene.addItem(item)
        self._page_items[idx] = item
        # Draw readable note card overlays for any annotations on this page
        self._draw_annotation_overlays(idx)

    def _draw_annotation_overlays(self, page_idx: int):
        """Create readable sticky-note card overlays for PDF annotations
        on *page_idx* so they're visible and selectable."""
        notes = self._annotations_by_page.get(page_idx)
        if not notes:
            return
        # Avoid re-drawing if overlays already exist for this page
        existing = getattr(self, '_annotation_overlays', None)
        if existing is None:
            self._annotation_overlays: dict[int, list] = {}
        if page_idx in self._annotation_overlays:
            return
        overlay_items = []
        for note in notes:
            rect_tuple = note.get("rect")
            content = note.get("content", "")
            author = note.get("author", "")
            ntype = note.get("type", "Text")
            if not rect_tuple or not content:
                continue
            x0, y0, x1, y1 = rect_tuple
            # Convert PDF-point position to scene coordinates
            sx, sy = self._pt_to_scene(page_idx, x0, y0)
            # Build a readable card
            note_w = 220.0
            lines = content.count('\n') + 1
            chars_per_line = max(1, int(note_w / 7))
            wrapped_lines = sum(
                max(1, (len(line) + chars_per_line - 1) // chars_per_line)
                for line in content.split('\n'))
            note_h = max(60.0, wrapped_lines * 18.0 + 36.0)
            # Background card — use setPos + local rect so transforms work
            card = QGraphicsRectItem(0, 0, note_w, note_h)
            card.setPos(sx, sy)
            card.setBrush(QBrush(QColor(255, 255, 200, 240)))
            card.setPen(QPen(QColor(200, 180, 60), 1.5))
            card.setZValue(50)
            card.setToolTip(content)
            self._scene.addItem(card)
            overlay_items.append(card)
            # Header bar — use setPos + local rect
            header_h = 22.0
            header = QGraphicsRectItem(0, 0, note_w, header_h)
            header.setPos(sx, sy)
            header.setBrush(QBrush(QColor(255, 230, 100)))
            header.setPen(QPen(Qt.PenStyle.NoPen))
            header.setZValue(51)
            self._scene.addItem(header)
            overlay_items.append(header)
            # Header text — show author or annotation type
            hdr_label = author if author else ntype
            hdr_txt = QGraphicsSimpleTextItem(hdr_label)
            hdr_txt.setFont(QFont("Segoe UI", 9, QFont.Weight.Bold))
            hdr_txt.setBrush(QBrush(QColor(100, 80, 0)))
            hdr_txt.setPos(sx + 6, sy + 3)
            hdr_txt.setZValue(52)
            self._scene.addItem(hdr_txt)
            overlay_items.append(hdr_txt)
            # Body text
            body = QGraphicsTextItem()
            body.setPlainText(content)
            body.setFont(QFont("Segoe UI", 9))
            body.setDefaultTextColor(QColor(60, 50, 0))
            body.setTextWidth(note_w - 12)
            body.setPos(sx + 6, sy + header_h + 4)
            body.setZValue(52)
            self._scene.addItem(body)
            overlay_items.append(body)
        self._annotation_overlays[page_idx] = overlay_items

    def _remove_annotation_overlays(self, page_idx: int):
        """Remove annotation card overlays for a page being unloaded."""
        overlays = getattr(self, '_annotation_overlays', None)
        if overlays is None:
            return
        items = overlays.pop(page_idx, [])
        for it in items:
            try:
                self._scene.removeItem(it)
            except Exception:
                pass

    def _unload_page_pixmap(self, idx):
        """Remove the in-memory pixmap for *idx*, keeping the disk cache."""
        item = self._page_items.pop(idx, None)
        if item is not None:
            self._scene.removeItem(item)
        self._remove_annotation_overlays(idx)

    def _load_cached_image(self, idx) -> QImage | None:
        """Read a previously cached page image from disk."""
        if idx not in self._cached_pages or not self._cache_dir:
            return None
        cache_path = os.path.join(self._cache_dir, f"{idx}.jpg")
        qimg = QImage(cache_path)
        if qimg.isNull():
            return None
        return qimg

    def _manage_viewport_pages(self):
        """Load pages near the viewport from cache, unload distant ones.

        Pages more than ``PAGE_BUFFER`` away from the current view are
        evicted from both the in-memory scene *and* the on-disk cache,
        so memory/disk usage stays bounded regardless of how large the
        PDF is.  The renderer will re-render them if the user scrolls
        back.
        """
        if not self._num_pages:
            return
        vis = self.get_visible_page()
        lo = max(0, vis - self.PAGE_BUFFER)
        hi = min(self._num_pages - 1, vis + self.PAGE_BUFFER)

        # Unload pixmaps outside the buffer window
        for idx in list(self._page_items.keys()):
            if idx < lo or idx > hi:
                self._unload_page_pixmap(idx)

        # Evict disk cache entries outside the buffer window
        for idx in list(self._cached_pages):
            if idx < lo or idx > hi:
                self._cached_pages.discard(idx)
                if self._cache_dir:
                    try:
                        os.unlink(os.path.join(self._cache_dir, f"{idx}.jpg"))
                    except OSError:
                        pass

        # Load pages inside the buffer window from cache
        for idx in range(lo, hi + 1):
            if idx not in self._page_items and idx in self._cached_pages:
                self._show_page_pixmap(idx)

    def _on_render_error(self, msg):
        print(f"[BoltPDF] Render warning: {msg}", file=sys.stderr)

    # -- Navigation / session helpers -------------------------------------
    def current_page_index(self) -> int:
        """0-based index of the page nearest the viewport centre."""
        try:
            return self._view._detect_current_page()
        except Exception:
            return 0

    def current_zoom(self) -> float:
        try:
            return float(getattr(self._view, "_current_scale", 1.0) or 1.0)
        except Exception:
            return 1.0

    def goto_page(self, idx: int, zoom: float | None = None):
        """Centre the view on 0-based page *idx* and refocus the
        renderer — mirrors the toolbar page-jump path so the on-disk
        cache and worker pool stay in lock-step."""
        if zoom is not None:
            try:
                self._view.set_scale(float(zoom))
            except Exception:
                pass
        if not self._page_positions or idx not in self._page_positions:
            return
        try:
            cx = (self._page_x.get(idx, 0.0)
                  + self._page_widths.get(idx, 0.0) / 2.0)
            self._view.centerOn(cx, self._page_positions[idx])
            if self._renderer and self._renderer.isRunning():
                self._renderer.set_focus(idx)
            if self._image_select_mode:
                self._ensure_image_detector_covers(idx)
            self._manage_viewport_pages()
        except Exception:
            pass

    def set_tint(self, mode: str):
        """Apply a reading tint (none/night/sepia/warm/dim) to this
        document's view."""
        try:
            self._view.set_tint(mode)
        except Exception:
            pass

    # -- Zoom -------------------------------------------------------------
    def zoom_in(self):
        self._view.set_scale(self._view._current_scale * 1.25)

    def zoom_out(self):
        self._view.set_scale(self._view._current_scale / 1.25)

    # -- Fit to Screen ---------------------------------------------------
    def set_fit_mode(self, checked: bool):
        self._view.set_fit_mode(checked)

    def refit_if_active(self):
        """Re-fit the current page if fit mode is on (called on resize)."""
        if self._view._fit_mode:
            self._view._fit_to_page(self._view._fit_current_page)

    # -- Rebuild as Images ------------------------------------------------
    def rebuild_as_images(self, main_window):
        if not self._doc_path:
            return

        # PDFium is not thread-safe — pause renderer while rebuild runs
        self._pause_renderer()

        base, ext = os.path.splitext(self._doc_path)
        output_path = f"{base}_optimised.pdf"
        if os.path.exists(output_path):
            n = 2
            while os.path.exists(f"{base}_optimised_{n}.pdf"):
                n += 1
            output_path = f"{base}_optimised_{n}.pdf"

        self._rebuild_progress = QProgressDialog(
            "Optimising PDF...", "Cancel", 0, self._num_pages, main_window
        )
        self._rebuild_progress.setWindowTitle("Optimise PDF")
        self._rebuild_progress.setMinimumDuration(0)
        self._rebuild_progress.setValue(0)
        self._rebuild_progress.setModal(True)

        self._rebuild_worker = RebuildWorker(
            self._doc_path, output_path, self._render_scale, parent=self
        )
        self._rebuild_worker.progress.connect(self._on_rebuild_progress)
        self._rebuild_worker.finished_ok.connect(
            lambda p: self._on_rebuild_done(p, main_window))
        self._rebuild_worker.error_occurred.connect(
            lambda m: self._on_rebuild_error(m, main_window))
        self._rebuild_progress.canceled.connect(self._rebuild_worker.terminate)
        self._rebuild_worker.start()

    def _on_rebuild_progress(self, page_num):
        if hasattr(self, '_rebuild_progress') and self._rebuild_progress:
            self._rebuild_progress.setValue(page_num)

    def _on_rebuild_done(self, output_path, main_window):
        if hasattr(self, '_rebuild_progress') and self._rebuild_progress:
            self._rebuild_progress.close()

        # Resume page rendering now that rebuild is finished
        self._resume_renderer()

        reply = QMessageBox.question(
            main_window, "Optimisation Complete",
            f"Saved to:\n{os.path.basename(output_path)}\n\n"
            "Do you want to open the new file?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply == QMessageBox.StandardButton.Yes:
            main_window.open_pdf_in_new_tab(output_path)

        self.status_changed.emit()

    def _on_rebuild_error(self, msg, main_window):
        if hasattr(self, '_rebuild_progress') and self._rebuild_progress:
            self._rebuild_progress.close()
        # Resume page rendering even after rebuild failure
        self._resume_renderer()
        QMessageBox.critical(main_window, "Rebuild Failed", f"Error:\n{msg}")
        self.status_changed.emit()

    # -- Export Images mode ------------------------------------------------
    def enter_image_select_mode(self):
        """Scan all pages for embedded images and overlay selectable rects."""
        if not self._doc_path or self._num_pages == 0:
            return
        if self._image_select_mode:
            return  # already active

        self._image_select_mode = True
        self._view.setCursor(Qt.CursorShape.CrossCursor)

        # Detect images on all pages in background.
        # Detection only reads page object metadata (bounds, pixel size) —
        # it does not render bitmaps — so it can coexist with the renderer.
        # Starts from the currently visible page and works outward so nearby
        # images are selectable almost immediately.
        vis = self.get_visible_page() if self._num_pages else 0
        all_pages = list(range(self._num_pages))
        self._img_detector = ImageDetectorWorker(
            self._doc_path, all_pages, focus=vis, parent=self)
        self._img_detector.page_images_found.connect(self._on_images_found)
        self._img_detector.error_occurred.connect(self._on_img_detect_error)
        self._img_detector.all_done.connect(self._on_img_detect_done)
        self._img_detector.start()
        self.status_changed.emit()

    def exit_image_select_mode(self):
        """Remove all image overlays and leave selection mode."""
        self._image_select_mode = False
        if self._img_detector and self._img_detector.isRunning():
            self._img_detector.cancel()
            self._img_detector.wait()
        for overlay in self._image_overlays:
            self._scene.removeItem(overlay)
        self._image_overlays.clear()
        self._image_overlays_by_page.clear()
        self._image_detected_pages.clear()
        self._view.unsetCursor()
        if self._ocr_active:
            self._view.setCursor(Qt.CursorShape.IBeamCursor)
        self.status_changed.emit()

    # === Edit Mode ========================================================
    # Edit Mode is a stateful toolbar — it never draws overlays over the
    # page in its idle state.  Actions are triggered one at a time:
    #   * Add Text    : click on a page to place a new text box
    #   * Add Image   : pick a file, then click on a page to drop it in
    #   * Replace Image: temporarily re-enters image-select mode so the
    #                    user can pick one of the real images already in
    #                    the PDF, then prompts for a replacement file
    #   * Save As     : writes every queued EditRecord to a new PDF
    #
    # Nothing in the original PDF is touched until Save As runs.
    # ======================================================================

    # -- coordinate conversion --------------------------------------------
    def _pt_to_scene(self, page_idx: int, x_pt: float, y_pt: float):
        """Convert PDF-point coordinates on *page_idx* to scene (x, y).
        Includes the per-page X origin so two-page-spread layout works
        for every overlay subsystem that routes through here."""
        scale = self._render_scale
        page_y = self._page_positions.get(page_idx, 0)
        page_x = self._page_x.get(page_idx, 0.0)
        return (page_x + x_pt * scale, page_y + y_pt * scale)

    def _scene_to_pt(self, page_idx: int, sx: float, sy: float):
        """Convert scene coordinates to PDF points on *page_idx*."""
        scale = self._render_scale or 1.0
        page_y = self._page_positions.get(page_idx, 0)
        page_x = self._page_x.get(page_idx, 0.0)
        return ((sx - page_x) / scale, (sy - page_y) / scale)

    def _scene_rect_to_pt(self, page_idx: int, rect: QRectF):
        """Translate a scene-space QRectF to a PDF-point tuple."""
        x0, y0 = self._scene_to_pt(page_idx, rect.x(), rect.y())
        x1, y1 = self._scene_to_pt(
            page_idx, rect.x() + rect.width(), rect.y() + rect.height())
        return (x0, y0, x1, y1)

    # -- enter / exit ------------------------------------------------------
    def enter_edit_mode(self):
        """Turn on the editor.  No overlays are drawn — the user picks
        an action from the toolbar and the next click does the work."""
        if not self._doc_path or self._num_pages == 0:
            return
        if self._edit_mode:
            return

        # PyMuPDF is required to save edits.  Bail out up-front with a
        # clear message if it isn't available in this build so the user
        # doesn't discover the problem mid-session.
        try:
            import fitz  # noqa: F401
        except Exception as e:
            QMessageBox.critical(
                self.window(), "Edit mode unavailable",
                "Editing requires PyMuPDF, which isn't available in this "
                "build of BoltPDF.\n\n"
                f"Details: {e}\n\n"
                "If you installed BoltPDF from boltpdf.co.uk, please "
                "update to the latest version.")
            return

        # Edit mode and image-select mode are mutually exclusive by
        # default.  (Replace Image explicitly re-opens image-select.)
        if self._image_select_mode:
            self.exit_image_select_mode()

        self._edit_mode = True
        self._edit_action = None
        self._edit_image_pick_mode = None
        # Baseline snapshot so the first edit becomes an undoable step.
        self._edit_history_reset()
        self._view.setCursor(Qt.CursorShape.ArrowCursor)
        self.status_changed.emit()

    def exit_edit_mode(self, *, discard_edits: bool = False):
        """Leave edit mode.  If *discard_edits* is True, any queued
        changes are thrown away.  Otherwise they're preserved so a
        subsequent Save As picks them up."""
        # Commit any inline text edit first so its changes are kept
        # (unless we're discarding everything anyway, in which case
        # drop it).
        if discard_edits:
            self._cancel_inline_text_edit()
        else:
            self._commit_inline_text_edit()

        # Dismiss any active transform selection (handles + outline).
        if self._active_transform is not None:
            try:
                self._active_transform.dismiss()
            except Exception:
                pass
            self._active_transform = None

        # If the user was in the middle of picking an image to replace,
        # close the image-select overlay as we leave.
        if self._edit_image_pick_mode and self._image_select_mode:
            self.exit_image_select_mode()

        self._edit_mode = False
        self._edit_action = None
        self._pending_add_image_path = None
        self._edit_image_pick_mode = None
        if discard_edits:
            self._edit_records.clear()
            self._clear_recovery()
            # Remove the live-preview items so the page reverts.
            self._teardown_edit_overlays()
        self._view.unsetCursor()
        if self._ocr_active:
            self._view.setCursor(Qt.CursorShape.IBeamCursor)
        self.status_changed.emit()

    # -- Undo/redo infrastructure (wired to UI in later stages) ----------
    def _teardown_edit_overlays(self):
        """Remove every live edit-preview item, cover and map entry from
        the scene WITHOUT touching _edit_records.  Factored verbatim from
        the old exit_edit_mode(discard) path so behaviour is identical;
        additionally clears _record_visuals (the old path leaked those
        dict entries — harmless, but the rebuild path needs a clean
        slate)."""
        for it in self._live_edit_items:
            try:
                self._scene.removeItem(it)
            except Exception:
                pass
        self._live_edit_items.clear()
        for cover in list(self._record_covers.values()):
            try:
                self._scene.removeItem(cover)
            except Exception:
                pass
        self._record_covers.clear()
        self._item_to_record.clear()
        self._record_visuals.clear()
        # Stale multi-select highlights point at now-removed items.
        self.clear_multi_selection()
        self._clear_snap_guides()

    def _snapshot_edit_records(self) -> list:
        """Deep copy so later mutations never alias a history entry."""
        return [copy.deepcopy(r) for r in self._edit_records]

    def _edit_checkpoint(self):
        """Record the current _edit_records state as a new undo step.
        No-op while restoring so undo/redo cannot recurse.  Wired to the
        mutation sites in Stage B/C."""
        if self._edit_restoring:
            return
        if self._edit_history_index < len(self._edit_history) - 1:
            del self._edit_history[self._edit_history_index + 1:]
        self._edit_history.append(self._snapshot_edit_records())
        if len(self._edit_history) > self._EDIT_HISTORY_CAP:
            self._edit_history.pop(0)
        self._edit_history_index = len(self._edit_history) - 1
        # Crash-recovery: persist pending edits shortly after a change.
        if self._doc_path:
            self._journal_timer.start()
        self.status_changed.emit()

    # -- E5: crash recovery ----------------------------------------------
    @staticmethod
    def _rec_to_dict(rec) -> dict:
        return dict(rec.__dict__)

    @staticmethod
    def _rec_from_dict(d: dict):
        r = EditRecord(kind=d.get("kind", "shape_add"),
                        page_idx=d.get("page_idx", 0))
        r.__dict__.update(d)
        return r

    def _recovery_path(self) -> str | None:
        if not self._doc_path:
            return None
        base = os.path.join(
            os.environ.get("LOCALAPPDATA", tempfile.gettempdir()),
            "BoltPDF", "recovery")
        try:
            os.makedirs(base, exist_ok=True)
        except OSError:
            return None
        key = hashlib.sha1(
            os.path.normcase(os.path.abspath(self._doc_path))
            .encode("utf-8", "replace")).hexdigest()
        return os.path.join(base, key + ".json")

    def _write_journal(self):
        p = self._recovery_path()
        if p is None:
            return
        try:
            if not self._edit_records:
                # Nothing pending → drop any stale journal.
                if os.path.isfile(p):
                    os.remove(p)
                return
            payload = {
                "doc": self._doc_path,
                "saved_at": time.time(),
                "records": [self._rec_to_dict(r)
                            for r in self._edit_records],
            }
            tmp = p + ".tmp"
            with open(tmp, "w", encoding="utf-8") as fh:
                json.dump(payload, fh)
            os.replace(tmp, p)
        except (OSError, TypeError, ValueError):
            pass

    def _clear_recovery(self):
        self._journal_timer.stop()
        p = self._recovery_path()
        if p and os.path.isfile(p):
            try:
                os.remove(p)
            except OSError:
                pass

    def _maybe_recover_edits(self):
        """If a crash-recovery journal exists for this document, offer
        to restore the unsaved edits.  Called at the end of load_pdf."""
        p = self._recovery_path()
        if not p or not os.path.isfile(p):
            return
        try:
            with open(p, "r", encoding="utf-8") as fh:
                payload = json.load(fh)
            recs = payload.get("records", [])
        except (OSError, ValueError):
            return
        if not recs:
            self._clear_recovery()
            return
        resp = QMessageBox.question(
            self.window(), "Recover unsaved edits",
            f"BoltPDF found {len(recs)} unsaved edit(s) for this "
            "document from a previous session.\n\nRestore them?",
            QMessageBox.StandardButton.Yes
            | QMessageBox.StandardButton.No)
        if resp != QMessageBox.StandardButton.Yes:
            self._clear_recovery()
            return
        try:
            self._edit_records = [self._rec_from_dict(d) for d in recs]
        except Exception:
            self._clear_recovery()
            return
        if not self._edit_mode:
            self.enter_edit_mode()
        self._rebuild_edit_overlays()
        self._edit_history_reset()
        self.status_changed.emit()

    # -- E5: copy / paste / duplicate ------------------------------------
    def copy_selected_edit(self) -> bool:
        rec = self.selected_record()
        if rec is None:
            return False
        global _EDIT_CLIPBOARD
        _EDIT_CLIPBOARD = copy.deepcopy(rec)
        return True

    def paste_edit(self) -> bool:
        global _EDIT_CLIPBOARD
        if _EDIT_CLIPBOARD is None or not self._edit_mode:
            return False
        clone = copy.deepcopy(_EDIT_CLIPBOARD)
        # Land the paste on the page the user is currently viewing
        # (enables cross-page paste); nudge it so it isn't hidden
        # exactly behind the original.
        try:
            clone.page_idx = self._view._detect_current_page()
        except Exception:
            pass
        off = 14.0
        if clone.new_rect:
            x0, y0, x1, y1 = clone.new_rect
            clone.new_rect = (x0 + off, y0 + off,
                              x1 + off, y1 + off)
        if getattr(clone, "orig_rect", None):
            clone.orig_rect = None  # a paste is a NEW addition
            if clone.kind in ("text_move", "text_edit"):
                clone.kind = "text_add"
            elif clone.kind in ("image_move", "image_replace"):
                clone.kind = "image_add"
        for attr in ("line_start", "line_end"):
            v = getattr(clone, attr, None)
            if v:
                setattr(clone, attr, (v[0] + off, v[1] + off))
        self._edit_records.append(clone)
        self._materialize_record(clone)
        self._reselect_record(clone)
        self._edit_checkpoint()
        self.status_changed.emit()
        return True

    def duplicate_selected_edit(self) -> bool:
        if not self.copy_selected_edit():
            return False
        return self.paste_edit()

    def _edit_history_reset(self):
        """Start a fresh history with the current state as baseline."""
        self._edit_history = [self._snapshot_edit_records()]
        self._edit_history_index = 0

    def can_edit_undo(self) -> bool:
        return self._edit_history_index > 0

    def can_edit_redo(self) -> bool:
        return 0 <= self._edit_history_index < len(self._edit_history) - 1

    def _restore_edit_state(self, snapshot: list):
        """Settle in-progress edits, tear down all overlays, restore the
        records, and rebuild the scene via the central renderer."""
        self._edit_restoring = True
        try:
            if self._active_transform is not None:
                try:
                    self._active_transform.dismiss()
                except Exception:
                    pass
                self._active_transform = None
            if self._inline_editor is not None:
                try:
                    self._cancel_inline_text_edit()
                except Exception:
                    pass
            self._teardown_edit_overlays()
            self._edit_records = [copy.deepcopy(r) for r in snapshot]
            self._rebuild_edit_overlays()
        finally:
            self._edit_restoring = False
        self.status_changed.emit()

    def edit_undo(self):
        if not self.can_edit_undo():
            return
        self._edit_history_index -= 1
        self._restore_edit_state(
            self._edit_history[self._edit_history_index])

    def edit_redo(self):
        if not self.can_edit_redo():
            return
        self._edit_history_index += 1
        self._restore_edit_state(
            self._edit_history[self._edit_history_index])

    def _rebuild_edit_overlays(self):
        """Re-create live previews for every record via the central
        renderer.  Per-kind bodies land in Stage B/C; unknown kinds are
        skipped without crashing."""
        for rec in self._edit_records:
            try:
                self._materialize_record(rec)
            except Exception as e:
                print(f"[BoltPDF] rebuild failed for "
                      f"{getattr(rec, 'kind', '?')}: {e}",
                      file=sys.stderr)

    def _pt_rect_to_scene(self, page_idx: int, pt_rect) -> QRectF:
        """Inverse of _scene_rect_to_pt: PDF-point (x0,y0,x1,y1) tuple →
        scene QRectF.  Exact linear inverse, so a record round-trips to
        a pixel-identical preview."""
        x0, y0, x1, y1 = pt_rect
        sx0, sy0 = self._pt_to_scene(page_idx, x0, y0)
        sx1, sy1 = self._pt_to_scene(page_idx, x1, y1)
        return QRectF(min(sx0, sx1), min(sy0, sy1),
                      abs(sx1 - sx0), abs(sy1 - sy0))

    def _register_visual(self, rec, item, visuals):
        """Add *item* to the scene and into all three tracking maps the
        rest of Edit Mode relies on."""
        self._scene.addItem(item)
        self._live_edit_items.append(item)
        self._item_to_record[item] = rec
        visuals.append(item)

    def selected_record(self):
        """The EditRecord the Transform tool currently has selected, or
        None.  This is what the Properties inspector edits."""
        sel = self._active_transform
        return sel.record if sel is not None else None

    def _reselect_record(self, record):
        """Re-wrap *record*'s (possibly rebuilt) visuals in fresh
        Transform handles.  Factored from try_begin_transform_at so the
        Properties panel can refresh the selection after an edit."""
        if self._active_transform is not None:
            try:
                self._active_transform.dismiss()
            except Exception:
                pass
            self._active_transform = None
        vis = self._record_visuals.get(id(record), [])
        if not vis:
            return
        combined = None
        for v in vis:
            try:
                br = v.sceneBoundingRect()
            except Exception:
                continue
            combined = br if combined is None else combined.united(br)
        if combined is None:
            return
        self._active_transform = TransformSelection(
            self, page_idx=record.page_idx, rect_scene=combined,
            record=record, visuals=list(vis),
            rotation=getattr(record, "rotation", 0.0) or 0.0)

    def _rebuild_one(self, record):
        """Drop *record*'s current visuals + cover and re-materialize
        them from the (mutated) record.  No reselect, no checkpoint —
        callers decide.  Single source of truth, reused by the
        Properties panel and the multi-select bulk operations."""
        for v in self._record_visuals.pop(id(record), []):
            try:
                self._scene.removeItem(v)
            except Exception:
                pass
            if v in self._live_edit_items:
                try:
                    self._live_edit_items.remove(v)
                except Exception:
                    pass
            self._item_to_record.pop(v, None)
        cov = self._record_covers.pop(id(record), None)
        if cov is not None:
            try:
                self._scene.removeItem(cov)
            except Exception:
                pass
            if cov in self._live_edit_items:
                try:
                    self._live_edit_items.remove(cov)
                except Exception:
                    pass
        self._materialize_record(record)

    def apply_property_edit(self, record):
        """Re-render *record* after the Properties panel mutated its
        fields, keep it selected, and push one undo step."""
        if record is None:
            return
        self._rebuild_one(record)
        self._reselect_record(record)
        self._edit_checkpoint()
        self.status_changed.emit()

    # -- E2: multi-select (dedicated "Select" tool) ----------------------
    def _record_scene_bounds(self, rec):
        vis = self._record_visuals.get(id(rec), [])
        b = None
        for v in vis:
            try:
                r = v.sceneBoundingRect()
            except Exception:
                continue
            b = r if b is None else b.united(r)
        return b

    def clear_multi_selection(self):
        for h in self._multi_hl:
            try:
                self._scene.removeItem(h)
            except Exception:
                pass
        self._multi_hl = []
        self._multi_sel = []

    def _draw_multi_highlights(self):
        for h in self._multi_hl:
            try:
                self._scene.removeItem(h)
            except Exception:
                pass
        self._multi_hl = []
        for rec in self._multi_sel:
            b = self._record_scene_bounds(rec)
            if b is None:
                continue
            hl = QGraphicsRectItem(b.adjusted(-3, -3, 3, 3))
            hl.setPen(QPen(QColor(51, 153, 255), 1.5,
                           Qt.PenStyle.DashLine))
            hl.setBrush(QBrush(QColor(51, 153, 255, 30)))
            hl.setZValue(205)
            self._scene.addItem(hl)
            self._multi_hl.append(hl)

    def set_multi_selection(self, records):
        # Dismiss the single-object Transform so the two systems never
        # fight over the same record.
        if self._active_transform is not None:
            try:
                self._active_transform.dismiss()
            except Exception:
                pass
            self._active_transform = None
        seen = set()
        uniq = []
        for r in records:
            if id(r) not in seen:
                seen.add(id(r))
                uniq.append(r)
        self._multi_sel = uniq
        self._draw_multi_highlights()
        self.status_changed.emit()

    def multiselect_from_band(self, rect_scene: QRectF):
        hits = []
        for rec in self._edit_records:
            b = self._record_scene_bounds(rec)
            if b is not None and b.intersects(rect_scene):
                hits.append(rec)
        self.set_multi_selection(hits)

    def has_multi_selection(self) -> bool:
        return len(self._multi_sel) > 0

    def nudge_multi(self, dx_pt: float, dy_pt: float) -> bool:
        if not self._multi_sel:
            return False
        for rec in self._multi_sel:
            if rec.new_rect:
                x0, y0, x1, y1 = rec.new_rect
                rec.new_rect = (x0 + dx_pt, y0 + dy_pt,
                                x1 + dx_pt, y1 + dy_pt)
            for attr in ("line_start", "line_end"):
                v = getattr(rec, attr, None)
                if v:
                    setattr(rec, attr,
                            (v[0] + dx_pt, v[1] + dy_pt))
            self._rebuild_one(rec)
        self._draw_multi_highlights()
        self._edit_checkpoint()
        self.status_changed.emit()
        return True

    def _drop_record(self, rec):
        """Remove a pure-addition record + its visuals entirely."""
        for v in self._record_visuals.pop(id(rec), []):
            try:
                self._scene.removeItem(v)
            except Exception:
                pass
            if v in self._live_edit_items:
                try:
                    self._live_edit_items.remove(v)
                except Exception:
                    pass
            self._item_to_record.pop(v, None)
        cov = self._record_covers.pop(id(rec), None)
        if cov is not None:
            try:
                self._scene.removeItem(cov)
            except Exception:
                pass
            if cov in self._live_edit_items:
                try:
                    self._live_edit_items.remove(cov)
                except Exception:
                    pass
        try:
            if rec in self._edit_records:
                self._edit_records.remove(rec)
        except Exception:
            pass

    def delete_multi(self) -> bool:
        if not self._multi_sel:
            return False
        for rec in list(self._multi_sel):
            if getattr(rec, "orig_rect", None) is None:
                self._drop_record(rec)
            else:
                if rec.kind.startswith("image"):
                    rec.kind = "image_delete"
                else:
                    rec.kind = "text_delete"
                rec.new_rect = None
                rec.text = None
                rec.html = None
                rec.image_path = None
                self._rebuild_one(rec)
        self.clear_multi_selection()
        self._edit_checkpoint()
        self.status_changed.emit()
        return True

    def align_multi(self, mode: str) -> bool:
        recs = [r for r in self._multi_sel if r.new_rect]
        if len(recs) < 2:
            return False
        xs0 = [r.new_rect[0] for r in recs]
        ys0 = [r.new_rect[1] for r in recs]
        xs1 = [r.new_rect[2] for r in recs]
        ys1 = [r.new_rect[3] for r in recs]
        L, T, R, B = min(xs0), min(ys0), max(xs1), max(ys1)
        cx, cy = (L + R) / 2.0, (T + B) / 2.0
        for r in recs:
            x0, y0, x1, y1 = r.new_rect
            w, h = x1 - x0, y1 - y0
            if mode == "left":
                x0, x1 = L, L + w
            elif mode == "right":
                x0, x1 = R - w, R
            elif mode == "top":
                y0, y1 = T, T + h
            elif mode == "bottom":
                y0, y1 = B - h, B
            elif mode == "centerx":
                x0 = cx - w / 2.0
                x1 = x0 + w
            elif mode == "centery":
                y0 = cy - h / 2.0
                y1 = y0 + h
            r.new_rect = (x0, y0, x1, y1)
            self._rebuild_one(r)
        self._draw_multi_highlights()
        self._edit_checkpoint()
        self.status_changed.emit()
        return True

    def distribute_multi(self, axis: str) -> bool:
        """Even out the spacing of the multi-selected objects along
        *axis* ('h' or 'v').  Outer two stay put; the rest are spaced so
        the gaps between bounding boxes are equal."""
        recs = [r for r in self._multi_sel if r.new_rect]
        if len(recs) < 3:
            return False
        i0, i1 = (0, 2) if axis == "h" else (1, 3)
        recs.sort(key=lambda r: r.new_rect[i0])
        span = recs[-1].new_rect[i1] - recs[0].new_rect[i0]
        total = sum(r.new_rect[i1] - r.new_rect[i0] for r in recs)
        gap = (span - total) / (len(recs) - 1)
        cursor = recs[0].new_rect[i0]
        for r in recs:
            x0, y0, x1, y1 = r.new_rect
            size = (x1 - x0) if axis == "h" else (y1 - y0)
            if axis == "h":
                r.new_rect = (cursor, y0, cursor + size, y1)
            else:
                r.new_rect = (x0, cursor, x1, cursor + size)
            cursor += size + gap
            self._rebuild_one(r)
        self._draw_multi_highlights()
        self._edit_checkpoint()
        self.status_changed.emit()
        return True

    def group_multi(self) -> bool:
        recs = list(self._multi_sel)
        if len(recs) < 2:
            return False
        self._next_group_id += 1
        gid = self._next_group_id
        for r in recs:
            if not isinstance(getattr(r, "extra", None), dict):
                r.extra = {}
            r.extra["group"] = gid
        self._edit_checkpoint()
        self.status_changed.emit()
        return True

    def ungroup_multi(self) -> bool:
        recs = list(self._multi_sel)
        if not recs:
            return False
        changed = False
        for r in recs:
            if isinstance(getattr(r, "extra", None), dict) \
                    and r.extra.pop("group", None) is not None:
                changed = True
        if changed:
            self._edit_checkpoint()
            self.status_changed.emit()
        return changed

    def select_group_of(self, record) -> bool:
        """If *record* is grouped, select the whole group as a
        multi-selection and return True; else False."""
        gid = None
        if isinstance(getattr(record, "extra", None), dict):
            gid = record.extra.get("group")
        if gid is None:
            return False
        members = [r for r in self._edit_records
                   if isinstance(getattr(r, "extra", None), dict)
                   and r.extra.get("group") == gid]
        if len(members) < 2:
            return False
        self.set_multi_selection(members)
        return True

    # -- E3: snapping & guide lines --------------------------------------
    def _clear_snap_guides(self):
        for g in self._snap_guide_items:
            try:
                self._scene.removeItem(g)
            except Exception:
                pass
        self._snap_guide_items = []

    def _draw_snap_guides(self, vxs, hys, page_idx):
        self._clear_snap_guides()
        from PyQt6.QtWidgets import QGraphicsLineItem as _GLI
        pw = self._page_widths.get(page_idx, 0.0)
        ph = self._page_heights.get(page_idx, 0.0)
        py = self._page_positions.get(page_idx, 0.0)
        pen = QPen(QColor(0, 200, 255), 0.8, Qt.PenStyle.DashLine)
        for x in vxs:
            ln = _GLI(x, py, x, py + ph)
            ln.setPen(pen)
            ln.setZValue(210)
            self._scene.addItem(ln)
            self._snap_guide_items.append(ln)
        for y in hys:
            ln = _GLI(0.0, y, pw, y)
            ln.setPen(pen)
            ln.setZValue(210)
            self._scene.addItem(ln)
            self._snap_guide_items.append(ln)

    def apply_snap(self, r: QRectF, page_idx: int, exclude) -> QRectF:
        """Snap the (translated) rect *r* to nearby object edges/centres
        and the page margins; draw guide lines for the matches.  Returns
        a possibly-adjusted copy.  Called only for move drags."""
        if not self._snap_enabled:
            self._clear_snap_guides()
            return r
        thr = 6.0
        pw = self._page_widths.get(page_idx, 0.0)
        ph = self._page_heights.get(page_idx, 0.0)
        py = self._page_positions.get(page_idx, 0.0)
        vx, hy = [], []
        if pw and ph:
            vx += [0.0, pw / 2.0, pw]
            hy += [py, py + ph / 2.0, py + ph]
        for rec in self._edit_records:
            if rec is exclude or rec.page_idx != page_idx:
                continue
            b = self._record_scene_bounds(rec)
            if b is None:
                continue
            vx += [b.left(), b.center().x(), b.right()]
            hy += [b.top(), b.center().y(), b.bottom()]
        out = QRectF(r)
        gv, gh = [], []
        bx = None
        for val, kind in ((out.left(), "l"), (out.center().x(), "c"),
                          (out.right(), "r")):
            for cx in vx:
                d = abs(val - cx)
                if d <= thr and (bx is None or d < bx[2]):
                    bx = (kind, cx, d)
        if bx is not None:
            kind, cx, _ = bx
            if kind == "l":
                out.moveLeft(cx)
            elif kind == "r":
                out.moveRight(cx)
            else:
                out.moveCenter(QPointF(cx, out.center().y()))
            gv.append(cx)
        by = None
        for val, kind in ((out.top(), "t"), (out.center().y(), "c"),
                          (out.bottom(), "b")):
            for cy in hy:
                d = abs(val - cy)
                if d <= thr and (by is None or d < by[2]):
                    by = (kind, cy, d)
        if by is not None:
            kind, cy, _ = by
            if kind == "t":
                out.moveTop(cy)
            elif kind == "b":
                out.moveBottom(cy)
            else:
                out.moveCenter(QPointF(out.center().x(), cy))
            gh.append(cy)
        self._draw_snap_guides(gv, gh, page_idx)
        return out

    def _make_cover(self, rec, scene_rect: QRectF, bg, z: float = 9.0):
        """Create + register a background-coloured cover over the
        original content position.  Covers are tracked separately from
        draggable visuals (in _record_covers, NOT _record_visuals /
        _item_to_record) so the Transform tool leaves them anchored
        while the live copy moves."""
        col = bg if bg else (255, 255, 255)
        cover = QGraphicsRectItem(QRectF(scene_rect))
        cover.setBrush(QBrush(QColor(*col)))
        cover.setPen(QPen(Qt.PenStyle.NoPen))
        cover.setZValue(z)
        self._scene.addItem(cover)
        self._live_edit_items.append(cover)
        self._record_covers[id(rec)] = cover
        return cover

    def _materialize_record(self, rec):
        """Central record → live-preview renderer and the single source
        of truth for every additive kind's visuals (the creators call
        this, and so does the undo/redo rebuild — so the create-path and
        the rebuild-path can never diverge).

        Stage B implements the self-contained additive kinds; Stage C
        adds the existing-content kinds (text/image move/edit/delete).
        Returns the list of scene items created."""
        kind = rec.kind
        page_idx = rec.page_idx
        visuals: list = []
        if rec.new_rect is None and kind not in (
                "text_delete", "image_delete"):
            return visuals
        rs = (self._pt_rect_to_scene(page_idx, rec.new_rect)
              if rec.new_rect is not None else None)
        ors = (self._pt_rect_to_scene(page_idx, rec.orig_rect)
               if rec.orig_rect is not None else None)

        if kind == "stamp_add":
            r, g, b = rec.color or (255, 0, 0)
            item = QGraphicsSimpleTextItem(rec.text or "")
            item.setFont(QFont("Helvetica", 20, QFont.Weight.Bold))
            item.setBrush(QBrush(QColor(r, g, b, 160)))
            item.setPos(rs.x(), rs.y())
            item.setZValue(100)
            self._register_visual(rec, item, visuals)
            border = QGraphicsRectItem(0, 0, rs.width(), rs.height())
            border.setPos(rs.x(), rs.y())
            border.setPen(QPen(QColor(r, g, b, 160), 2.0))
            border.setBrush(QBrush(Qt.GlobalColor.transparent))
            border.setZValue(99)
            self._register_visual(rec, border, visuals)

        elif kind == "note_add":
            text = rec.text or ""
            note_w = 220.0
            lines = text.count('\n') + 1  # noqa: F841 (kept for parity)
            chars_per_line = max(1, int(note_w / 7))
            wrapped = sum(
                max(1, (len(ln) + chars_per_line - 1) // chars_per_line)
                for ln in text.split('\n'))
            note_h = max(60.0, wrapped * 18.0 + 36.0)
            ox, oy = rs.x(), rs.y()
            card = QGraphicsRectItem(0, 0, note_w, note_h)
            card.setPos(ox, oy)
            card.setBrush(QBrush(QColor(255, 255, 200, 240)))
            card.setPen(QPen(QColor(200, 180, 60), 1.5))
            card.setZValue(100)
            card.setToolTip(text)
            self._register_visual(rec, card, visuals)
            header_h = 22.0
            header = QGraphicsRectItem(0, 0, note_w, header_h)
            header.setPos(ox, oy)
            header.setBrush(QBrush(QColor(255, 230, 100)))
            header.setPen(QPen(Qt.PenStyle.NoPen))
            header.setZValue(101)
            self._register_visual(rec, header, visuals)
            hdr_txt = QGraphicsSimpleTextItem("Sticky Note")
            hdr_txt.setFont(QFont("Segoe UI", 9, QFont.Weight.Bold))
            hdr_txt.setBrush(QBrush(QColor(100, 80, 0)))
            hdr_txt.setPos(ox + 6, oy + 3)
            hdr_txt.setZValue(102)
            self._register_visual(rec, hdr_txt, visuals)
            body = QGraphicsTextItem()
            body.setDefaultTextColor(QColor(40, 40, 40))
            body.setFont(QFont("Segoe UI", 10))
            body.setTextWidth(note_w - 12)
            body.setPlainText(text)
            body.setPos(ox + 6, oy + header_h + 4)
            body.setZValue(102)
            self._register_visual(rec, body, visuals)

        elif kind == "redact_add":
            item = QGraphicsRectItem(0, 0, rs.width(), rs.height())
            item.setPos(rs.x(), rs.y())
            item.setBrush(QBrush(QColor(255, 255, 255)))
            item.setPen(QPen(QColor(200, 200, 200), 1.0,
                             Qt.PenStyle.DashLine))
            item.setZValue(100)
            item.setToolTip("Whiteout — content will be removed on save")
            self._register_visual(rec, item, visuals)

        elif kind == "highlight_add":
            item = QGraphicsRectItem(0, 0, rs.width(), rs.height())
            item.setPos(rs.x(), rs.y())
            item.setBrush(QBrush(QColor(255, 235, 60, 90)))
            item.setPen(QPen(Qt.PenStyle.NoPen))
            item.setZValue(100)
            item.setToolTip(
                "Highlight — saved as a real PDF highlight annotation")
            self._register_visual(rec, item, visuals)

        elif kind == "shape_add":
            r, g, b = rec.stroke_color or (0, 0, 0)
            pen = QPen(QColor(r, g, b), rec.stroke_width or 2.0)
            sstyle = getattr(rec, "stroke_style", "solid")
            if sstyle == "dash":
                pen.setStyle(Qt.PenStyle.DashLine)
            elif sstyle == "dot":
                pen.setStyle(Qt.PenStyle.DotLine)
            brush = QBrush(Qt.GlobalColor.transparent)
            if rec.fill_color:
                fr, fg, fb = rec.fill_color
                fa = max(0, min(255, getattr(rec, "fill_opacity", 80)))
                brush = QBrush(QColor(fr, fg, fb, fa))
            st = rec.shape_type
            if st == "rect":
                item = QGraphicsRectItem(0, 0, rs.width(), rs.height())
                item.setPos(rs.x(), rs.y())
                item.setPen(pen)
                item.setBrush(brush)
            elif st == "circle":
                item = QGraphicsEllipseItem(0, 0, rs.width(), rs.height())
                item.setPos(rs.x(), rs.y())
                item.setPen(pen)
                item.setBrush(brush)
            elif st in ("line", "arrow"):
                from PyQt6.QtWidgets import QGraphicsLineItem as _GLI
                sx0, sy0 = self._pt_to_scene(
                    page_idx, rec.line_start[0], rec.line_start[1])
                sx1, sy1 = self._pt_to_scene(
                    page_idx, rec.line_end[0], rec.line_end[1])
                item = _GLI(sx0, sy0, sx1, sy1)
                item.setPen(pen)
                if st == "arrow":
                    import math
                    hl = min(15, (rec.stroke_width or 2.0) * 5)
                    heads = [(sx1, sy1,
                              math.atan2(sy1 - sy0, sx1 - sx0))]
                    if getattr(rec, "arrow_both", False):
                        heads.append(
                            (sx0, sy0,
                             math.atan2(sy0 - sy1, sx0 - sx1)))
                    for hx, hy, ang in heads:
                        for da in (math.pi * 0.85, -math.pi * 0.85):
                            h = _GLI(hx, hy,
                                     hx + hl * math.cos(ang + da),
                                     hy + hl * math.sin(ang + da))
                            h.setPen(pen)
                            h.setZValue(100)
                            self._register_visual(rec, h, visuals)
            else:
                item = QGraphicsRectItem(0, 0, rs.width(), rs.height())
                item.setPos(rs.x(), rs.y())
                item.setPen(pen)
                item.setBrush(brush)
            item.setZValue(100)
            # Keep the primary item first in the visuals list, matching
            # the original _record_visuals = [item] + extra ordering.
            self._scene.addItem(item)
            self._live_edit_items.append(item)
            self._item_to_record[item] = rec
            visuals.insert(0, item)

        elif kind in ("text_move", "text_edit"):
            # Cover the ORIGINAL block (anchored at orig_rect); the live
            # text copy sits at new_rect.  text_move always covers (z=9);
            # text_edit covers at z=10 but honours a transparent
            # background (None → no cover, page shows through).
            if kind == "text_move":
                if ors is not None:
                    self._make_cover(rec, ors,
                                     rec.background_color, z=9.0)
            elif rec.background_color is not None and ors is not None:
                self._make_cover(rec, ors,
                                 rec.background_color, z=10.0)
            ex = getattr(rec, "extra", {}) or {}
            ti = QGraphicsTextItem()
            fnt = QFont(ex.get("family", "Helvetica"))
            fnt.setBold(bool(ex.get("bold")))
            fnt.setItalic(bool(ex.get("italic")))
            fnt.setPointSizeF(
                max(6.0, (rec.font_size or 12.0)
                    * self._render_scale * 0.72))
            ti.setFont(fnt)
            cr, cg, cb = rec.color or (0, 0, 0)
            ti.setDefaultTextColor(QColor(cr, cg, cb))
            if rec.html:
                ti.setHtml(rec.html)
            else:
                ti.setPlainText(rec.text or "")
            if rs is not None:
                ti.setTextWidth(rs.width())
                ti.setPos(rs.x(), rs.y())
            ti.setZValue(11)
            self._register_visual(rec, ti, visuals)

        elif kind in ("image_move", "image_replace"):
            if (kind == "image_move" and ors is not None):
                self._make_cover(rec, ors, rec.background_color, z=9.0)
            try:
                pm = QPixmap(rec.image_path) if rec.image_path else None
            except Exception:
                pm = None
            if pm is not None and not pm.isNull() and rs is not None:
                w = max(1, int(round(rs.width())))
                h = max(1, int(round(rs.height())))
                item = QGraphicsPixmapItem(pm.scaled(
                    w, h,
                    Qt.AspectRatioMode.IgnoreAspectRatio,
                    Qt.TransformationMode.SmoothTransformation))
                item.setPos(rs.x(), rs.y())
                item.setZValue(10)
                self._register_visual(rec, item, visuals)

        elif kind in ("text_delete", "image_delete"):
            # The original region is erased on save; only the cover
            # remains so the user sees it "gone".
            if ors is not None:
                bg = rec.background_color
                if bg is None:
                    try:
                        bg = self._sample_background_color(
                            page_idx, rec.orig_rect)
                    except Exception:
                        bg = None
                self._make_cover(rec, ors, bg, z=9.0)

        # text_add / image_add: no live preview at creation (parity with
        # the original handlers).  Existing-content kinds: Stage C.
        if visuals:
            self._record_visuals[id(rec)] = list(visuals)
        return visuals

    def set_edit_action(self, action):
        """Arm the next click — 'add_text', 'add_image', 'transform',
        'stamp', 'note', 'redact', 'shape_rect', 'shape_circle',
        'shape_line', 'shape_arrow', or None."""
        # Leaving transform mode? Dismiss any active selection so the
        # handles don't linger on screen when the user switches tools.
        if (self._edit_action == "transform" and action != "transform"
                and self._active_transform is not None):
            try:
                self._active_transform.dismiss()
            except Exception:
                pass
            self._active_transform = None
        # Clean up any drag rubber band from shape/redact
        old_rubber = getattr(self, '_drag_rubber', None)
        if old_rubber is not None:
            try:
                self._scene.removeItem(old_rubber)
            except Exception:
                pass
            self._drag_rubber = None
        self._drag_start = None
        self._edit_action = action
        drag_tools = ("redact", "highlight", "multiselect", "shape_rect",
                      "shape_circle", "shape_line", "shape_arrow")
        click_tools = ("add_text", "add_image", "stamp", "note")
        if action in click_tools or action in drag_tools:
            self._view.setCursor(_get_edit_cursor(action))
        elif action == "transform":
            self._view.setCursor(Qt.CursorShape.PointingHandCursor)
        else:
            self._view.setCursor(Qt.CursorShape.ArrowCursor)

    # -- click-driven add actions -----------------------------------------
    def handle_edit_mode_click(self, scene_pos: QPointF) -> bool:
        """Called from PDFGraphicsView when the user left-clicks while
        edit mode is active.  If an 'add' action is armed, creates a
        new EditRecord at the click position and returns True."""
        if not self._edit_mode:
            return False
        if self._edit_action == "add_text":
            self._add_new_text_record_at(scene_pos)
            return True
        if self._edit_action == "add_image":
            self._add_new_image_record_at(scene_pos)
            return True
        if self._edit_action == "stamp":
            self._add_stamp_at(scene_pos)
            return True
        if self._edit_action == "note":
            self._add_note_at(scene_pos)
            return True
        if self._edit_action in ("redact", "highlight", "multiselect"):
            # Redact / highlight / select use drag — record the start
            self._drag_start = scene_pos
            return True
        if self._edit_action and self._edit_action.startswith("shape_"):
            self._drag_start = scene_pos
            return True
        return False

    # -- Stamp tool ----------------------------------------------------------
    def _add_stamp_at(self, scene_pos: QPointF):
        """Place the pending stamp text at the clicked position."""
        page_idx = self._page_at_scene_point(scene_pos)
        if page_idx is None:
            return
        stamp_text = getattr(self, '_pending_stamp_text', None)
        stamp_color = getattr(self, '_pending_stamp_color', (255, 0, 0))
        if not stamp_text:
            return
        # Stamp size: ~200 x 50, centred on click
        w, h = 200.0, 50.0
        x = scene_pos.x() - w / 2
        y = scene_pos.y() - h / 2
        rect_scene = QRectF(x, y, w, h)
        new_rect = self._scene_rect_to_pt(page_idx, rect_scene)
        rec = EditRecord(
            kind="stamp_add", page_idx=page_idx,
            new_rect=new_rect, text=stamp_text,
            font_size=28.0, color=stamp_color,
            rotation=0.0)
        self._edit_records.append(rec)
        self._materialize_record(rec)
        self._edit_checkpoint()
        self.status_changed.emit()

    # -- Sticky Note tool ----------------------------------------------------
    def _add_note_at(self, scene_pos: QPointF):
        """Place a sticky note annotation at the clicked position."""
        page_idx = self._page_at_scene_point(scene_pos)
        if page_idx is None:
            return
        text, ok = QInputDialog.getMultiLineText(
            self.window(), "Sticky Note",
            "Note text:", "")
        if not ok or not text:
            return
        # Note anchor for PDF annotation: small point
        anchor_w, anchor_h = 24.0, 24.0
        anchor_rect = QRectF(scene_pos.x(), scene_pos.y(),
                             anchor_w, anchor_h)
        new_rect = self._scene_rect_to_pt(page_idx, anchor_rect)
        rec = EditRecord(
            kind="note_add", page_idx=page_idx,
            new_rect=new_rect, text=text)
        self._edit_records.append(rec)
        self._materialize_record(rec)
        self._edit_checkpoint()
        self.status_changed.emit()

    def add_note_on_current_page(self, text: str) -> int:
        """Add a sticky note near the top-left of the currently visible
        page (used by the Notes panel's Add Note button — no click
        needed).  Returns the 0-based page index, or -1 on failure."""
        if not text or not self._num_pages:
            return -1
        page_idx = self.get_visible_page() if self._num_pages else 0
        page_top = self._page_positions.get(page_idx, 0.0)
        anchor = QRectF(36.0, page_top + 36.0, 24.0, 24.0)
        new_rect = self._scene_rect_to_pt(page_idx, anchor)
        rec = EditRecord(kind="note_add", page_idx=page_idx,
                         new_rect=new_rect, text=text)
        self._edit_records.append(rec)
        self._materialize_record(rec)
        self._edit_checkpoint()
        self.status_changed.emit()
        return page_idx

    # -- Shape / Redact drag support -----------------------------------------
    def handle_edit_drag_release(self, scene_pos: QPointF) -> bool:
        """Called from PDFGraphicsView on mouse-release after a drag
        for shape or redact tools.  Returns True if handled."""
        start = getattr(self, '_drag_start', None)
        if start is None:
            return False
        self._drag_start = None
        page_idx = self._page_at_scene_point(start)
        if page_idx is None:
            return False

        # Shift-constrain: perfect square/circle, or 45°-snapped line.
        act = self._edit_action or ""
        if act.startswith("shape_"):
            from PyQt6.QtWidgets import QApplication
            if (QApplication.keyboardModifiers()
                    & Qt.KeyboardModifier.ShiftModifier):
                stype = act.replace("shape_", "")
                dx = scene_pos.x() - start.x()
                dy = scene_pos.y() - start.y()
                if stype in ("rect", "circle"):
                    side = max(abs(dx), abs(dy))
                    scene_pos = QPointF(
                        start.x() + (side if dx >= 0 else -side),
                        start.y() + (side if dy >= 0 else -side))
                elif stype in ("line", "arrow"):
                    import math
                    length = math.hypot(dx, dy)
                    step = math.pi / 4
                    snapped = round(math.atan2(dy, dx) / step) * step
                    scene_pos = QPointF(
                        start.x() + length * math.cos(snapped),
                        start.y() + length * math.sin(snapped))

        x0, y0 = min(start.x(), scene_pos.x()), min(start.y(), scene_pos.y())
        x1, y1 = max(start.x(), scene_pos.x()), max(start.y(), scene_pos.y())
        # Minimum size guard
        if abs(x1 - x0) < 5 and abs(y1 - y0) < 5:
            return False
        rect_scene = QRectF(x0, y0, x1 - x0, y1 - y0)
        new_rect = self._scene_rect_to_pt(page_idx, rect_scene)

        if self._edit_action == "multiselect":
            self.multiselect_from_band(rect_scene)
            return True

        if self._edit_action == "redact":
            rec = EditRecord(
                kind="redact_add", page_idx=page_idx,
                new_rect=new_rect)
            self._edit_records.append(rec)
            self._materialize_record(rec)
            self._edit_checkpoint()
            self.status_changed.emit()
            return True

        if self._edit_action == "highlight":
            rec = EditRecord(
                kind="highlight_add", page_idx=page_idx,
                new_rect=new_rect)
            self._edit_records.append(rec)
            self._materialize_record(rec)
            self._edit_checkpoint()
            self.status_changed.emit()
            return True

        if self._edit_action and self._edit_action.startswith("shape_"):
            shape_type = self._edit_action.replace("shape_", "")
            s_color = getattr(self, '_shape_stroke_color', (0, 0, 0))
            s_width = getattr(self, '_shape_stroke_width', 2.0)
            f_color = getattr(self, '_shape_fill_color', None)

            line_start = None
            line_end = None
            if shape_type in ("line", "arrow"):
                line_start = self._scene_to_pt(
                    page_idx, start.x(), start.y())
                line_end = self._scene_to_pt(
                    page_idx, scene_pos.x(), scene_pos.y())

            rec = EditRecord(
                kind="shape_add", page_idx=page_idx,
                new_rect=new_rect,
                shape_type=shape_type,
                stroke_color=s_color,
                stroke_width=s_width,
                fill_color=f_color,
                line_start=line_start,
                line_end=line_end)
            self._edit_records.append(rec)
            self._materialize_record(rec)
            self._edit_checkpoint()
            self.status_changed.emit()
            return True

        return False

    def update_drag_preview(self, scene_pos: QPointF):
        """Draw a rubber-band preview while the user drags for shape/redact."""
        start = getattr(self, '_drag_start', None)
        if start is None:
            return
        # Remove old rubber band
        old = getattr(self, '_drag_rubber', None)
        if old is not None:
            try:
                self._scene.removeItem(old)
            except Exception:
                pass
        x0, y0 = min(start.x(), scene_pos.x()), min(start.y(), scene_pos.y())
        x1, y1 = max(start.x(), scene_pos.x()), max(start.y(), scene_pos.y())
        rect = QRectF(x0, y0, x1 - x0, y1 - y0)

        if self._edit_action == "redact":
            item = QGraphicsRectItem(rect)
            item.setBrush(QBrush(QColor(255, 255, 255, 150)))
            item.setPen(QPen(QColor(180, 180, 180), 1.0,
                             Qt.PenStyle.DashLine))
        elif self._edit_action and self._edit_action.startswith("shape_"):
            stype = self._edit_action.replace("shape_", "")
            s_color = getattr(self, '_shape_stroke_color', (0, 0, 0))
            r, g, b = s_color
            pen = QPen(QColor(r, g, b, 150), 2.0, Qt.PenStyle.DashLine)
            if stype == "circle":
                from PyQt6.QtWidgets import QGraphicsEllipseItem
                item = QGraphicsEllipseItem(rect)
                item.setPen(pen)
                item.setBrush(QBrush(Qt.GlobalColor.transparent))
            elif stype in ("line", "arrow"):
                from PyQt6.QtWidgets import QGraphicsLineItem
                item = QGraphicsLineItem(
                    start.x(), start.y(),
                    scene_pos.x(), scene_pos.y())
                item.setPen(pen)
            else:
                item = QGraphicsRectItem(rect)
                item.setPen(pen)
                item.setBrush(QBrush(Qt.GlobalColor.transparent))
        else:
            return

        item.setZValue(200)
        self._scene.addItem(item)
        self._drag_rubber = item

    def _get_edit_text_bboxes(self, page_idx: int) -> list:
        """Return a cached list of (x0, y0, x1, y1) bboxes (in PDF
        points) for every text block on *page_idx*.  Populated on
        demand the first time the hover handler or click handler asks
        about the page, so load_pdf stays fast."""
        if page_idx in self._edit_text_bboxes:
            return self._edit_text_bboxes[page_idx]
        bboxes: list = []
        try:
            import fitz
        except Exception:
            self._edit_text_bboxes[page_idx] = bboxes
            return bboxes
        if not self._doc_path:
            self._edit_text_bboxes[page_idx] = bboxes
            return bboxes
        doc = None
        try:
            doc = fitz.open(self._doc_path)
            if 0 <= page_idx < len(doc):
                text_dict = doc[page_idx].get_text("dict")
                for block in text_dict.get("blocks", []):
                    if block.get("type", 0) != 0:
                        continue
                    bbox = block.get("bbox")
                    if bbox and len(bbox) == 4:
                        bboxes.append(tuple(bbox))
        except Exception:
            pass
        finally:
            if doc is not None:
                try:
                    doc.close()
                except Exception:
                    pass
        self._edit_text_bboxes[page_idx] = bboxes
        return bboxes

    def scene_point_over_editable_text(self, scene_pos: QPointF) -> bool:
        """True if the given scene position lands inside any text
        block's bbox on the underlying page.  Used by the view to
        flip the cursor to an I-beam on hover while in edit mode."""
        if not self._edit_mode:
            return False
        page_idx = self._page_at_scene_point(scene_pos)
        if page_idx is None:
            return False
        try:
            pt_x, pt_y = self._scene_to_pt(
                page_idx, scene_pos.x(), scene_pos.y())
        except Exception:
            return False
        for (x0, y0, x1, y1) in self._get_edit_text_bboxes(page_idx):
            if x0 <= pt_x <= x1 and y0 <= pt_y <= y1:
                return True
        return False

    def _page_at_scene_point(self, scene_pos: QPointF):
        """Return the 0-based page index whose slot contains *scene_pos*,
        or None if the point is outside every page."""
        y = scene_pos.y()
        for idx in range(self._num_pages):
            py = self._page_positions.get(idx, 0)
            ph = self._page_heights.get(idx, 0)
            if py <= y <= py + ph:
                return idx
        return None

    def _page_at_scene_rect_center(self, rect: QRectF):
        """Page index whose scene box contains the centre of *rect*
        (both X and Y, so it is correct in two-page-spread too).
        Returns None if the centre is in a gap / outside every page."""
        cx = rect.center().x()
        cy = rect.center().y()
        for idx in range(self._num_pages):
            px = self._page_x.get(idx, 0.0)
            py = self._page_positions.get(idx, 0.0)
            pw = self._page_widths.get(idx, 0.0)
            ph = self._page_heights.get(idx, 0.0)
            if px <= cx <= px + pw and py <= cy <= py + ph:
                return idx
        return None

    def _add_new_text_record_at(self, scene_pos: QPointF):
        """Queue an EditRecord for a brand-new text box at the clicked
        point.  Nothing is drawn on the page until Save As runs."""
        page_idx = self._page_at_scene_point(scene_pos)
        if page_idx is None:
            self.set_edit_action(None)
            return
        text, ok = QInputDialog.getMultiLineText(
            self.window(), "Add Text",
            "Text to insert:", "")
        if not ok or not text:
            self.set_edit_action(None)
            return
        # Default size: ~200 x 40 scene px, centred on the click.
        w = 200.0
        h = 40.0
        x = scene_pos.x() - w / 2
        y = scene_pos.y() - h / 2
        rect_scene = QRectF(x, y, w, h)
        new_rect = self._scene_rect_to_pt(page_idx, rect_scene)
        self._edit_records.append(EditRecord(
            kind="text_add",
            page_idx=page_idx,
            orig_rect=None,
            new_rect=new_rect,
            text=text,
            font_size=12.0,
            color=(0, 0, 0),
        ))
        self.set_edit_action(None)
        QMessageBox.information(
            self.window(), "Text Queued",
            "The text will appear in the saved PDF after you use "
            "Save As.")
        self._edit_checkpoint()
        self.status_changed.emit()

    def _add_new_image_record_at(self, scene_pos: QPointF):
        """Queue an EditRecord for a brand-new image at the clicked
        point using the file the user selected before clicking."""
        path = self._pending_add_image_path
        if not path or not os.path.isfile(path):
            self.set_edit_action(None)
            return
        page_idx = self._page_at_scene_point(scene_pos)
        if page_idx is None:
            self.set_edit_action(None)
            return
        iw, ih = 200.0, 200.0
        try:
            from PIL import Image as _Image
            with _Image.open(path) as im:
                iw, ih = im.size
        except Exception:
            pass
        max_w = 240.0
        scale = min(1.0, max_w / max(iw, 1.0))
        w = iw * scale
        h = ih * scale
        x = scene_pos.x() - w / 2
        y = scene_pos.y() - h / 2
        rect_scene = QRectF(x, y, w, h)
        new_rect = self._scene_rect_to_pt(page_idx, rect_scene)
        self._edit_records.append(EditRecord(
            kind="image_add",
            page_idx=page_idx,
            orig_rect=None,
            new_rect=new_rect,
            image_path=path,
        ))
        self._pending_add_image_path = None
        self.set_edit_action(None)
        QMessageBox.information(
            self.window(), "Image Queued",
            "The image will appear in the saved PDF after you use "
            "Save As.")
        self._edit_checkpoint()
        self.status_changed.emit()

    # -- Replace Image (via image-select pick) ----------------------------
    def begin_image_replace_pick(self):
        """Flip into a temporary image-picking state.  The existing
        image-select overlays let the user click one of the real images
        already in the PDF; ``handle_image_overlay_click`` will route
        the pick back into the replace flow."""
        if not self._edit_mode:
            return
        self._edit_image_pick_mode = "replace"
        # Make sure the image detector has run so there are overlays to
        # click on.  ``enter_image_select_mode`` handles both first-time
        # activation and re-entry.
        if not self._image_select_mode:
            self.enter_image_select_mode()
        QMessageBox.information(
            self.window(), "Replace Image",
            "Click the image in the PDF you want to replace.")

    def _handle_image_replace_pick(self, overlay) -> bool:
        """User clicked *overlay* while in replace-pick mode.  Prompts
        for a replacement file, shows it live on the page, and queues
        an EditRecord for Save As."""
        path, _ = QFileDialog.getOpenFileName(
            self.window(), "Choose Replacement Image", "",
            "Image Files (*.png *.jpg *.jpeg *.bmp *.tif *.tiff)")
        if not path:
            # Stay in pick mode so they can try again.
            return True
        rect_scene = QRectF(overlay.rect().translated(overlay.pos()))
        orig_rect = self._scene_rect_to_pt(overlay.page_idx, rect_scene)
        record = EditRecord(
            kind="image_replace",
            page_idx=overlay.page_idx,
            orig_rect=orig_rect,
            new_rect=orig_rect,
            image_path=path,
        )
        self._edit_records.append(record)
        # Render the new image directly into the scene on top of the
        # old one so the user sees it immediately.  Passing the record
        # registers the item so the Transform tool can grab it later.
        self._add_live_image_replacement(rect_scene, path, record=record)
        # Done picking — leave image-select mode and return to the
        # normal edit-mode idle state.
        self._edit_image_pick_mode = None
        self.exit_image_select_mode()
        self.status_changed.emit()
        return True

    def _add_live_image_replacement(self, rect_scene: QRectF, path: str,
                                    record: "EditRecord | None" = None):
        """Draw *path* scaled into *rect_scene* as a live preview item
        on top of the page pixmap.  If *record* is supplied the new
        item is registered in the item→record map so the Transform
        tool can re-select it on a later click."""
        try:
            pixmap = QPixmap(path)
            if pixmap.isNull():
                return None
            w = max(1, int(round(rect_scene.width())))
            h = max(1, int(round(rect_scene.height())))
            scaled = pixmap.scaled(
                w, h,
                Qt.AspectRatioMode.IgnoreAspectRatio,
                Qt.TransformationMode.SmoothTransformation,
            )
            item = QGraphicsPixmapItem(scaled)
            item.setPos(rect_scene.x(), rect_scene.y())
            # Sit above the page pixmap (Z=0) and word overlays (Z~5)
            # but below the image-select overlays (Z=15).
            item.setZValue(10)
            self._scene.addItem(item)
            self._live_edit_items.append(item)
            if record is not None:
                self._item_to_record[item] = record
            return item
        except Exception as e:
            print(f"[BoltPDF] live image replace failed: {e}",
                  file=sys.stderr)
            return None

    # -- Transform tool ---------------------------------------------------
    def try_begin_transform_at(self, scene_pos: QPointF) -> bool:
        """Pick the thing under *scene_pos* and wrap it in transform
        handles.

        Hit-test order: live preview items we've already tracked,
        then existing text blocks, then existing images.  Returns
        True when a selection was started (or already active)."""
        if not self._edit_mode:
            return False

        page_idx = self._page_at_scene_point(scene_pos)
        if page_idx is None:
            return False

        # Dismiss any existing selection first so we never have two
        # handle sets drawn at once.
        if self._active_transform is not None:
            try:
                self._active_transform.dismiss()
            except Exception:
                pass
            self._active_transform = None

        # 1) click lands on something we drew for a prior edit?
        for it in list(self._item_to_record.keys()):
            try:
                br = it.sceneBoundingRect()
            except Exception:
                continue
            if br.contains(scene_pos):
                record = self._item_to_record[it]
                # Grouped object → select the whole group instead of a
                # single-object transform (move/align/delete together).
                if self.select_group_of(record):
                    return True
                # Gather ALL visuals for this record so they move as
                # a group and Delete removes them all.
                all_vis = self._record_visuals.get(id(record), [it])
                # Compute bounding rect of all visuals
                combined = QRectF(br)
                for v in all_vis:
                    try:
                        combined = combined.united(v.sceneBoundingRect())
                    except Exception:
                        pass
                self._active_transform = TransformSelection(
                    self,
                    page_idx=record.page_idx,
                    rect_scene=combined,
                    record=record,
                    visuals=list(all_vis),
                    rotation=getattr(record, "rotation", 0.0) or 0.0,
                )
                return True

        # 2) click lands on an existing text block or image in the PDF?
        try:
            import fitz
        except Exception:
            return False

        click_pt_x, click_pt_y = self._scene_to_pt(
            page_idx, scene_pos.x(), scene_pos.y())

        doc = None
        try:
            doc = fitz.open(self._doc_path)
            if page_idx < 0 or page_idx >= len(doc):
                return False
            page = doc[page_idx]

            # 2a) existing image?
            try:
                images = page.get_images(full=True)
            except Exception:
                images = []
            image_hit = None
            for img in images:
                xref = img[0]
                try:
                    rects = page.get_image_rects(xref)
                except Exception:
                    rects = []
                for r in rects:
                    if (r.x0 <= click_pt_x <= r.x1
                            and r.y0 <= click_pt_y <= r.y1):
                        image_hit = (xref, tuple(r))
                        break
                if image_hit:
                    break
            if image_hit is not None:
                xref, (ix0, iy0, ix1, iy1) = image_hit
                # Crop the rendered page pixmap to the image bbox and
                # write it out as a temp PNG so the save worker has a
                # real file to re-insert at the new rect.
                tmp_path = self._crop_existing_image_to_tempfile(
                    page_idx, (ix0, iy0, ix1, iy1))
                if tmp_path is None:
                    return False
                sx0, sy0 = self._pt_to_scene(page_idx, ix0, iy0)
                sx1, sy1 = self._pt_to_scene(page_idx, ix1, iy1)
                rect_scene = QRectF(sx0, sy0,
                                    max(14.0, sx1 - sx0),
                                    max(14.0, sy1 - sy0))
                record = EditRecord(
                    kind="image_move",
                    page_idx=page_idx,
                    orig_rect=(ix0, iy0, ix1, iy1),
                    new_rect=(ix0, iy0, ix1, iy1),
                    image_path=tmp_path,
                    image_xref=xref,
                )
                self._edit_records.append(record)

                # Paint a background-coloured cover over the orig rect
                # so the user sees the original image "pick up".
                bg = self._sample_background_color(
                    page_idx, (ix0, iy0, ix1, iy1)) or (255, 255, 255)
                # Persist the sampled cover colour so undo/redo can
                # faithfully rebuild the cover from the record alone.
                record.background_color = bg
                cover = QGraphicsRectItem(QRectF(rect_scene))
                cover.setBrush(QBrush(QColor(*bg)))
                cover.setPen(QPen(Qt.PenStyle.NoPen))
                cover.setZValue(9)
                self._scene.addItem(cover)
                self._live_edit_items.append(cover)
                self._record_covers[id(record)] = cover

                # Live pixmap that the user will drag around.
                visual = self._add_live_image_replacement(
                    rect_scene, tmp_path, record=record)
                if visual is None:
                    return False

                self._active_transform = TransformSelection(
                    self,
                    page_idx=page_idx,
                    rect_scene=rect_scene,
                    record=record,
                    visuals=[visual],
                    rotation=0.0,
                )
                self.status_changed.emit()
                return True

            # 2b) existing text block?
            text_dict = page.get_text("dict")
            best_block = None
            best_area = float("inf")
            for block in text_dict.get("blocks", []):
                if block.get("type", 0) != 0:
                    continue
                bbox = block.get("bbox")
                if not bbox or len(bbox) != 4:
                    continue
                x0, y0, x1, y1 = bbox
                if (x0 <= click_pt_x <= x1
                        and y0 <= click_pt_y <= y1):
                    area = max(1.0, (x1 - x0) * (y1 - y0))
                    if area < best_area:
                        best_area = area
                        best_block = block
            if best_block is None:
                return False

            bbox = best_block["bbox"]
            lines_text = []
            for line in best_block.get("lines", []):
                lines_text.append("".join(
                    span.get("text", "")
                    for span in line.get("spans", [])))
            text = "\n".join(lines_text).rstrip()

            meta = self._detect_paragraph_font(best_block)
            bg = self._sample_background_color(page_idx, bbox)

            sx0, sy0 = self._pt_to_scene(page_idx, bbox[0], bbox[1])
            sx1, sy1 = self._pt_to_scene(page_idx, bbox[2], bbox[3])
            rect_scene = QRectF(sx0, sy0,
                                max(14.0, sx1 - sx0),
                                max(14.0, sy1 - sy0))

            record = EditRecord(
                kind="text_move",
                page_idx=page_idx,
                orig_rect=tuple(bbox),
                new_rect=tuple(bbox),
                text=text,
                font_size=meta["size"],
                font_name=meta.get("pdf_font"),
                color=meta["color"],
                background_color=bg,
            )
            record.extra = {
                "family": meta.get("family", "Helvetica"),
                "bold": bool(meta.get("bold")),
                "italic": bool(meta.get("italic")),
            }
            self._edit_records.append(record)

            # Cover the original block with its sampled background.
            bg_fill = bg or (255, 255, 255)
            cover = QGraphicsRectItem(QRectF(rect_scene))
            cover.setBrush(QBrush(QColor(*bg_fill)))
            cover.setPen(QPen(Qt.PenStyle.NoPen))
            cover.setZValue(9)
            self._scene.addItem(cover)
            self._live_edit_items.append(cover)
            self._record_covers[id(record)] = cover

            # Live text item that the user will drag around.
            text_item = QGraphicsTextItem(text)
            fnt = QFont(meta.get("family", "Helvetica"))
            fnt.setBold(bool(meta.get("bold")))
            fnt.setItalic(bool(meta.get("italic")))
            fnt.setPointSizeF(
                max(6.0, meta["size"] * self._render_scale * 0.72))
            text_item.setFont(fnt)
            r, g, b = meta.get("color", (0, 0, 0))
            text_item.setDefaultTextColor(QColor(r, g, b))
            text_item.setTextWidth(rect_scene.width())
            text_item.setPos(rect_scene.x(), rect_scene.y())
            text_item.setZValue(11)
            self._scene.addItem(text_item)
            self._live_edit_items.append(text_item)
            self._item_to_record[text_item] = record

            self._active_transform = TransformSelection(
                self,
                page_idx=page_idx,
                rect_scene=rect_scene,
                record=record,
                visuals=[text_item],
                rotation=0.0,
            )
            self.status_changed.emit()
            return True
        except Exception as e:
            print(f"[BoltPDF] transform begin failed: {e}",
                  file=sys.stderr)
            return False
        finally:
            if doc is not None:
                try:
                    doc.close()
                except Exception:
                    pass

    def _crop_existing_image_to_tempfile(self, page_idx: int,
                                         bbox_pt: tuple) -> str | None:
        """Crop the rendered page pixmap to *bbox_pt* and save it to a
        temporary PNG.  Returns the path, or None on failure."""
        try:
            item = self._page_items.get(page_idx)
            if item is None:
                return None
            pix = item.pixmap()
            if pix is None or pix.isNull():
                return None
            sx0, sy0 = self._pt_to_scene(
                page_idx, bbox_pt[0], bbox_pt[1])
            sx1, sy1 = self._pt_to_scene(
                page_idx, bbox_pt[2], bbox_pt[3])
            page_y = self._page_positions.get(page_idx, 0)
            # Convert scene coords back to pixmap-local pixels.
            local_x = int(round(sx0 - item.x()))
            local_y = int(round(sy0 - item.y()))
            local_w = max(1, int(round(sx1 - sx0)))
            local_h = max(1, int(round(sy1 - sy0)))
            cropped = pix.copy(local_x, local_y, local_w, local_h)
            if cropped.isNull():
                return None
            fd, path = tempfile.mkstemp(
                prefix="boltpdf_xform_", suffix=".png")
            os.close(fd)
            cropped.save(path, "PNG")
            return path
        except Exception as e:
            print(f"[BoltPDF] crop existing image failed: {e}",
                  file=sys.stderr)
            return None

    # -- Delete selected item ---------------------------------------------
    def delete_active_selection(self) -> bool:
        """Remove whatever the Transform tool currently has selected —
        an existing PDF text block, an existing image, or a pending
        live-edit item — from the document.

        For content that exists in the source PDF (anything with an
        ``orig_rect``) the record's kind is flipped to
        ``text_delete`` / ``image_delete`` so the save worker redacts
        the original region and inserts nothing in its place.  For pure
        additions (no ``orig_rect``) the record is simply dropped from
        the queue.

        Returns True if something was deleted."""
        if not self._edit_mode:
            return False
        sel = self._active_transform
        if sel is None:
            return False

        record = sel.record
        visuals = list(sel.visuals)
        page_idx = sel.page_idx

        # Teardown the handles + body first so nothing lingers.
        try:
            sel.dismiss()
        except Exception:
            pass
        self._active_transform = None

        # Remove any live preview visuals from the scene.  These are
        # the text item / pixmap the Transform tool was steering.
        for v in visuals:
            try:
                self._scene.removeItem(v)
            except Exception:
                pass
            try:
                if v in self._live_edit_items:
                    self._live_edit_items.remove(v)
            except Exception:
                pass
            # Drop the item→record mapping so clicks don't find a
            # dangling reference.
            try:
                if v in self._item_to_record:
                    del self._item_to_record[v]
            except Exception:
                pass

        # Also remove any extra visuals tracked in _record_visuals that
        # might not have been in sel.visuals.
        if record is not None:
            extra = self._record_visuals.pop(id(record), [])
            for ev in extra:
                try:
                    if ev not in visuals:
                        self._scene.removeItem(ev)
                        if ev in self._live_edit_items:
                            self._live_edit_items.remove(ev)
                        if ev in self._item_to_record:
                            del self._item_to_record[ev]
                except Exception:
                    pass

        if record is None:
            self.status_changed.emit()
            return True

        # Ensure the original area is visually covered so the user
        # sees the deletion immediately.  If the record already has a
        # cover rect (text_move / image_move paths add one), great;
        # otherwise paint one now using the sampled / stored
        # background colour.
        cover = self._record_covers.get(id(record))
        if cover is None and record.orig_rect is not None:
            try:
                ox0, oy0, ox1, oy1 = record.orig_rect
                sx0, sy0 = self._pt_to_scene(page_idx, ox0, oy0)
                sx1, sy1 = self._pt_to_scene(page_idx, ox1, oy1)
                rect_scene = QRectF(
                    sx0, sy0, max(2.0, sx1 - sx0), max(2.0, sy1 - sy0))
                bg = record.background_color
                if bg is None:
                    bg = self._sample_background_color(
                        page_idx, record.orig_rect) or (255, 255, 255)
                new_cover = QGraphicsRectItem(rect_scene)
                new_cover.setBrush(QBrush(QColor(*bg)))
                new_cover.setPen(QPen(Qt.PenStyle.NoPen))
                new_cover.setZValue(9)
                self._scene.addItem(new_cover)
                self._live_edit_items.append(new_cover)
                self._record_covers[id(record)] = new_cover
            except Exception as e:
                print(f"[BoltPDF] delete cover paint failed: {e}",
                      file=sys.stderr)

        # Decide the fate of the record itself.
        if record.orig_rect is None:
            # Pure addition — drop it from the queue entirely, along
            # with its cover rect (if any).
            cover = self._record_covers.pop(id(record), None)
            if cover is not None:
                try:
                    self._scene.removeItem(cover)
                except Exception:
                    pass
                try:
                    if cover in self._live_edit_items:
                        self._live_edit_items.remove(cover)
                except Exception:
                    pass
            try:
                if record in self._edit_records:
                    self._edit_records.remove(record)
            except Exception:
                pass
        else:
            # Existing PDF content — convert into a delete record so
            # the save worker erases the original region without
            # inserting anything in its place.
            if record.kind.startswith("image"):
                record.kind = "image_delete"
            else:
                record.kind = "text_delete"
            record.new_rect = None
            record.text = None
            record.html = None
            record.image_path = None
            if record not in self._edit_records:
                self._edit_records.append(record)

        self._edit_checkpoint()
        self.status_changed.emit()
        return True

    # -- Live text editing ------------------------------------------------
    _FONT_FLAG_ITALIC = 1 << 1
    _FONT_FLAG_SERIF = 1 << 2
    _FONT_FLAG_MONO = 1 << 3
    _FONT_FLAG_BOLD = 1 << 4

    def _detect_paragraph_font(self, block: dict) -> dict:
        """Pull the dominant font / size / colour out of a PyMuPDF text
        block.  Returns a dict with ``size`` (pt), ``color`` (r,g,b),
        ``bold``, ``italic``, ``serif``, ``mono``, ``pdf_font`` (the
        PyMuPDF builtin font name, e.g. 'helv-b', 'tiro-i') and
        ``family`` (best-guess Qt font family)."""
        size = 11.0
        color = (0, 0, 0)
        bold = italic = serif = mono = False
        font_name = ""
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                size = float(span.get("size", 11.0) or 11.0)
                c = span.get("color", 0) or 0
                color = ((c >> 16) & 0xFF, (c >> 8) & 0xFF, c & 0xFF)
                flags = int(span.get("flags", 0) or 0)
                italic = bool(flags & self._FONT_FLAG_ITALIC)
                serif = bool(flags & self._FONT_FLAG_SERIF)
                mono = bool(flags & self._FONT_FLAG_MONO)
                bold = bool(flags & self._FONT_FLAG_BOLD)
                font_name = span.get("font", "") or ""
                # Name-based fallback: PyMuPDF's flags are sometimes
                # zero on embedded subset fonts, but the name almost
                # always tells us.
                low = font_name.lower()
                if "bold" in low or "-b" in low:
                    bold = True
                if "italic" in low or "oblique" in low or "-i" in low:
                    italic = True
                if ("times" in low or "serif" in low or "tiro" in low
                        or "roman" in low):
                    serif = True
                if "cour" in low or "mono" in low:
                    mono = True
                # Use the first span's metadata as representative.
                break
            if font_name:
                break

        # Map to a PyMuPDF builtin font name.
        if mono:
            base = "cour"
        elif serif:
            base = "tiro"
        else:
            base = "helv"
        suffix = ""
        if bold and italic:
            suffix = "bi"
        elif bold:
            suffix = "b"
        elif italic:
            suffix = "i"
        pdf_font = f"{base}-{suffix}" if suffix else base

        # Pick a Qt font family that looks roughly similar.
        if mono:
            family = "Courier New"
        elif serif:
            family = "Times New Roman"
        else:
            family = "Helvetica"

        return {
            "size": size,
            "color": color,
            "bold": bold,
            "italic": italic,
            "serif": serif,
            "mono": mono,
            "pdf_font": pdf_font,
            "family": family,
            "orig_font_name": font_name,
        }

    def _sample_background_color(self, page_idx: int,
                                 bbox_pt: tuple):
        """Sample the rendered page image near *bbox_pt* (in PDF
        points) and return a (r, g, b) tuple that represents the
        dominant colour just outside the block's edges.  Returns None
        if no sample could be taken (e.g. the page isn't loaded)."""
        item = self._page_items.get(page_idx)
        if item is None:
            return None
        try:
            pixmap = item.pixmap()
            qimg = pixmap.toImage()
        except Exception:
            return None
        if qimg.isNull():
            return None

        scale = self._render_scale or 1.0
        x0 = int(round(bbox_pt[0] * scale))
        y0 = int(round(bbox_pt[1] * scale))
        x1 = int(round(bbox_pt[2] * scale))
        y1 = int(round(bbox_pt[3] * scale))
        w = qimg.width()
        h = qimg.height()

        def clamp_x(v):
            return max(0, min(w - 1, v))

        def clamp_y(v):
            return max(0, min(h - 1, v))

        samples = []
        cx = (x0 + x1) // 2
        margin = 4
        # Above, below, left, right of the block.
        candidates = [
            (cx, y0 - margin),
            (cx, y1 + margin),
            (x0 - margin, (y0 + y1) // 2),
            (x1 + margin, (y0 + y1) // 2),
            (x0 - margin, y0 - margin),
            (x1 + margin, y0 - margin),
            (x0 - margin, y1 + margin),
            (x1 + margin, y1 + margin),
        ]
        for px, py in candidates:
            px = clamp_x(px)
            py = clamp_y(py)
            try:
                pixel = qimg.pixel(px, py)
            except Exception:
                continue
            r = (pixel >> 16) & 0xFF
            g = (pixel >> 8) & 0xFF
            b = pixel & 0xFF
            samples.append((r, g, b))
        if not samples:
            return None
        # Pick the most common sample (mode), falling back to average.
        from collections import Counter
        counter = Counter(samples)
        most, _ = counter.most_common(1)[0]
        return most

    def try_begin_inline_text_reedit_at(self, scene_pos: QPointF) -> bool:
        """If the click landed on a text item we already drew for a
        previous edit, reopen the inline editor over it so the user
        can keep editing the same paragraph.  Returns True when the
        click was consumed."""
        if not self._edit_mode:
            return False
        if self._inline_editor is not None:
            self._commit_inline_text_edit()

        # Find a text item that (a) has a record mapping and (b)
        # contains the click point.  Image records are skipped — those
        # belong to the Transform tool, not the text editor.
        target_item = None
        target_record = None
        for it, rec in list(self._item_to_record.items()):
            if not isinstance(it, QGraphicsTextItem):
                continue
            if rec is None or rec.kind.startswith("image"):
                continue
            if rec.kind in ("text_delete",):
                continue
            try:
                br = it.sceneBoundingRect()
            except Exception:
                continue
            if br.contains(scene_pos):
                target_item = it
                target_record = rec
                break
        if target_item is None or target_record is None:
            return False

        # Dismiss any active transform selection so its handles don't
        # fight with the inline editor we're about to open.
        if self._active_transform is not None:
            try:
                self._active_transform.dismiss()
            except Exception:
                pass
            self._active_transform = None

        # Compute the rect the editor should sit in.  Prefer the live
        # item's current sceneBoundingRect so that if the user moved
        # or resized it via the Transform tool, the re-edit opens at
        # the new position.
        try:
            rect_scene = QRectF(target_item.sceneBoundingRect())
        except Exception:
            return False
        if rect_scene.width() < 40.0:
            rect_scene.setWidth(40.0)
        if rect_scene.height() < 16.0:
            rect_scene.setHeight(16.0)

        # Pull the cover rect paired with this record (if any) so we
        # can hide it while the editor is open and restore / replace
        # it on commit / cancel.
        existing_cover = self._record_covers.get(id(target_record))
        saved_visuals = [target_item]
        if existing_cover is not None:
            saved_visuals.append(existing_cover)
        # Hide the live preview items while the editor is open so the
        # user doesn't see a stale copy of the text behind the widget.
        for v in saved_visuals:
            try:
                v.setVisible(False)
            except Exception:
                pass

        page_idx = target_record.page_idx
        orig_text = target_record.text or ""
        current_html = target_record.html
        meta_family = None
        meta_size = target_record.font_size or 11.0
        meta_color = target_record.color or (0, 0, 0)
        meta_bold = False
        meta_italic = False
        bg = target_record.background_color

        # Best-effort: pick up the live item's font so bold/italic/
        # family survive a re-edit even if they weren't persisted on
        # the record's scalar fields.
        try:
            live_font = target_item.font()
            meta_family = live_font.family()
            meta_bold = live_font.bold()
            meta_italic = live_font.italic()
            if live_font.pointSizeF() > 0:
                meta_size = live_font.pointSizeF() / (
                    self._render_scale * 0.72)
        except Exception:
            pass
        if not meta_family:
            meta_family = "Helvetica"

        try:
            editor = QTextEdit()
            if current_html:
                editor.setHtml(current_html)
            else:
                editor.setPlainText(orig_text)
            editor.setFrameStyle(0)
            editor.setVerticalScrollBarPolicy(
                Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            editor.setHorizontalScrollBarPolicy(
                Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            editor.setAcceptRichText(True)

            bg_css = ("#ffffff" if bg is None
                      else f"rgb({bg[0]},{bg[1]},{bg[2]})")
            fg_css = "rgb({},{},{})".format(*meta_color)
            editor.setStyleSheet(
                "QTextEdit {{"
                " background: {bg}; color: {fg};"
                " border: 2px solid #3399ff;"
                " padding: 1px;"
                "}}".format(bg=bg_css, fg=fg_css))
            editor.viewport().setAutoFillBackground(True)

            fnt = QFont(meta_family)
            fnt.setBold(meta_bold)
            fnt.setItalic(meta_italic)
            fnt.setPointSizeF(
                max(6.0, meta_size * self._render_scale * 0.72))
            editor.setFont(fnt)

            proxy = self._scene.addWidget(editor)
            proxy.setPos(rect_scene.x(), rect_scene.y())
            proxy.resize(rect_scene.width(), rect_scene.height())
            proxy.setZValue(100)

            editor.setFocus(Qt.FocusReason.OtherFocusReason)
            editor.selectAll()

            self._inline_editor = proxy
            self._inline_editor_ctx = {
                "page_idx": page_idx,
                "orig_rect": target_record.orig_rect,
                "orig_text": orig_text,
                "rect_scene": rect_scene,
                "font_size": meta_size,
                "color": meta_color,
                "background_color": bg,
                "pdf_font": target_record.font_name,
                "family": meta_family,
                "bold": meta_bold,
                "italic": meta_italic,
                "cover_item": None,
                # Re-edit bookkeeping — commit / cancel use these to
                # update (or restore) the existing record instead of
                # creating a fresh one.
                "reedit_record": target_record,
                "reedit_visuals": saved_visuals,
                "reedit_cover": existing_cover,
            }

            # Formatting toolbar — same as the first-time editor path.
            try:
                tb = InlineFormatToolbar(self._view.viewport())
                tb.bind(editor)
                tb.apply_initial_style(
                    meta_family,
                    meta_size * self._render_scale * 0.72,
                    meta_bold,
                    meta_italic,
                    meta_color,
                    bg,
                )
                tb.done_clicked.connect(self._commit_inline_text_edit)
                tb.cancel_clicked.connect(self._cancel_inline_text_edit)

                top_left_view = self._view.mapFromScene(
                    rect_scene.topLeft())
                tb.adjustSize()
                tb_w = tb.width()
                tb_h = tb.height()
                viewport_w = self._view.viewport().width()
                tx = max(4, min(top_left_view.x(),
                                viewport_w - tb_w - 4))
                if top_left_view.y() - tb_h - 6 >= 4:
                    ty = top_left_view.y() - tb_h - 6
                else:
                    ty = top_left_view.y() + 6
                tb.move(int(tx), int(ty))
                tb.show()
                tb.raise_()
                editor.setFocus(Qt.FocusReason.OtherFocusReason)
                editor.selectAll()
                self._inline_editor_toolbar = tb
            except Exception as e:
                print(f"[BoltPDF] re-edit toolbar setup failed: {e}",
                      file=sys.stderr)
                self._inline_editor_toolbar = None
            return True
        except Exception as e:
            print(f"[BoltPDF] inline re-edit setup failed: {e}",
                  file=sys.stderr)
            # Restore the hidden visuals so the previous edit remains
            # visible if we failed to spin up the editor.
            for v in saved_visuals:
                try:
                    v.setVisible(True)
                except Exception:
                    pass
            self._inline_editor = None
            self._inline_editor_ctx = None
            return False

    def try_begin_inline_text_edit_at(self, scene_pos: QPointF) -> bool:
        """If the click landed on a real text block, open an inline
        QTextEdit over the whole paragraph so the user can retype it.
        Returns True if the click was consumed."""
        if not self._edit_mode:
            return False
        # If another inline editor is already open, commit it first.
        if self._inline_editor is not None:
            self._commit_inline_text_edit()
        page_idx = self._page_at_scene_point(scene_pos)
        if page_idx is None:
            return False

        try:
            import fitz
        except Exception:
            return False

        click_pt_x, click_pt_y = self._scene_to_pt(
            page_idx, scene_pos.x(), scene_pos.y())

        doc = None
        try:
            doc = fitz.open(self._doc_path)
            if page_idx < 0 or page_idx >= len(doc):
                return False
            page = doc[page_idx]
            text_dict = page.get_text("dict")
        except Exception as e:
            print(f"[BoltPDF] inline text probe failed: {e}",
                  file=sys.stderr)
            return False
        finally:
            if doc is not None:
                try:
                    doc.close()
                except Exception:
                    pass

        # Find the smallest text BLOCK containing the click — that
        # gives us paragraph granularity instead of single-line spans.
        best_block = None
        best_area = float("inf")
        for block in text_dict.get("blocks", []):
            if block.get("type", 0) != 0:
                continue  # not a text block
            bbox = block.get("bbox")
            if not bbox or len(bbox) != 4:
                continue
            x0, y0, x1, y1 = bbox
            if (x0 <= click_pt_x <= x1
                    and y0 <= click_pt_y <= y1):
                area = max(1.0, (x1 - x0) * (y1 - y0))
                if area < best_area:
                    best_area = area
                    best_block = block
        if best_block is None:
            return False

        bbox = best_block["bbox"]
        # Reconstruct the paragraph text: join spans within each line,
        # join lines with newlines.  This preserves the user's mental
        # model of "editing a paragraph".
        lines_text = []
        for line in best_block.get("lines", []):
            line_str = "".join(
                span.get("text", "") for span in line.get("spans", []))
            lines_text.append(line_str)
        text = "\n".join(lines_text).rstrip()

        meta = self._detect_paragraph_font(best_block)
        bg = self._sample_background_color(page_idx, bbox)

        sx0, sy0 = self._pt_to_scene(page_idx, bbox[0], bbox[1])
        sx1, sy1 = self._pt_to_scene(page_idx, bbox[2], bbox[3])
        # Pad a few pixels so the editor's opaque background fully
        # overlaps any anti-aliased glyph fringes from the rendered
        # page pixmap sitting underneath.
        pad = 3.0
        rect_scene = QRectF(sx0 - pad, sy0 - pad,
                            max(40.0, (sx1 - sx0) + pad * 2.0),
                            max(16.0, (sy1 - sy0) + pad * 2.0))

        # No separate cover rect — the editor itself is fully opaque
        # (filled with the sampled paragraph background, or white if we
        # couldn't sample one) so it completely hides the rendered text
        # underneath.  The user is visually editing the original text
        # in place instead of overlaying a new layer.
        cover_item = None

        try:
            editor = QTextEdit()
            editor.setPlainText(text)
            editor.setFrameStyle(0)
            editor.setVerticalScrollBarPolicy(
                Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            editor.setHorizontalScrollBarPolicy(
                Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            # Rich text IS now allowed so the formatting toolbar can
            # apply bold / italic / colours directly to the editor.
            editor.setAcceptRichText(True)

            # Opaque background matching the sampled paragraph colour
            # (falling back to white) so the underlying rendered text is
            # fully hidden while the user is editing.  Store the sampled
            # colour on the ctx below so the save worker and the Bg
            # picker still know what the page really looked like.
            bg_css = ("#ffffff" if bg is None
                      else f"rgb({bg[0]},{bg[1]},{bg[2]})")
            fg_css = "rgb({},{},{})".format(*meta["color"])
            editor.setStyleSheet(
                "QTextEdit {{"
                " background: {bg}; color: {fg};"
                " border: 2px solid #3399ff;"
                " padding: 1px;"
                "}}".format(bg=bg_css, fg=fg_css))
            editor.viewport().setAutoFillBackground(True)

            fnt = QFont(meta["family"])
            fnt.setBold(meta["bold"])
            fnt.setItalic(meta["italic"])
            fnt.setPointSizeF(
                max(6.0, meta["size"] * self._render_scale * 0.72))
            editor.setFont(fnt)

            proxy = self._scene.addWidget(editor)
            proxy.setPos(rect_scene.x(), rect_scene.y())
            proxy.resize(rect_scene.width(), rect_scene.height())
            proxy.setZValue(100)

            editor.setFocus(Qt.FocusReason.OtherFocusReason)
            editor.selectAll()

            self._inline_editor = proxy
            self._inline_editor_ctx = {
                "page_idx": page_idx,
                "orig_rect": tuple(bbox),
                "orig_text": text,
                "rect_scene": rect_scene,
                "font_size": meta["size"],
                "color": meta["color"],
                "background_color": bg,
                "pdf_font": meta["pdf_font"],
                "family": meta["family"],
                "bold": meta["bold"],
                "italic": meta["italic"],
                "cover_item": cover_item,
            }

            # ---- formatting toolbar -------------------------------
            # Floating QFrame parented to the view's viewport so it
            # doesn't rotate / scale with the scene.  Positioned just
            # above the editor if there's room, otherwise just below.
            try:
                tb = InlineFormatToolbar(self._view.viewport())
                tb.bind(editor)
                tb.apply_initial_style(
                    meta["family"],
                    meta["size"] * self._render_scale * 0.72,
                    meta["bold"],
                    meta["italic"],
                    meta["color"],
                    bg,
                )
                tb.done_clicked.connect(self._commit_inline_text_edit)
                tb.cancel_clicked.connect(self._cancel_inline_text_edit)

                # Convert the editor's scene rect to viewport coords so
                # we can position the toolbar next to it.
                top_left_view = self._view.mapFromScene(
                    rect_scene.topLeft())
                tb.adjustSize()
                tb_w = tb.width()
                tb_h = tb.height()
                viewport_w = self._view.viewport().width()
                # Horizontal: clamp so the toolbar stays on-screen.
                tx = max(4, min(top_left_view.x(),
                                viewport_w - tb_w - 4))
                # Vertical: above the editor if there's room, else below.
                if top_left_view.y() - tb_h - 6 >= 4:
                    ty = top_left_view.y() - tb_h - 6
                else:
                    ty = top_left_view.y() + int(
                        self._view.mapFromScene(
                            rect_scene.bottomLeft()).y()
                        - top_left_view.y()) + 6
                tb.move(int(tx), int(ty))
                tb.show()
                tb.raise_()
                # Re-apply focus to the editor after the toolbar steals
                # focus during its show.
                editor.setFocus(Qt.FocusReason.OtherFocusReason)
                editor.selectAll()
                self._inline_editor_toolbar = tb
            except Exception as e:
                print(f"[BoltPDF] format toolbar setup failed: {e}",
                      file=sys.stderr)
                self._inline_editor_toolbar = None

            return True
        except Exception as e:
            print(f"[BoltPDF] inline editor setup failed: {e}",
                  file=sys.stderr)
            self._inline_editor = None
            self._inline_editor_ctx = None
            if self._inline_editor_toolbar is not None:
                try:
                    self._inline_editor_toolbar.deleteLater()
                except Exception:
                    pass
                self._inline_editor_toolbar = None
            return False

    def _commit_inline_text_edit(self):
        """Apply the inline editor's text to the scene (live preview)
        and queue a text_edit EditRecord for Save As."""
        proxy = self._inline_editor
        ctx = self._inline_editor_ctx
        toolbar = self._inline_editor_toolbar
        # Clear state BEFORE teardown so any signal cascades during
        # removeItem can't re-enter us.
        self._inline_editor = None
        self._inline_editor_ctx = None
        self._inline_editor_toolbar = None
        if proxy is None or ctx is None:
            # Still clean up the toolbar if it somehow exists.
            if toolbar is not None:
                try:
                    toolbar.deleteLater()
                except Exception:
                    pass
            return

        new_text = ""
        new_html = None
        try:
            widget = proxy.widget()
            if isinstance(widget, QTextEdit):
                new_text = widget.toPlainText()
                new_html = widget.toHtml()
            elif isinstance(widget, QLineEdit):
                new_text = widget.text()
            else:
                new_text = ""
        except Exception:
            new_text = ""

        # Pull the toolbar's final background choice (None → transparent
        # → don't paint a cover rect at all) before tearing it down.
        bg_override = None
        bg_transparent = False
        if toolbar is not None:
            try:
                chosen = toolbar.bg_color()
                if chosen is None:
                    bg_transparent = True
                else:
                    bg_override = (chosen.red(), chosen.green(),
                                   chosen.blue())
            except Exception:
                pass
            try:
                toolbar.deleteLater()
            except Exception:
                pass

        try:
            self._scene.removeItem(proxy)
        except Exception:
            pass

        # Remove the temporary "hide-original-text" cover that we placed
        # beneath the editor.  The commit path below will add its own
        # permanent cover rect (or leave the page bare if the user chose
        # a transparent background).
        pre_cover = ctx.get("cover_item")
        if pre_cover is not None:
            try:
                self._scene.removeItem(pre_cover)
            except Exception:
                pass

        reedit_record = ctx.get("reedit_record")
        reedit_visuals = ctx.get("reedit_visuals") or []
        reedit_cover = ctx.get("reedit_cover")

        if (new_text == ctx.get("orig_text", "")
                and bg_override is None and not bg_transparent
                and reedit_record is None):
            return  # nothing to do

        rect_scene: QRectF = ctx["rect_scene"]
        # Prefer the toolbar's explicit background override, then fall
        # back to the auto-sampled colour, then white.
        if bg_transparent:
            bg_for_save = None
        elif bg_override is not None:
            bg_for_save = bg_override
        else:
            bg_for_save = ctx.get("background_color")

        # Re-edit: drop the old hidden preview items entirely — we'll
        # rebuild them below with the new text.
        if reedit_record is not None:
            for v in reedit_visuals:
                try:
                    self._scene.removeItem(v)
                except Exception:
                    pass
                if v in self._live_edit_items:
                    try:
                        self._live_edit_items.remove(v)
                    except Exception:
                        pass
                if v in self._item_to_record:
                    try:
                        del self._item_to_record[v]
                    except Exception:
                        pass
            if reedit_cover is not None:
                try:
                    self._record_covers.pop(id(reedit_record), None)
                except Exception:
                    pass

        text_item = None
        # Live preview: fill the original paragraph's rect with the
        # chosen background (unless transparent) and draw the new text
        # on top using the matched font.  If rich HTML is available we
        # render it via setHtml so formatting shows in the preview.
        new_cover = None
        try:
            if not bg_transparent:
                bg_fill = bg_for_save or (255, 255, 255)
                new_cover = QGraphicsRectItem(QRectF(rect_scene))
                new_cover.setBrush(QBrush(QColor(*bg_fill)))
                new_cover.setPen(QPen(Qt.PenStyle.NoPen))
                new_cover.setZValue(10)
                self._scene.addItem(new_cover)
                self._live_edit_items.append(new_cover)

            text_item = QGraphicsTextItem()
            fnt = QFont(ctx.get("family", "Helvetica"))
            fnt.setBold(bool(ctx.get("bold")))
            fnt.setItalic(bool(ctx.get("italic")))
            fnt.setPointSizeF(
                max(6.0, ctx["font_size"] * self._render_scale * 0.72))
            text_item.setFont(fnt)
            r, g, b = ctx.get("color", (0, 0, 0))
            text_item.setDefaultTextColor(QColor(r, g, b))
            if new_html:
                text_item.setHtml(new_html)
            else:
                text_item.setPlainText(new_text)
            text_item.setPos(rect_scene.x(), rect_scene.y())
            text_item.setTextWidth(rect_scene.width())
            text_item.setZValue(11)
            self._scene.addItem(text_item)
            self._live_edit_items.append(text_item)
        except Exception as e:
            print(f"[BoltPDF] live text paint failed: {e}",
                  file=sys.stderr)

        if reedit_record is not None:
            # Update the existing record in place so Save As picks up
            # the latest text/formatting without creating a duplicate.
            record = reedit_record
            record.text = new_text
            record.html = new_html
            record.font_size = ctx["font_size"]
            record.font_name = ctx.get("pdf_font")
            record.color = ctx["color"]
            record.background_color = bg_for_save
            if record.kind not in ("text_edit", "text_add"):
                record.kind = "text_edit"
        else:
            record = EditRecord(
                kind="text_edit",
                page_idx=ctx["page_idx"],
                orig_rect=ctx["orig_rect"],
                new_rect=ctx["orig_rect"],
                text=new_text,
                html=new_html,
                font_size=ctx["font_size"],
                font_name=ctx.get("pdf_font"),
                color=ctx["color"],
                background_color=bg_for_save,
            )
            self._edit_records.append(record)
        # Persist the matched font so undo/redo can rebuild the live
        # text preview identically from the record alone.
        record.extra = {
            "family": ctx.get("family", "Helvetica"),
            "bold": bool(ctx.get("bold")),
            "italic": bool(ctx.get("italic")),
        }
        # Let the Transform tool find this record by clicking on the
        # live text preview item we just added above.
        if text_item is not None:
            try:
                self._item_to_record[text_item] = record
            except Exception:
                pass
        if new_cover is not None:
            try:
                self._record_covers[id(record)] = new_cover
            except Exception:
                pass
        self._edit_checkpoint()
        self.status_changed.emit()

    def _cancel_inline_text_edit(self):
        """Dismiss the inline editor without saving or live-drawing."""
        proxy = self._inline_editor
        ctx = self._inline_editor_ctx
        toolbar = self._inline_editor_toolbar
        self._inline_editor = None
        self._inline_editor_ctx = None
        self._inline_editor_toolbar = None
        if toolbar is not None:
            try:
                toolbar.deleteLater()
            except Exception:
                pass
        # Drop the "hide-original-text" cover rect so the underlying
        # PDF text reappears exactly as it was.
        if ctx is not None:
            pre_cover = ctx.get("cover_item")
            if pre_cover is not None:
                try:
                    self._scene.removeItem(pre_cover)
                except Exception:
                    pass
            # Re-edit cancel: restore the live preview items we hid so
            # the earlier edit remains untouched.
            reedit_visuals = ctx.get("reedit_visuals") or []
            for v in reedit_visuals:
                try:
                    v.setVisible(True)
                except Exception:
                    pass
        if proxy is None:
            return
        try:
            self._scene.removeItem(proxy)
        except Exception:
            pass

    # -- Save As ----------------------------------------------------------
    def _collect_pending_edit_records(self) -> list:
        """Return the queued EditRecords.  (Overlay-derived records were
        removed when the overlay system was dropped.)"""
        return list(self._edit_records)

    def save_edited_pdf_as(self, main_window):
        """Prompt for an output path and kick off the save worker."""
        if not self._edit_mode:
            return

        records = self._collect_pending_edit_records()
        if not records:
            QMessageBox.information(
                main_window, "Save Edited PDF",
                "No edits to save yet.  Use Add Text, Add Image, or "
                "Replace Image first.")
            return

        default_dir = os.path.dirname(self._doc_path) if self._doc_path else ""
        stem = os.path.splitext(os.path.basename(self._doc_path or "edited"))[0]
        default_path = os.path.join(default_dir, f"{stem}_edited.pdf")
        out_path, _ = QFileDialog.getSaveFileName(
            main_window, "Save Edited PDF", default_path, "PDF Files (*.pdf)")
        if not out_path:
            return
        if not out_path.lower().endswith(".pdf"):
            out_path += ".pdf"

        # Pause the renderer so PyMuPDF has the document to itself.
        self._pause_renderer()

        self._save_progress = QProgressDialog(
            "Saving edited PDF...", "Cancel", 0, 100, main_window)
        self._save_progress.setWindowTitle("Save Edited PDF")
        self._save_progress.setMinimumDuration(0)
        self._save_progress.setValue(0)
        self._save_progress.setModal(True)

        self._save_worker = SaveEditedPdfWorker(
            self._doc_path, out_path, records, parent=self)
        self._save_worker.progress.connect(self._on_save_progress)
        self._save_worker.finished_ok.connect(
            lambda p: self._on_save_done(p, main_window))
        self._save_worker.error_occurred.connect(
            lambda m: self._on_save_error(m, main_window))
        self._save_progress.canceled.connect(self._save_worker.cancel)
        self._save_worker.start()

    def _on_save_progress(self, pct):
        if self._save_progress is not None:
            self._save_progress.setValue(pct)

    def _on_save_done(self, output_path, main_window):
        if self._save_progress is not None:
            self._save_progress.close()
            self._save_progress = None
        self._resume_renderer()
        reply = QMessageBox.question(
            main_window, "Save Complete",
            f"Saved to:\n{os.path.basename(output_path)}\n\n"
            "Open the edited file in a new tab?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        # Clear the record queue now that it has been flushed to disk,
        # and drop the crash-recovery journal (edits are now saved).
        self._edit_records.clear()
        self._clear_recovery()
        # Tidy up the live preview items — the saved file now contains
        # the real edits, and leaving phantom previews on the source
        # tab is misleading.
        for it in self._live_edit_items:
            try:
                self._scene.removeItem(it)
            except Exception:
                pass
        self._live_edit_items.clear()
        if reply == QMessageBox.StandardButton.Yes:
            main_window.open_pdf_in_new_tab(output_path)
        self.status_changed.emit()

    def _on_save_error(self, msg, main_window):
        if self._save_progress is not None:
            self._save_progress.close()
            self._save_progress = None
        self._resume_renderer()
        QMessageBox.critical(
            main_window, "Save Failed",
            f"Could not save edited PDF:\n{msg}")
        self.status_changed.emit()
    # === end Edit Mode ====================================================

    def _on_images_found(self, page_idx, images):
        """Create ImageOverlay items for each detected image on the page."""
        # Mark this page as covered even if it has zero images — we
        # don't want to keep re-scanning empty pages.
        self._image_detected_pages.add(page_idx)
        page_y = self._page_positions.get(page_idx, 0)
        page_list = self._image_overlays_by_page.setdefault(page_idx, [])
        for info in images:
            # Convert PDF user-space coords to scene coords.
            # PDF origin is bottom-left; scene origin is top-left.
            page_h = info["page_h"]
            scale = self._render_scale
            sx = info["left"] * scale
            sy = (page_h - info["top"]) * scale + page_y
            sw = (info["right"] - info["left"]) * scale
            sh = (info["top"] - info["bottom"]) * scale

            if sw < 5 or sh < 5:
                continue  # skip tiny artifacts

            overlay = ImageOverlay(
                sx, sy, sw, sh, page_idx, info["obj_index"])
            self._scene.addItem(overlay)
            self._image_overlays.append(overlay)
            page_list.append(overlay)

    def _on_img_detect_error(self, msg):
        print(f"[BoltPDF] Image detection error: {msg}", file=sys.stderr)

    def _on_img_detect_done(self):
        # Only bail out with "no images found" if we've scanned every
        # page of the document without producing a single overlay.
        # A partial scan that covered a sub-range of pages (because we
        # restarted the detector for a newly visible window) should
        # never trigger the bail-out.
        if (not self._image_overlays
                and len(self._image_detected_pages) >= self._num_pages):
            QMessageBox.information(
                self.window(), "Export Images",
                "No embedded images were found in this PDF.")
            self.exit_image_select_mode()
        self.status_changed.emit()

    def _ensure_image_detector_covers(self, page_idx: int):
        """Make sure pages around *page_idx* have had their overlays
        built.  Called whenever image-select mode is on and the user's
        visible window changes.

        Behaviour:
          * If every page in the current window is already detected,
            does nothing.
          * If the background detector is still running, just refocuses
            it so the missing pages jump to the front of its queue.
          * If the detector has already finished, spins up a fresh
            detector whose queue is only the still-undetected window
            pages — keeps the work bounded and fast.
        """
        if not self._image_select_mode or not self._num_pages:
            return
        lo = max(0, page_idx - self.PAGE_BUFFER)
        hi = min(self._num_pages - 1, page_idx + self.PAGE_BUFFER)
        missing = [
            i for i in range(lo, hi + 1)
            if i not in self._image_detected_pages
        ]
        if not missing:
            return
        if self._img_detector and self._img_detector.isRunning():
            # Existing detector still has work to do — re-prioritise
            # so the missing window pages scan next.
            self._img_detector.set_focus(page_idx)
            return
        # Detector has exited; spin up a new one to cover just the
        # undetected pages.  A short, targeted queue means the new
        # detector finishes almost immediately.
        self._img_detector = ImageDetectorWorker(
            self._doc_path, missing, focus=page_idx, parent=self)
        self._img_detector.page_images_found.connect(self._on_images_found)
        self._img_detector.error_occurred.connect(self._on_img_detect_error)
        self._img_detector.all_done.connect(self._on_img_detect_done)
        self._img_detector.start()

    def _page_at_scene_y(self, y: float) -> int:
        """Return the page index whose vertical range contains *y*."""
        for idx in range(self._num_pages):
            py = self._page_positions.get(idx, 0)
            ph = self._page_heights.get(idx, 0)
            if py <= y <= py + ph:
                return idx
        return 0

    def handle_image_overlay_click(self, scene_pos):
        """Toggle the smallest ImageOverlay under the click (ignores
        full-page background images when smaller images overlap).

        Uses the page-indexed overlay dict so only overlays on the
        clicked page are checked — O(images-on-page) instead of
        O(all-images-in-document).
        """
        page_idx = self._page_at_scene_y(scene_pos.y())
        # If the clicked page hasn't been detected yet (detector was
        # already finished by the time the user scrolled here), kick off
        # a fresh detector run so this page's overlays get built.  The
        # first click on a newly revealed page won't find anything, but
        # a split second later the overlays appear and subsequent clicks
        # work normally.
        if (self._image_select_mode
                and page_idx not in self._image_detected_pages):
            self._ensure_image_detector_covers(page_idx)
        best = None
        best_area = float("inf")
        # Check the clicked page and its immediate neighbours (overlays
        # that straddle a page boundary are rare but possible).
        for p in (page_idx - 1, page_idx, page_idx + 1):
            for overlay in self._image_overlays_by_page.get(p, ()):
                r = overlay.rect().translated(overlay.pos())
                if r.contains(scene_pos):
                    area = r.width() * r.height()
                    if area < best_area:
                        best_area = area
                        best = overlay
        if best is not None:
            # If we're in the middle of an Edit-Mode "replace image"
            # pick, route the click into the replace flow instead of
            # toggling the overlay's export selection.
            if self._edit_mode and self._edit_image_pick_mode == "replace":
                return self._handle_image_replace_pick(best)
            best.toggle_selected()
            self.status_changed.emit()
            return True
        return False

    def export_selected_images(self, main_window):
        """Export all currently-selected image overlays as JPEG files."""
        selected = [o for o in self._image_overlays if o.selected]
        if not selected:
            QMessageBox.information(
                main_window, "Export Images", "No images selected.")
            return

        # PDFium is not thread-safe — pause renderer while export runs
        self._pause_renderer()

        pdf_dir = os.path.dirname(self._doc_path)
        pdf_stem = os.path.splitext(os.path.basename(self._doc_path))[0]
        export_dir = os.path.join(pdf_dir, f"{pdf_stem}_exports")

        image_list = [{"page_idx": o.page_idx, "obj_index": o.obj_index}
                      for o in selected]

        self._export_progress = QProgressDialog(
            "Exporting images...", "Cancel", 0, len(image_list), main_window)
        self._export_progress.setWindowTitle("Export Images")
        self._export_progress.setMinimumDuration(0)
        self._export_progress.setValue(0)
        self._export_progress.setModal(True)

        self._export_worker = ImageExportWorker(
            self._doc_path, image_list, export_dir, pdf_stem,
            parent=self)
        self._export_worker.progress.connect(self._on_export_progress)
        self._export_worker.finished_ok.connect(
            lambda p: self._on_export_done(p, main_window))
        self._export_worker.error_occurred.connect(
            lambda m: self._on_export_error(m, main_window))
        self._export_progress.canceled.connect(self._export_worker.terminate)
        self._export_worker.start()

    def _on_export_progress(self, count):
        if hasattr(self, '_export_progress') and self._export_progress:
            self._export_progress.setValue(count)

    def _on_export_done(self, output_dir, main_window):
        if hasattr(self, '_export_progress') and self._export_progress:
            self._export_progress.close()

        # Resume page rendering now that export is finished
        self._resume_renderer()

        reply = QMessageBox.question(
            main_window, "Export Complete",
            f"Saved to:\n{output_dir}\n\n"
            "Do you want to open the folder?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply == QMessageBox.StandardButton.Yes:
            if sys.platform == "win32":
                os.startfile(output_dir)
            else:
                subprocess.Popen(["xdg-open", output_dir])

        # Leave image-select mode after export
        self.exit_image_select_mode()

    def _on_export_error(self, msg, main_window):
        if hasattr(self, '_export_progress') and self._export_progress:
            self._export_progress.close()
        # Resume page rendering even after export failure
        self._resume_renderer()
        QMessageBox.critical(main_window, "Export Failed", f"Error:\n{msg}")

    # =====================================================================
    # OCR
    # =====================================================================
    def get_visible_page(self):
        viewport_rect = self._view.mapToScene(
            self._view.viewport().rect()).boundingRect()
        center_y = viewport_rect.center().y()
        best_page, best_dist = 0, float("inf")
        for idx in range(self._num_pages):
            page_center = self._page_positions[idx] + self._page_heights[idx] / 2
            d = abs(center_y - page_center)
            if d < best_dist:
                best_dist = d
                best_page = idx
        return best_page

    def _get_buffer_pages(self, center):
        pages = []
        for off in [-1, 0, 1, 2]:
            p = center + off
            if 0 <= p < self._num_pages and p not in self._ocr_done_pages:
                pages.append(p)
        return pages

    def toggle_ocr(self):
        if not self._doc_path or self._num_pages == 0:
            return
        if self._ocr_active:
            self._ocr_active = False
            self._clear_ocr()
            self._view.set_ocr_mode(False)
            self._disconnect_scroll_listener()
            self.status_changed.emit()
            return

        self._ocr_active = True
        self._view.set_ocr_mode(True)
        self._run_ocr_for_visible()
        self._connect_scroll_listener()
        self.status_changed.emit()

    def _run_ocr_for_visible(self):
        if not self._ocr_active or not self._doc_path:
            return
        pages = self._get_buffer_pages(self.get_visible_page())
        if not pages:
            self.status_changed.emit()
            return
        if self._ocr_worker and self._ocr_worker.isRunning():
            self._ocr_worker.cancel()
            self._ocr_worker.wait()

        self._ocr_worker = OCRWorker(
            self._doc_path, pages, self._render_scale, self)
        self._ocr_worker.page_ocr_done.connect(self._on_page_ocr_done)
        self._ocr_worker.status_update.connect(lambda _: self.status_changed.emit())
        self._ocr_worker.error_occurred.connect(self._on_ocr_error)
        self._ocr_worker.all_done.connect(self._on_ocr_finished)
        self._ocr_worker.start()
        self.status_changed.emit()

    def _on_page_ocr_done(self, page_idx, lines):
        page_y = self._page_positions.get(page_idx, 0)
        overlays = []

        for line_words in lines:
            first_in_line = True
            for wd in line_words:
                text = wd.get("text", "").strip()
                if not text:
                    continue
                idx = self._next_word_index
                self._next_word_index += 1

                if first_in_line:
                    self._line_breaks.add(idx)
                    first_in_line = False

                item = WordOverlay(
                    wd["x"], page_y + wd["y"], wd["w"], wd["h"],
                    text, line_id=f"{page_idx}_{id(line_words)}",
                    word_index=idx,
                )
                self._scene.addItem(item)
                overlays.append(item)
                self._all_words.append(item)

        self._page_word_overlays[page_idx] = overlays
        self._ocr_done_pages.add(page_idx)
        self._all_words.sort(key=lambda w: w.word_index)

    def _on_ocr_error(self, msg):
        print(f"[BoltPDF] OCR error: {msg}", file=sys.stderr)
        if "not found" in msg.lower() or "powershell" in msg.lower():
            QMessageBox.warning(self.window(), "OCR Error",
                f"Windows OCR encountered an issue:\n\n{msg}\n\n"
                "Make sure you are running Windows 10 or later\n"
                "and have at least one language pack installed\n"
                "(Settings > Time & Language > Language).")

    def _on_ocr_finished(self):
        if self._ocr_active:
            self.status_changed.emit()

    def _clear_ocr(self):
        for overlays in self._page_word_overlays.values():
            for item in overlays:
                self._scene.removeItem(item)
        self._page_word_overlays.clear()
        self._all_words.clear()
        self._line_breaks.clear()
        self._next_word_index = 0
        self._ocr_done_pages.clear()
        self._view.clear_selection()

    def select_all_ocr(self):
        if not self._ocr_active or not self._all_words:
            return
        any_sel = any(w.selected for w in self._all_words)
        for w in self._all_words:
            w.set_selected(not any_sel)
        if not any_sel:
            self._view._anchor_idx = self._all_words[0].word_index
            self._view._extent_idx = self._all_words[-1].word_index

    def copy_selected_text(self):
        selected = [w for w in self._all_words if w.selected]
        if not selected:
            return

        parts = []
        for w in selected:
            if parts and w.word_index in self._line_breaks:
                parts.append("\n")
            elif parts:
                parts.append(" ")
            parts.append(w.word_text)

        QApplication.clipboard().setText("".join(parts))

    # -- Scroll auto-OCR -------------------------------------------------
    def _connect_scroll_listener(self):
        if self._scroll_ocr_timer is None:
            self._scroll_ocr_timer = QTimer(self)
            self._scroll_ocr_timer.setSingleShot(True)
            self._scroll_ocr_timer.setInterval(300)
            self._scroll_ocr_timer.timeout.connect(self._on_scroll_settled)
        vsb = self._view.verticalScrollBar()
        if vsb:
            vsb.valueChanged.connect(self._on_scroll_changed)

    def _disconnect_scroll_listener(self):
        vsb = self._view.verticalScrollBar()
        if vsb:
            try:
                vsb.valueChanged.disconnect(self._on_scroll_changed)
            except TypeError:
                pass
        if self._scroll_ocr_timer:
            self._scroll_ocr_timer.stop()

    def _on_scroll_changed(self):
        if self._scroll_ocr_timer:
            self._scroll_ocr_timer.start()

    def _on_scroll_settled(self):
        if self._ocr_active:
            self._run_ocr_for_visible()

    # -- Scroll-based render reprioritisation --------------------------------
    def _connect_render_scroll_listener(self):
        """Reprioritise the renderer whenever the user scrolls."""
        if not hasattr(self, '_render_scroll_timer') or self._render_scroll_timer is None:
            self._render_scroll_timer = QTimer(self)
            self._render_scroll_timer.setSingleShot(True)
            # 10 ms is low enough that jump-to-page feels instant, but
            # still coalesces high-frequency wheel events.
            self._render_scroll_timer.setInterval(10)
            self._render_scroll_timer.timeout.connect(self._on_render_scroll_settled)
            # Only connect the scroll signal once (guard is the timer check above)
            vsb = self._view.verticalScrollBar()
            if vsb:
                vsb.valueChanged.connect(self._on_render_scroll_changed)

    def _on_render_scroll_changed(self):
        if hasattr(self, '_render_scroll_timer') and self._render_scroll_timer:
            self._render_scroll_timer.start()

    def _on_render_scroll_settled(self):
        vis = self.get_visible_page()
        if self._renderer and self._renderer.isRunning():
            self._renderer.set_focus(vis)
        # Refresh search highlights with a separate 100 ms throttle so
        # rapid scrolling doesn't call refresh_search_highlights on
        # every 10 ms tick.
        if self._search_matches:
            if not hasattr(self, '_search_scroll_timer') or self._search_scroll_timer is None:
                self._search_scroll_timer = QTimer(self)
                self._search_scroll_timer.setSingleShot(True)
                self._search_scroll_timer.setInterval(100)
                self._search_scroll_timer.timeout.connect(self.refresh_search_highlights)
            self._search_scroll_timer.start()
        # Make sure image-select mode has overlays for every page in
        # the new viewport window — reprioritises the running
        # detector if it's still working, or spins up a fresh one for
        # the undetected pages if it has already finished.
        if self._image_select_mode:
            self._ensure_image_detector_covers(vis)
        # Load/unload cached pages based on current viewport
        self._manage_viewport_pages()

    # -- Helpers ----------------------------------------------------------
    def _stop_workers(self):
        # Workers that expose a cooperative .cancel() flag get asked to
        # stop first, then joined with a short timeout.
        for attr in ("_renderer", "_ocr_worker", "_img_detector",
                     "_search_worker"):
            w = getattr(self, attr, None)
            if w is None:
                continue
            try:
                if w.isRunning():
                    if hasattr(w, "cancel"):
                        w.cancel()
                    w.wait(2000)
                    if w.isRunning():
                        # Last resort — QThread.terminate is unsafe in
                        # general, but by this point we're tearing down
                        # the tab and would otherwise leak a worker.
                        w.terminate()
                        w.wait(500)
            except Exception:
                pass

        # Short-lived workers (export / rebuild) don't have a cancel
        # flag but we still need to wait for them before the window
        # disappears, otherwise the multiprocessing Pools they own
        # leave lingering BoltPDF.exe subprocesses behind.
        for attr in ("_export_worker", "_rebuild_worker", "_save_worker"):
            w = getattr(self, attr, None)
            if w is None:
                continue
            try:
                if w.isRunning():
                    w.wait(2000)
                    if w.isRunning():
                        w.terminate()
                        w.wait(500)
            except Exception:
                pass

    # -- PDFium thread safety helpers ----------------------------------------
    # PDFium is NOT thread-safe at the C level, even with separate document
    # handles.  Any worker that opens its own PdfDocument (ImageDetectorWorker,
    # RebuildWorker, ImageExportWorker) must NOT run concurrently with the
    # PageRenderer.  These helpers pause/resume the renderer around such ops.

    def _pause_renderer(self):
        """Stop the page renderer so another worker can use PDFium safely."""
        if self._renderer and self._renderer.isRunning():
            self._renderer.cancel()
            self._renderer.wait()

    def _resume_renderer(self):
        """Restart the page renderer, skipping pages already cached."""
        if not self._doc_path or not self._num_pages:
            return
        # Don't restart if renderer is somehow still running
        if self._renderer and self._renderer.isRunning():
            return
        self._renderer = PageRenderer(
            self._doc_path, self._render_scale, self._num_pages, self)
        # Pre-populate the done set so already-rendered pages are skipped
        with self._renderer._lock:
            self._renderer._done = set(self._cached_pages)
        vis = self.get_visible_page() if self._num_pages else 0
        self._renderer.set_focus(vis)
        self._renderer.page_ready.connect(self._on_page_rendered)
        self._renderer.error_occurred.connect(self._on_render_error)
        self._renderer.start()

    def _cleanup_cache(self):
        """Remove the disk cache directory and all cached page images."""
        if self._cache_dir and os.path.isdir(self._cache_dir):
            import shutil
            try:
                shutil.rmtree(self._cache_dir, ignore_errors=True)
            except Exception:
                pass
        self._cache_dir = None
        self._cached_pages.clear()

    # -- Search ---------------------------------------------------------------
    # Minimum query length before a search is dispatched.  Single-character
    # queries can match thousands of times and create just as many scene
    # items, so we skip them entirely.
    _SEARCH_MIN_QUERY_LEN = 2

    # How many pages either side of the viewport to draw highlights for.
    # Matches on pages further away are still tracked (for Next/Prev and
    # the counter) but their QGraphicsRectItems are deferred until the
    # user scrolls near them.
    _SEARCH_HIGHLIGHT_WINDOW = 3

    def start_search(self, query: str):
        """Kick off a streaming async search.  Results arrive in
        batches via ``_on_search_batch`` so the user can jump to the
        first hit immediately while the rest of the document is still
        being scanned."""
        self._cancel_search_worker()
        self._clear_search_highlights()
        self._search_matches.clear()
        self._search_page_index.clear()
        self._search_drawn_pages.clear()
        self._search_current = -1
        self._search_query = query
        self._search_done = False
        self._search_highlight_hwm = 0
        self._search_status_pending = False

        if (not query or len(query) < self._SEARCH_MIN_QUERY_LEN
                or not self._doc_path):
            self._search_done = True
            self.status_changed.emit()
            return

        worker = SearchWorker(
            self._doc_path, query,
            text_cache=self._search_text_cache,
            parent=self)
        worker.batch_ready.connect(self._on_search_batch)
        worker.finished_search.connect(self._on_search_finished)
        worker.text_cache_ready.connect(self._on_search_text_cache)
        self._search_worker = worker
        worker.start()

    def _cancel_search_worker(self):
        if self._search_worker is not None:
            try:
                self._search_worker.cancel()
                self._search_worker.batch_ready.disconnect()
                self._search_worker.finished_search.disconnect()
                self._search_worker.text_cache_ready.disconnect()
            except Exception:
                pass
            self._search_worker = None

    def _emit_search_status(self):
        """Throttled status emission — at most every 300 ms during
        active search to avoid flooding _sync_toolbar."""
        if self._search_done:
            # Search finished — emit immediately
            self._search_status_pending = False
            self.status_changed.emit()
            return
        if self._search_status_pending:
            return  # timer already scheduled
        self._search_status_pending = True
        QTimer.singleShot(300, self._flush_search_status)

    def _flush_search_status(self):
        self._search_status_pending = False
        self.status_changed.emit()

    def _on_search_text_cache(self, cache: dict):
        self._search_text_cache = cache

    def _on_search_batch(self, query: str, matches: list):
        """Called when the worker flushes a batch of new matches."""
        if query != self._search_query:
            return
        first_batch = len(self._search_matches) == 0
        base = len(self._search_matches)
        self._search_matches.extend(matches)

        # Update page index for the new matches
        for i, (pg, _) in enumerate(matches, start=base):
            self._search_page_index.setdefault(pg, []).append(i)

        # On the very first batch, jump to the closest match
        if first_batch and self._search_matches:
            visible = self.get_visible_page()
            best_idx = 0
            best_dist = abs(self._search_matches[0][0] - visible)
            for i, (pg, _) in enumerate(self._search_matches):
                d = abs(pg - visible)
                if d < best_dist:
                    best_dist = d
                    best_idx = i
            self._search_current = best_idx
            self._scroll_to_match(self._search_current)

        # Draw highlights only for newly arrived matches that fall
        # within the current viewport window.
        self._draw_new_highlights_fast(matches, base)
        self._emit_search_status()

    def _on_search_finished(self, query: str):
        if query != self._search_query:
            return
        self._search_worker = None
        self._search_done = True
        self._emit_search_status()

    def search_next(self) -> int:
        if not self._search_matches:
            return -1
        old = self._search_current
        self._search_current = (
            (self._search_current + 1) % len(self._search_matches))
        self._scroll_to_match(self._search_current)
        self._update_current_highlight(old)
        return self._search_current

    def search_prev(self) -> int:
        if not self._search_matches:
            return -1
        old = self._search_current
        self._search_current = (
            (self._search_current - 1) % len(self._search_matches))
        self._scroll_to_match(self._search_current)
        self._update_current_highlight(old)
        return self._search_current

    def clear_search(self):
        self._cancel_search_worker()
        self._clear_search_highlights()
        self._search_matches.clear()
        self._search_page_index.clear()
        self._search_drawn_pages.clear()
        self._search_current = -1
        self._search_query = ""
        self._search_done = True
        self._search_highlight_hwm = 0

    def refresh_search_highlights(self):
        """Called on scroll — incrementally add/remove highlights for
        pages entering/leaving the viewport window."""
        if not self._search_matches:
            return
        visible = self.get_visible_page()
        lo = visible - self._SEARCH_HIGHLIGHT_WINDOW
        hi = visible + self._SEARCH_HIGHLIGHT_WINDOW
        want = set(range(max(0, lo), min(self._num_pages, hi + 1)))
        # Also always include the page of the current match
        if 0 <= self._search_current < len(self._search_matches):
            want.add(self._search_matches[self._search_current][0])

        # Pages to remove (no longer in window)
        to_remove = self._search_drawn_pages - want
        if to_remove:
            keep = []
            for item in self._search_highlights:
                pg = item.data(0)  # we store page_idx in data(0)
                if pg in to_remove:
                    self._scene.removeItem(item)
                else:
                    keep.append(item)
            self._search_highlights = keep
            self._search_drawn_pages -= to_remove

        # Pages to add (newly entered window)
        to_add = want - self._search_drawn_pages
        if to_add:
            for pg in to_add:
                for mi in self._search_page_index.get(pg, []):
                    item = self._make_highlight_item_fast(mi, visible)
                    if item is not None:
                        self._scene.addItem(item)
                        self._search_highlights.append(item)
            self._search_drawn_pages |= to_add

    def _clear_search_highlights(self):
        for item in self._search_highlights:
            try:
                self._scene.removeItem(item)
            except Exception:
                pass
        self._search_highlights.clear()
        self._search_drawn_pages.clear()
        self._search_highlight_hwm = 0

    def _make_highlight_item_fast(self, match_idx: int,
                                   visible: int) -> QGraphicsRectItem | None:
        """Create a highlight rect.  *visible* is pre-computed to avoid
        calling get_visible_page() per item."""
        page_idx, rect = self._search_matches[match_idx]
        is_current = (match_idx == self._search_current)
        x0, y0, x1, y1 = rect
        sx0, sy0 = self._pt_to_scene(page_idx, x0, y0)
        sx1, sy1 = self._pt_to_scene(page_idx, x1, y1)
        r = QRectF(sx0, sy0, sx1 - sx0, sy1 - sy0)
        item = QGraphicsRectItem(r)
        if is_current:
            item.setBrush(QBrush(QColor(255, 140, 0, 120)))
            item.setPen(QPen(QColor(255, 140, 0), 2.0))
        else:
            item.setBrush(QBrush(QColor(255, 255, 0, 70)))
            item.setPen(QPen(Qt.PenStyle.NoPen))
        item.setZValue(50)
        # Store page index for fast removal by page
        item.setData(0, page_idx)
        # Store match index for surgical current-highlight updates
        item.setData(1, match_idx)
        return item

    def _draw_new_highlights_fast(self, new_matches: list, base: int):
        """Draw highlights only for newly arrived matches that are
        within the viewport window.  O(new_matches) not O(all)."""
        visible = self.get_visible_page()
        lo = visible - self._SEARCH_HIGHLIGHT_WINDOW
        hi = visible + self._SEARCH_HIGHLIGHT_WINDOW
        for i, (pg, _) in enumerate(new_matches, start=base):
            is_current = (i == self._search_current)
            if not is_current and not (lo <= pg <= hi):
                continue
            item = self._make_highlight_item_fast(i, visible)
            if item is not None:
                self._scene.addItem(item)
                self._search_highlights.append(item)
                self._search_drawn_pages.add(pg)

    def _update_current_highlight(self, old_idx: int):
        """Surgically update just the old and new current-match items
        instead of a full redraw."""
        new_idx = self._search_current
        # Update existing items that match old/new indices
        old_found = False
        new_found = False
        for item in self._search_highlights:
            mi = item.data(1)
            if mi == old_idx and old_idx >= 0:
                # Demote old current → normal yellow
                item.setBrush(QBrush(QColor(255, 255, 0, 70)))
                item.setPen(QPen(Qt.PenStyle.NoPen))
                old_found = True
            elif mi == new_idx:
                # Promote new current → orange
                item.setBrush(QBrush(QColor(255, 140, 0, 120)))
                item.setPen(QPen(QColor(255, 140, 0), 2.0))
                new_found = True
            if old_found and new_found:
                break

        # If the new current wasn't drawn (outside viewport), add it
        if not new_found and 0 <= new_idx < len(self._search_matches):
            visible = self.get_visible_page()
            item = self._make_highlight_item_fast(new_idx, visible)
            if item is not None:
                self._scene.addItem(item)
                self._search_highlights.append(item)
                self._search_drawn_pages.add(
                    self._search_matches[new_idx][0])

    def _scroll_to_match(self, idx: int):
        """Scroll the view so that match *idx* is visible and centred."""
        if idx < 0 or idx >= len(self._search_matches):
            return
        page_idx, rect = self._search_matches[idx]
        x0, y0, x1, y1 = rect
        cx = (x0 + x1) / 2.0
        cy = (y0 + y1) / 2.0
        sx, sy = self._pt_to_scene(page_idx, cx, cy)
        self._view.centerOn(sx, sy)

    @property
    def search_match_count(self) -> int:
        return len(self._search_matches)

    @property
    def search_current_index(self) -> int:
        return self._search_current

    def cleanup(self):
        """Called when the tab is about to be closed."""
        self._stop_workers()
        self._cleanup_cache()


# ---------------------------------------------------------------------------
# Export to Word — dialog + worker
# ---------------------------------------------------------------------------
def _parse_page_spec(spec: str, max_page: int) -> list[int] | str:
    """Parse a page specification like '1-5,8,12-15' into a sorted list
    of 0-based page indices.  Returns an error string on bad input."""
    pages = set()
    for part in spec.replace(" ", "").split(","):
        if not part:
            continue
        if "-" in part:
            pieces = part.split("-", 1)
            try:
                start, end = int(pieces[0]), int(pieces[1])
            except ValueError:
                return f"Invalid range: {part}"
            if start < 1 or end < 1:
                return f"Page numbers must be positive: {part}"
            if start > max_page or end > max_page:
                return f"Page {max(start, end)} exceeds document ({max_page} pages)"
            if start > end:
                start, end = end, start
            pages.update(range(start - 1, end))  # convert to 0-based
        else:
            try:
                p = int(part)
            except ValueError:
                return f"Invalid page number: {part}"
            if p < 1:
                return f"Page numbers must be positive: {part}"
            if p > max_page:
                return f"Page {p} exceeds document ({max_page} pages)"
            pages.add(p - 1)  # convert to 0-based
    if not pages:
        return "No pages specified"
    return sorted(pages)


class ExportWordDialog(QDialog):
    """Dialog letting the user choose which pages to export as a Word doc."""

    ENTIRE_FILE = 0
    CURRENT_PAGE = 1
    PAGE_RANGE = 2

    def __init__(self, num_pages: int, current_page: int, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Export To Docx")
        self.setMinimumWidth(400)
        self.num_pages = num_pages
        self.current_page = current_page   # 0-based

        layout = QVBoxLayout(self)

        group = QGroupBox("Pages to export")
        g_layout = QVBoxLayout(group)

        self._btn_group = QButtonGroup(self)

        self._rb_all = QRadioButton(f"Entire file ({num_pages} pages)")
        self._rb_all.setChecked(True)
        self._btn_group.addButton(self._rb_all, self.ENTIRE_FILE)
        g_layout.addWidget(self._rb_all)

        self._rb_current = QRadioButton(
            f"Current page (page {current_page + 1})")
        self._btn_group.addButton(self._rb_current, self.CURRENT_PAGE)
        g_layout.addWidget(self._rb_current)

        self._rb_range = QRadioButton("Page range:")
        self._btn_group.addButton(self._rb_range, self.PAGE_RANGE)
        range_row = QHBoxLayout()
        range_row.addWidget(self._rb_range)
        self._range_edit = QLineEdit()
        self._range_edit.setPlaceholderText("e.g. 1-26  or  1,2,5,9,76")
        self._range_edit.setEnabled(False)
        range_row.addWidget(self._range_edit, 1)
        g_layout.addLayout(range_row)

        self._rb_range.toggled.connect(self._range_edit.setEnabled)
        self._range_edit.textChanged.connect(
            lambda: self._rb_range.setChecked(True))

        layout.addWidget(group)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok
            | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

        self.setStyleSheet("""
            QDialog { background: #2d2d2d; }
            QGroupBox { color: #e0e0e0; border: 1px solid #555;
                        border-radius: 4px; margin-top: 12px;
                        padding: 16px 12px 12px; font-size: 13px; }
            QGroupBox::title { subcontrol-origin: margin;
                               padding: 0 6px; }
            QRadioButton { color: #e0e0e0; font-size: 13px;
                           spacing: 6px; padding: 4px 0; }
            QLineEdit { background: #1e1e1e; color: #e0e0e0;
                        border: 1px solid #555; border-radius: 4px;
                        padding: 4px 8px; font-size: 13px; }
            QLineEdit:disabled { background: #353535; color: #777; }
            QPushButton { background: #3d3d3d; color: #e0e0e0;
                          border: 1px solid #555; border-radius: 4px;
                          padding: 6px 16px; font-size: 13px; }
            QPushButton:hover { background: #4a4a4a; }
        """)

    def get_page_indices(self) -> list[int] | str:
        """Return a sorted list of 0-based page indices, or an error string."""
        mode = self._btn_group.checkedId()
        if mode == self.ENTIRE_FILE:
            return list(range(self.num_pages))
        elif mode == self.CURRENT_PAGE:
            return [self.current_page]
        else:
            return _parse_page_spec(self._range_edit.text(), self.num_pages)


def _mp_render_page_for_docx(args):
    """Render one PDF page to a full-resolution PNG for docx embedding.

    Returning the rendered bitmap's pixel dimensions (rather than the
    page's mediabox dimensions) means rotated pages come out the right
    way up with matching section sizes — pypdfium2 auto-applies the
    page rotation when rendering, so the PNG dimensions already reflect
    the visible page.
    """
    doc_path, page_idx, dpi, out_dir = args
    import pypdfium2 as _pdf
    doc = _pdf.PdfDocument(doc_path)
    page = doc[page_idx]
    scale = dpi / 72.0
    bitmap = page.render(scale=scale)
    pil = bitmap.to_pil()
    px_w, px_h = pil.size
    out_path = os.path.join(out_dir, f"page_{page_idx:05d}.png")
    # PNG with optimize=False keeps encoding fast; the temp file is
    # deleted right after the docx is packed.
    pil.save(out_path, "PNG", optimize=False)
    try:
        bitmap.close()
    except Exception:
        pass
    page.close()
    doc.close()
    return (page_idx, out_path, px_w, px_h)


class ExportWordWorker(QThread):
    """Rasterise every selected PDF page and build a Word document
    where each page is a pixel-perfect copy of the original.

    The docx is constructed with one section per PDF page.  The
    section's page size is set to the exact dimensions of the PDF
    page (so landscape, portrait, and mixed-orientation documents all
    render correctly), margins are zeroed, and a single full-page
    image is inserted.  Because the image IS the page, every element —
    text, fonts, positioning, tables, figures, colours, headers,
    footers, page numbers — is preserved exactly as it appears in the
    original PDF.
    """
    progress = pyqtSignal(int)
    finished_ok = pyqtSignal(str)
    error_occurred = pyqtSignal(str)

    # Render DPI — 200 gives a crisp, print-quality result without
    # making the docx enormous.  Bump to 300 for print-master output.
    _RENDER_DPI = 200

    def __init__(self, doc_path, page_indices, output_path, parent=None):
        super().__init__(parent)
        self.doc_path = doc_path
        self.page_indices = page_indices
        self.output_path = output_path

    def run(self):
        import tempfile
        import shutil

        tmpdir = None
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches, Pt, Emu
            from docx.enum.section import WD_SECTION, WD_ORIENT

            n = len(self.page_indices)
            if n == 0:
                self.error_occurred.emit("No pages selected for export.")
                return

            # Staging area for per-page PNGs.  Wiped on the way out.
            tmpdir = tempfile.mkdtemp(prefix="boltpdf_docx_")

            # -- Phase 1: parallel rasterisation --------------------------
            workers = max(1, min(n, os.cpu_count() or 1))
            args = [(self.doc_path, idx, self._RENDER_DPI, tmpdir)
                    for idx in self.page_indices]

            rendered: dict = {}   # page_idx -> (png_path, px_w, px_h)
            with multiprocessing.Pool(processes=workers) as pool:
                for page_idx, png_path, px_w, px_h in pool.imap_unordered(
                        _mp_render_page_for_docx, args):
                    rendered[page_idx] = (png_path, px_w, px_h)
                    done = len(rendered)
                    # Rasterisation is the heavy step — give it 80%.
                    self.progress.emit(int((done / n) * 80))

            # -- Phase 2: build the docx ----------------------------------
            doc = DocxDocument()

            # Zero out Normal's spacing so the picture sits flush against
            # the page edges.  Without this, python-docx's default 1.15
            # line spacing + 8pt after-paragraph would push the image
            # down and force a second blank page.
            normal = doc.styles['Normal']
            normal.paragraph_format.space_before = Pt(0)
            normal.paragraph_format.space_after = Pt(0)
            normal.paragraph_format.line_spacing = 1.0

            for i, page_idx in enumerate(self.page_indices):
                data = rendered.get(page_idx)
                if data is None:
                    continue
                png_path, px_w, px_h = data

                # Pixels → inches using the render DPI.  This is the
                # physical size of the rendered image and therefore the
                # size the docx page should be.
                w_in = px_w / float(self._RENDER_DPI)
                h_in = px_h / float(self._RENDER_DPI)

                # First page uses the document's default section; every
                # subsequent page is its own NEW_PAGE section so page
                # sizes can vary independently across the document.
                if i == 0:
                    section = doc.sections[0]
                else:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)

                # Match each PDF page's orientation individually.  The
                # orientation flag is set *before* the dimensions because
                # python-docx's orientation setter can swap width/height
                # — by writing the dimensions last we guarantee the
                # section carries the exact rendered size regardless of
                # any internal swap.  This is what makes mixed-orientation
                # PDFs (e.g. page 3 landscape, page 4 portrait) come
                # through correctly.
                landscape = w_in > h_in
                section.orientation = (
                    WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT)
                section.page_width = Inches(w_in)
                section.page_height = Inches(h_in)
                section.left_margin = Inches(0)
                section.right_margin = Inches(0)
                section.top_margin = Inches(0)
                section.bottom_margin = Inches(0)
                section.header_distance = Inches(0)
                section.footer_distance = Inches(0)
                try:
                    section.gutter = Inches(0)
                except Exception:
                    pass

                # One paragraph, one full-page image.  Setting only the
                # width makes python-docx compute the height from the
                # image's own aspect ratio — which (because we rendered
                # at a uniform DPI) matches the section's page height
                # exactly.
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0
                run = p.add_run()
                run.add_picture(png_path, width=Inches(w_in))

                self.progress.emit(80 + int(((i + 1) / n) * 20))

            doc.save(self.output_path)
            self.progress.emit(100)
            self.finished_ok.emit(self.output_path)

        except Exception as e:
            self.error_occurred.emit(str(e))
        finally:
            # Always clean up the rendered PNGs, even on error.
            if tmpdir and os.path.isdir(tmpdir):
                try:
                    shutil.rmtree(tmpdir, ignore_errors=True)
                except Exception:
                    pass


# ---------------------------------------------------------------------------
# Export Pages as JPEGs
# ---------------------------------------------------------------------------
class ExportPagesDialog(QDialog):
    """Dialog letting the user choose which pages to export as JPEG images."""

    ENTIRE_FILE = 0
    CURRENT_PAGE = 1
    PAGE_RANGE = 2

    def __init__(self, num_pages: int, current_page: int, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Export Pages")
        self.setMinimumWidth(400)
        self.num_pages = num_pages
        self.current_page = current_page

        layout = QVBoxLayout(self)

        group = QGroupBox("Pages to export as JPEG")
        g_layout = QVBoxLayout(group)

        self._btn_group = QButtonGroup(self)

        self._rb_all = QRadioButton(f"Entire file ({num_pages} pages)")
        self._rb_all.setChecked(True)
        self._btn_group.addButton(self._rb_all, self.ENTIRE_FILE)
        g_layout.addWidget(self._rb_all)

        self._rb_current = QRadioButton(
            f"Current page (page {current_page + 1})")
        self._btn_group.addButton(self._rb_current, self.CURRENT_PAGE)
        g_layout.addWidget(self._rb_current)

        self._rb_range = QRadioButton("Page range:")
        self._btn_group.addButton(self._rb_range, self.PAGE_RANGE)
        range_row = QHBoxLayout()
        range_row.addWidget(self._rb_range)
        self._range_edit = QLineEdit()
        self._range_edit.setPlaceholderText("e.g. 1-26  or  1,2,5,9,76")
        self._range_edit.setEnabled(False)
        range_row.addWidget(self._range_edit, 1)
        g_layout.addLayout(range_row)

        self._rb_range.toggled.connect(self._range_edit.setEnabled)
        self._range_edit.textChanged.connect(
            lambda: self._rb_range.setChecked(True))

        layout.addWidget(group)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok
            | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

        self.setStyleSheet("""
            QDialog { background: #2d2d2d; }
            QGroupBox { color: #e0e0e0; border: 1px solid #555;
                        border-radius: 4px; margin-top: 12px;
                        padding: 16px 12px 12px; font-size: 13px; }
            QGroupBox::title { subcontrol-origin: margin;
                               padding: 0 6px; }
            QRadioButton { color: #e0e0e0; font-size: 13px;
                           spacing: 6px; padding: 4px 0; }
            QLineEdit { background: #1e1e1e; color: #e0e0e0;
                        border: 1px solid #555; border-radius: 4px;
                        padding: 4px 8px; font-size: 13px; }
            QLineEdit:disabled { background: #353535; color: #777; }
            QPushButton { background: #3d3d3d; color: #e0e0e0;
                          border: 1px solid #555; border-radius: 4px;
                          padding: 6px 16px; font-size: 13px; }
            QPushButton:hover { background: #4a4a4a; }
        """)

    def get_page_indices(self) -> list | str:
        """Return a sorted list of 0-based page indices, or an error string."""
        mode = self._btn_group.checkedId()
        if mode == self.ENTIRE_FILE:
            return list(range(self.num_pages))
        elif mode == self.CURRENT_PAGE:
            return [self.current_page]
        else:
            return _parse_page_spec(self._range_edit.text(), self.num_pages)


def _mp_render_page_jpeg(args):
    """Render one PDF page to a JPEG file at high quality."""
    doc_path, page_idx, dpi, out_dir, name_prefix = args
    import pypdfium2 as _pdf
    doc = _pdf.PdfDocument(doc_path)
    page = doc[page_idx]
    scale = dpi / 72.0
    bitmap = page.render(scale=scale)
    pil = bitmap.to_pil()
    out_path = os.path.join(
        out_dir, f"{name_prefix}_page_{page_idx + 1:04d}.jpg")
    pil.save(out_path, "JPEG", quality=95)
    try:
        bitmap.close()
    except Exception:
        pass
    page.close()
    doc.close()
    return out_path


class ExportPagesWorker(QThread):
    """Render selected PDF pages as JPEG images into a folder."""
    progress = pyqtSignal(int)
    finished_ok = pyqtSignal(str)          # output folder path
    error_occurred = pyqtSignal(str)

    def __init__(self, doc_path: str, page_indices: list,
                 out_folder: str, name_prefix: str,
                 dpi: int = 200, parent=None):
        super().__init__(parent)
        self.doc_path = doc_path
        self.page_indices = page_indices
        self.out_folder = out_folder
        self.name_prefix = name_prefix
        self.dpi = dpi

    def run(self):
        try:
            os.makedirs(self.out_folder, exist_ok=True)
            import multiprocessing as mp
            n = len(self.page_indices)
            tasks = [
                (self.doc_path, idx, self.dpi, self.out_folder,
                 self.name_prefix)
                for idx in self.page_indices
            ]
            done = 0
            pool = mp.Pool(processes=min(mp.cpu_count(), n, 8))
            try:
                for _ in pool.imap_unordered(_mp_render_page_jpeg, tasks):
                    done += 1
                    self.progress.emit(int(done / n * 100))
            finally:
                pool.close()
                pool.join()
            self.progress.emit(100)
            self.finished_ok.emit(self.out_folder)
        except Exception as e:
            self.error_occurred.emit(str(e))


class PreviewWorker(QThread):
    """Render every page as a low-res JPEG and assemble them into a
    lightweight preview PDF saved next to the original file."""
    progress = pyqtSignal(int)
    finished_ok = pyqtSignal(str)          # output file path
    error_occurred = pyqtSignal(str)

    def __init__(self, doc_path: str, out_path: str,
                 target_width: int, parent=None):
        super().__init__(parent)
        self.doc_path = doc_path
        self.out_path = out_path
        self.target_width = target_width
        self._cancelled = False

    def cancel(self):
        self._cancelled = True

    def run(self):
        try:
            import fitz  # PyMuPDF
            src = fitz.open(self.doc_path)
            out = fitz.open()  # new empty PDF
            n = len(src)
            for i in range(n):
                if self._cancelled:
                    src.close()
                    out.close()
                    return
                page = src[i]
                # Calculate scale to hit target_width at 72 DPI
                pw = page.rect.width  # in points (72 dpi)
                scale = self.target_width / pw if pw > 0 else 1.0
                mat = fitz.Matrix(scale, scale)
                pix = page.get_pixmap(matrix=mat)
                # Create a new page matching the pixmap dimensions
                # (pixmap size is in pixels; at 72 DPI, 1px = 1pt)
                img_w = pix.width
                img_h = pix.height
                new_page = out.new_page(width=img_w, height=img_h)
                # Insert the pixmap as a JPEG image onto the page
                img_bytes = pix.tobytes("jpeg", jpg_quality=60)
                new_page.insert_image(
                    fitz.Rect(0, 0, img_w, img_h),
                    stream=img_bytes)
                self.progress.emit(int((i + 1) / n * 100))
            src.close()
            out.save(self.out_path, deflate=True, garbage=3)
            out.close()
            self.progress.emit(100)
            self.finished_ok.emit(self.out_path)
        except Exception as e:
            self.error_occurred.emit(str(e))


# ---------------------------------------------------------------------------
# Auto-update — checks GitHub for new versions, downloads + swaps exe
# ---------------------------------------------------------------------------
def _parse_version(v: str):
    """Turn '1.2.3' into a comparable tuple (1, 2, 3)."""
    return tuple(int(x) for x in v.strip().split("."))


class UpdateChecker(QThread):
    """Background thread that checks for a newer version.

    Signals:
        update_available(remote_version: str, download_url: str, changelog: str)
        no_update()
        check_failed(error: str)
    """
    update_available = pyqtSignal(str, str, str)   # version, url, changelog
    no_update = pyqtSignal()
    check_failed = pyqtSignal(str)

    def run(self):
        try:
            import urllib.request
            req = urllib.request.Request(
                _UPDATE_URL,
                headers={"User-Agent": f"BoltPDF/{__version__}"})
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            remote = data.get("version", "0.0.0")
            dl_url = data.get("download_url", "")
            changelog = data.get("changelog", "")
            if _parse_version(remote) > _parse_version(__version__):
                self.update_available.emit(remote, dl_url, changelog)
            else:
                self.no_update.emit()
        except Exception as e:
            self.check_failed.emit(str(e))


class UpdateDownloader(QThread):
    """Downloads the new release (a .zip of the onedir bundle) to a temp file.

    Signals:
        progress(percent: int)
        finished_ok(temp_path: str)
        error_occurred(error: str)
    """
    progress = pyqtSignal(int)
    finished_ok = pyqtSignal(str)
    error_occurred = pyqtSignal(str)

    def __init__(self, url: str, parent=None):
        super().__init__(parent)
        self.url = url

    def run(self):
        try:
            import urllib.request
            req = urllib.request.Request(
                self.url,
                headers={"User-Agent": f"BoltPDF/{__version__}"})
            resp = urllib.request.urlopen(req, timeout=120)
            total = int(resp.headers.get("Content-Length", 0))
            # Download into BoltPDF temp folder
            dl_dir = os.path.join(
                os.environ.get("LOCALAPPDATA", tempfile.gettempdir()),
                "BoltPDF", "update")
            os.makedirs(dl_dir, exist_ok=True)
            # Keep the remote extension so the updater knows whether
            # to swap a single exe or extract a zip.
            ext = ".zip" if self.url.lower().endswith(".zip") else ".exe"
            tmp_path = os.path.join(dl_dir, f"BoltPDF_update{ext}")
            downloaded = 0
            with open(tmp_path, "wb") as f:
                while True:
                    chunk = resp.read(65536)
                    if not chunk:
                        break
                    f.write(chunk)
                    downloaded += len(chunk)
                    if total > 0:
                        self.progress.emit(int(downloaded / total * 100))
            self.progress.emit(100)
            self.finished_ok.emit(tmp_path)
        except Exception as e:
            self.error_occurred.emit(str(e))


def _apply_update_and_restart(new_file_path: str):
    """Create a .bat that waits for BoltPDF to exit, then applies the
    downloaded update and relaunches.

    For onedir builds the update is a .zip containing the full bundle
    (exe + DLLs + _internal/...).  The updater extracts it into the
    install directory, replacing everything.  For backwards
    compatibility with a single-exe payload the old swap-and-replace
    path is still supported.
    """
    if not getattr(sys, 'frozen', False):
        # Dev mode — nothing to swap
        return
    current_exe = sys.executable
    install_dir = os.path.dirname(current_exe)
    bat_dir = os.path.join(
        os.environ.get("LOCALAPPDATA", tempfile.gettempdir()),
        "BoltPDF", "update")
    os.makedirs(bat_dir, exist_ok=True)
    bat_path = os.path.join(bat_dir, "_updater.bat")

    is_zip = new_file_path.lower().endswith(".zip")

    if is_zip:
        # Extract the zip over the install dir.  PowerShell's
        # Expand-Archive needs an existing (empty-ish) destination, so
        # we wipe the install dir first.  Files that are still locked
        # (while the exe is shutting down) get retried briefly.
        script = f"""@echo off
:: Wait for BoltPDF to fully exit
timeout /t 3 /nobreak >nul
:: Retry loop — wait until the main exe unlocks
:retry
del "{current_exe}" >nul 2>&1
if exist "{current_exe}" (
    timeout /t 2 /nobreak >nul
    goto retry
)
:: Remove the rest of the old install (safe now that the exe is gone)
powershell -NoProfile -Command "Get-ChildItem -LiteralPath '{install_dir}' -Force | Remove-Item -Recurse -Force -ErrorAction SilentlyContinue"
:: Extract the new onedir bundle over the install directory
powershell -NoProfile -Command "Expand-Archive -LiteralPath '{new_file_path}' -DestinationPath '{install_dir}' -Force"
:: If the zip contained a top-level 'BoltPDF' folder, lift its contents up
if exist "{install_dir}\\BoltPDF\\BoltPDF.exe" (
    powershell -NoProfile -Command "Move-Item -LiteralPath '{install_dir}\\BoltPDF\\*' -Destination '{install_dir}' -Force; Remove-Item -LiteralPath '{install_dir}\\BoltPDF' -Recurse -Force"
)
:: Relaunch
start "" "{current_exe}"
:: Clean up the downloaded zip and this script
del "{new_file_path}" >nul 2>&1
del "%~f0" >nul 2>&1
"""
    else:
        # Legacy single-exe swap
        script = f"""@echo off
timeout /t 3 /nobreak >nul
:retry
del "{current_exe}" >nul 2>&1
if exist "{current_exe}" (
    timeout /t 2 /nobreak >nul
    goto retry
)
move /Y "{new_file_path}" "{current_exe}" >nul
start "" "{current_exe}"
del "%~f0" >nul 2>&1
"""

    with open(bat_path, "w") as f:
        f.write(script)

    # Launch the bat detached, then exit
    subprocess.Popen(
        ["cmd.exe", "/c", bat_path],
        creationflags=subprocess.DETACHED_PROCESS | subprocess.CREATE_NO_WINDOW,
        close_fds=True)
    QApplication.instance().quit()


# ---------------------------------------------------------------------------
# Ad Banner — Adsterra banner via QWebEngineView at the bottom of the window
# ---------------------------------------------------------------------------
_AD_SETTINGS_FILE = os.path.join(
    os.environ.get("LOCALAPPDATA", os.path.expanduser("~")),
    "BoltPDF", "ad_settings.json")


def _ads_enabled() -> bool:
    """Return True unless the user has explicitly disabled ads."""
    try:
        with open(_AD_SETTINGS_FILE, "r") as f:
            return json.load(f).get("ads_enabled", True)
    except Exception:
        return True


def _set_ads_enabled(enabled: bool):
    """Persist the user's ad preference."""
    os.makedirs(os.path.dirname(_AD_SETTINGS_FILE), exist_ok=True)
    try:
        with open(_AD_SETTINGS_FILE, "w") as f:
            json.dump({"ads_enabled": enabled}, f)
    except Exception:
        pass


def _build_ad_html() -> str:
    """Return a self-contained HTML page that loads an Adsterra banner."""
    return """<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<style>
  html, body {
    margin: 0; padding: 0; overflow: hidden;
    background: #2d2d2d;
    display: flex; justify-content: center; align-items: center;
    height: 100%;
  }
</style>
</head><body>
<script type="text/javascript">
  atOptions = {
    'key' : 'a480a2b06ea7bacda0d8ad45cb03d488',
    'format' : 'iframe',
    'height' : 90,
    'width' : 728,
    'params' : {}
  };
</script>
<script type="text/javascript"
  src="https://www.highperformanceformat.com/a480a2b06ea7bacda0d8ad45cb03d488/invoke.js">
</script>
</body></html>"""


class AdBanner(QWidget):
    """Thin bottom banner showing an Adsterra ad via QWebEngineView.

    - WebEngine is initialized on the next event-loop tick (zero delay)
      so the banner appears as soon as the main window is visible.
    - Loads ad HTML from a temp file URL so third-party scripts work.
    - Graceful fallback: if PyQtWebEngine isn't installed or there's
      no internet, the banner stays hidden.
    - Dismissable via a small x button (hides for the session).
    - Respects persistent ads_enabled preference.
    """

    def __init__(self, parent=None):
        super().__init__(parent)
        self._web = None
        self._tmp_html = None
        self.setFixedHeight(0)
        self.setStyleSheet("background: #2d2d2d;")

        self._layout = QHBoxLayout(self)
        self._layout.setContentsMargins(0, 0, 0, 0)
        self._layout.setSpacing(0)

        self._close_btn = QPushButton("\u00d7", self)
        self._close_btn.setFixedSize(20, 20)
        self._close_btn.setStyleSheet(
            "QPushButton { color: #888; background: #2d2d2d; border: none;"
            " font-size: 16px; font-weight: bold; }"
            "QPushButton:hover { color: #fff; }")
        self._close_btn.clicked.connect(self._dismiss)
        self._close_btn.hide()

        if not _ads_enabled():
            return
        # Kick off WebEngine init on the next event-loop tick so the
        # main window paints first, but with zero artificial delay.
        QTimer.singleShot(0, self._init_web)

    def _init_web(self):
        if not _ads_enabled():
            return
        try:
            from PyQt6.QtWebEngineWidgets import QWebEngineView
            from PyQt6.QtWebEngineCore import (QWebEnginePage,
                                               QWebEngineProfile,
                                               QWebEngineSettings)
        except Exception as e:
            print(f"[AdBanner] WebEngine import failed: {e}")
            return

        print("[AdBanner] WebEngine loaded OK, setting up ad...")

        # Write ad HTML to a temp file so the browser navigates to a real
        # file:// URL — this lets third-party scripts (Adsterra) load
        # properly, unlike setHtml() which restricts remote resources.
        try:
            ad_dir = os.path.join(
                os.environ.get("LOCALAPPDATA", tempfile.gettempdir()),
                "BoltPDF", "adcache")
            os.makedirs(ad_dir, exist_ok=True)
            self._tmp_html = os.path.join(ad_dir, "ad_banner.html")
            with open(self._tmp_html, "w", encoding="utf-8") as f:
                f.write(_build_ad_html())
        except Exception:
            return

        profile = QWebEngineProfile("boltpdf_ads", self)
        profile.setHttpCacheType(
            QWebEngineProfile.HttpCacheType.MemoryHttpCache)
        page = QWebEnginePage(profile, self)

        settings = page.settings()
        settings.setAttribute(
            QWebEngineSettings.WebAttribute.JavascriptEnabled, True)
        settings.setAttribute(
            QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls,
            True)

        self._web = QWebEngineView(self)
        self._web.setPage(page)
        self._web.setFixedHeight(90)
        self._web.setStyleSheet("background: #2d2d2d;")
        self._web.setContextMenuPolicy(
            Qt.ContextMenuPolicy.NoContextMenu)

        self._layout.addWidget(self._web, 1)
        self._layout.addWidget(self._close_btn, 0,
                               Qt.AlignmentFlag.AlignTop)

        # Navigate to the temp file — real URL lets remote JS load
        file_url = QUrl.fromLocalFile(self._tmp_html)
        self._web.load(file_url)
        self._close_btn.show()
        self.setFixedHeight(90)

    def _dismiss(self):
        self.setFixedHeight(0)
        self._close_btn.hide()
        if self._web:
            self._web.setUrl(QUrl("about:blank"))
            self._web.hide()

    def set_enabled(self, enabled: bool):
        _set_ads_enabled(enabled)
        if enabled and self._web is None:
            self._init_web()
        elif not enabled:
            self._dismiss()


# ---------------------------------------------------------------------------
# Search bar — Ctrl+F opens a floating search panel at the top of the window
# ---------------------------------------------------------------------------
class SearchBar(QFrame):
    """Compact inline search widget designed to live inside the toolbar."""
    search_changed = pyqtSignal(str)   # emitted on every keystroke
    next_clicked = pyqtSignal()
    prev_clicked = pyqtSignal()
    closed = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("SearchBar")
        self.setStyleSheet(
            "#SearchBar {"
            "  background: transparent;"
            "  border: none; padding: 0;"
            "}"
            "#SearchBar QLineEdit {"
            "  background: #3c3c3c; color: #e0e0e0;"
            "  border: 1px solid #555; border-radius: 3px;"
            "  padding: 2px 6px; font-size: 12px;"
            "  selection-background-color: #3399ff;"
            "}"
            "#SearchBar QLineEdit:focus { border: 1px solid #7a7aff; }"
            "#SearchBar QLabel {"
            "  color: #aaa; font-size: 11px; padding: 0 2px;"
            "}"
            "#SearchBar QPushButton {"
            "  background: transparent; color: #e0e0e0;"
            "  border: 1px solid transparent; border-radius: 3px;"
            "  padding: 2px 6px; font-size: 11px;"
            "}"
            "#SearchBar QPushButton:hover {"
            "  background: #4a4a4a; border-color: #555;"
            "}"
            "#SearchBar QPushButton#closeSearchBtn {"
            "  color: #aaa; font-size: 14px; font-weight: bold;"
            "  padding: 0 4px;"
            "}"
            "#SearchBar QPushButton#closeSearchBtn:hover {"
            "  color: #fff;"
            "}")
        layout = QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(3)

        self._input = QLineEdit()
        self._input.setPlaceholderText("Search...")
        self._input.setFixedWidth(160)
        self._input.setClearButtonEnabled(True)
        self._input.textChanged.connect(self.search_changed.emit)
        self._input.returnPressed.connect(self.next_clicked.emit)
        layout.addWidget(self._input)

        self._match_label = QLabel("")
        self._match_label.setFixedWidth(75)
        layout.addWidget(self._match_label)

        btn_prev = QPushButton("\u25B2")
        btn_prev.setToolTip("Previous match (Shift+Enter)")
        btn_prev.setFixedWidth(22)
        btn_prev.clicked.connect(self.prev_clicked.emit)
        layout.addWidget(btn_prev)

        btn_next = QPushButton("\u25BC")
        btn_next.setToolTip("Next match (Enter)")
        btn_next.setFixedWidth(22)
        btn_next.clicked.connect(self.next_clicked.emit)
        layout.addWidget(btn_next)

        close_btn = QPushButton("\u2715")
        close_btn.setObjectName("closeSearchBtn")
        close_btn.setToolTip("Clear search (Esc)")
        close_btn.setFixedWidth(20)
        close_btn.clicked.connect(self.closed.emit)
        layout.addWidget(close_btn)

    def open_bar(self):
        """Focus the input and select any existing text."""
        self._input.setFocus(Qt.FocusReason.ShortcutFocusReason)
        self._input.selectAll()

    def close_bar(self):
        """Clear the search text and release focus."""
        self._input.clear()
        self._match_label.setText("")
        self._input.clearFocus()

    def text(self) -> str:
        return self._input.text()

    def set_match_info(self, current: int, total: int,
                       searching: bool = False):
        if total == 0:
            txt = self._input.text()
            if searching:
                self._match_label.setText("Searching...")
            else:
                self._match_label.setText(
                    "No matches" if txt else "")
        else:
            suffix = "..." if searching else ""
            self._match_label.setText(
                f"{current + 1} of {total}{suffix}")

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_Escape:
            self.closed.emit()
            return
        if (event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter)
                and event.modifiers() & Qt.KeyboardModifier.ShiftModifier):
            self.prev_clicked.emit()
            return
        super().keyPressEvent(event)


class SearchWorker(QThread):
    """Run fitz search_for() off the main thread, streaming results in
    batches so the UI stays responsive even for huge documents.

    Results are flushed to the main thread either when the first match
    is found (for instant feedback) or every ~200 ms thereafter, and
    finally on completion.

    Signals:
        batch_ready(query, list[(page_idx, (x0,y0,x1,y1))])
            Emitted with a batch of new matches.
        finished_search(query)
            Emitted once when the entire document has been searched.
        text_cache_ready(dict)
            Emitted once with the per-page plain text index.
    """
    batch_ready = pyqtSignal(str, list)        # query, matches
    finished_search = pyqtSignal(str)           # query
    text_cache_ready = pyqtSignal(dict)         # page_idx → text

    def __init__(self, doc_path: str, query: str,
                 text_cache: dict | None = None, parent=None):
        super().__init__(parent)
        self.doc_path = doc_path
        self.query = query
        self._text_cache = text_cache
        self._cancelled = False

    def cancel(self):
        self._cancelled = True

    def run(self):
        if not self.query or not self.doc_path:
            self.finished_search.emit(self.query)
            return
        try:
            import fitz
        except ImportError:
            self.finished_search.emit(self.query)
            return
        try:
            doc = fitz.open(self.doc_path)
        except Exception:
            self.finished_search.emit(self.query)
            return

        query_lower = self.query.lower()
        building_cache = self._text_cache is None
        new_cache: dict = {} if building_cache else {}
        pending: list = []
        first_emitted = False
        last_flush = time.monotonic()

        try:
            for page_idx in range(len(doc)):
                if self._cancelled:
                    break
                page = doc[page_idx]

                if building_cache:
                    page_text = page.get_text("text") or ""
                    new_cache[page_idx] = page_text
                else:
                    page_text = self._text_cache.get(page_idx, "")

                if query_lower not in page_text.lower():
                    continue

                rects = page.search_for(self.query)
                for r in rects:
                    pending.append(
                        (page_idx, (r.x0, r.y0, r.x1, r.y1)))

                if self._cancelled:
                    break

                # Flush strategy: immediately on first results, then
                # every 200 ms to avoid flooding the main thread.
                now = time.monotonic()
                if pending and (
                        not first_emitted
                        or (now - last_flush) >= 0.2):
                    self.batch_ready.emit(self.query, pending)
                    pending = []
                    last_flush = now
                    first_emitted = True
        except Exception:
            pass
        finally:
            doc.close()

        if not self._cancelled:
            # Flush any remaining matches
            if pending:
                self.batch_ready.emit(self.query, pending)
            if building_cache and new_cache:
                self.text_cache_ready.emit(new_cache)
            self.finished_search.emit(self.query)


# ---------------------------------------------------------------------------
# Update banner — non-intrusive bar shown at the top of the window when a
# new version is detected on startup.
# ---------------------------------------------------------------------------
class UpdateBanner(QFrame):
    """Dismissible banner that tells the user a new version is available."""
    update_clicked = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("UpdateBanner")
        self.setStyleSheet(
            "#UpdateBanner {"
            "  background: qlineargradient(x1:0,y1:0,x2:1,y2:0,"
            "    stop:0 #1a5276, stop:1 #2e86c1);"
            "  border-bottom: 1px solid #1a5276;"
            "  padding: 4px 8px;"
            "}"
            "#UpdateBanner QLabel { color: #ffffff; font-size: 13px; }"
            "#UpdateBanner QPushButton {"
            "  background: #ffffff; color: #1a5276;"
            "  border: none; border-radius: 3px;"
            "  padding: 4px 14px; font-weight: bold; font-size: 12px;"
            "}"
            "#UpdateBanner QPushButton:hover { background: #d6eaf8; }"
            "#UpdateBanner QPushButton#dismissBtn {"
            "  background: transparent; color: #ffffff;"
            "  font-size: 16px; font-weight: bold; padding: 2px 6px;"
            "}"
            "#UpdateBanner QPushButton#dismissBtn:hover {"
            "  background: rgba(255,255,255,0.15);"
            "}")
        layout = QHBoxLayout(self)
        layout.setContentsMargins(8, 4, 8, 4)
        layout.setSpacing(10)

        self._label = QLabel("")
        self._label.setWordWrap(False)
        layout.addWidget(self._label, 1)

        self._update_btn = QPushButton("Update Now")
        self._update_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self._update_btn.clicked.connect(self.update_clicked.emit)
        layout.addWidget(self._update_btn)

        dismiss = QPushButton("\u2715")
        dismiss.setObjectName("dismissBtn")
        dismiss.setCursor(Qt.CursorShape.PointingHandCursor)
        dismiss.setToolTip("Dismiss")
        dismiss.clicked.connect(self._dismiss)
        layout.addWidget(dismiss)

        self.hide()

    def show_update(self, remote_version: str, changelog: str = ""):
        txt = (f"A new version of BoltPDF is available: "
               f"<b>v{remote_version}</b>  (you have v{__version__})")
        if changelog:
            # Show first line of changelog only in the banner
            first_line = changelog.strip().split("\n")[0].strip()
            if first_line:
                txt += f" — {first_line}"
        self._label.setText(txt)
        self.show()

    def _dismiss(self):
        self.hide()


# ---------------------------------------------------------------------------
# Preview Panel — left sidebar showing preview-PDF thumbnails
# ---------------------------------------------------------------------------
class _ClickableThumb(QLabel):
    """A QLabel thumbnail that emits *clicked* with its page index."""
    clicked = pyqtSignal(int)

    def __init__(self, page_index: int, parent=None):
        super().__init__(parent)
        self._page_index = page_index
        self.setCursor(Qt.CursorShape.PointingHandCursor)

    def mousePressEvent(self, ev):
        self.clicked.emit(self._page_index)
        super().mousePressEvent(ev)


# ---------------------------------------------------------------------------
# Notes Panel — right-side panel listing PDF annotations for the current page
# ---------------------------------------------------------------------------
class _ClickableCard(QFrame):
    """A note card that emits *clicked* with its 0-based page index so
    the Notes panel can jump the main view to that page."""
    clicked = pyqtSignal(int)

    def __init__(self, page_index: int, parent=None):
        super().__init__(parent)
        self._page_index = page_index
        self.setCursor(Qt.CursorShape.PointingHandCursor)

    def mousePressEvent(self, ev):
        if ev.button() == Qt.MouseButton.LeftButton:
            self.clicked.emit(self._page_index)
        super().mousePressEvent(ev)


class NotesPanel(QWidget):
    """Right-side panel that lists all comment annotations on the currently
    visible page.  Automatically updates as the user scrolls.  Clicking a
    note jumps the main document to that note's page."""

    note_clicked = pyqtSignal(int)       # 0-based page index
    add_note_requested = pyqtSignal()    # user wants to add a note

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setMinimumWidth(260)
        self.setMaximumWidth(400)

        root_layout = QVBoxLayout(self)
        root_layout.setContentsMargins(0, 0, 0, 0)
        root_layout.setSpacing(0)

        # Header
        self._header = QLabel("Notes")
        self._header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._header.setStyleSheet(
            "QLabel { background: #2b2b2b; color: #e0e0e0;"
            " font-weight: bold; font-size: 14px;"
            " padding: 8px; border-bottom: 1px solid #444; }")
        root_layout.addWidget(self._header)

        # Add Note button
        self._add_btn = QPushButton("+ Add Note")
        self._add_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self._add_btn.setStyleSheet(
            "QPushButton { background: #3a6ea5; color: #fff;"
            " border: none; padding: 7px; font-weight: bold;"
            " font-size: 13px; }"
            " QPushButton:hover { background: #4a82c0; }")
        self._add_btn.clicked.connect(
            lambda: self.add_note_requested.emit())
        root_layout.addWidget(self._add_btn)

        # Scrollable area for note cards
        self._scroll = QScrollArea()
        self._scroll.setWidgetResizable(True)
        self._scroll.setHorizontalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self._scroll.setStyleSheet(
            "QScrollArea { background: #1e1e1e; border: none; }")
        root_layout.addWidget(self._scroll)

        # Inner container
        self._container = QWidget()
        self._container.setStyleSheet("background: #1e1e1e;")
        self._card_layout = QVBoxLayout(self._container)
        self._card_layout.setContentsMargins(8, 8, 8, 8)
        self._card_layout.setSpacing(8)
        self._card_layout.addStretch()
        self._scroll.setWidget(self._container)

        self._current_page = -1
        self._loaded = False
        self._page_sections: dict[int, list[QFrame]] = {}

    def show_notes(self, current_page: int,
                   all_annotations: dict[int, list[dict]]):
        """Populate the panel with all annotations, grouped by page.
        *current_page* (0-based) is highlighted and scrolled into view.
        *all_annotations* maps page index → list of note dicts."""
        if current_page == self._current_page and self._loaded:
            return  # already showing with this page highlighted
        self._current_page = current_page

        # Full rebuild only on first call or when the document changes
        if not self._loaded:
            self._rebuild_cards(all_annotations)
            self._loaded = True

        # Update header
        total = sum(len(v) for v in all_annotations.values())
        self._header.setText(
            f"Notes ({total}) — Page {current_page + 1}")

        # Highlight the current page's section and scroll to it
        for pg, widgets in self._page_sections.items():
            for w in widgets:
                if pg == current_page:
                    w.setStyleSheet(
                        w._base_style + " border-left: 4px solid #4a90d9;")
                else:
                    w.setStyleSheet(w._base_style)

        # Scroll to current page's first card
        if current_page in self._page_sections:
            cards = self._page_sections[current_page]
            if cards:
                self._scroll.ensureWidgetVisible(cards[0], 0, 20)

    def _rebuild_cards(self, all_annotations):
        """Rebuild all note cards from scratch."""
        # Clear everything
        while self._card_layout.count() > 1:  # keep the stretch
            item = self._card_layout.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
        self._page_sections.clear()

        if not all_annotations:
            empty = QLabel("No comments in this document.")
            empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
            empty.setStyleSheet(
                "color: #888; font-size: 13px; padding: 20px;")
            empty._base_style = "color: #888; font-size: 13px; padding: 20px;"
            self._card_layout.insertWidget(0, empty)
            return

        insert_idx = 0
        for page_idx in sorted(all_annotations.keys()):
            notes = all_annotations[page_idx]
            if not notes:
                continue

            # Page header label
            pg_label = QLabel(f"Page {page_idx + 1}")
            pg_label.setStyleSheet(
                "color: #aaa; font-weight: bold; font-size: 13px;"
                " padding: 8px 4px 4px 4px; background: transparent;"
                " border: none;")
            pg_label._base_style = (
                "color: #aaa; font-weight: bold; font-size: 13px;"
                " padding: 8px 4px 4px 4px; background: transparent;"
                " border: none;")
            self._card_layout.insertWidget(insert_idx, pg_label)
            insert_idx += 1

            page_cards = []
            for note in notes:
                card = _ClickableCard(page_idx)
                card.clicked.connect(self.note_clicked)
                card.setFrameShape(QFrame.Shape.StyledPanel)
                base = (
                    "QFrame { background: #fffde7;"
                    " border: 1px solid #d4c84a;"
                    " border-radius: 8px; }"
                    " QLabel { background: transparent; border: none; }")
                card.setStyleSheet(base)
                card._base_style = base

                clayout = QVBoxLayout(card)
                clayout.setContentsMargins(14, 12, 14, 12)
                clayout.setSpacing(6)

                # Author + type header
                header_parts = []
                if note.get("author"):
                    header_parts.append(
                        f"<b>{note['author']}</b>")
                if note.get("type"):
                    header_parts.append(
                        f"<i>{note['type']}</i>")
                if header_parts:
                    hlbl = QLabel(" · ".join(header_parts))
                    hlbl.setStyleSheet(
                        "font-size: 12px; color: #777;")
                    hlbl.setWordWrap(True)
                    clayout.addWidget(hlbl)

                # Subject
                if note.get("subject"):
                    slbl = QLabel(f"<b>{note['subject']}</b>")
                    slbl.setStyleSheet(
                        "font-size: 14px; color: #333;")
                    slbl.setWordWrap(True)
                    clayout.addWidget(slbl)

                # Content
                if note.get("content"):
                    clbl = QLabel(note["content"])
                    clbl.setStyleSheet(
                        "font-size: 14px; color: #222;")
                    clbl.setWordWrap(True)
                    clayout.addWidget(clbl)

                self._card_layout.insertWidget(insert_idx, card)
                page_cards.append(card)
                insert_idx += 1

            self._page_sections[page_idx] = page_cards

    def clear(self):
        """Remove all notes and reset."""
        self._current_page = -1
        self._loaded = False
        self._page_sections.clear()
        while self._card_layout.count() > 1:
            item = self._card_layout.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
        self._header.setText("Notes")


class _ThumbRenderer(QThread):
    """Single in-process thread that renders small page thumbnails on
    demand.  Deliberately *not* a process pool — thumbnails are tiny
    (~150 px) and must never compete with the main multi-process
    PageRenderer for cores.  Renders nearest-to-focus first and only
    what was requested, so it stays lazy and cheap."""

    thumb_ready = pyqtSignal(int, QImage)

    def __init__(self, doc_path: str, thumb_px: int = 150, parent=None):
        super().__init__(parent)
        self._doc_path = doc_path
        self._thumb_px = thumb_px
        self._lock = threading.Lock()
        self._queue: list[int] = []
        self._done: set[int] = set()
        self._focus = 0
        self._cancel = False

    def request(self, indices, focus: int):
        with self._lock:
            self._focus = focus
            self._queue = sorted(
                {i for i in indices if i not in self._done},
                key=lambda i: abs(i - focus))

    def forget(self, indices):
        """Allow re-render of evicted thumbs if scrolled back to."""
        with self._lock:
            self._done.difference_update(indices)

    def cancel(self):
        self._cancel = True

    def _next(self):
        with self._lock:
            while self._queue:
                p = self._queue.pop(0)
                if p not in self._done:
                    return p
        return None

    def run(self):
        try:
            doc = pdfium.PdfDocument(self._doc_path)
        except Exception:
            return
        try:
            while not self._cancel:
                idx = self._next()
                if idx is None:
                    self.msleep(40)
                    continue
                try:
                    page = doc[idx]
                    try:
                        w_pt = page.get_width() or 1
                        scale = max(0.05, min(self._thumb_px / w_pt, 1.5))
                        pil = page.render(
                            scale=scale, draw_annots=False).to_pil()
                    finally:
                        page.close()
                    if pil.mode != "RGBA":
                        pil = pil.convert("RGBA")
                    data = pil.tobytes("raw", "RGBA")
                    qimg = QImage(data, pil.width, pil.height,
                                  QImage.Format.Format_RGBA8888).copy()
                    with self._lock:
                        self._done.add(idx)
                    if not self._cancel:
                        self.thumb_ready.emit(idx, qimg)
                except Exception:
                    with self._lock:
                        self._done.add(idx)
        finally:
            try:
                doc.close()
            except Exception:
                pass


class ThumbnailPanel(QWidget):
    """Lazy page-thumbnail sidebar for the *main* document.

    Hidden by default (so startup stays instant).  Builds cheap empty
    placeholders for every page, then renders only the thumbnails in the
    visible scroll window via _ThumbRenderer, evicting pixmaps beyond a
    cap to bound memory.  Clicking a thumbnail jumps the main view."""

    page_clicked = pyqtSignal(int)  # 0-based

    _GAP = 6
    _CAP = 160          # max pixmaps kept resident
    _BUFFER = 6         # extra thumbs rendered above/below the viewport

    def __init__(self, parent=None):
        super().__init__(parent)
        self._thumbs: list[_ClickableThumb] = []
        self._loaded: set[int] = set()
        self._doc_path: str | None = None
        self._num = 0
        self._thumb_w = 150
        self._thumb_h = 200
        self._cur = -1
        self._renderer: _ThumbRenderer | None = None

        self.setMinimumWidth(120)
        self.setMaximumWidth(260)

        self._scroll = QScrollArea(self)
        self._scroll.setWidgetResizable(True)
        self._scroll.setHorizontalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self._container = QWidget()
        self._col = QVBoxLayout(self._container)
        self._col.setContentsMargins(6, 6, 6, 6)
        self._col.setSpacing(self._GAP)
        self._scroll.setWidget(self._container)

        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(0)
        header = QLabel("Pages")
        header.setStyleSheet(
            "font-weight:bold; padding:4px; text-align:center; "
            "background:#2c3e50; color:#ecf0f1;")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        outer.addWidget(header)
        outer.addWidget(self._scroll, 1)

        self._scroll_timer = QTimer(self)
        self._scroll_timer.setSingleShot(True)
        self._scroll_timer.setInterval(90)
        self._scroll_timer.timeout.connect(self._request_visible)
        sb = self._scroll.verticalScrollBar()
        if sb:
            sb.valueChanged.connect(lambda _v: self._scroll_timer.start())

    # -- lifecycle --------------------------------------------------------
    def set_document(self, doc_path: str, num_pages: int, aspect: float):
        """Point the panel at a document.  Only called when the panel is
        actually visible, so hidden = zero cost."""
        if doc_path == self._doc_path and num_pages == self._num:
            return
        self.stop()
        self.clear()
        self._doc_path = doc_path
        self._num = num_pages
        aspect = aspect if aspect and aspect > 0 else 0.75
        self._thumb_w = max(90, self.width() - 28)
        self._thumb_h = max(40, int(self._thumb_w / aspect))
        for i in range(num_pages):
            t = _ClickableThumb(i)
            t.setFixedSize(self._thumb_w, self._thumb_h)
            t.setAlignment(Qt.AlignmentFlag.AlignCenter)
            t.setText(str(i + 1))
            t.setStyleSheet(
                "background:#3a3a3a; color:#888; border:1px solid #555;")
            t.clicked.connect(self.page_clicked)
            self._col.addWidget(t)
            self._thumbs.append(t)
        self._col.addStretch(1)
        if num_pages:
            self._renderer = _ThumbRenderer(doc_path, self._thumb_w, self)
            self._renderer.thumb_ready.connect(self._on_thumb_ready)
            self._renderer.start()
            QTimer.singleShot(0, self._request_visible)

    def stop(self):
        if self._renderer is not None:
            try:
                self._renderer.cancel()
                self._renderer.wait(1500)
            except Exception:
                pass
            self._renderer = None

    def clear(self):
        for t in self._thumbs:
            self._col.removeWidget(t)
            t.deleteLater()
        self._thumbs.clear()
        self._loaded.clear()
        self._cur = -1
        while self._col.count():
            it = self._col.takeAt(0)
            if it.widget():
                it.widget().deleteLater()
        self._doc_path = None
        self._num = 0

    # -- lazy viewport rendering -----------------------------------------
    def _visible_range(self):
        if not self._thumbs:
            return (0, -1)
        step = self._thumb_h + self._GAP
        sb = self._scroll.verticalScrollBar()
        top = sb.value() if sb else 0
        vh = self._scroll.viewport().height()
        lo = max(0, top // step - self._BUFFER)
        hi = min(self._num - 1,
                 (top + vh) // step + self._BUFFER)
        return (int(lo), int(hi))

    def _request_visible(self):
        if self._renderer is None or not self._thumbs:
            return
        lo, hi = self._visible_range()
        if hi < lo:
            return
        want = list(range(lo, hi + 1))
        # Evict pixmaps far outside the window to bound memory.
        if len(self._loaded) > self._CAP:
            keep_lo, keep_hi = lo - self._CAP, hi + self._CAP
            evict = [i for i in self._loaded
                     if i < keep_lo or i > keep_hi]
            for i in evict:
                if 0 <= i < len(self._thumbs):
                    t = self._thumbs[i]
                    t.clear()
                    t.setText(str(i + 1))
                self._loaded.discard(i)
            if evict and self._renderer:
                self._renderer.forget(evict)
        missing = [i for i in want if i not in self._loaded]
        if missing:
            self._renderer.request(missing, (lo + hi) // 2)

    def _on_thumb_ready(self, idx: int, qimg: QImage):
        if not (0 <= idx < len(self._thumbs)) or qimg.isNull():
            return
        pm = QPixmap.fromImage(qimg)
        if pm.isNull():
            return
        t = self._thumbs[idx]
        t.setText("")
        t.setPixmap(pm)
        self._loaded.add(idx)
        self._restyle(idx)

    # -- current-page highlight ------------------------------------------
    def _restyle(self, idx: int):
        if not (0 <= idx < len(self._thumbs)):
            return
        border = ("2px solid #7a7aff" if idx == self._cur
                  else "1px solid #555")
        self._thumbs[idx].setStyleSheet(
            f"background:#3a3a3a; color:#888; border:{border};")

    def highlight(self, idx: int):
        if idx == self._cur:
            return
        prev = self._cur
        self._cur = idx
        for i in (prev, idx):
            self._restyle(i)
        if 0 <= idx < len(self._thumbs):
            self._scroll.ensureWidgetVisible(self._thumbs[idx])


class PreviewPanel(QWidget):
    """Scrollable sidebar that displays every page of a preview PDF as a
    small thumbnail.  Sized so roughly 10 preview pages fit alongside one
    full page of the main document."""

    page_clicked = pyqtSignal(int)  # 0-based page index

    # Target: ~10 preview pages visible per main-page height.
    # At 792 pt (US Letter), that's ~79 pt per thumbnail + gap.
    THUMB_GAP = 4  # px between thumbnails

    def __init__(self, parent=None):
        super().__init__(parent)
        self._pixmaps: list[QPixmap] = []
        self._doc = None  # fitz document kept open for rendering
        self._labels: list[_ClickableThumb] = []

        self.setMinimumWidth(100)
        self.setMaximumWidth(300)

        # Scroll area wrapping a vertical column of QLabels
        self._scroll = QScrollArea(self)
        self._scroll.setWidgetResizable(True)
        self._scroll.setHorizontalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        self._container = QWidget()
        self._col = QVBoxLayout(self._container)
        self._col.setContentsMargins(4, 4, 4, 4)
        self._col.setSpacing(self.THUMB_GAP)
        self._scroll.setWidget(self._container)

        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(0)
        # Clickable header — toggles thumbnail visibility
        self._header = QPushButton("Preview  \u25BC")
        self._header.setFlat(True)
        self._header.setCursor(Qt.CursorShape.PointingHandCursor)
        self._header.setStyleSheet(
            "font-weight:bold; padding:4px; text-align:center; "
            "background:#2c3e50; color:#ecf0f1; border:none;")
        self._header.clicked.connect(self.toggle_thumbnails)
        self._thumbnails_visible = True
        outer.addWidget(self._header)
        outer.addWidget(self._scroll, 1)

    def toggle_thumbnails(self):
        """Show or hide the thumbnail scroll area."""
        self._thumbnails_visible = not self._thumbnails_visible
        self._scroll.setVisible(self._thumbnails_visible)
        arrow = "\u25BC" if self._thumbnails_visible else "\u25B6"
        self._header.setText(f"Preview  {arrow}")

    def load_preview(self, preview_path: str, main_page_height: float):
        """Load a preview PDF and render every page as a thumbnail.

        *main_page_height* is the height (in points) of a typical page
        in the main document, used to calculate thumbnail size so that
        ~10 preview pages fit in that height.
        """
        self.clear()
        try:
            import fitz
            doc = fitz.open(preview_path)
        except Exception as e:
            print(f"[PreviewPanel] Failed to open {preview_path}: {e}",
                  file=sys.stderr)
            return

        # Target thumbnail height: main_page_height / 10, minus gaps
        thumb_h = max((main_page_height - self.THUMB_GAP * 9) / 10, 40)
        panel_w = self.width() - 16  # account for margins + scrollbar

        for i in range(len(doc)):
            page = doc[i]
            pw, ph = page.rect.width, page.rect.height
            if ph <= 0:
                continue
            # Scale to fit thumb_h, but also cap at panel width
            scale_h = thumb_h / ph
            scale_w = panel_w / pw if pw > 0 else scale_h
            scale = min(scale_h, scale_w)
            mat = fitz.Matrix(scale, scale)
            pix = page.get_pixmap(matrix=mat)
            qimg = QImage(pix.samples, pix.width, pix.height,
                          pix.stride, QImage.Format.Format_RGB888)
            pm = QPixmap.fromImage(qimg)
            self._pixmaps.append(pm)

            lbl = _ClickableThumb(i)
            lbl.setPixmap(pm)
            lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lbl.setStyleSheet("background: white; border:1px solid #bbb;")
            lbl.clicked.connect(self.page_clicked)
            self._col.addWidget(lbl)
            self._labels.append(lbl)

        doc.close()
        self._col.addStretch(1)

    def clear(self):
        """Remove all thumbnails."""
        for lbl in self._labels:
            self._col.removeWidget(lbl)
            lbl.deleteLater()
        self._labels.clear()
        self._pixmaps.clear()
        # Remove any stretch item
        while self._col.count():
            item = self._col.takeAt(0)
            if item.widget():
                item.widget().deleteLater()


# ---------------------------------------------------------------------------
# Main Window
# ---------------------------------------------------------------------------
class PropertiesPanel(QWidget):
    """Right-side inspector for the edit object the Transform tool has
    selected.  Edits position/size/rotation, colour, stroke width, fill
    and font size; each change is applied through
    DocumentTab.apply_property_edit() so it re-renders faithfully and is
    a single undo step.  Hidden by default (zero cost until shown)."""

    def __init__(self, parent=None):
        super().__init__(parent)
        from PyQt6.QtWidgets import (QDoubleSpinBox, QFormLayout,
                                     QCheckBox, QComboBox)
        self._tab = None
        self._record = None
        self._loading = False
        self.setMinimumWidth(150)
        self.setMaximumWidth(300)

        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(0)
        header = QLabel("Properties")
        header.setStyleSheet(
            "font-weight:bold; padding:4px; text-align:center; "
            "background:#2c3e50; color:#ecf0f1;")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        outer.addWidget(header)

        self._hint = QLabel(
            "Select an object with the\nTransform tool to edit its\n"
            "properties.")
        self._hint.setWordWrap(True)
        self._hint.setStyleSheet("padding:8px; color:#aaa;")
        outer.addWidget(self._hint)

        self._form_host = QWidget()
        form = QFormLayout(self._form_host)
        form.setContentsMargins(8, 8, 8, 8)

        def _spin(lo, hi, step=1.0, dec=1):
            s = QDoubleSpinBox()
            s.setRange(lo, hi)
            s.setDecimals(dec)
            s.setSingleStep(step)
            s.valueChanged.connect(self._schedule_apply)
            return s

        self._sx = _spin(-100000, 100000)
        self._sy = _spin(-100000, 100000)
        self._sw = _spin(1, 100000)
        self._sh = _spin(1, 100000)
        self._srot = _spin(-360, 360, 1.0, 1)
        self._sstroke = _spin(0.1, 100, 0.5, 1)
        self._sfont = _spin(4, 400, 1.0, 1)
        form.addRow("X (pt)", self._sx)
        form.addRow("Y (pt)", self._sy)
        form.addRow("W (pt)", self._sw)
        form.addRow("H (pt)", self._sh)
        form.addRow("Rotation°", self._srot)
        form.addRow("Stroke w", self._sstroke)
        form.addRow("Font size", self._sfont)

        self._btn_colour = QPushButton("Colour…")
        self._btn_colour.clicked.connect(lambda: self._pick("colour"))
        form.addRow("Colour", self._btn_colour)
        self._btn_fill = QPushButton("Fill…")
        self._btn_fill.clicked.connect(lambda: self._pick("fill"))
        form.addRow("Fill", self._btn_fill)
        self._chk_nofill = QCheckBox("No fill")
        self._chk_nofill.stateChanged.connect(self._schedule_apply)
        form.addRow("", self._chk_nofill)
        self._sfopac = _spin(0, 255, 5, 0)
        form.addRow("Fill opacity", self._sfopac)
        self._cmb_style = QComboBox()
        self._cmb_style.addItems(["Solid", "Dash", "Dot"])
        self._cmb_style.currentIndexChanged.connect(self._schedule_apply)
        form.addRow("Stroke style", self._cmb_style)
        self._chk_arrowboth = QCheckBox("Arrow both ends")
        self._chk_arrowboth.stateChanged.connect(self._schedule_apply)
        form.addRow("", self._chk_arrowboth)
        outer.addWidget(self._form_host)
        outer.addStretch(1)

        self._apply_timer = QTimer(self)
        self._apply_timer.setSingleShot(True)
        self._apply_timer.setInterval(280)
        self._apply_timer.timeout.connect(self._apply_now)
        self._form_host.setVisible(False)

    # -- refresh from the current selection -------------------------------
    def refresh(self, tab):
        self._tab = tab
        rec = tab.selected_record() if tab is not None else None
        self._record = rec
        if rec is None:
            self._hint.setVisible(True)
            self._form_host.setVisible(False)
            return
        self._hint.setVisible(False)
        self._form_host.setVisible(True)
        self._loading = True
        try:
            kind = rec.kind
            shape = rec.shape_type
            is_line = (kind == "shape_add"
                       and shape in ("line", "arrow"))
            has_geom = (rec.new_rect is not None and not is_line)
            if rec.new_rect:
                x0, y0, x1, y1 = rec.new_rect
                self._sx.setValue(x0)
                self._sy.setValue(y0)
                self._sw.setValue(max(1.0, x1 - x0))
                self._sh.setValue(max(1.0, y1 - y0))
            for w in (self._sx, self._sy, self._sw, self._sh):
                w.setEnabled(has_geom)
            self._srot.setValue(getattr(rec, "rotation", 0.0) or 0.0)
            self._srot.setEnabled(has_geom)
            is_shape = (kind == "shape_add")
            is_text = kind in ("stamp_add", "text_add",
                               "text_move", "text_edit")
            self._sstroke.setValue(rec.stroke_width or 2.0)
            self._sstroke.setEnabled(is_shape)
            self._sfont.setValue(rec.font_size or 12.0)
            self._sfont.setEnabled(is_text)
            self._btn_colour.setEnabled(is_shape or is_text)
            self._btn_fill.setEnabled(is_shape)
            self._chk_nofill.setEnabled(is_shape)
            self._chk_nofill.setChecked(
                is_shape and not rec.fill_color)
            self._sfopac.setValue(getattr(rec, "fill_opacity", 80))
            self._sfopac.setEnabled(is_shape)
            sty = getattr(rec, "stroke_style", "solid")
            self._cmb_style.setCurrentIndex(
                {"solid": 0, "dash": 1, "dot": 2}.get(sty, 0))
            self._cmb_style.setEnabled(is_shape)
            is_arrow = (is_shape and shape == "arrow")
            self._chk_arrowboth.setChecked(
                is_arrow and bool(getattr(rec, "arrow_both", False)))
            self._chk_arrowboth.setEnabled(is_arrow)
        finally:
            self._loading = False

    # -- editing ----------------------------------------------------------
    def _schedule_apply(self, *_):
        if not self._loading and self._record is not None:
            self._apply_timer.start()

    def _pick(self, which):
        if self._record is None:
            return
        cur = (self._record.fill_color if which == "fill"
               else (self._record.stroke_color
                     if self._record.kind == "shape_add"
                     else self._record.color)) or (0, 0, 0)
        col = QColorDialog.getColor(
            QColor(*cur), self, "Pick colour")
        if not col.isValid():
            return
        rgb = (col.red(), col.green(), col.blue())
        if which == "fill":
            self._record.fill_color = rgb
            self._chk_nofill.setChecked(False)
        elif self._record.kind == "shape_add":
            self._record.stroke_color = rgb
        else:
            self._record.color = rgb
        self._apply_now()

    def _apply_now(self):
        tab, rec = self._tab, self._record
        if tab is None or rec is None:
            return
        try:
            if (rec.new_rect is not None
                    and self._sx.isEnabled()):
                x = self._sx.value()
                y = self._sy.value()
                rec.new_rect = (x, y,
                                x + self._sw.value(),
                                y + self._sh.value())
            if self._srot.isEnabled():
                rec.rotation = self._srot.value()
            if self._sstroke.isEnabled():
                rec.stroke_width = self._sstroke.value()
            if self._sfont.isEnabled():
                rec.font_size = self._sfont.value()
            if (self._chk_nofill.isEnabled()
                    and self._chk_nofill.isChecked()):
                rec.fill_color = None
            if self._sfopac.isEnabled():
                rec.fill_opacity = int(self._sfopac.value())
            if self._cmb_style.isEnabled():
                rec.stroke_style = ("solid", "dash", "dot")[
                    self._cmb_style.currentIndex()]
            if self._chk_arrowboth.isEnabled():
                rec.arrow_both = self._chk_arrowboth.isChecked()
        except Exception as e:
            print(f"[BoltPDF] property apply failed: {e}",
                  file=sys.stderr)
            return
        tab.apply_property_edit(rec)


class _AppState:
    """Tiny JSON-backed persistence for recent files, the last session,
    and per-document bookmarks.

    Stored in %LOCALAPPDATA%/BoltPDF/state.json (the same base dir the
    OCR helper already uses).  Every operation is best-effort and must
    never raise into the UI: a missing or corrupt file just resets to
    defaults.  Writes are atomic (temp + os.replace) so a crash mid-save
    cannot corrupt the store.
    """

    _MAX_RECENT = 15

    def __init__(self):
        base = os.path.join(
            os.environ.get("LOCALAPPDATA", tempfile.gettempdir()),
            "BoltPDF")
        try:
            os.makedirs(base, exist_ok=True)
        except OSError:
            base = tempfile.gettempdir()
        self._path = os.path.join(base, "state.json")
        self._data = {"recent": [], "session": [], "bookmarks": {},
                      "prefs": {}}
        self._load()

    @staticmethod
    def _key(path: str) -> str:
        try:
            return os.path.normcase(os.path.abspath(path))
        except Exception:
            return path or ""

    def _load(self):
        try:
            with open(self._path, "r", encoding="utf-8") as fh:
                d = json.load(fh)
            if isinstance(d, dict):
                self._data["recent"] = [
                    p for p in d.get("recent", []) if isinstance(p, str)]
                self._data["session"] = [
                    s for s in d.get("session", [])
                    if isinstance(s, dict)
                    and isinstance(s.get("path"), str)]
                bm = d.get("bookmarks", {})
                if isinstance(bm, dict):
                    self._data["bookmarks"] = {
                        k: v for k, v in bm.items() if isinstance(v, list)}
                pr = d.get("prefs", {})
                if isinstance(pr, dict):
                    self._data["prefs"] = pr
        except (OSError, ValueError, TypeError):
            pass  # missing or corrupt → keep defaults

    def _save(self):
        try:
            tmp = self._path + ".tmp"
            with open(tmp, "w", encoding="utf-8") as fh:
                json.dump(self._data, fh, indent=2)
            os.replace(tmp, self._path)
        except OSError:
            pass

    # -- recent files -----------------------------------------------------
    def add_recent(self, path: str):
        if not path or not os.path.isfile(path):
            return
        k = self._key(path)
        rec = [p for p in self._data["recent"] if self._key(p) != k]
        rec.insert(0, path)
        self._data["recent"] = rec[:self._MAX_RECENT]
        self._save()

    def recent(self) -> list:
        live = [p for p in self._data["recent"] if os.path.isfile(p)]
        if live != self._data["recent"]:
            self._data["recent"] = live
            self._save()
        return list(live)

    def clear_recent(self):
        self._data["recent"] = []
        self._save()

    # -- last session -----------------------------------------------------
    def set_session(self, entries: list):
        self._data["session"] = entries
        self._save()

    def session(self) -> list:
        return [s for s in self._data["session"]
                if os.path.isfile(s.get("path", ""))]

    # -- per-document bookmarks ------------------------------------------
    def bookmarks(self, path: str) -> list:
        return list(self._data["bookmarks"].get(self._key(path), []))

    def set_bookmarks(self, path: str, marks: list):
        k = self._key(path)
        if marks:
            self._data["bookmarks"][k] = marks
        else:
            self._data["bookmarks"].pop(k, None)
        self._save()

    # -- generic preferences ---------------------------------------------
    def pref(self, key: str, default=None):
        return self._data.get("prefs", {}).get(key, default)

    def set_pref(self, key: str, value):
        self._data.setdefault("prefs", {})[key] = value
        self._save()


# App-wide clipboard for copy/paste/duplicate of edit objects across
# tabs and pages.  Holds a deep-copied EditRecord (or None).
_EDIT_CLIPBOARD = None


_APP_STATE: "_AppState | None" = None


def app_state() -> _AppState:
    """Lazy singleton accessor.

    Deliberately lazy so the spawn-based PageRenderer worker processes —
    which re-import this whole module — never construct the store or
    touch disk.  Only the GUI process ever calls this.
    """
    global _APP_STATE
    if _APP_STATE is None:
        _APP_STATE = _AppState()
    return _APP_STATE


class BoltPDFReader(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("BoltPDF")
        self.resize(900, 1000)
        self.setStyleSheet(self._stylesheet())

        # Tab widget
        self._tabs = QTabWidget(self)
        self._tabs.setTabsClosable(True)
        self._tabs.setMovable(True)
        self._tabs.tabCloseRequested.connect(self._close_tab)
        self._tabs.currentChanged.connect(self._on_tab_changed)
        self._tabs.setDocumentMode(True)

        # Update banner at the top (hidden until an update is found)
        self._update_banner = UpdateBanner(self)
        self._update_banner.update_clicked.connect(
            self._banner_update_clicked)

        # Ad banner at the bottom
        self._ad_banner = AdBanner(self)

        # Notes panel (right sidebar, hidden until annotations are found).
        # Clicking a note jumps the main view to that note's page.
        self._notes_panel = NotesPanel(self)
        self._notes_panel.setVisible(False)
        self._notes_panel.note_clicked.connect(self._on_note_clicked)
        self._notes_panel.add_note_requested.connect(
            self._on_add_note_requested)

        # Splitter: thumbnails | tab area | notes | properties
        self._splitter = QSplitter(Qt.Orientation.Horizontal, self)
        self._splitter.setChildrenCollapsible(False)

        # Centre: update banner + tabs + ad banner
        centre = QWidget(self)
        centre_vbox = QVBoxLayout(centre)
        centre_vbox.setContentsMargins(0, 0, 0, 0)
        centre_vbox.setSpacing(0)
        centre_vbox.addWidget(self._update_banner, 0)
        centre_vbox.addWidget(self._tabs, 1)
        centre_vbox.addWidget(self._ad_banner, 0)

        # Page-thumbnail sidebar (far left, hidden until toggled).
        self._thumb_panel = ThumbnailPanel(self)
        self._thumb_panel.setVisible(False)
        self._thumb_panel.page_clicked.connect(self._on_thumb_page_clicked)

        # Properties inspector (far right, hidden until toggled in
        # Edit Mode).
        self._props_panel = PropertiesPanel(self)
        self._props_panel.setVisible(False)

        self._splitter.addWidget(self._thumb_panel)
        self._splitter.addWidget(centre)
        self._splitter.addWidget(self._notes_panel)
        self._splitter.addWidget(self._props_panel)
        # Initial widths: thumbnails | tabs | notes | props
        self._splitter.setSizes([170, 800, 250, 210])

        self.setCentralWidget(self._splitter)

        self._build_toolbar()
        self._is_fullscreen = False
        self._edit_tb_visible_before_fs = False
        self._notes_visible_before_fs = False

        # Hide tab bar when 0 or 1 tabs
        self._update_tab_bar_visibility()

        # -- DEFERRED INITIALIZATION ------------------------------------
        # Push non-critical setup (icon load, drop filter walk, update
        # check) off the first-paint path so the window appears fast.
        self._update_checker = None
        self._update_downloader = None
        self._edit_beta_shown = False
        QTimer.singleShot(0, self._post_show_init)
        QTimer.singleShot(10000, self._check_for_updates)

    def _post_show_init(self):
        """Run non-critical setup after the window has painted once."""
        # App icon — works both for frozen exe and dev mode
        base = getattr(sys, '_MEIPASS',
                       os.path.dirname(os.path.abspath(__file__)))
        for icon_name in ("boltpdf_icon.png", "boltpdf_icon.ico"):
            icon_path = os.path.join(base, icon_name)
            if os.path.isfile(icon_path):
                app_icon = QIcon(icon_path)
                self.setWindowIcon(app_icon)
                QApplication.instance().setWindowIcon(app_icon)
                break

        # Global drag-and-drop: install event filter on all child widgets
        # so PDFs can be dropped anywhere in the window.  This walks the
        # full widget tree so we defer it off the first-paint path.
        self._install_drop_filter(self)

    # -- public: open a PDF in a new tab ----------------------------------
    def open_pdf_in_new_tab(self, path):
        """Create a new DocumentTab, load the PDF, and switch to it."""
        tab = DocumentTab(self)
        tab._view.zoom_changed.connect(self._on_zoom_changed)
        tab.status_changed.connect(self._sync_toolbar)
        name = os.path.basename(path)
        idx = self._tabs.addTab(tab, name)
        self._tabs.setCurrentIndex(idx)
        # Build the initial layout in the saved reading mode (spread
        # is decided inside load_pdf from tab._reading_mode).
        tab._reading_mode = app_state().pref(
            "reading_mode", "continuous")
        tab.load_pdf(path)
        app_state().add_recent(path)
        tab.set_tint(app_state().pref("display_tint", "none"))
        self._refresh_thumbs()
        if app_state().pref("reading_mode", "continuous") == "single":
            QTimer.singleShot(
                120, lambda t=tab: self._apply_reading_mode(t))
        self._update_tab_bar_visibility()
        self._sync_toolbar()
        # Show notes panel if the document has annotations
        self._show_notes_panel_for_tab(tab)
        # Connect scroll events to update the notes panel as the user
        # scrolls through pages (throttled via a 200 ms single-shot timer)
        self._connect_notes_scroll(tab)

    # -- Toolbar ----------------------------------------------------------
    def _build_toolbar(self):
        tb = QToolBar("Main", self)
        tb.setMovable(False)
        tb.setIconSize(QSize(20, 20))
        tb.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextBesideIcon)
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, tb)

        act_open = QAction("Open", self)
        act_open.setShortcut(QKeySequence("Ctrl+O"))
        act_open.triggered.connect(self._open_file)
        tb.addAction(act_open)

        # Recent files + Bookmarks dropdowns (instant-popup menus,
        # populated lazily on aboutToShow so they cost nothing until used)
        from PyQt6.QtWidgets import QMenu
        self._recent_btn = QToolButton(self)
        self._recent_btn.setText("Recent")
        self._recent_btn.setToolTip("Reopen a recently viewed PDF")
        self._recent_btn.setPopupMode(
            QToolButton.ToolButtonPopupMode.InstantPopup)
        self._recent_btn.setToolButtonStyle(
            Qt.ToolButtonStyle.ToolButtonTextOnly)
        self._recent_menu = QMenu(self._recent_btn)
        self._recent_menu.aboutToShow.connect(self._populate_recent_menu)
        self._recent_btn.setMenu(self._recent_menu)
        tb.addWidget(self._recent_btn)

        self._bookmark_btn = QToolButton(self)
        self._bookmark_btn.setText("Bookmarks")
        self._bookmark_btn.setToolTip(
            "Add or jump to a bookmark in this document")
        self._bookmark_btn.setPopupMode(
            QToolButton.ToolButtonPopupMode.InstantPopup)
        self._bookmark_btn.setToolButtonStyle(
            Qt.ToolButtonStyle.ToolButtonTextOnly)
        self._bookmark_menu = QMenu(self._bookmark_btn)
        self._bookmark_menu.aboutToShow.connect(
            self._populate_bookmark_menu)
        self._bookmark_btn.setMenu(self._bookmark_menu)
        tb.addWidget(self._bookmark_btn)

        self._display_btn = QToolButton(self)
        self._display_btn.setText("Display")
        self._display_btn.setToolTip(
            "Reading tint — night (invert), sepia, warm, dim")
        self._display_btn.setPopupMode(
            QToolButton.ToolButtonPopupMode.InstantPopup)
        self._display_btn.setToolButtonStyle(
            Qt.ToolButtonStyle.ToolButtonTextOnly)
        self._display_menu = QMenu(self._display_btn)
        self._display_menu.aboutToShow.connect(
            self._populate_display_menu)
        self._display_btn.setMenu(self._display_menu)
        tb.addWidget(self._display_btn)

        self._act_thumbs = QAction("Pages", self)
        self._act_thumbs.setCheckable(True)
        self._act_thumbs.setToolTip(
            "Show the page-thumbnail sidebar (lazy — hidden by default)")
        self._act_thumbs.toggled.connect(self._toggle_thumbs)
        tb.addAction(self._act_thumbs)

        self._layout_btn = QToolButton(self)
        self._layout_btn.setText("Layout")
        self._layout_btn.setToolTip(
            "Reading layout — continuous scroll or single page")
        self._layout_btn.setPopupMode(
            QToolButton.ToolButtonPopupMode.InstantPopup)
        self._layout_btn.setToolButtonStyle(
            Qt.ToolButtonStyle.ToolButtonTextOnly)
        self._layout_menu = QMenu(self._layout_btn)
        self._layout_menu.aboutToShow.connect(self._populate_layout_menu)
        self._layout_btn.setMenu(self._layout_menu)
        tb.addWidget(self._layout_btn)

        tb.addSeparator()

        act_zout = QAction("\u2212", self)
        act_zout.setToolTip("Zoom Out")
        act_zout.setShortcut(QKeySequence("Ctrl+-"))
        act_zout.triggered.connect(self._zoom_out)
        tb.addAction(act_zout)

        self._zoom_label = QLabel("100%")
        self._zoom_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._zoom_label.setFixedWidth(55)
        tb.addWidget(self._zoom_label)

        act_zin = QAction("+", self)
        act_zin.setToolTip("Zoom In")
        act_zin.setShortcut(QKeySequence("Ctrl+="))
        act_zin.triggered.connect(self._zoom_in)
        tb.addAction(act_zin)

        tb.addSeparator()

        act_fs = QAction("Fullscreen", self)
        act_fs.setShortcut(QKeySequence("F11"))
        act_fs.triggered.connect(self._toggle_fullscreen)
        tb.addAction(act_fs)

        self._act_fit = QAction("Fit Page", self)
        self._act_fit.setShortcut(QKeySequence("Ctrl+0"))
        self._act_fit.setToolTip("Fit page to screen, scroll one page at a time")
        self._act_fit.setCheckable(True)
        self._act_fit.triggered.connect(self._toggle_fit)
        self._act_fit.setEnabled(False)
        tb.addAction(self._act_fit)

        tb.addSeparator()

        self._act_ocr = QAction("Detect Text", self)
        self._act_ocr.setShortcut(QKeySequence("Ctrl+T"))
        self._act_ocr.setToolTip(
            "Detect text — click and drag to select, Shift+drag to extend")
        self._act_ocr.triggered.connect(self._toggle_ocr)
        self._act_ocr.setEnabled(False)
        tb.addAction(self._act_ocr)

        self._act_copy = QAction("Copy", self)
        self._act_copy.setShortcut(QKeySequence("Ctrl+C"))
        self._act_copy.setToolTip("Copy selected text to clipboard")
        self._act_copy.triggered.connect(self._copy_selected_text)
        self._act_copy.setVisible(False)
        tb.addAction(self._act_copy)

        self._act_select_all = QAction("Select All", self)
        self._act_select_all.setShortcut(QKeySequence("Ctrl+A"))
        self._act_select_all.triggered.connect(self._select_all_ocr)
        self._act_select_all.setVisible(False)
        tb.addAction(self._act_select_all)

        tb.addSeparator()

        self._act_combine = QAction("Combine", self)
        self._act_combine.setShortcut(QKeySequence("Ctrl+M"))
        self._act_combine.setToolTip(
            "Select multiple PDF files and merge them into one")
        self._act_combine.triggered.connect(self._combine_pdfs)
        tb.addAction(self._act_combine)

        self._act_word = QAction("Export To Docx", self)
        self._act_word.setShortcut(QKeySequence("Ctrl+W"))
        self._act_word.setToolTip(
            "Export pages to a Word document (.docx)")
        self._act_word.triggered.connect(self._export_word)
        self._act_word.setEnabled(False)
        tb.addAction(self._act_word)

        self._act_rebuild = QAction("Optimise PDF", self)
        self._act_rebuild.setToolTip(
            "Snapshot every page as JPEG and rebuild an optimised image-based PDF")
        self._act_rebuild.triggered.connect(self._rebuild_as_images)
        self._act_rebuild.setEnabled(False)
        tb.addAction(self._act_rebuild)

        self._act_export = QAction("Export Images", self)
        self._act_export.setToolTip(
            "Enter image selection mode — click images to select, then export")
        self._act_export.triggered.connect(self._toggle_image_select)
        self._act_export.setEnabled(False)
        tb.addAction(self._act_export)

        self._act_export_sel = QAction("Export Selected", self)
        self._act_export_sel.setToolTip("Export selected images as JPEG files")
        self._act_export_sel.triggered.connect(self._export_selected_images)
        self._act_export_sel.setVisible(False)
        tb.addAction(self._act_export_sel)

        self._act_export_pages = QAction("Export Pages", self)
        self._act_export_pages.setToolTip(
            "Export selected pages as JPEG images into a folder")
        self._act_export_pages.triggered.connect(self._export_pages)
        self._act_export_pages.setEnabled(False)
        tb.addAction(self._act_export_pages)


        self._act_notes = QAction("Notes", self)
        self._act_notes.setToolTip(
            "Show or hide the annotations / comments panel")
        self._act_notes.setCheckable(True)
        self._act_notes.triggered.connect(self._toggle_notes_panel)
        self._act_notes.setEnabled(False)
        tb.addAction(self._act_notes)

        tb.addSeparator()

        # Edit Mode — top-level toggle
        self._act_edit = QAction("Edit", self)
        self._act_edit.setShortcut(QKeySequence("Ctrl+E"))
        self._act_edit.setToolTip(
            "Enter edit mode — move/edit text boxes and images, "
            "add new text or images, and save as a new PDF")
        self._act_edit.setCheckable(True)
        self._act_edit.triggered.connect(self._toggle_edit_mode)
        self._act_edit.setEnabled(False)
        tb.addAction(self._act_edit)

        # ── Right-side vertical toolbar for edit tools ──────────────
        etb = QToolBar("Edit Tools", self)
        etb.setMovable(False)
        etb.setIconSize(QSize(20, 20))
        etb.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextBesideIcon)
        etb.setStyleSheet(
            "QToolBar { spacing: 2px; padding: 4px; }"
            "QToolButton { min-width: 90px; text-align: left; }")
        self.addToolBar(Qt.ToolBarArea.RightToolBarArea, etb)
        self._edit_toolbar = etb
        etb.setVisible(False)  # hidden until edit mode is entered

        # Undo / Redo — history lives on DocumentTab; Ctrl+Z / Ctrl+Y
        # and Ctrl+Shift+Z are also wired in keyPressEvent.
        self._act_edit_undo = QAction("Undo", self)
        self._act_edit_undo.setToolTip("Undo the last edit (Ctrl+Z)")
        self._act_edit_undo.triggered.connect(self._edit_undo)
        self._act_edit_undo.setEnabled(False)
        etb.addAction(self._act_edit_undo)

        self._act_edit_redo = QAction("Redo", self)
        self._act_edit_redo.setToolTip(
            "Redo the last undo (Ctrl+Shift+Z or Ctrl+Y)")
        self._act_edit_redo.triggered.connect(self._edit_redo)
        self._act_edit_redo.setEnabled(False)
        etb.addAction(self._act_edit_redo)

        self._act_edit_props = QAction("Properties", self)
        self._act_edit_props.setCheckable(True)
        self._act_edit_props.setToolTip(
            "Show the Properties inspector for the selected object")
        self._act_edit_props.toggled.connect(self._toggle_props)
        etb.addAction(self._act_edit_props)

        etb.addSeparator()

        # Edit Mode sub-actions
        self._act_edit_add_text = QAction("Add Text", self)
        self._act_edit_add_text.setToolTip(
            "Click on a page to drop a new text box")
        self._act_edit_add_text.setCheckable(True)
        self._act_edit_add_text.triggered.connect(self._edit_add_text)
        etb.addAction(self._act_edit_add_text)

        self._act_edit_add_image = QAction("Add Image", self)
        self._act_edit_add_image.setToolTip(
            "Pick an image file, then click on a page to drop it")
        self._act_edit_add_image.setCheckable(True)
        self._act_edit_add_image.triggered.connect(self._edit_add_image)
        etb.addAction(self._act_edit_add_image)

        self._act_edit_transform = QAction("Transform", self)
        self._act_edit_transform.setToolTip(
            "Click a text box or image to show move / scale / rotate "
            "handles.  Drag the middle to move, the corners/edges to "
            "resize, and the green knob above the top to rotate "
            "(hold Shift to snap rotation to 15°).")
        self._act_edit_transform.setCheckable(True)
        self._act_edit_transform.triggered.connect(self._edit_transform)
        etb.addAction(self._act_edit_transform)

        self._act_edit_select = QAction("Select", self)
        self._act_edit_select.setToolTip(
            "Rubber-band multi-select: drag a box over objects, then "
            "Align / nudge with arrows / Delete them together.")
        self._act_edit_select.setCheckable(True)
        self._act_edit_select.triggered.connect(self._edit_multiselect)
        etb.addAction(self._act_edit_select)

        from PyQt6.QtWidgets import QToolButton as _QTB, QMenu as _QM
        self._act_edit_align_btn = _QTB(self)
        self._act_edit_align_btn.setText("Align")
        self._act_edit_align_btn.setToolTip(
            "Align the multi-selected objects")
        self._act_edit_align_btn.setPopupMode(
            _QTB.ToolButtonPopupMode.InstantPopup)
        _amenu = _QM(self._act_edit_align_btn)
        for _label, _mode in (("Left", "left"), ("Right", "right"),
                              ("Top", "top"), ("Bottom", "bottom"),
                              ("Center H", "centerx"),
                              ("Center V", "centery")):
            _aa = _amenu.addAction(_label)
            _aa.triggered.connect(
                lambda _c=False, m=_mode: self._edit_align(m))
        _amenu.addSeparator()
        for _label, _ax in (("Distribute horizontally", "h"),
                            ("Distribute vertically", "v")):
            _da = _amenu.addAction(_label)
            _da.triggered.connect(
                lambda _c=False, a=_ax: self._edit_distribute(a))
        self._act_edit_align_btn.setMenu(_amenu)
        etb.addWidget(self._act_edit_align_btn)

        self._act_edit_group = QAction("Group", self)
        self._act_edit_group.setToolTip(
            "Group the multi-selected objects so they select & move "
            "together (Transform-click any member to select the group)")
        self._act_edit_group.triggered.connect(self._edit_group)
        etb.addAction(self._act_edit_group)

        self._act_edit_ungroup = QAction("Ungroup", self)
        self._act_edit_ungroup.setToolTip(
            "Ungroup the selected group")
        self._act_edit_ungroup.triggered.connect(self._edit_ungroup)
        etb.addAction(self._act_edit_ungroup)

        self._act_edit_snap = QAction("Snap", self)
        self._act_edit_snap.setCheckable(True)
        self._act_edit_snap.setChecked(True)
        self._act_edit_snap.setToolTip(
            "Snap a moved object to other objects' edges/centres and "
            "the page margins, with live guide lines")
        self._act_edit_snap.toggled.connect(self._edit_toggle_snap)
        etb.addAction(self._act_edit_snap)

        self._act_edit_delete = QAction("Delete", self)
        self._act_edit_delete.setToolTip(
            "Delete the currently selected text box or image from the "
            "document.  Pick something with the Transform tool first, "
            "then click Delete (or press Delete / Backspace) to remove "
            "it.  The change is written when you hit Save As.")
        self._act_edit_delete.triggered.connect(self._edit_delete_selection)
        etb.addAction(self._act_edit_delete)

        self._act_edit_duplicate = QAction("Duplicate", self)
        self._act_edit_duplicate.setToolTip(
            "Duplicate the selected object (Ctrl+D). Ctrl+C/Ctrl+V "
            "copy/paste — paste lands on the page you're viewing.")
        self._act_edit_duplicate.triggered.connect(
            self._edit_duplicate)
        etb.addAction(self._act_edit_duplicate)

        self._act_edit_replace_img = QAction("Replace Image", self)
        self._act_edit_replace_img.setToolTip(
            "Replace the currently selected image with a file from disk")
        self._act_edit_replace_img.triggered.connect(self._edit_replace_image)
        etb.addAction(self._act_edit_replace_img)

        etb.addSeparator()

        # --- Stamp, Shapes, Whiteout, Sticky Note ---
        self._act_edit_stamp = QAction("Stamp", self)
        self._act_edit_stamp.setToolTip(
            "Place a stamp (APPROVED, DRAFT, etc.) on the page")
        self._act_edit_stamp.triggered.connect(self._edit_stamp)
        etb.addAction(self._act_edit_stamp)

        self._act_edit_shape_rect = QAction("Rect", self)
        self._act_edit_shape_rect.setToolTip("Draw a rectangle")
        self._act_edit_shape_rect.setCheckable(True)
        self._act_edit_shape_rect.triggered.connect(
            lambda: self._edit_shape("shape_rect"))
        etb.addAction(self._act_edit_shape_rect)

        self._act_edit_shape_circle = QAction("Circle", self)
        self._act_edit_shape_circle.setToolTip("Draw a circle / ellipse")
        self._act_edit_shape_circle.setCheckable(True)
        self._act_edit_shape_circle.triggered.connect(
            lambda: self._edit_shape("shape_circle"))
        etb.addAction(self._act_edit_shape_circle)

        self._act_edit_shape_line = QAction("Line", self)
        self._act_edit_shape_line.setToolTip("Draw a straight line")
        self._act_edit_shape_line.setCheckable(True)
        self._act_edit_shape_line.triggered.connect(
            lambda: self._edit_shape("shape_line"))
        etb.addAction(self._act_edit_shape_line)

        self._act_edit_shape_arrow = QAction("Arrow", self)
        self._act_edit_shape_arrow.setToolTip("Draw an arrow")
        self._act_edit_shape_arrow.setCheckable(True)
        self._act_edit_shape_arrow.triggered.connect(
            lambda: self._edit_shape("shape_arrow"))
        etb.addAction(self._act_edit_shape_arrow)

        self._act_edit_redact = QAction("Whiteout", self)
        self._act_edit_redact.setToolTip(
            "Draw a white rectangle to hide content")
        self._act_edit_redact.setCheckable(True)
        self._act_edit_redact.triggered.connect(self._edit_redact)
        etb.addAction(self._act_edit_redact)

        self._act_edit_highlight = QAction("Highlight", self)
        self._act_edit_highlight.setToolTip(
            "Drag over content to add a real PDF highlight annotation")
        self._act_edit_highlight.setCheckable(True)
        self._act_edit_highlight.triggered.connect(self._edit_highlight)
        etb.addAction(self._act_edit_highlight)

        self._act_edit_note = QAction("Sticky Note", self)
        self._act_edit_note.setToolTip(
            "Click on the page to add a sticky note comment")
        self._act_edit_note.setCheckable(True)
        self._act_edit_note.triggered.connect(self._edit_note)
        etb.addAction(self._act_edit_note)

        etb.addSeparator()

        self._act_edit_save = QAction("Save As...", self)
        self._act_edit_save.setToolTip(
            "Apply all pending edits and write the result as a new PDF")
        self._act_edit_save.triggered.connect(self._edit_save_as)
        etb.addAction(self._act_edit_save)

        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding,
                             QSizePolicy.Policy.Preferred)
        tb.addWidget(spacer)

        # Search bar — inline in the toolbar, hidden until Ctrl+F
        self._search_bar = SearchBar(self)
        self._search_bar.search_changed.connect(self._on_search_changed)
        self._search_bar.next_clicked.connect(self._on_search_next)
        self._search_bar.prev_clicked.connect(self._on_search_prev)
        self._search_bar.closed.connect(self._on_search_closed)
        tb.addWidget(self._search_bar)

        # Check for Updates — always visible at the right end
        self._act_check_update = QAction("Check for Updates", self)
        self._act_check_update.setToolTip(
            "Check if a newer version of BoltPDF is available")
        self._act_check_update.triggered.connect(
            self._manual_check_for_updates)
        tb.addAction(self._act_check_update)

        # About / Licenses — always visible (GNU AGPL v3 §5: Appropriate
        # Legal Notices must be conveniently and prominently accessible).
        self._act_about = QAction("About", self)
        self._act_about.setToolTip(
            "Copyright, licence (GNU AGPL v3) and third-party notices")
        self._act_about.triggered.connect(self._show_about)
        tb.addAction(self._act_about)

        tb.addSeparator()

        # Page jump box — type a number and press Enter to jump to that page
        self._page_jump = QLineEdit()
        self._page_jump.setFixedWidth(50)
        self._page_jump.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._page_jump.setPlaceholderText("#")
        self._page_jump.setToolTip("Enter a page number and press Enter")
        self._page_jump.setStyleSheet(
            "QLineEdit { background: #3c3c3c; color: #e0e0e0; border: 1px solid #555;"
            " border-radius: 3px; padding: 2px; font-size: 12px; }"
            "QLineEdit:focus { border: 1px solid #7a7aff; }")
        self._page_jump.returnPressed.connect(self._on_page_jump)
        tb.addWidget(self._page_jump)

        self._page_label = QLabel("")
        tb.addWidget(self._page_label)

    # -- About / Licenses (GNU AGPL v3 §5: Appropriate Legal Notices) -----
    def _show_about(self):
        """Display copyright, the no-warranty disclaimer, redistribution
        terms, and convenient access to the full AGPL-3.0 text and the
        bundled third-party licence notices."""
        candidates = []
        if getattr(sys, 'frozen', False):
            candidates.append(getattr(sys, '_MEIPASS', ''))
            candidates.append(os.path.dirname(sys.executable))
        else:
            candidates.append(os.path.dirname(os.path.abspath(__file__)))

        def _find(name):
            for d in candidates:
                if d:
                    p = os.path.join(d, name)
                    if os.path.isfile(p):
                        return p
            return None

        license_path = _find("LICENSE")
        tpl_path = _find("THIRD_PARTY_LICENSES.txt")

        dlg = QDialog(self)
        dlg.setWindowTitle("About BoltPDF")
        dlg.resize(660, 540)
        lay = QVBoxLayout(dlg)

        header = QLabel()
        header.setTextFormat(Qt.TextFormat.RichText)
        header.setWordWrap(True)
        header.setOpenExternalLinks(True)
        header.setText(
            f"<h2 style='margin-bottom:2px'>BoltPDF v{__version__}</h2>"
            "<p>Copyright &copy; 2026 BoltPDF</p>"
            "<p>BoltPDF is <b>free software</b>, licensed under the "
            "<b>GNU Affero General Public License, version 3</b> "
            "(AGPL-3.0-or-later).</p>"
            "<p>This program comes with <b>ABSOLUTELY NO WARRANTY</b>. "
            "This is free software, and you are welcome to redistribute it "
            "under the terms of the GNU AGPL v3. BoltPDF links the "
            "AGPL-licensed PyMuPDF and the GPL-3.0-licensed PyQt6; the "
            "combined work is conveyed under the AGPL-3.0.</p>"
            "<p>Complete corresponding source code: "
            "<a href='https://github.com/HBDPN/BoltPDF'>"
            "github.com/HBDPN/BoltPDF</a></p>"
        )
        lay.addWidget(header)

        viewer = QTextEdit()
        viewer.setReadOnly(True)
        viewer.setLineWrapMode(QTextEdit.LineWrapMode.NoWrap)
        lay.addWidget(viewer, 1)

        def _load(path, fallback):
            if path:
                try:
                    with open(path, "r", encoding="utf-8",
                              errors="replace") as fh:
                        viewer.setPlainText(fh.read())
                    return
                except OSError:
                    pass
            viewer.setPlainText(fallback)

        agpl_fallback = ("The full GNU Affero General Public License v3 "
                         "text is available at "
                         "https://www.gnu.org/licenses/agpl-3.0.txt")

        row = QHBoxLayout()
        b_lic = QPushButton("Licence (AGPL-3.0)")
        b_tpl = QPushButton("Third-Party Licences")
        b_lic.clicked.connect(lambda: _load(license_path, agpl_fallback))
        b_tpl.clicked.connect(lambda: _load(
            tpl_path, "THIRD_PARTY_LICENSES.txt was not found in this "
                      "installation. See https://github.com/HBDPN/BoltPDF"))
        row.addWidget(b_lic)
        row.addWidget(b_tpl)
        row.addStretch(1)
        bb = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        bb.rejected.connect(dlg.reject)
        bb.accepted.connect(dlg.accept)
        row.addWidget(bb)
        lay.addLayout(row)

        _load(license_path, agpl_fallback)
        dlg.exec()

    # -- current tab helper -----------------------------------------------
    def _current_tab(self) -> DocumentTab | None:
        w = self._tabs.currentWidget()
        return w if isinstance(w, DocumentTab) else None

    # -- Toolbar sync (called when active tab changes or tab status
    #    changes) ---------------------------------------------------------
    def _sync_toolbar(self):
        tab = self._current_tab()
        if tab is None:
            self.setWindowTitle("BoltPDF")
            self._page_label.setText("")
            self._zoom_label.setText("100%")
            self._act_ocr.setEnabled(False)
            self._act_ocr.setText("Detect Text")
            self._act_copy.setVisible(False)
            self._act_select_all.setVisible(False)
            self._act_fit.setEnabled(False)
            self._act_fit.setChecked(False)
            self._act_rebuild.setEnabled(False)
            self._act_word.setEnabled(False)
            self._act_export.setEnabled(False)
            self._act_export.setText("Export Images")
            self._act_export_sel.setVisible(False)
            self._act_export_pages.setEnabled(False)
            self._act_edit.setEnabled(False)
            self._act_edit.setChecked(False)
            self._edit_toolbar.setVisible(False)
            self._act_notes.setEnabled(False)
            self._act_notes.setChecked(False)
            self._notes_panel.setVisible(False)
            return

        # Window title
        name = os.path.basename(tab.doc_path) if tab.doc_path else ""
        self.setWindowTitle(f"BoltPDF \u2014 {name}" if name else "BoltPDF")

        # Zoom
        self._zoom_label.setText(f"{int(tab.view._current_scale * 100)}%")

        # Page label
        if tab.ocr_active:
            if tab._ocr_worker and tab._ocr_worker.isRunning():
                self._page_label.setText("Running OCR...")
            else:
                self._page_label.setText(
                    f"{tab.num_pages} pages \u2022 OCR active "
                    f"({len(tab._ocr_done_pages)} scanned)")
        else:
            self._page_label.setText(
                f"{tab.num_pages} pages" if tab.num_pages else "")

        # OCR button
        has_doc = tab.doc_path is not None and tab.num_pages > 0
        self._act_ocr.setEnabled(has_doc)
        self._act_ocr.setText("Clear Text" if tab.ocr_active else "Detect Text")
        self._act_copy.setVisible(tab.ocr_active)
        self._act_select_all.setVisible(tab.ocr_active)

        # Fit mode
        self._act_fit.setEnabled(has_doc)
        self._act_fit.setChecked(tab.view._fit_mode)

        # Rebuild / Export / Word
        self._act_rebuild.setEnabled(has_doc)
        self._act_word.setEnabled(has_doc)
        self._act_export.setEnabled(has_doc)
        self._act_export_pages.setEnabled(has_doc)
        self._act_notes.setEnabled(has_doc)
        if tab.image_select_mode:
            self._act_export.setText("Cancel Export")
            self._act_export_sel.setVisible(True)
            sel_count = sum(1 for o in tab._image_overlays if o.selected)
            total = len(tab._image_overlays)
            self._act_export_sel.setText(
                f"Export Selected ({sel_count})" if sel_count
                else "Export Selected")
            self._act_export_sel.setEnabled(sel_count > 0)
        else:
            self._act_export.setText("Export Images")
            self._act_export_sel.setVisible(False)

        # Edit mode
        self._act_edit.setEnabled(has_doc)
        in_edit = bool(tab.edit_mode)
        self._act_edit.setChecked(in_edit)
        self._act_edit.setText("Exit Edit" if in_edit else "Edit")
        self._edit_toolbar.setVisible(in_edit)
        if in_edit:
            ea = tab.edit_action
            self._act_edit_add_text.setChecked(ea == 'add_text')
            self._act_edit_add_image.setChecked(ea == 'add_image')
            self._act_edit_transform.setChecked(ea == 'transform')
            self._act_edit_shape_rect.setChecked(ea == 'shape_rect')
            self._act_edit_shape_circle.setChecked(ea == 'shape_circle')
            self._act_edit_shape_line.setChecked(ea == 'shape_line')
            self._act_edit_shape_arrow.setChecked(ea == 'shape_arrow')
            self._act_edit_redact.setChecked(ea == 'redact')
            self._act_edit_highlight.setChecked(ea == 'highlight')
            self._act_edit_select.setChecked(ea == 'multiselect')
            self._act_edit_note.setChecked(ea == 'note')
            # Delete is only meaningful when something is actually
            # selected by the Transform tool.
            self._act_edit_delete.setEnabled(
                tab._active_transform is not None)
            # Replace Image is always available in edit mode — clicking
            # it drops the user into the image-pick flow.
            self._act_edit_replace_img.setEnabled(True)
            # Only let the user save when there's at least one queued
            # EditRecord to apply.
            self._act_edit_save.setEnabled(
                bool(getattr(tab, '_edit_records', None)))
            # Undo/redo availability tracks the history cursor.
            self._act_edit_undo.setEnabled(tab.can_edit_undo())
            self._act_edit_redo.setEnabled(tab.can_edit_redo())
            # Keep the Properties inspector in sync with the current
            # Transform selection (no-op work when it's hidden).
            if self._props_panel.isVisible():
                self._props_panel.refresh(tab)

        # Search bar — keep match counter in sync when results arrive
        if self._search_bar.text():
            still_searching = not getattr(tab, '_search_done', True)
            self._search_bar.set_match_info(
                tab.search_current_index, tab.search_match_count,
                searching=still_searching)

        # Notes panel — show/hide and update for current page
        self._update_notes_panel(tab)

    def _gather_all_notes(self, tab):
        """Collect all notes from both PDF annotations and user-added
        edit records into a single dict: page_idx → list of dicts."""
        combined: dict[int, list[dict]] = {}
        # 1. PDF annotations
        for pg, notes in tab._annotations_by_page.items():
            combined.setdefault(pg, []).extend(notes)
        # 2. User-added sticky notes from edit records
        for rec in getattr(tab, '_edit_records', []):
            if rec.kind == 'note_add' and rec.text:
                combined.setdefault(rec.page_idx, []).append({
                    "author": "You",
                    "content": rec.text,
                    "subject": "",
                    "type": "Sticky Note",
                })
        return combined

    def _update_notes_panel(self, tab=None):
        """Refresh the notes panel content for the visible page."""
        if tab is None:
            tab = self._current_tab()
        if tab is None:
            return
        self._act_notes.setChecked(self._notes_panel.isVisible())
        if not self._notes_panel.isVisible():
            return  # panel is hidden, nothing to update
        # Ensure the splitter gives the panel enough width
        sizes = self._splitter.sizes()
        if len(sizes) == 3 and sizes[2] < 200:
            total = sum(sizes)
            notes_w = 280
            sizes[2] = notes_w
            sizes[1] = total - sizes[0] - notes_w
            self._splitter.setSizes(sizes)
        vis_page = tab.get_visible_page() if tab._num_pages else 0
        all_notes = self._gather_all_notes(tab)
        # Rebuild when the per-page note distribution changes.  Total
        # count alone misses a note *moved* to another page (same
        # count, different page) — which would leave a stale jump
        # target in the panel.
        sig = tuple(sorted((pg, len(v)) for pg, v in all_notes.items()))
        if sig != getattr(self, '_last_notes_sig', None):
            self._notes_panel.clear()
            self._last_notes_sig = sig
        self._notes_panel.show_notes(vis_page, all_notes)

    def _toggle_notes_panel(self):
        """Toggle the notes panel visibility."""
        tab = self._current_tab()
        if tab is None:
            return
        if self._notes_panel.isVisible():
            self._notes_panel.setVisible(False)
            self._act_notes.setChecked(False)
        else:
            self._notes_panel.setVisible(True)
            self._act_notes.setChecked(True)
            # Force a full rebuild since the user explicitly toggled
            self._notes_panel.clear()
            self._update_notes_panel(tab)

    def _show_notes_panel_for_tab(self, tab):
        """Auto-show the notes panel when a tab with notes loads."""
        all_notes = self._gather_all_notes(tab)
        if all_notes:
            self._notes_panel.setVisible(True)
            self._act_notes.setChecked(True)
            self._notes_panel.clear()
            self._update_notes_panel(tab)

    def _connect_notes_scroll(self, tab):
        """Connect a throttled scroll listener so the notes panel updates
        as the user scrolls through pages."""
        if not hasattr(self, '_notes_scroll_timer') or self._notes_scroll_timer is None:
            self._notes_scroll_timer = QTimer(self)
            self._notes_scroll_timer.setSingleShot(True)
            self._notes_scroll_timer.setInterval(200)
            self._notes_scroll_timer.timeout.connect(self._update_notes_panel)
        vsb = tab.view.verticalScrollBar()
        if vsb:
            vsb.valueChanged.connect(self._on_notes_scroll_tick)

    def _on_notes_scroll_tick(self):
        """Restart the notes-panel debounce timer on each scroll event."""
        if hasattr(self, '_notes_scroll_timer') and self._notes_scroll_timer:
            self._notes_scroll_timer.start()

    def _on_page_jump(self):
        """Jump to the page number typed in the page-jump box."""
        tab = self._current_tab()
        if tab is None or not tab._page_positions:
            return
        text = self._page_jump.text().strip()
        if not text.isdigit():
            self._page_jump.clear()
            return
        page = int(text)
        if page < 1:
            page = 1
        if page > tab.num_pages:
            page = tab.num_pages
        idx = page - 1  # 0-based
        if idx in tab._page_positions:
            y = tab._page_positions[idx]
            tab.view.centerOn(0, y)
            # Bypass the scroll debounce — the user has explicitly
            # jumped, so refocus the renderer and the viewport cache
            # immediately instead of waiting for the debounce timer.
            if tab._renderer and tab._renderer.isRunning():
                tab._renderer.set_focus(idx)
            if tab._image_select_mode:
                tab._ensure_image_detector_covers(idx)
            tab._manage_viewport_pages()
        self._page_jump.clear()
        self._page_jump.clearFocus()

    # -- Auto-update ------------------------------------------------------
    # _pending_update stores (version, url, changelog) when a new version
    # is detected.  Both the banner and manual-check dialog use it.
    _pending_update: tuple | None = None

    def _check_for_updates(self):
        """Automatic startup check — shows the banner bar, not a dialog."""
        self._update_checker = UpdateChecker(self)
        self._update_checker.update_available.connect(
            self._on_startup_update_available)
        self._update_checker.start()

    def _on_startup_update_available(self, version, url, changelog):
        """Startup path: slide the non-intrusive banner into view."""
        self._pending_update = (version, url, changelog)
        self._update_banner.show_update(version, changelog)

    def _banner_update_clicked(self):
        """User clicked 'Update Now' on the banner bar."""
        if self._pending_update is None:
            return
        _, url, _ = self._pending_update
        self._update_banner.hide()
        self._start_download(url)

    # -- Manual "Check for Updates" (toolbar button) -----------------------
    def _manual_check_for_updates(self):
        """User explicitly clicked 'Check for Updates' in the toolbar."""
        # If we already know about an update, show a dialog immediately.
        if self._pending_update is not None:
            self._show_update_dialog(*self._pending_update)
            return
        # Otherwise fire a fresh check and show a small "Checking..." dialog
        # that auto-closes when the result arrives.
        self._manual_check_dlg = QProgressDialog(
            "Checking for updates...", "Cancel", 0, 0, self)
        self._manual_check_dlg.setWindowTitle("Check for Updates")
        self._manual_check_dlg.setMinimumDuration(0)
        self._manual_check_dlg.setModal(True)

        checker = UpdateChecker(self)
        checker.update_available.connect(self._on_manual_update_found)
        checker.no_update.connect(self._on_manual_no_update)
        checker.check_failed.connect(self._on_manual_check_failed)
        self._manual_check_dlg.canceled.connect(checker.terminate)
        self._manual_checker = checker  # prevent GC
        checker.start()

    def _on_manual_update_found(self, version, url, changelog):
        try:
            self._manual_check_dlg.close()
        except Exception:
            pass
        self._pending_update = (version, url, changelog)
        self._update_banner.show_update(version, changelog)
        self._show_update_dialog(version, url, changelog)

    def _on_manual_no_update(self):
        try:
            self._manual_check_dlg.close()
        except Exception:
            pass
        QMessageBox.information(
            self, "No Updates",
            f"You're running the latest version of BoltPDF (v{__version__}).")

    def _on_manual_check_failed(self, error):
        try:
            self._manual_check_dlg.close()
        except Exception:
            pass
        QMessageBox.warning(
            self, "Check Failed",
            f"Could not check for updates:\n{error}\n\n"
            "You can check manually at https://github.com/HBDPN/BoltPDF")

    def _show_update_dialog(self, version, url, changelog):
        """Show a modal dialog offering to download and install."""
        msg = QMessageBox(self)
        msg.setWindowTitle("Update Available")
        msg.setIcon(QMessageBox.Icon.Information)
        text = f"BoltPDF v{version} is available (you have v{__version__})."
        if changelog:
            text += f"\n\nWhat's new:\n{changelog}"
        msg.setText(text)
        msg.setStandardButtons(
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        msg.button(QMessageBox.StandardButton.Yes).setText("Update Now")
        msg.button(QMessageBox.StandardButton.No).setText("Later")
        if msg.exec() != QMessageBox.StandardButton.Yes:
            return
        self._update_banner.hide()
        self._start_download(url)

    def _start_download(self, url):
        """Download the new exe with a progress dialog."""
        self._dl_progress = QProgressDialog(
            "Downloading update...", "Cancel", 0, 100, self)
        self._dl_progress.setWindowTitle("Updating BoltPDF")
        self._dl_progress.setMinimumDuration(0)
        self._dl_progress.setValue(0)

        self._update_downloader = UpdateDownloader(url, self)
        self._update_downloader.progress.connect(self._dl_progress.setValue)
        self._update_downloader.finished_ok.connect(self._on_download_done)
        self._update_downloader.error_occurred.connect(self._on_download_error)
        self._dl_progress.canceled.connect(self._update_downloader.terminate)
        self._update_downloader.start()

    def _on_download_done(self, tmp_path):
        self._dl_progress.close()
        msg = QMessageBox(self)
        msg.setWindowTitle("Update Ready")
        msg.setIcon(QMessageBox.Icon.Question)
        msg.setText(
            "Update downloaded successfully.\n\n"
            "BoltPDF will close and restart with the new version.")
        msg.setStandardButtons(
            QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel)
        msg.button(QMessageBox.StandardButton.Ok).setText("Restart Now")
        if msg.exec() == QMessageBox.StandardButton.Ok:
            _apply_update_and_restart(tmp_path)

    def _on_download_error(self, error):
        self._dl_progress.close()
        QMessageBox.warning(
            self, "Update Failed",
            f"Could not download update:\n{error}\n\n"
            "You can download manually from https://github.com/HBDPN/BoltPDF")

    # -- Search handlers ----------------------------------------------------
    def _on_search_changed(self, query: str):
        """Debounce: wait 300ms after the last keystroke before searching."""
        # Reset the debounce timer on every keystroke
        if not hasattr(self, '_search_debounce'):
            self._search_debounce = QTimer(self)
            self._search_debounce.setSingleShot(True)
            self._search_debounce.timeout.connect(self._fire_search)
        self._search_debounce.start(300)

    def _fire_search(self):
        """Actually start the async search after the debounce delay."""
        query = self._search_bar.text()
        tab = self._current_tab()
        if not tab:
            self._search_bar.set_match_info(0, 0)
            return
        if not query:
            tab.clear_search()
            self._search_bar.set_match_info(0, 0)
            return
        if len(query) < DocumentTab._SEARCH_MIN_QUERY_LEN:
            tab.clear_search()
            self._search_bar._match_label.setText(
                f"Type {DocumentTab._SEARCH_MIN_QUERY_LEN}+ chars")
            return
        self._search_bar._match_label.setText("Searching...")
        tab.start_search(query)

    def _on_search_status_update(self):
        """Called when any tab emits status_changed — if the search bar
        has a query, refresh the match counter."""
        if not self._search_bar.text():
            return
        tab = self._current_tab()
        if not tab:
            return
        still_searching = not getattr(tab, '_search_done', True)
        self._search_bar.set_match_info(
            tab.search_current_index, tab.search_match_count,
            searching=still_searching)

    def _on_search_next(self):
        tab = self._current_tab()
        if not tab:
            return
        tab.search_next()
        still_searching = not getattr(tab, '_search_done', True)
        self._search_bar.set_match_info(
            tab.search_current_index, tab.search_match_count,
            searching=still_searching)

    def _on_search_prev(self):
        tab = self._current_tab()
        if not tab:
            return
        tab.search_prev()
        still_searching = not getattr(tab, '_search_done', True)
        self._search_bar.set_match_info(
            tab.search_current_index, tab.search_match_count,
            searching=still_searching)

    def _on_search_closed(self):
        self._search_bar.close_bar()
        tab = self._current_tab()
        if tab:
            tab.clear_search()

    def _on_tab_changed(self, index):
        self._sync_toolbar()
        # If the search bar has a query, re-run on the new tab
        if self._search_bar.text():
            self._fire_search()
        # Update preview panel for the newly active tab
        tab = self._tabs.widget(index) if index >= 0 else None
        if isinstance(tab, DocumentTab):
            self._show_notes_panel_for_tab(tab)
            self._refresh_thumbs()
        else:
            self._notes_panel.clear()
            self._notes_panel.setVisible(False)
            if self._thumb_panel.isVisible():
                self._thumb_panel.clear()

    # -- Tab management ---------------------------------------------------
    def _close_tab(self, index):
        tab = self._tabs.widget(index)
        if isinstance(tab, DocumentTab):
            tab.cleanup()
        self._tabs.removeTab(index)
        self._update_tab_bar_visibility()
        self._sync_toolbar()
        # Refresh side panels for whatever tab is now active
        new_tab = self._current_tab()
        if new_tab:
            self._show_notes_panel_for_tab(new_tab)
        else:
            self._notes_panel.clear()
            self._notes_panel.setVisible(False)

    # -- Window close: make sure every worker / subprocess is gone -------
    def closeEvent(self, event):
        """Shut down every background worker and subprocess this app
        owns before Qt tears the window down.

        Without this hook the PageRenderer's multiprocessing workers
        (plus any Pools owned by short-lived export / rebuild threads)
        can stay alive after the main window has closed, leaving a
        pile of BoltPDF.exe entries in Task Manager until they're
        killed manually.
        """
        # 0. Persist the session (open docs + page + zoom) BEFORE any
        #    teardown, so it can be restored next launch.  Best-effort;
        #    must never block or fail the close.
        try:
            entries = []
            for i in range(self._tabs.count()):
                t = self._tabs.widget(i)
                if isinstance(t, DocumentTab) and t.doc_path:
                    entries.append({
                        "path": t.doc_path,
                        "page": t.current_page_index(),
                        "zoom": t.current_zoom(),
                    })
            app_state().set_session(entries)
        except Exception:
            pass

        # 0a. Mark a clean shutdown so the next plain launch opens
        #     blank (a crash leaves this False → recover instead).
        try:
            app_state().set_pref("clean_exit", True)
        except Exception:
            pass

        # 0b. Stop the thumbnail render thread (not owned by any tab).
        try:
            self._thumb_panel.stop()
        except Exception:
            pass

        # 1. Stop update-related QThreads on the main window itself.
        for attr in ("_update_checker", "_update_downloader"):
            w = getattr(self, attr, None)
            if w is None:
                continue
            try:
                if w.isRunning():
                    if hasattr(w, "cancel"):
                        w.cancel()
                    w.wait(1000)
                    if w.isRunning():
                        w.terminate()
                        w.wait(500)
            except Exception:
                pass

        # 2. Stop any top-level workers that live on the window (combine
        #    PDFs, export-to-Word) — these each own their own mp.Pool.
        for attr in ("_combine_worker", "_word_worker", "_pages_worker"):
            w = getattr(self, attr, None)
            if w is None:
                continue
            try:
                if w.isRunning():
                    w.wait(2000)
                    if w.isRunning():
                        w.terminate()
                        w.wait(500)
            except Exception:
                pass

        # 3. Cleanup every open tab: stops PageRenderer workers, OCR,
        #    image detector, image export, rebuild, and clears the
        #    on-disk cache for each.
        for i in range(self._tabs.count()):
            tab = self._tabs.widget(i)
            if isinstance(tab, DocumentTab):
                try:
                    tab.cleanup()
                except Exception:
                    pass

        # 4. Belt-and-braces: kill any multiprocessing child that slipped
        #    through the cooperative shutdown above.  These are what
        #    show up in Task Manager as extra BoltPDF.exe entries.
        try:
            import multiprocessing as _mp
            alive = _mp.active_children()
            for p in alive:
                try:
                    p.terminate()
                except Exception:
                    pass
            # Give them half a second to go, then hard-kill anything
            # that's still breathing.
            deadline = time.monotonic() + 0.5
            while time.monotonic() < deadline:
                if not any(p.is_alive() for p in alive):
                    break
                time.sleep(0.02)
            for p in _mp.active_children():
                try:
                    if p.is_alive():
                        if hasattr(p, "kill"):
                            p.kill()
                        else:
                            p.terminate()
                except Exception:
                    pass
        except Exception:
            pass

        super().closeEvent(event)

    def _update_tab_bar_visibility(self):
        self._tabs.tabBar().setVisible(self._tabs.count() > 1)

    # -- File loading (toolbar Open) --------------------------------------
    def _open_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Open PDF", "", "PDF Files (*.pdf)")
        if path:
            self.open_pdf_in_new_tab(path)

    # -- Recent files -----------------------------------------------------
    def _populate_recent_menu(self):
        m = self._recent_menu
        m.clear()
        recent = app_state().recent()
        if not recent:
            a = m.addAction("(no recent files)")
            a.setEnabled(False)
            return
        for p in recent:
            act = m.addAction(os.path.basename(p))
            act.setToolTip(p)
            act.triggered.connect(
                lambda _c=False, path=p: self.open_pdf_in_new_tab(path))
        m.addSeparator()
        clr = m.addAction("Clear recent")
        clr.triggered.connect(lambda: app_state().clear_recent())

    # -- Per-document bookmarks ------------------------------------------
    def _populate_bookmark_menu(self):
        m = self._bookmark_menu
        m.clear()
        tab = self._current_tab()
        if tab is None or not tab.doc_path:
            a = m.addAction("(open a PDF first)")
            a.setEnabled(False)
            return
        cur = tab.current_page_index()
        add = m.addAction(f"Add bookmark — page {cur + 1}")
        add.triggered.connect(
            lambda _c=False, t=tab, p=cur: self._add_bookmark(t, p))
        marks = app_state().bookmarks(tab.doc_path)
        if marks:
            m.addSeparator()
            for bm in sorted(marks, key=lambda b: b.get("page", 0)):
                pg = int(bm.get("page", 0))
                label = bm.get("label") or f"Page {pg + 1}"
                act = m.addAction(f"p.{pg + 1}  —  {label}")
                act.triggered.connect(
                    lambda _c=False, t=tab, p=pg: t.goto_page(p))
            m.addSeparator()
            clr = m.addAction("Remove all bookmarks")
            clr.triggered.connect(
                lambda _c=False, t=tab: app_state().set_bookmarks(
                    t.doc_path, []))

    def _add_bookmark(self, tab, page_idx):
        label, ok = QInputDialog.getText(
            self, "Add bookmark",
            f"Label for page {page_idx + 1}:",
            text=f"Page {page_idx + 1}")
        if not ok:
            return
        marks = app_state().bookmarks(tab.doc_path)
        marks = [b for b in marks
                 if int(b.get("page", -1)) != page_idx]
        marks.append({"page": page_idx,
                      "label": label.strip() or f"Page {page_idx + 1}"})
        app_state().set_bookmarks(tab.doc_path, marks)

    # -- Display / reading tint ------------------------------------------
    def _populate_display_menu(self):
        m = self._display_menu
        m.clear()
        cur = app_state().pref("display_tint", "none")
        for key, text in (("none", "Normal"),
                          ("night", "Night (invert)"),
                          ("sepia", "Sepia"),
                          ("warm", "Warm"),
                          ("dim", "Dim")):
            act = m.addAction(text)
            act.setCheckable(True)
            act.setChecked(key == cur)
            act.triggered.connect(
                lambda _c=False, k=key: self._set_tint(k))

    def _set_tint(self, mode):
        app_state().set_pref("display_tint", mode)
        for i in range(self._tabs.count()):
            t = self._tabs.widget(i)
            if isinstance(t, DocumentTab):
                t.set_tint(mode)

    # -- Thumbnails sidebar ----------------------------------------------
    def _toggle_thumbs(self, checked: bool):
        self._thumb_panel.setVisible(checked)
        if checked:
            self._refresh_thumbs()
        else:
            # Free the render thread + pixmaps when hidden.
            self._thumb_panel.stop()
            self._thumb_panel.clear()

    def _refresh_thumbs(self):
        """Point the thumbnail panel at the active document.  No-op when
        the panel is hidden, so it costs nothing until the user asks."""
        if not self._thumb_panel.isVisible():
            return
        tab = self._current_tab()
        if tab is None or not tab.doc_path or not tab.num_pages:
            self._thumb_panel.clear()
            return
        try:
            w = tab._page_widths.get(0, 0.0)
            h = tab._page_heights.get(0, 0.0)
            aspect = (w / h) if (w and h) else 0.75
        except Exception:
            aspect = 0.75
        self._thumb_panel.set_document(
            tab.doc_path, tab.num_pages, aspect)
        self._thumb_panel.highlight(tab.current_page_index())

    def _on_thumb_page_clicked(self, idx: int):
        tab = self._current_tab()
        if tab is not None:
            tab.goto_page(idx)
            self._thumb_panel.highlight(idx)

    def _on_note_clicked(self, page_idx: int):
        """Jump the main document to the page of the clicked note."""
        tab = self._current_tab()
        if tab is not None:
            tab.goto_page(page_idx)

    def _on_add_note_requested(self):
        """Add Note button in the Notes panel: prompt for text and drop
        a sticky note on the current page (no click needed).  It shows
        in the list immediately and is written on Save As."""
        tab = self._current_tab()
        if tab is None or not tab.doc_path:
            QMessageBox.information(
                self, "Add Note", "Open a PDF first.")
            return
        text, ok = QInputDialog.getMultiLineText(
            self, "Add Note", "Note text:", "")
        if not ok or not text.strip():
            return
        page_idx = tab.add_note_on_current_page(text)
        if page_idx < 0:
            return
        tab.goto_page(page_idx)
        # Refresh the panel so the new note appears right away.
        self._notes_panel.clear()
        self._update_notes_panel(tab)
        self._sync_toolbar()

    # -- Reading layout ---------------------------------------------------
    def _populate_layout_menu(self):
        m = self._layout_menu
        m.clear()
        cur = app_state().pref("reading_mode", "continuous")
        for key, text in (("continuous", "Continuous scroll"),
                          ("single", "Single page"),
                          ("spread", "Two-page spread")):
            act = m.addAction(text)
            act.setCheckable(True)
            act.setChecked(key == cur)
            act.triggered.connect(
                lambda _c=False, k=key: self._set_reading_mode(k))

    def _set_reading_mode(self, mode: str):
        if mode not in ("continuous", "single", "spread"):
            mode = "continuous"
        prev = app_state().pref("reading_mode", "continuous")
        if mode == prev:
            return
        tab = self._current_tab()
        # Toggling spread changes the actual page layout, so the
        # cleanest + safest path is a full document reload — every
        # overlay subsystem rebuilds through the normal (now X-origin
        # aware) code paths instead of being repositioned in place.
        need_reload = ("spread" in (mode, prev)
                       and tab is not None and tab.doc_path)
        if (need_reload and getattr(tab, "_edit_records", None)):
            ans = QMessageBox.question(
                self, "Change layout?",
                "Switching to/from two-page spread reloads the "
                "document and will discard unsaved edits in this "
                "tab.\n\nContinue?",
                QMessageBox.StandardButton.Yes
                | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No)
            if ans != QMessageBox.StandardButton.Yes:
                return
        app_state().set_pref("reading_mode", mode)
        if tab is None:
            return
        tab._reading_mode = mode
        if need_reload:
            page = tab.current_page_index()
            # The reload intentionally drops pending edits; clear the
            # crash-recovery journal too, otherwise load_pdf's recovery
            # prompt would offer to restore what we just discarded.
            try:
                tab._clear_recovery()
            except Exception:
                pass
            tab.load_pdf(tab.doc_path)
            QTimer.singleShot(
                160,
                lambda t=tab, p=page: (t.goto_page(p),
                                       self._apply_reading_mode(t)))
        else:
            self._apply_reading_mode(tab)

    def _apply_reading_mode(self, tab):
        """Single = the proven fit-one-page engine; continuous & spread
        = normal scrolling (spread's side-by-side layout is built in
        load_pdf and flows through _pt_to_scene, so OCR / search / edit
        overlays need no special-casing)."""
        mode = app_state().pref("reading_mode", "continuous")
        tab._reading_mode = mode
        single = (mode == "single")
        try:
            tab.set_fit_mode(single)
        except Exception:
            pass
        if tab is self._current_tab():
            self._act_fit.blockSignals(True)
            self._act_fit.setChecked(single)
            self._act_fit.blockSignals(False)

    # -- Session restore --------------------------------------------------
    def restore_last_session(self):
        """Reopen the documents (and page/zoom) from the previous run.
        Called deferred at startup only when no PDF was passed on the
        command line, so normal startup stays instant."""
        try:
            sess = app_state().session()
        except Exception:
            return
        for entry in sess:
            path = entry.get("path")
            if not path or not os.path.isfile(path):
                continue
            self.open_pdf_in_new_tab(path)
            tab = self._current_tab()
            if tab is None:
                continue
            page = int(entry.get("page", 0) or 0)
            zoom = entry.get("zoom")
            # Defer the jump so the async layout/render settles first.
            QTimer.singleShot(
                150, lambda t=tab, p=page, z=zoom: t.goto_page(p, z))

    def reopen_crashed_docs(self):
        """Open only the PDF(s) that still had unsaved edits when the
        app last exited abnormally (crash), so the user can recover
        them.  A normal launch otherwise starts with no document open.
        Opening each doc triggers load_pdf's own recovery prompt."""
        try:
            base = os.path.join(
                os.environ.get("LOCALAPPDATA", tempfile.gettempdir()),
                "BoltPDF", "recovery")
            if not os.path.isdir(base):
                return
            seen, opened = set(), 0
            for name in sorted(os.listdir(base)):
                if not name.endswith(".json"):
                    continue
                try:
                    with open(os.path.join(base, name), "r",
                              encoding="utf-8") as fh:
                        payload = json.load(fh)
                except (OSError, ValueError):
                    continue
                doc = payload.get("doc")
                recs = payload.get("records") or []
                if not doc or not recs or not os.path.isfile(doc):
                    continue
                key = os.path.normcase(os.path.abspath(doc))
                if key in seen:
                    continue
                seen.add(key)
                self.open_pdf_in_new_tab(doc)
                opened += 1
                if opened >= 5:        # safety cap
                    break
        except Exception:
            pass

    # -- Toolbar action delegates -----------------------------------------
    def _zoom_in(self):
        tab = self._current_tab()
        if tab:
            tab.zoom_in()

    def _zoom_out(self):
        tab = self._current_tab()
        if tab:
            tab.zoom_out()

    def _on_zoom_changed(self, scale):
        # Only update label if the signal came from the active tab
        tab = self._current_tab()
        if tab and self.sender() is tab._view:
            self._zoom_label.setText(f"{int(scale * 100)}%")

    def _toggle_fullscreen(self):
        if self._is_fullscreen:
            self.showNormal()
            # Restore edit toolbar if it was visible before fullscreen
            if self._edit_tb_visible_before_fs:
                self._edit_toolbar.setVisible(True)
            # Restore notes panel if it was visible before fullscreen
            if self._notes_visible_before_fs:
                self._notes_panel.setVisible(True)
        else:
            # Remember whether panels were showing
            self._edit_tb_visible_before_fs = self._edit_toolbar.isVisible()
            self._edit_toolbar.setVisible(False)
            self._notes_visible_before_fs = self._notes_panel.isVisible()
            self._notes_panel.setVisible(False)
            self.showFullScreen()
        self._is_fullscreen = not self._is_fullscreen

    def _toggle_fit(self, checked: bool):
        tab = self._current_tab()
        if tab:
            tab.set_fit_mode(checked)

    def _toggle_ocr(self):
        tab = self._current_tab()
        if tab:
            tab.toggle_ocr()

    def _copy_selected_text(self):
        tab = self._current_tab()
        if tab:
            tab.copy_selected_text()

    def _select_all_ocr(self):
        tab = self._current_tab()
        if tab:
            tab.select_all_ocr()

    def _combine_pdfs(self):
        """Open a file dialog, select multiple PDFs, merge them in name
        order into a single file using all CPU cores."""
        files, _ = QFileDialog.getOpenFileNames(
            self, "Select PDFs to Combine", "",
            "PDF Files (*.pdf);;All Files (*)")
        if not files or len(files) < 2:
            if files and len(files) == 1:
                QMessageBox.information(
                    self, "Combine",
                    "Please select at least 2 PDF files to combine.")
            return

        # Sort by filename (case-insensitive)
        files.sort(key=lambda p: os.path.basename(p).lower())

        # Ask where to save
        default_name = "Combined.pdf"
        save_path, _ = QFileDialog.getSaveFileName(
            self, "Save Combined PDF",
            os.path.join(os.path.dirname(files[0]), default_name),
            "PDF Files (*.pdf)")
        if not save_path:
            return

        # Progress dialog
        self._combine_progress = QProgressDialog(
            f"Combining {len(files)} PDFs...", "Cancel", 0, 100, self)
        self._combine_progress.setWindowTitle("Combine PDFs")
        self._combine_progress.setMinimumDuration(0)
        self._combine_progress.setValue(0)

        # Worker thread
        self._combine_worker = CombineWorker(files, save_path, parent=self)
        self._combine_worker.progress.connect(self._combine_progress.setValue)
        self._combine_worker.finished_ok.connect(self._on_combine_done)
        self._combine_worker.error_occurred.connect(self._on_combine_error)
        self._combine_progress.canceled.connect(self._combine_worker.terminate)
        self._combine_worker.start()

    def _on_combine_done(self, output_path):
        self._combine_progress.close()
        reply = QMessageBox.question(
            self, "Combine Complete",
            f"Combined PDF saved to:\n{os.path.basename(output_path)}\n\n"
            "Open it now?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.open_pdf_in_new_tab(output_path)

    def _on_combine_error(self, msg):
        self._combine_progress.close()
        QMessageBox.critical(self, "Combine Error", f"Failed to combine:\n{msg}")

    def _export_word(self):
        """Open the Export to Word dialog and start the export."""
        tab = self._current_tab()
        if not tab or not tab.doc_path:
            return

        current_page = tab.get_visible_page()
        dlg = ExportWordDialog(tab.num_pages, current_page, self)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        result = dlg.get_page_indices()
        if isinstance(result, str):
            QMessageBox.warning(self, "Export to Word", result)
            return

        page_indices = result

        # Ask where to save
        base_name = os.path.splitext(os.path.basename(tab.doc_path))[0]
        default_name = f"{base_name}.docx"
        save_path, _ = QFileDialog.getSaveFileName(
            self, "Save Word Document",
            os.path.join(os.path.dirname(tab.doc_path), default_name),
            "Word Documents (*.docx)")
        if not save_path:
            return

        # Progress dialog
        self._word_progress = QProgressDialog(
            f"Exporting {len(page_indices)} page(s) to Word...",
            "Cancel", 0, 100, self)
        self._word_progress.setWindowTitle("Export to Word")
        self._word_progress.setMinimumDuration(0)
        self._word_progress.setValue(0)

        # Worker thread
        self._word_worker = ExportWordWorker(
            tab.doc_path, page_indices, save_path, parent=self)
        self._word_worker.progress.connect(self._word_progress.setValue)
        self._word_worker.finished_ok.connect(self._on_word_export_done)
        self._word_worker.error_occurred.connect(self._on_word_export_error)
        self._word_progress.canceled.connect(self._word_worker.terminate)
        self._word_worker.start()

    def _on_word_export_done(self, output_path):
        self._word_progress.close()
        QMessageBox.information(
            self, "Export Complete",
            f"Word document saved to:\n{os.path.basename(output_path)}")

    def _on_word_export_error(self, msg):
        self._word_progress.close()
        QMessageBox.critical(
            self, "Export Error", f"Failed to export:\n{msg}")

    def _rebuild_as_images(self):
        tab = self._current_tab()
        if tab:
            self._act_rebuild.setEnabled(False)
            tab.rebuild_as_images(self)

    def _toggle_image_select(self):
        tab = self._current_tab()
        if not tab:
            return
        if tab.image_select_mode:
            tab.exit_image_select_mode()
        else:
            tab.enter_image_select_mode()

    def _export_selected_images(self):
        tab = self._current_tab()
        if tab:
            tab.export_selected_images(self)

    # -- Export Pages as JPEG ------------------------------------------------
    def _export_pages(self):
        """Open the Export Pages dialog and render selected pages as JPEGs."""
        tab = self._current_tab()
        if not tab or not tab.doc_path:
            return

        current_page = tab.get_visible_page()
        dlg = ExportPagesDialog(tab.num_pages, current_page, self)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        result = dlg.get_page_indices()
        if isinstance(result, str):
            QMessageBox.warning(self, "Export Pages", result)
            return

        page_indices = result

        # Build the output folder next to the PDF:
        #   <pdf_dir>/<pdf_stem>_pages export/
        pdf_dir = os.path.dirname(tab.doc_path)
        stem = os.path.splitext(os.path.basename(tab.doc_path))[0]
        out_folder = os.path.join(pdf_dir, f"{stem}_pages export")

        self._pages_progress = QProgressDialog(
            f"Exporting {len(page_indices)} page(s) as JPEG...",
            "Cancel", 0, 100, self)
        self._pages_progress.setWindowTitle("Export Pages")
        self._pages_progress.setMinimumDuration(0)
        self._pages_progress.setValue(0)

        self._pages_worker = ExportPagesWorker(
            tab.doc_path, page_indices, out_folder, stem,
            dpi=200, parent=self)
        self._pages_worker.progress.connect(self._pages_progress.setValue)
        self._pages_worker.finished_ok.connect(self._on_pages_export_done)
        self._pages_worker.error_occurred.connect(
            self._on_pages_export_error)
        self._pages_progress.canceled.connect(self._pages_worker.terminate)
        self._pages_worker.start()

    def _on_pages_export_done(self, folder_path):
        self._pages_progress.close()
        QMessageBox.information(
            self, "Export Complete",
            f"Pages exported to:\n{folder_path}")

    def _on_pages_export_error(self, msg):
        self._pages_progress.close()
        QMessageBox.critical(
            self, "Export Error", f"Failed to export pages:\n{msg}")

    # -- Edit Mode ---------------------------------------------------------
    def _toggle_edit_mode(self):
        tab = self._current_tab()
        if not tab:
            return
        if tab.edit_mode:
            # Exiting with pending edits: keep them in this session
            # (preview stays visible, Save As still available) unless
            # the user explicitly discards.
            dirty = bool(getattr(tab, '_edit_records', None))
            if dirty:
                mb = QMessageBox(self)
                mb.setWindowTitle("Exit edit mode")
                mb.setIcon(QMessageBox.Icon.Question)
                mb.setText(
                    "Keep your edits for this session?\n\n"
                    "Keep — leave edit mode but keep the changes "
                    "visible (you can still Save As later; they're "
                    "dropped only when the file is closed).\n"
                    "Discard — throw the edits away now.")
                keep_btn = mb.addButton(
                    "Keep", QMessageBox.ButtonRole.AcceptRole)
                disc_btn = mb.addButton(
                    "Discard", QMessageBox.ButtonRole.DestructiveRole)
                cancel_btn = mb.addButton(
                    "Cancel", QMessageBox.ButtonRole.RejectRole)
                mb.setDefaultButton(keep_btn)
                mb.exec()
                clicked = mb.clickedButton()
                if clicked is cancel_btn or clicked is None:
                    self._act_edit.setChecked(True)
                    return
                tab.exit_edit_mode(
                    discard_edits=(clicked is disc_btn))
            else:
                tab.exit_edit_mode(discard_edits=False)
        else:
            # Show a one-time beta notice the first time the user
            # enters edit mode in this session.
            if not self._edit_beta_shown:
                self._edit_beta_shown = True
                msg = QMessageBox(self)
                msg.setWindowTitle("Edit Tools — Beta")
                msg.setIcon(QMessageBox.Icon.Information)
                msg.setText(
                    "The editing tools are currently in <b>beta</b>.\n\n"
                    "They work well for most documents, but you may "
                    "encounter occasional quirks with complex layouts "
                    "or unusual fonts.\n\n"
                    "We're actively improving them and updates will "
                    "be included in future releases. Your feedback is "
                    "welcome!")
                msg.setStandardButtons(QMessageBox.StandardButton.Ok)
                msg.button(QMessageBox.StandardButton.Ok).setText(
                    "Got it")
                msg.exec()
            # Exiting other mutually-exclusive modes
            if tab.image_select_mode:
                tab.exit_image_select_mode()
            tab.enter_edit_mode()
        self._sync_toolbar()

    def _edit_add_text(self):
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        # Toggle: clicking an active sub-action disarms it
        if tab.edit_action == 'add_text':
            tab.set_edit_action(None)
        else:
            tab.set_edit_action('add_text')
        self._sync_toolbar()

    def _edit_transform(self):
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if tab.edit_action == 'transform':
            tab.set_edit_action(None)
        else:
            tab.set_edit_action('transform')
        self._sync_toolbar()

    def _edit_undo(self):
        tab = self._current_tab()
        if tab and tab.edit_mode:
            tab.edit_undo()
            self._sync_toolbar()

    def _edit_redo(self):
        tab = self._current_tab()
        if tab and tab.edit_mode:
            tab.edit_redo()
            self._sync_toolbar()

    def _toggle_props(self, checked: bool):
        self._props_panel.setVisible(checked)
        if checked:
            self._props_panel.refresh(self._current_tab())

    def _edit_duplicate(self):
        """Duplicate the Transform-selected object (Ctrl+D)."""
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if not tab.duplicate_selected_edit():
            QMessageBox.information(
                self, "Duplicate",
                "Select an object with the Transform tool first, "
                "then Duplicate (Ctrl+D).")
            return
        self._sync_toolbar()

    def _edit_multiselect(self):
        """Arm the rubber-band multi-select tool."""
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if tab.edit_action == 'multiselect':
            tab.set_edit_action(None)
            tab.clear_multi_selection()
        else:
            tab.set_edit_action('multiselect')
        self._sync_toolbar()

    def _edit_toggle_snap(self, checked: bool):
        tab = self._current_tab()
        if tab is not None:
            tab._snap_enabled = bool(checked)
            if not checked:
                tab._clear_snap_guides()

    def _edit_align(self, mode: str):
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if not tab.align_multi(mode):
            QMessageBox.information(
                self, "Align",
                "Use the Select tool to rubber-band at least two "
                "objects first, then choose an Align option.")
            return
        self._sync_toolbar()

    def _edit_distribute(self, axis: str):
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if not tab.distribute_multi(axis):
            QMessageBox.information(
                self, "Distribute",
                "Select at least three objects with the Select tool "
                "first, then choose a Distribute option.")
            return
        self._sync_toolbar()

    def _edit_group(self):
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if not tab.group_multi():
            QMessageBox.information(
                self, "Group",
                "Select at least two objects with the Select tool "
                "first, then Group.")
            return
        self._sync_toolbar()

    def _edit_ungroup(self):
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if not tab.ungroup_multi():
            QMessageBox.information(
                self, "Ungroup",
                "Select a grouped object (Transform-click it, or "
                "rubber-band its members) first, then Ungroup.")
            return
        self._sync_toolbar()

    def _edit_delete_selection(self):
        """Delete the text box or image currently held by the
        Transform tool's active selection."""
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if tab._active_transform is None:
            QMessageBox.information(
                self, "Delete",
                "Select a text box or image first with the Transform "
                "tool, then press Delete to remove it.")
            return
        tab.delete_active_selection()
        self._sync_toolbar()

    def _edit_add_image(self):
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if tab.edit_action == 'add_image':
            tab.set_edit_action(None)
            self._sync_toolbar()
            return
        # Prompt for an image file up front so we have it ready when
        # the user clicks on a page.
        path, _ = QFileDialog.getOpenFileName(
            self, "Choose image to add", "",
            "Images (*.png *.jpg *.jpeg *.bmp *.gif *.webp)")
        if not path:
            tab.set_edit_action(None)
            self._sync_toolbar()
            return
        tab._pending_add_image_path = path
        tab.set_edit_action('add_image')
        self._sync_toolbar()

    def _edit_replace_image(self):
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        # Hand off to the tab — it puts the viewer into image-select
        # mode so the user can click one of the real images already
        # in the PDF, then prompts for a replacement file.
        tab.begin_image_replace_pick()
        self._sync_toolbar()

    def _edit_stamp(self):
        """Show a menu of preset stamps, then arm the stamp tool."""
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        from PyQt6.QtWidgets import QMenu
        menu = QMenu("Stamp", self)
        presets = [
            ("APPROVED", (0, 150, 0)),
            ("DRAFT", (180, 0, 0)),
            ("CONFIDENTIAL", (200, 0, 0)),
            ("REVISED", (0, 0, 200)),
            ("COPY", (100, 100, 100)),
            ("FINAL", (0, 100, 0)),
        ]
        for text, color in presets:
            act = menu.addAction(text)
            act.setData((text, color))
        # Saved custom-stamp library (persisted across sessions).
        library = app_state().pref("stamp_library", []) or []
        if library:
            menu.addSeparator()
            for entry in library:
                try:
                    t = entry.get("text", "")
                    c = tuple(entry.get("color", (255, 0, 0)))
                except Exception:
                    continue
                if t:
                    menu.addAction(t).setData((t, c))
        menu.addSeparator()
        menu.addAction("Date / Time stamp").setData(
            ("__date__", (60, 60, 60)))
        menu.addAction("Custom...").setData(
            ("__custom__", (255, 0, 0)))
        menu.addAction("Save text to library...").setData(
            ("__savelib__", None))
        menu.addAction("Bates / page numbering...").setData(
            ("__bates__", None))
        chosen = menu.exec(self.cursor().pos())
        if not chosen or chosen.data() is None:
            return
        text, color = chosen.data()
        if text == "__bates__":
            self._bates_number(tab)
            return
        if text == "__savelib__":
            t, ok = QInputDialog.getText(
                self, "Save stamp", "Stamp text:")
            if not ok or not t:
                return
            qc = QColorDialog.getColor(
                QColor(255, 0, 0), self, "Stamp colour")
            c = ((qc.red(), qc.green(), qc.blue())
                 if qc.isValid() else (255, 0, 0))
            lib = list(app_state().pref("stamp_library", []) or [])
            lib.append({"text": t.upper(), "color": list(c)})
            app_state().set_pref("stamp_library", lib)
            return
        if text == "__date__":
            text = time.strftime("%Y-%m-%d %H:%M")
        elif text == "__custom__":
            text, ok = QInputDialog.getText(
                self, "Custom Stamp", "Stamp text:")
            if not ok or not text:
                return
            qc = QColorDialog.getColor(
                QColor(255, 0, 0), self, "Stamp Colour")
            if qc.isValid():
                color = (qc.red(), qc.green(), qc.blue())
            text = text.upper()
        else:
            text = text.upper()
        tab._pending_stamp_text = text
        tab._pending_stamp_color = color
        tab.set_edit_action('stamp')
        self._sync_toolbar()

    def _bates_number(self, tab):
        """Stamp sequential Bates/page numbers in the bottom-centre of
        every page (generates one undoable stamp record per page, so it
        round-trips through Save As like any other stamp)."""
        prefix, ok = QInputDialog.getText(
            self, "Bates / page numbering",
            "Prefix (e.g. 'ABC-' or blank for plain page numbers):")
        if not ok:
            return
        start, ok = QInputDialog.getInt(
            self, "Bates / page numbering",
            "Start number:", 1, 0, 1_000_000)
        if not ok:
            return
        digits, ok = QInputDialog.getInt(
            self, "Bates / page numbering",
            "Zero-pad to how many digits? (0 = none):", 4, 0, 12)
        if not ok:
            return
        scale = tab._render_scale or 1.0
        made = 0
        for i in range(tab.num_pages):
            num = start + i
            label = f"{prefix}{str(num).zfill(digits) if digits else num}"
            pw = tab._page_widths.get(i, 0.0) / scale
            ph = tab._page_heights.get(i, 0.0) / scale
            if pw <= 0 or ph <= 0:
                continue
            box_w, box_h = 200.0, 24.0
            x0 = (pw - box_w) / 2.0
            y0 = ph - box_h - 12.0
            rec = EditRecord(
                kind="stamp_add", page_idx=i,
                new_rect=(x0, y0, x0 + box_w, y0 + box_h),
                text=label, font_size=12.0, color=(60, 60, 60),
                rotation=0.0)
            tab._edit_records.append(rec)
            tab._materialize_record(rec)
            made += 1
        if made:
            tab._edit_checkpoint()
            tab.status_changed.emit()
            QMessageBox.information(
                self, "Bates / page numbering",
                f"Added numbering to {made} page(s). Use Save As to "
                "write them into the PDF.")

    def _edit_shape(self, shape_action: str):
        """Arm one of the shape drawing tools."""
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if tab.edit_action == shape_action:
            tab.set_edit_action(None)
        else:
            tab.set_edit_action(shape_action)
        self._sync_toolbar()

    def _edit_redact(self):
        """Arm the whiteout / redact tool."""
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if tab.edit_action == 'redact':
            tab.set_edit_action(None)
        else:
            tab.set_edit_action('redact')
        self._sync_toolbar()

    def _edit_highlight(self):
        """Arm the highlight tool (saves as a real PDF highlight)."""
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if tab.edit_action == 'highlight':
            tab.set_edit_action(None)
        else:
            tab.set_edit_action('highlight')
        self._sync_toolbar()

    def _edit_note(self):
        """Arm the sticky note tool."""
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        if tab.edit_action == 'note':
            tab.set_edit_action(None)
        else:
            tab.set_edit_action('note')
        self._sync_toolbar()

    def _edit_save_as(self):
        tab = self._current_tab()
        if not tab or not tab.edit_mode:
            return
        tab.save_edited_pdf_as(self)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        tab = self._current_tab()
        if tab:
            tab.refit_if_active()

    # -- Drag & Drop (global — works on every widget) -----------------------
    def _install_drop_filter(self, widget):
        """Recursively enable drops and install this window as event filter
        on every child widget so PDF drops work anywhere."""
        from PyQt6.QtCore import QEvent
        widget.setAcceptDrops(True)
        widget.installEventFilter(self)
        for child in widget.findChildren(QWidget):
            child.setAcceptDrops(True)
            child.installEventFilter(self)

    def eventFilter(self, obj, event):
        from PyQt6.QtCore import QEvent
        if event.type() == QEvent.Type.DragEnter:
            if event.mimeData().hasUrls():
                for url in event.mimeData().urls():
                    if url.toLocalFile().lower().endswith(".pdf"):
                        event.acceptProposedAction()
                        return True
        elif event.type() == QEvent.Type.DragMove:
            if event.mimeData().hasUrls():
                event.acceptProposedAction()
                return True
        elif event.type() == QEvent.Type.Drop:
            paths = []
            for url in event.mimeData().urls():
                path = url.toLocalFile()
                if path.lower().endswith(".pdf"):
                    paths.append(path)
            if paths:
                for path in paths:
                    self.open_pdf_in_new_tab(path)
                return True
        elif event.type() == QEvent.Type.ChildAdded:
            # New child widgets (e.g. new tabs) also need the filter
            child = event.child()
            if isinstance(child, QWidget):
                child.setAcceptDrops(True)
                child.installEventFilter(self)
        return super().eventFilter(obj, event)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                if url.toLocalFile().lower().endswith(".pdf"):
                    event.acceptProposedAction()
                    return
        event.ignore()

    def dropEvent(self, event):
        paths = []
        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if path.lower().endswith(".pdf"):
                paths.append(path)
        for path in paths:
            self.open_pdf_in_new_tab(path)

    # -- Keyboard ---------------------------------------------------------
    def keyPressEvent(self, event):
        # Ctrl+F → open / focus the search bar
        if (event.key() == Qt.Key.Key_F
                and event.modifiers() & Qt.KeyboardModifier.ControlModifier):
            self._search_bar.open_bar()
            event.accept()
            return

        # Ctrl+Z / Ctrl+Shift+Z / Ctrl+Y → undo / redo edit-mode changes
        if (event.key() in (Qt.Key.Key_Z, Qt.Key.Key_Y)
                and event.modifiers()
                & Qt.KeyboardModifier.ControlModifier):
            tab = self._current_tab()
            if tab and tab.edit_mode:
                shift = bool(event.modifiers()
                             & Qt.KeyboardModifier.ShiftModifier)
                if event.key() == Qt.Key.Key_Y or (
                        event.key() == Qt.Key.Key_Z and shift):
                    tab.edit_redo()
                else:
                    tab.edit_undo()
                self._sync_toolbar()
                event.accept()
                return

        # Ctrl+C / Ctrl+V / Ctrl+D → copy / paste / duplicate the
        # Transform-selected edit object (works across pages & tabs).
        if (event.key() in (Qt.Key.Key_C, Qt.Key.Key_V, Qt.Key.Key_D)
                and event.modifiers()
                & Qt.KeyboardModifier.ControlModifier
                and not (event.modifiers()
                         & Qt.KeyboardModifier.ShiftModifier)):
            tab = self._current_tab()
            if tab and tab.edit_mode:
                k = event.key()
                handled = False
                if k == Qt.Key.Key_C:
                    handled = tab.copy_selected_edit()
                elif k == Qt.Key.Key_V:
                    handled = tab.paste_edit()
                else:
                    handled = tab.duplicate_selected_edit()
                if handled:
                    self._sync_toolbar()
                    event.accept()
                    return

        if event.key() == Qt.Key.Key_Escape and self._is_fullscreen:
            self._toggle_fullscreen()
            return

        # Multi-select (E2): arrow-nudge / Delete the rubber-band group.
        tab = self._current_tab()
        if (tab and tab.edit_mode and tab.has_multi_selection()
                and tab._inline_editor is None):
            step = (10.0 if (event.modifiers()
                             & Qt.KeyboardModifier.ShiftModifier)
                    else 2.0)
            k = event.key()
            if k == Qt.Key.Key_Left:
                tab.nudge_multi(-step, 0)
                event.accept()
                return
            if k == Qt.Key.Key_Right:
                tab.nudge_multi(step, 0)
                event.accept()
                return
            if k == Qt.Key.Key_Up:
                tab.nudge_multi(0, -step)
                event.accept()
                return
            if k == Qt.Key.Key_Down:
                tab.nudge_multi(0, step)
                event.accept()
                return
            if k in (Qt.Key.Key_Delete, Qt.Key.Key_Backspace):
                if tab.delete_multi():
                    self._sync_toolbar()
                    event.accept()
                    return

        # Delete / Backspace while a Transform selection is active →
        # remove the selected text box or image from the document.
        if event.key() in (Qt.Key.Key_Delete, Qt.Key.Key_Backspace):
            tab = self._current_tab()
            if (tab and tab.edit_mode
                    and tab._active_transform is not None
                    and tab._inline_editor is None):
                if tab.delete_active_selection():
                    self._sync_toolbar()
                    event.accept()
                    return

        tab = self._current_tab()
        if tab and tab.view._fit_mode:
            if event.key() in (Qt.Key.Key_PageDown, Qt.Key.Key_Down,
                               Qt.Key.Key_Space, Qt.Key.Key_Right):
                tab.view.fit_step(1)
                return
            elif event.key() in (Qt.Key.Key_PageUp, Qt.Key.Key_Up,
                                 Qt.Key.Key_Left):
                tab.view.fit_step(-1)
                return
            elif event.key() == Qt.Key.Key_Home:
                tab.view.fit_go_to_page(0)
                return
            elif event.key() == Qt.Key.Key_End:
                tab.view.fit_go_to_page(tab.num_pages - 1)
                return

        super().keyPressEvent(event)

    # -- Stylesheet -------------------------------------------------------
    @staticmethod
    def _stylesheet():
        return """
            QMainWindow { background: #1e1e1e; }
            QToolBar {
                background: #2d2d2d; border: none;
                padding: 4px 8px; spacing: 6px;
            }
            QToolBar QToolButton {
                color: #e0e0e0; background: transparent;
                border: 1px solid transparent; border-radius: 4px;
                padding: 4px 10px; font-size: 13px;
            }
            QToolBar QToolButton:hover {
                background: #3d3d3d; border-color: #555;
            }
            QToolBar QToolButton:pressed { background: #4a4a4a; }
            QLabel { color: #b0b0b0; font-size: 12px; }
            QGraphicsView { border: none; }
            QTabWidget::pane { border: none; }
            QTabBar {
                background: #2d2d2d; border: none;
            }
            QTabBar::tab {
                background: #353535; color: #b0b0b0;
                padding: 6px 16px; margin-right: 2px;
                border: 1px solid #444; border-bottom: none;
                border-top-left-radius: 4px; border-top-right-radius: 4px;
                font-size: 12px;
            }
            QTabBar::tab:selected {
                background: #1e1e1e; color: #e0e0e0;
                border-color: #555;
            }
            QTabBar::tab:hover:!selected {
                background: #3d3d3d;
            }
            QScrollBar:vertical {
                background: #2d2d2d; width: 10px; border: none;
            }
            QScrollBar::handle:vertical {
                background: #555; border-radius: 5px; min-height: 30px;
            }
            QScrollBar::handle:vertical:hover { background: #777; }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0;
            }
            QScrollBar:horizontal {
                background: #2d2d2d; height: 10px; border: none;
            }
            QScrollBar::handle:horizontal {
                background: #555; border-radius: 5px; min-width: 30px;
            }
            QScrollBar::handle:horizontal:hover { background: #777; }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
                width: 0;
            }
        """


# ---------------------------------------------------------------------------
# Installation & default-reader registration (Windows only, frozen exe)
# ---------------------------------------------------------------------------
_PROG_ID = "BoltPDF.Document"
_APP_NAME = "BoltPDF"
_INSTALL_DIR_NAME = "BoltPDF"
_UNINSTALL_REG = rf"Software\Microsoft\Windows\CurrentVersion\Uninstall\{_APP_NAME}"


def _get_install_dir() -> str:
    """Return the target install directory under C:\\Program Files."""
    return os.path.join(os.environ.get("PROGRAMFILES", r"C:\Program Files"), _INSTALL_DIR_NAME)


def _get_installed_exe() -> str:
    """Return the full path to the installed exe."""
    return os.path.join(_get_install_dir(), "BoltPDF.exe")


def _is_installed() -> bool:
    """Return True if BoltPDF is installed in its permanent location."""
    return os.path.isfile(_get_installed_exe())


def _is_running_from_install_dir() -> bool:
    """Return True if the current exe is the installed copy."""
    if not getattr(sys, 'frozen', False):
        return False
    try:
        return (os.path.normcase(os.path.abspath(sys.executable))
                == os.path.normcase(os.path.abspath(_get_installed_exe())))
    except Exception:
        return False


def _is_admin() -> bool:
    """Return True if running with admin privileges."""
    try:
        import ctypes
        return ctypes.windll.shell32.IsUserAnAdmin() != 0
    except Exception:
        return False


def _elevate_and_install():
    """Re-launch this exe with admin rights for the install step."""
    try:
        import ctypes
        ret = ctypes.windll.shell32.ShellExecuteW(
            None, "runas", sys.executable,
            "--install " + " ".join(
                f'"{a}"' for a in sys.argv[1:] if a.lower().endswith(".pdf")),
            None, 1)
        if ret > 32:
            sys.exit(0)  # elevation launched OK, exit this non-admin copy
    except Exception:
        pass


def _install_app() -> bool:
    """Copy the entire onedir bundle to C:\\Program Files\\BoltPDF,
    create Start Menu and Desktop shortcuts, and register with
    Apps & Features.  Requires admin rights — will request UAC
    elevation if needed.

    Because BoltPDF is built with PyInstaller in onedir mode, the
    exe has a lot of sibling files (Qt DLLs, Python runtime, the
    _internal/ subfolder, ocr_helper.ps1, etc.).  Installation
    means copying the whole directory, not just BoltPDF.exe.

    Returns True if installation succeeded.
    """
    if sys.platform != "win32" or not getattr(sys, 'frozen', False):
        return False

    # Request elevation if we don't have admin rights
    if not _is_admin():
        _elevate_and_install()
        return False  # if elevation was declined, continue running from current location

    import shutil
    import winreg

    src_exe = os.path.abspath(sys.executable)
    src_dir = os.path.dirname(src_exe)
    install_dir = _get_install_dir()
    dst = _get_installed_exe()

    # --- Copy the whole onedir bundle ---------------------------------------
    # If we're already running from the install dir, nothing to copy.
    if os.path.normcase(src_dir) == os.path.normcase(install_dir):
        pass
    else:
        try:
            # Wipe any stale install first so we don't leave orphan DLLs
            # from an older version.
            if os.path.isdir(install_dir):
                # Skip files currently locked by a running instance.
                for name in os.listdir(install_dir):
                    p = os.path.join(install_dir, name)
                    try:
                        if os.path.isdir(p):
                            shutil.rmtree(p, ignore_errors=True)
                        else:
                            os.unlink(p)
                    except Exception:
                        pass
            os.makedirs(install_dir, exist_ok=True)
            # Copy everything from the source onedir to the target
            shutil.copytree(src_dir, install_dir, dirs_exist_ok=True)
        except Exception as e:
            print(f"[BoltPDF] Install copy failed: {e}", file=sys.stderr)
            return False

    # --- Start Menu + Desktop shortcuts -------------------------------------
    try:
        _create_shortcut(dst, _APP_NAME)
    except Exception as e:
        print(f"[BoltPDF] Shortcut creation warning: {e}", file=sys.stderr)

    # --- Apps & Features (Uninstall entry) ----------------------------------
    try:
        with winreg.CreateKey(winreg.HKEY_CURRENT_USER, _UNINSTALL_REG) as key:
            winreg.SetValueEx(key, "DisplayName", 0,
                              winreg.REG_SZ, _APP_NAME)
            winreg.SetValueEx(key, "DisplayIcon", 0,
                              winreg.REG_SZ, f'"{dst}",0')
            winreg.SetValueEx(key, "UninstallString", 0,
                              winreg.REG_SZ, f'"{dst}" --uninstall')
            winreg.SetValueEx(key, "InstallLocation", 0,
                              winreg.REG_SZ, install_dir)
            winreg.SetValueEx(key, "Publisher", 0,
                              winreg.REG_SZ, "BoltPDF")
            winreg.SetValueEx(key, "NoModify", 0, winreg.REG_DWORD, 1)
            winreg.SetValueEx(key, "NoRepair", 0, winreg.REG_DWORD, 1)
            # Estimate total install size in KB (walk the install dir)
            try:
                total = 0
                for root, _dirs, files in os.walk(install_dir):
                    for f in files:
                        try:
                            total += os.path.getsize(os.path.join(root, f))
                        except OSError:
                            pass
                winreg.SetValueEx(key, "EstimatedSize", 0,
                                  winreg.REG_DWORD, total // 1024)
            except Exception:
                pass
    except Exception as e:
        print(f"[BoltPDF] Uninstall reg warning: {e}", file=sys.stderr)

    return True


def _create_shortcut(target_exe: str, name: str):
    """Create Start Menu AND Desktop shortcuts (.lnk) using PowerShell."""
    start_menu = os.path.join(
        os.environ.get("APPDATA", ""),
        r"Microsoft\Windows\Start Menu\Programs")
    start_lnk = os.path.join(start_menu, f"{name}.lnk")

    # Desktop shortcut — use the Public Desktop so it works for all users,
    # fall back to current user's desktop
    desktop = os.path.join(os.environ.get("USERPROFILE", ""), "Desktop")
    desktop_lnk = os.path.join(desktop, f"{name}.lnk")

    working_dir = os.path.dirname(target_exe)
    ps = f"""
$ws = New-Object -ComObject WScript.Shell

$sc = $ws.CreateShortcut('{start_lnk}')
$sc.TargetPath = '{target_exe}'
$sc.WorkingDirectory = '{working_dir}'
$sc.Description = 'BoltPDF - Lightweight PDF Reader'
$sc.Save()

$dc = $ws.CreateShortcut('{desktop_lnk}')
$dc.TargetPath = '{target_exe}'
$dc.WorkingDirectory = '{working_dir}'
$dc.Description = 'BoltPDF - Lightweight PDF Reader'
$dc.Save()
"""
    with tempfile.NamedTemporaryFile(
            mode="w", suffix=".ps1", delete=False) as f:
        f.write(ps)
        ps_path = f.name
    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass",
             "-File", ps_path],
            capture_output=True, timeout=10,
            creationflags=(subprocess.CREATE_NO_WINDOW
                           if sys.platform == "win32" else 0),
        )
    finally:
        try:
            os.unlink(ps_path)
        except OSError:
            pass


def _uninstall_app():
    """Remove the installed files, shortcuts, and registry entries, then exit."""
    if sys.platform != "win32":
        sys.exit(0)

    import winreg

    install_dir = _get_install_dir()

    # Remove Start Menu shortcut
    start_menu = os.path.join(
        os.environ.get("APPDATA", ""),
        r"Microsoft\Windows\Start Menu\Programs")
    lnk = os.path.join(start_menu, f"{_APP_NAME}.lnk")
    try:
        if os.path.isfile(lnk):
            os.unlink(lnk)
    except Exception:
        pass

    # Remove Desktop shortcut
    desktop = os.path.join(os.environ.get("USERPROFILE", ""), "Desktop")
    desktop_lnk = os.path.join(desktop, f"{_APP_NAME}.lnk")
    try:
        if os.path.isfile(desktop_lnk):
            os.unlink(desktop_lnk)
    except Exception:
        pass

    # Remove registry entries
    for key_path in [
        _UNINSTALL_REG,
        rf"Software\Classes\{_PROG_ID}",
        r"Software\Classes\.pdf",
        r"Software\BoltPDF",
    ]:
        try:
            _reg_delete_tree(winreg.HKEY_CURRENT_USER, key_path)
        except Exception:
            pass

    # Remove from RegisteredApplications
    try:
        with winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                r"Software\RegisteredApplications",
                0, winreg.KEY_SET_VALUE) as key:
            try:
                winreg.DeleteValue(key, _APP_NAME)
            except FileNotFoundError:
                pass
    except Exception:
        pass

    _notify_shell()

    # Schedule deletion of install directory after exit (files are locked)
    bat = os.path.join(tempfile.gettempdir(), "boltpdf_cleanup.bat")
    try:
        with open(bat, "w") as f:
            f.write(f'@echo off\n')
            f.write(f'timeout /t 2 /nobreak >nul\n')
            f.write(f'rmdir /s /q "{install_dir}"\n')
            f.write(f'del /q "%~f0"\n')
        subprocess.Popen(
            ["cmd", "/c", bat],
            creationflags=subprocess.CREATE_NO_WINDOW | subprocess.DETACHED_PROCESS,
        )
    except Exception:
        pass

    sys.exit(0)


def _reg_delete_tree(hive, subkey):
    """Recursively delete a registry key tree."""
    import winreg
    try:
        winreg.DeleteKey(hive, subkey)
    except OSError:
        try:
            with winreg.OpenKey(hive, subkey, 0,
                                winreg.KEY_READ) as key:
                i = 0
                children = []
                while True:
                    try:
                        children.append(winreg.EnumKey(key, i))
                        i += 1
                    except OSError:
                        break
            for child in children:
                _reg_delete_tree(hive, rf"{subkey}\{child}")
            winreg.DeleteKey(hive, subkey)
        except Exception:
            pass


def _is_default_pdf_reader() -> bool | None:
    """Return True if BoltPDF is the default .pdf handler (per UserChoice),
    False if not, or None if we can't determine (non-Windows / registry error).

    Only the UserChoice key is authoritative — Windows uses it to pick the
    default file handler.  We intentionally do NOT fall back to checking
    HKCU\\Software\\Classes\\.pdf because _register_prog_id() writes our
    ProgId there during install, which would make us appear to be the
    default even before the user has actually accepted the prompt.
    """
    if sys.platform != "win32":
        return None
    try:
        import winreg
        key_path = (r"Software\Microsoft\Windows\CurrentVersion"
                    r"\Explorer\FileExts\.pdf\UserChoice")
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path) as key:
            prog_id, _ = winreg.QueryValueEx(key, "ProgId")
            return prog_id == _PROG_ID
    except FileNotFoundError:
        # UserChoice missing — not default yet
        return False
    except Exception:
        return False


def _register_prog_id():
    """Register BoltPDF's ProgId and .pdf association in the current user's
    registry.  No admin rights required.  Always points to the installed
    exe path so associations survive even if the original download is moved."""
    if sys.platform != "win32":
        return
    try:
        import winreg
        # Always use the installed location so file associations are stable
        exe_path = _get_installed_exe() if _is_installed() else sys.executable

        # ProgId
        prog_key = rf"Software\Classes\{_PROG_ID}"
        with winreg.CreateKey(winreg.HKEY_CURRENT_USER, prog_key) as key:
            winreg.SetValueEx(key, "", 0, winreg.REG_SZ, "BoltPDF Document")

        # DefaultIcon
        icon_key = rf"Software\Classes\{_PROG_ID}\DefaultIcon"
        with winreg.CreateKey(winreg.HKEY_CURRENT_USER, icon_key) as key:
            winreg.SetValueEx(key, "", 0, winreg.REG_SZ, f'"{exe_path}",0')

        # shell\open\command
        cmd_key = rf"Software\Classes\{_PROG_ID}\shell\open\command"
        with winreg.CreateKey(winreg.HKEY_CURRENT_USER, cmd_key) as key:
            winreg.SetValueEx(key, "", 0, winreg.REG_SZ,
                              f'"{exe_path}" "%1"')

        # Per-user .pdf → our ProgId (fallback when UserChoice is absent)
        with winreg.CreateKey(
                winreg.HKEY_CURRENT_USER,
                r"Software\Classes\.pdf") as key:
            winreg.SetValueEx(key, "", 0, winreg.REG_SZ, _PROG_ID)

        # OpenWithProgids
        owp_key = r"Software\Classes\.pdf\OpenWithProgids"
        with winreg.CreateKey(winreg.HKEY_CURRENT_USER, owp_key) as key:
            winreg.SetValueEx(key, _PROG_ID, 0, winreg.REG_NONE, b"")

        # RegisteredApplications
        try:
            with winreg.OpenKey(
                    winreg.HKEY_CURRENT_USER,
                    r"Software\RegisteredApplications",
                    0, winreg.KEY_SET_VALUE) as key:
                winreg.SetValueEx(key, "BoltPDF", 0, winreg.REG_SZ,
                                  r"Software\BoltPDF\Capabilities")
        except Exception:
            with winreg.CreateKey(
                    winreg.HKEY_CURRENT_USER,
                    r"Software\RegisteredApplications") as key:
                winreg.SetValueEx(key, "BoltPDF", 0, winreg.REG_SZ,
                                  r"Software\BoltPDF\Capabilities")

        # Capabilities
        cap = r"Software\BoltPDF\Capabilities"
        with winreg.CreateKey(winreg.HKEY_CURRENT_USER, cap) as key:
            winreg.SetValueEx(key, "ApplicationName", 0,
                              winreg.REG_SZ, "BoltPDF")
            winreg.SetValueEx(key, "ApplicationDescription", 0,
                              winreg.REG_SZ, "Lightweight PDF Reader")
        with winreg.CreateKey(
                winreg.HKEY_CURRENT_USER,
                rf"Software\BoltPDF\Capabilities\FileAssociations") as key:
            winreg.SetValueEx(key, ".pdf", 0, winreg.REG_SZ, _PROG_ID)

    except Exception as e:
        print(f"[BoltPDF] Registry registration warning: {e}",
              file=sys.stderr)


def _get_user_sid() -> str:
    """Return the current user's SID string via whoami."""
    try:
        r = subprocess.run(
            ["whoami", "/user", "/fo", "csv", "/nh"],
            capture_output=True, text=True, timeout=10,
            creationflags=(subprocess.CREATE_NO_WINDOW
                           if sys.platform == "win32" else 0),
        )
        # Output: "DOMAIN\user","S-1-5-21-..."
        return r.stdout.strip().split(",")[1].strip('"')
    except Exception as e:
        print(f"[BoltPDF] SID lookup failed: {e}", file=sys.stderr)
        return ""


def _get_current_filetime() -> int:
    """Return the current system time as a Windows FILETIME (int)."""
    import ctypes
    class FILETIME(ctypes.Structure):
        _fields_ = [("lo", ctypes.c_uint32), ("hi", ctypes.c_uint32)]
    ft = FILETIME()
    ctypes.windll.kernel32.GetSystemTimeAsFileTime(ctypes.byref(ft))
    return (ft.hi << 32) | ft.lo


def _compute_user_choice_hash(ext: str, sid: str, progid: str,
                              filetime: int) -> str:
    """Compute the Windows 10/11 UserChoice Hash value.

    The algorithm is reverse-engineered from the Windows shell (matching the
    tools_setfta / PS-SFTA / Mozilla Firefox implementations).  It combines
    the file extension, user SID, ProgId, a timestamp (FILETIME rounded to
    the minute), and the hard-coded "experience" string, runs them through
    MD5, then two independent scramble passes whose results are XOR'd and
    Base64-encoded.
    """
    import struct
    import base64
    import hashlib

    M = 0xFFFFFFFF
    TICKS_PER_MIN = 600_000_000
    ft_round = (filetime // TICKS_PER_MIN) * TICKS_PER_MIN

    experience = ("user choice set via windows user experience "
                  "{d18b6dd5-6124-4341-9318-804003bafa0b}")
    regdate = f"{ft_round:016x}"
    # Input: extension + SID + ProgId + timestamp-hex + experience + NUL
    input_str = f"{ext}{sid}{progid}{regdate}{experience}\0".lower()
    data = input_str.encode("utf-16-le")

    # --- MD5 of input bytes ------------------------------------------------
    md5 = hashlib.md5(data).digest()

    # --- Prepare DWORD arrays ---------------------------------------------
    num_dw = len(data) // 4
    length = num_dw if num_dw % 2 == 0 else num_dw - 1
    if length <= 1:
        return ""
    dwords = struct.unpack(f"<{length}I", data[: length * 4])
    md5_dw = struct.unpack("<4I", md5)
    n_pairs = length // 2

    # === Pass 1 (sub_1) ====================================================
    s1_c0 = ((md5_dw[0] | 1) + 0x69FB0000) & M
    s1_c1 = ((md5_dw[1] | 1) + 0x13DB0000) & M
    s1_r = 0       # running result
    s1_acc = 0     # accumulator
    idx = 0
    for _ in range(n_pairs):
        v11 = (dwords[idx] + s1_r) & M
        idx += 2
        temp = (s1_c0 * v11 - 0x10FA9605 * (v11 >> 16)) & M
        v12 = (0x79F8A395 * temp + 0x689B6B9F * (temp >> 16)) & M
        v13 = (0xEA970001 * v12 - 0x3C101569 * (v12 >> 16)) & M
        v14 = (v13 + s1_acc) & M
        t = (dwords[idx - 1] + v13) & M
        v15 = (s1_c1 * t - 0x3CE8EC25 * (t >> 16)) & M
        tmp2 = (0x59C3AF2D * v15 - 0x2232E0F1 * (v15 >> 16)) & M
        s1_r = (0x1EC90001 * tmp2 + 0x35BD1EC9 * (tmp2 >> 16)) & M
        s1_acc = (s1_r + v14) & M
    hash1 = struct.pack("<II", s1_r, s1_acc)

    # === Pass 2 (sub_2) ====================================================
    s2_k0 = (md5_dw[0] | 1) & M
    s2_k1 = (md5_dw[1] | 1) & M
    s2_c0 = (0xB1110000 * s2_k0) & M
    s2_c1 = (0x16F50000 * s2_k1) & M
    s2_r = 0       # running result (v5)
    s2_acc = 0     # accumulator (v7)
    idx = 0
    for _ in range(n_pairs):
        idx += 2
        inp1 = dwords[idx - 2]
        inp2 = dwords[idx - 1]
        a1 = (inp1 + s2_r) & M
        v9 = (a1 * s2_c0 - 0x30674EEF * (((s2_k0 * a1) & M) >> 16)) & M
        v10 = v9 >> 16
        inner = (0x5B9F0000 * v9 - 0x78F7A461 * v10) & M
        v11 = (0xE9B30000 * v10 + 0x12CEB96D * (inner >> 16)) & M
        v12 = (0x1D830000 * v11 + 0x257E1D83 * (v11 >> 16)) & M
        a2 = (v12 + inp2) & M
        x_val = (a2 * s2_c1 - 0x5D8BE90B * (((s2_k1 * a2) & M) >> 16)) & M
        v13 = x_val >> 16
        v14 = ((0x96FF0000 * x_val - 0x2C7C6901 * v13) & M) >> 16
        inner2 = (0x7C932B89 * v14 - 0x5C890000 * v13) & M
        s2_r = (0xF2310000 * v14 - 0x405B6097 * (inner2 >> 16)) & M
        s2_acc = (s2_acc + s2_r + v12) & M
    hash2 = struct.pack("<II", s2_r, s2_acc)

    # === XOR the two passes and Base64-encode ==============================
    final = bytes(a ^ b for a, b in zip(hash1, hash2))
    return base64.b64encode(final).decode("ascii")


def _delete_user_choice_key() -> bool:
    """Take ownership of the protected UserChoice key and delete it.

    The UserChoice key has special ACLs — we must take ownership before
    we can remove it.  This is the *only* step that needs PowerShell.
    """
    ps_script = r"""
$ErrorActionPreference = 'Stop'
$subKey = 'Software\Microsoft\Windows\CurrentVersion\Explorer\FileExts\.pdf\UserChoice'
try {
    $reg = [Microsoft.Win32.Registry]::CurrentUser.OpenSubKey(
        $subKey,
        [Microsoft.Win32.RegistryKeyPermissionCheck]::ReadWriteSubTree,
        [System.Security.AccessControl.RegistryRights]::TakeOwnership)
    if (-not $reg) { Write-Output 'OK'; exit 0 }

    $acl = $reg.GetAccessControl(
        [System.Security.AccessControl.AccessControlSections]::None)
    $me  = [System.Security.Principal.WindowsIdentity]::GetCurrent().User
    $acl.SetOwner($me)
    $reg.SetAccessControl($acl)

    $acl  = $reg.GetAccessControl()
    $rule = [System.Security.AccessControl.RegistryAccessRule]::new(
        $me, 'FullControl',
        [System.Security.AccessControl.InheritanceFlags]::None,
        [System.Security.AccessControl.PropagationFlags]::None,
        'Allow')
    $acl.SetAccessRule($rule)
    $reg.SetAccessControl($acl)
    $reg.Close()

    [Microsoft.Win32.Registry]::CurrentUser.DeleteSubKeyTree($subKey, $false)
    Write-Output 'OK'
} catch {
    Write-Output "FAIL:$($_.Exception.Message)"
}
"""
    try:
        # Save script to a temp file to avoid quoting issues
        with tempfile.NamedTemporaryFile(
                mode="w", suffix=".ps1", delete=False) as f:
            f.write(ps_script)
            ps_path = f.name
        try:
            r = subprocess.run(
                ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass",
                 "-File", ps_path],
                capture_output=True, text=True, timeout=15,
                creationflags=(subprocess.CREATE_NO_WINDOW
                               if sys.platform == "win32" else 0),
            )
            out = r.stdout.strip()
            if out == "OK":
                return True
            print(f"[BoltPDF] UserChoice delete: {out}", file=sys.stderr)
            if r.stderr.strip():
                print(f"[BoltPDF] PS stderr: {r.stderr.strip()}",
                      file=sys.stderr)
            return False
        finally:
            try:
                os.unlink(ps_path)
            except OSError:
                pass
    except Exception as e:
        print(f"[BoltPDF] UserChoice delete error: {e}", file=sys.stderr)
        return False


def _write_user_choice() -> bool:
    """Delete the existing UserChoice key and recreate it with the correct
    ProgId *and* computed Hash so Windows 10/11 accepts it.

    Flow:
      1. Delete the old key (needs PowerShell for ACL ownership trick).
      2. Create the new empty key, then query its actual LastWriteTime.
      3. Compute hash using that real timestamp (avoids timing mismatches).
      4. Write ProgId + Hash via winreg.
    """
    if sys.platform != "win32":
        return False

    import winreg
    import ctypes

    # Step 1 — delete old key
    if not _delete_user_choice_key():
        print("[BoltPDF] Could not delete old UserChoice key",
              file=sys.stderr)
        return False

    # Step 2 — create empty key, then read its actual timestamp
    key_path = (r"Software\Microsoft\Windows\CurrentVersion"
                r"\Explorer\FileExts\.pdf\UserChoice")
    try:
        hkey = winreg.CreateKey(winreg.HKEY_CURRENT_USER, key_path)
    except Exception as e:
        print(f"[BoltPDF] UserChoice create failed: {e}", file=sys.stderr)
        return False

    # Query the key's real LastWriteTime via RegQueryInfoKeyW
    sid = _get_user_sid()
    if not sid:
        winreg.CloseKey(hkey)
        return False

    filetime = ctypes.c_int64(0)
    try:
        ctypes.windll.advapi32.RegQueryInfoKeyW(
            int(hkey), None, None, None, None, None,
            None, None, None, None, None,
            ctypes.byref(filetime))
    except Exception:
        pass
    ft_val = filetime.value if filetime.value else _get_current_filetime()

    # Step 3 — compute hash using the key's real timestamp
    hash_val = _compute_user_choice_hash(".pdf", sid, _PROG_ID, ft_val)

    # Step 4 — write both values
    try:
        winreg.SetValueEx(hkey, "ProgId", 0, winreg.REG_SZ, _PROG_ID)
        winreg.SetValueEx(hkey, "Hash", 0, winreg.REG_SZ, hash_val)
        winreg.CloseKey(hkey)
        return True
    except Exception as e:
        winreg.CloseKey(hkey)
        print(f"[BoltPDF] UserChoice write failed: {e}", file=sys.stderr)
        return False


def _build_setfta_script(result_file: str) -> str:
    """Return the content of the all-in-one elevated PowerShell script.

    The script:
      1. Disables UCPD if present.
      2. Takes ownership of the old UserChoice key and deletes it.
      3. Compiles the tools_setfta C# hash code via Add-Type.
      4. Creates the new key, reads its real timestamp, computes the
         hash, and writes ProgId + Hash.
      5. Re-enables UCPD.
      6. Writes the result to *result_file* so the caller can check.
    """
    # Escape backslashes in the result path for PowerShell
    rp = result_file.replace("'", "''")
    return r"""
$ErrorActionPreference = 'Stop'

$cs = @'
using System;
using System.Text;
using System.Runtime.InteropServices;
using Microsoft.Win32;

public class UserChoiceHash {
    public static string ComputeAndWrite(string extension, string progid) {
        string regpath = @"SOFTWARE\Microsoft\Windows\CurrentVersion\Explorer\FileExts\" + extension + @"\UserChoice";
        try { Registry.CurrentUser.DeleteSubKey(regpath, false); } catch {}
        RegistryKey regnode = Registry.CurrentUser.CreateSubKey(regpath);

        System.Security.Principal.WindowsIdentity user = System.Security.Principal.WindowsIdentity.GetCurrent();
        string sid = user.User.Value;

        long ftLastWriteTime;
        RegQueryInfoKey(regnode.Handle.DangerousGetHandle(), IntPtr.Zero, IntPtr.Zero, IntPtr.Zero,
            IntPtr.Zero, IntPtr.Zero, IntPtr.Zero, IntPtr.Zero, IntPtr.Zero, IntPtr.Zero, IntPtr.Zero,
            out ftLastWriteTime);

        DateTime time = DateTime.FromFileTime(ftLastWriteTime);
        time = time.AddTicks(-(time.Ticks % 600000000));
        ftLastWriteTime = time.ToFileTime();

        string regdate = ftLastWriteTime.ToString("x16");
        string experience = "user choice set via windows user experience {d18b6dd5-6124-4341-9318-804003bafa0b}";
        byte[] bytes = Encoding.Unicode.GetBytes((extension + sid + progid + regdate + experience + "\0").ToLower());

        System.Security.Cryptography.MD5 md5Hash = System.Security.Cryptography.MD5.Create();
        byte[] md5 = md5Hash.ComputeHash(bytes);

        byte[] mshash1 = sub_1(bytes, md5);
        byte[] mshash2 = sub_2(bytes, md5);
        byte[] finalraw = new byte[8];
        for (int i = 0; i < 8; i++) finalraw[i] = (byte)(mshash1[i] ^ mshash2[i]);

        string hash = System.Convert.ToBase64String(finalraw);

        regnode.SetValue("ProgId", progid);
        regnode.SetValue("Hash", hash);
        regnode.Close();
        return "OK:" + hash;
    }

    static byte[] sub_1(byte[] data, byte[] md5) {
        byte[] retval = new byte[8];
        uint length = (uint)(((((data.Length) >> 2) & 1) < 1 ? 1 : 0) + ((data.Length) >> 2) - 1);
        uint[] dword_data = new uint[length];
        uint[] dword_md5 = new uint[4];
        for (int i = 0; i < dword_data.Length; i++) dword_data[i] = BitConverter.ToUInt32(data, i * 4);
        for (int i = 0; i < 4; i++) dword_md5[i] = BitConverter.ToUInt32(md5, i * 4);
        if (length <= 1 || (length & 1) == 1) return retval;
        uint v5 = 0, v6 = 0;
        uint v7 = (length - 2) >> 1;
        uint v18 = v7++;
        uint v8 = v7;
        uint result = 0;
        uint v9 = (dword_md5[1] | 1) + 0x13DB0000u;
        uint v10 = (dword_md5[0] | 1) + 0x69FB0000u;
        do {
            uint v11 = dword_data[v6] + result; v6 += 2;
            uint v12 = 0x79F8A395u * (v10 * v11 - 0x10FA9605u * (v11 >> 16)) + 0x689B6B9Fu * ((v10 * v11 - 0x10FA9605u * (v11 >> 16)) >> 16);
            uint v13 = 0xEA970001u * v12 - 0x3C101569u * (v12 >> 16);
            uint v14 = v13 + v5;
            uint v15 = v9 * (dword_data[v6 - 1] + v13) - 0x3CE8EC25u * ((dword_data[v6 - 1] + v13) >> 16);
            result = 0x1EC90001u * (0x59C3AF2Du * v15 - 0x2232E0F1u * (v15 >> 16)) + 0x35BD1EC9u * ((0x59C3AF2Du * v15 - 0x2232E0F1u * (v15 >> 16)) >> 16);
            v5 = result + v14;
            --v8;
        } while (v8 != 0);
        BitConverter.GetBytes(result).CopyTo(retval, 0);
        BitConverter.GetBytes(v5).CopyTo(retval, 4);
        return retval;
    }

    static byte[] sub_2(byte[] data, byte[] md5) {
        byte[] retval = new byte[8];
        uint length = (uint)(((((data.Length) >> 2) & 1) < 1 ? 1 : 0) + ((data.Length) >> 2) - 1);
        uint[] dword_data = new uint[length];
        uint[] dword_md5 = new uint[4];
        for (int i = 0; i < dword_data.Length; i++) dword_data[i] = BitConverter.ToUInt32(data, i * 4);
        for (int i = 0; i < 4; i++) dword_md5[i] = BitConverter.ToUInt32(md5, i * 4);
        if (length <= 1 || (length & 1) == 1) return retval;
        uint v5 = 0, v6 = 0, v7 = 0;
        uint v25 = (length - 2) >> 1;
        uint v21 = dword_md5[0] | 1;
        uint v22 = dword_md5[1] | 1;
        uint v23 = 0xB1110000u * v21;
        uint v24 = 0x16F50000u * v22;
        uint v8 = v25 + 1;
        do {
            v6 += 2;
            uint v9 = (dword_data[v6 - 2] + v5) * v23 - 0x30674EEFu * (v21 * (dword_data[v6 - 2] + v5) >> 16);
            uint v10 = v9 >> 16;
            uint v11 = 0xE9B30000u * v10 + 0x12CEB96Du * ((0x5B9F0000u * v9 - 0x78F7A461u * v10) >> 16);
            uint v12 = 0x1D830000u * v11 + 0x257E1D83u * (v11 >> 16);
            uint v13 = ((v12 + dword_data[v6 - 1]) * v24 - 0x5D8BE90Bu * ((v22 * (v12 + dword_data[v6 - 1])) >> 16)) >> 16;
            uint v14 = 0x96FF0000u * ((v12 + dword_data[v6 - 1]) * v24 - 0x5D8BE90Bu * ((v22 * (v12 + dword_data[v6 - 1])) >> 16)) - 0x2C7C6901u * v13 >> 16;
            v5 = 0xF2310000u * v14 - 0x405B6097u * ((0x7C932B89u * v14 - 0x5C890000u * v13) >> 16);
            v7 += v5 + v12;
            --v8;
        } while (v8 != 0);
        BitConverter.GetBytes(v5).CopyTo(retval, 0);
        BitConverter.GetBytes(v7).CopyTo(retval, 4);
        return retval;
    }

    [DllImport("advapi32.dll", EntryPoint = "RegQueryInfoKey", CallingConvention = CallingConvention.Winapi, SetLastError = true)]
    extern static int RegQueryInfoKey(IntPtr handle, IntPtr a, IntPtr b, IntPtr c, IntPtr d, IntPtr e, IntPtr f, IntPtr g, IntPtr h, IntPtr i, IntPtr j, out long lpftLastWriteTime);
}
'@

try {
    # --- Disable UCPD if present -------------------------------------------
    $ucpdDisabled = $false
    try {
        $svc = Get-Service -Name 'UCPD' -ErrorAction SilentlyContinue
        if ($svc) {
            Set-Service -Name 'UCPD' -StartupType Disabled -ErrorAction SilentlyContinue
            Stop-Service -Name 'UCPD' -Force -ErrorAction SilentlyContinue
            Start-Sleep -Seconds 2
            $ucpdDisabled = $true
        }
    } catch {}
    try {
        Disable-ScheduledTask -TaskName '\Microsoft\Windows\AppxDeploymentClient\UCPD velocity' `
            -ErrorAction SilentlyContinue | Out-Null
    } catch {}

    # --- Take ownership of old key and delete it ---------------------------
    $subKey = 'Software\Microsoft\Windows\CurrentVersion\Explorer\FileExts\.pdf\UserChoice'
    try {
        $reg = [Microsoft.Win32.Registry]::CurrentUser.OpenSubKey(
            $subKey,
            [Microsoft.Win32.RegistryKeyPermissionCheck]::ReadWriteSubTree,
            [System.Security.AccessControl.RegistryRights]::TakeOwnership)
        if ($reg) {
            $acl = $reg.GetAccessControl(
                [System.Security.AccessControl.AccessControlSections]::None)
            $me  = [System.Security.Principal.WindowsIdentity]::GetCurrent().User
            $acl.SetOwner($me)
            $reg.SetAccessControl($acl)
            $acl  = $reg.GetAccessControl()
            $rule = [System.Security.AccessControl.RegistryAccessRule]::new(
                $me, 'FullControl',
                [System.Security.AccessControl.InheritanceFlags]::None,
                [System.Security.AccessControl.PropagationFlags]::None,
                'Allow')
            $acl.SetAccessRule($rule)
            $reg.SetAccessControl($acl)
            $reg.Close()
            [Microsoft.Win32.Registry]::CurrentUser.DeleteSubKeyTree($subKey, $false)
        }
    } catch {}

    # --- Compile and run the C# hash computation --------------------------
    Add-Type -TypeDefinition $cs -Language CSharp
    $result = [UserChoiceHash]::ComputeAndWrite('.pdf', '""" + _PROG_ID + r"""')
    $result | Out-File -Encoding ascii '""" + rp + r"""'

    # --- Re-enable UCPD (will take effect on next boot) --------------------
    if ($ucpdDisabled) {
        try {
            Set-Service -Name 'UCPD' -StartupType Automatic -ErrorAction SilentlyContinue
        } catch {}
        try {
            Enable-ScheduledTask -TaskName '\Microsoft\Windows\AppxDeploymentClient\UCPD velocity' `
                -ErrorAction SilentlyContinue | Out-Null
        } catch {}
    }
} catch {
    "FAIL:$($_.Exception.Message)" | Out-File -Encoding ascii '""" + rp + r"""'
}
"""


def _run_elevated(ps_path: str) -> bool:
    """Launch a PowerShell script elevated (UAC) using ShellExecuteExW
    and wait for it to finish.  Returns True if the process ran."""
    import ctypes
    import ctypes.wintypes

    class SHELLEXECUTEINFOW(ctypes.Structure):
        _fields_ = [
            ("cbSize", ctypes.wintypes.DWORD),
            ("fMask", ctypes.c_ulong),
            ("hwnd", ctypes.wintypes.HWND),
            ("lpVerb", ctypes.c_wchar_p),
            ("lpFile", ctypes.c_wchar_p),
            ("lpParameters", ctypes.c_wchar_p),
            ("lpDirectory", ctypes.c_wchar_p),
            ("nShow", ctypes.c_int),
            ("hInstApp", ctypes.wintypes.HINSTANCE),
            ("lpIDList", ctypes.c_void_p),
            ("lpClass", ctypes.c_wchar_p),
            ("hkeyClass", ctypes.wintypes.HKEY),
            ("dwHotKey", ctypes.wintypes.DWORD),
            ("hIconOrMonitor", ctypes.wintypes.HANDLE),
            ("hProcess", ctypes.wintypes.HANDLE),
        ]

    SEE_MASK_NOCLOSEPROCESS = 0x00000040
    SW_HIDE = 0
    INFINITE = 0xFFFFFFFF

    sei = SHELLEXECUTEINFOW()
    sei.cbSize = ctypes.sizeof(sei)
    sei.fMask = SEE_MASK_NOCLOSEPROCESS
    sei.lpVerb = "runas"
    sei.lpFile = "powershell.exe"
    sei.lpParameters = (
        f'-NoProfile -ExecutionPolicy Bypass -WindowStyle Hidden '
        f'-File "{ps_path}"'
    )
    sei.nShow = SW_HIDE

    ok = ctypes.windll.shell32.ShellExecuteExW(ctypes.byref(sei))
    if not ok:
        print("[BoltPDF] ShellExecuteExW failed", file=sys.stderr)
        return False

    if sei.hProcess:
        ctypes.windll.kernel32.WaitForSingleObject(sei.hProcess, 60000)
        ctypes.windll.kernel32.CloseHandle(sei.hProcess)
    return True


def _write_user_choice_elevated() -> bool:
    """Set .pdf UserChoice via an elevated PowerShell process that embeds
    the C# hash computation so timing is always correct.

    Uses ShellExecuteExW (no nested PowerShell quoting issues) and
    communicates the result via a temp file.
    """
    result_path = os.path.join(tempfile.gettempdir(), "boltpdf_setfta.txt")
    ps_content = _build_setfta_script(result_path)

    ps_path = os.path.join(tempfile.gettempdir(), "boltpdf_setfta.ps1")
    try:
        with open(ps_path, "w", encoding="utf-8") as f:
            f.write(ps_content)
    except Exception as e:
        print(f"[BoltPDF] Could not write PS script: {e}", file=sys.stderr)
        return False

    try:
        # Clean any stale result
        try:
            os.unlink(result_path)
        except OSError:
            pass

        ok = _run_elevated(ps_path)
        if not ok:
            return False

        # Read the result file the elevated script wrote
        import time
        time.sleep(0.5)
        try:
            with open(result_path, "r", encoding="utf-8-sig") as rf:
                out = rf.read().strip()
            print(f"[BoltPDF] Elevated result: {out}", file=sys.stderr)
            if out.startswith("OK:"):
                return True
        except FileNotFoundError:
            print("[BoltPDF] No result from elevated process",
                  file=sys.stderr)

        # Fallback: just check registry
        return _is_default_pdf_reader() is True
    finally:
        for p in (ps_path, result_path):
            try:
                os.unlink(p)
            except OSError:
                pass


def _notify_shell():
    """Tell Explorer that file associations have changed."""
    try:
        import ctypes
        SHCNE_ASSOCCHANGED = 0x08000000
        SHCNF_IDLIST = 0x0000
        ctypes.windll.shell32.SHChangeNotify(
            SHCNE_ASSOCCHANGED, SHCNF_IDLIST, None, None)
    except Exception:
        pass


def _set_default_pdf_reader():
    """Set BoltPDF as the default PDF reader.

    Strategy (each step is tried only if the previous one failed):
      1. Non-elevated: delete UserChoice, write correct hash, set ProgId.
         Works on systems without UCPD or older Windows 10.
      2. Elevated: launch PowerShell as admin via ShellExecuteExW to
         disable UCPD, compute the hash in C# (with correct key
         timestamp), and write the key.  Shows a one-time UAC prompt.
    Returns True if successful."""
    _register_prog_id()

    # --- Attempt 1: non-elevated ------------------------------------------
    ok = _write_user_choice()
    _notify_shell()
    if ok and _is_default_pdf_reader() is True:
        return True

    # --- Attempt 2: elevated (handles UCPD) --------------------------------
    ok = _write_user_choice_elevated()
    _notify_shell()
    return ok


def _prompt_set_default(parent=None):
    """Check if BoltPDF is the default PDF reader; if not, offer to set it.
    If the user accepts, apply the change fully — no further input needed."""
    status = _is_default_pdf_reader()
    if status is True or status is None:
        return  # already default or can't determine

    reply = QMessageBox.question(
        parent,
        "Default PDF Reader",
        "BoltPDF is not your default PDF reader.\n\n"
        "Would you like to set it as the default?",
        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
    )
    if reply != QMessageBox.StandardButton.Yes:
        return

    success = _set_default_pdf_reader()

    if success:
        QMessageBox.information(
            parent,
            "Default PDF Reader",
            "BoltPDF is now your default PDF reader.",
        )
    else:
        # Both approaches failed — open Settings as last resort
        _automate_settings_fallback(parent)


def _automate_settings_fallback(parent=None):
    """Last resort: open the Windows Default Apps settings page.
    The user will need to select BoltPDF manually."""
    try:
        subprocess.Popen(
            ["cmd", "/c", "start",
             "ms-settings:defaultapps"],
            creationflags=(subprocess.CREATE_NO_WINDOW
                           if sys.platform == "win32" else 0),
        )
    except Exception:
        try:
            os.startfile("ms-settings:defaultapps")
        except Exception:
            pass

    QMessageBox.information(
        parent,
        "Almost Done",
        "Please select BoltPDF in the Settings window\n"
        "that has opened to finish the setup.",
    )


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def main():
    # Handle --uninstall before creating a GUI
    if "--uninstall" in sys.argv:
        _uninstall_app()  # exits

    # Handle --install (elevated re-launch for first-run install)
    if "--install" in sys.argv and getattr(sys, 'frozen', False):
        installed = _install_app()
        if installed and _is_installed():
            _register_prog_id()
            # Relaunch from installed location (non-elevated, normal user)
            args = [_get_installed_exe()] + [
                a for a in sys.argv[1:]
                if a.lower().endswith(".pdf")]
            try:
                subprocess.Popen(args)
            except Exception:
                pass
            sys.exit(0)
        sys.exit(1)

    # Must be set BEFORE QApplication is created so QWebEngineView works later
    from PyQt6.QtCore import Qt as _Qt
    QApplication.setAttribute(_Qt.ApplicationAttribute.AA_ShareOpenGLContexts)

    app = QApplication(sys.argv)
    app.setApplicationName("BoltPDF")

    # Set taskbar icon (Windows needs AppUserModelID for correct grouping)
    if sys.platform == "win32":
        try:
            import ctypes
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(
                "BoltPDF.PDFReader.1")
        except Exception:
            pass

    # --- First-run install (frozen exe on Windows) --------------------------
    # If we're a frozen exe and NOT already running from the install dir,
    # install ourselves, then relaunch from the installed location so all
    # registry paths are correct from the start.
    if (getattr(sys, 'frozen', False)
            and sys.platform == "win32"
            and not _is_running_from_install_dir()):
        installed = _install_app()
        if installed and _is_installed():
            # Register ProgId pointing to the installed exe
            _register_prog_id()
            # Relaunch from installed location, passing any PDF args through
            args = [_get_installed_exe()] + [
                a for a in sys.argv[1:] if a.lower().endswith(".pdf")]
            try:
                subprocess.Popen(args)
            except Exception:
                pass
            sys.exit(0)

    window = BoltPDFReader()
    window.setAcceptDrops(True)
    window.show()

    # Distinguish a clean quit from a crash: read last run's flag, then
    # immediately arm "not clean" for THIS run (closeEvent flips it back
    # to clean on a normal quit).  Default True so a first-ever run is
    # treated as clean.
    _was_clean = True
    try:
        _was_clean = bool(app_state().pref("clean_exit", True))
        app_state().set_pref("clean_exit", False)
    except Exception:
        pass

    # Open any PDFs passed on the command line
    opened_from_args = False
    for arg in sys.argv[1:]:
        if arg.lower().endswith(".pdf"):
            window.open_pdf_in_new_tab(arg)
            opened_from_args = True

    # A plain launch opens with NO document.  Only auto-reopen a PDF if
    # the previous run ended abnormally (crash) with unsaved edits, so
    # the user can recover that work.  (Deferred so the window paints
    # first.)
    if not opened_from_args and not _was_clean:
        QTimer.singleShot(0, window.reopen_crashed_docs)

    # Prompt for default reader (only for installed frozen exe).
    # Deferred so the window appears first, then the dialog pops up
    # over a fully-rendered UI instead of blocking startup.
    if (getattr(sys, 'frozen', False)
            and sys.platform == "win32"
            and _is_running_from_install_dir()):
        QTimer.singleShot(500, lambda: _prompt_set_default(window))

    exit_code = app.exec()

    # Final sweep: even after closeEvent has run, hard-kill any
    # multiprocessing child that is somehow still alive so the exe
    # process tree terminates cleanly.  Without this, stray PageRenderer
    # or Pool workers can keep BoltPDF.exe entries in Task Manager
    # after the main window has closed.
    try:
        import multiprocessing as _mp
        for p in _mp.active_children():
            try:
                p.terminate()
            except Exception:
                pass
        for p in _mp.active_children():
            try:
                if p.is_alive():
                    if hasattr(p, "kill"):
                        p.kill()
                    else:
                        p.terminate()
                p.join(timeout=0.2)
            except Exception:
                pass
    except Exception:
        pass

    sys.exit(exit_code)


if __name__ == "__main__":
    multiprocessing.freeze_support()   # required for PyInstaller on Windows
    main()
